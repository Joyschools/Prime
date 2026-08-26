from __future__ import annotations
import re
import hashlib
import secrets

import csv
import json
import io
import base64
import os
import sqlite3
import uuid
import threading
import urllib.error
import urllib.parse
import urllib.request
import mimetypes
import smtplib
import ssl
from email.message import EmailMessage
from collections.abc import MutableMapping
from datetime import datetime, timedelta
from functools import wraps
from pathlib import Path
from typing import Any, Callable, Iterable

from flask import (
    Flask,
    abort,
    flash as flask_flash,
    g,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session as flask_session,
    url_for,
)
from werkzeug.security import check_password_hash, generate_password_hash
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import qrcode
from PIL import Image
from pypdf import PdfReader, PdfWriter

BASE_DIR = Path(__file__).resolve().parent
INSTANCE_DIR = BASE_DIR / "instance"
_render_data_dir = Path(os.environ.get("PERSISTENT_DATA_DIR", "/var/data"))
if os.environ.get("DATA_DIR"):
    DATA_DIR = Path(os.environ["DATA_DIR"]).expanduser()
elif os.environ.get("RENDER") and _render_data_dir.exists():
    DATA_DIR = _render_data_dir
else:
    DATA_DIR = INSTANCE_DIR
DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / "school.db"
SECRET_FILE = DATA_DIR / "secret.key"
UPLOAD_DIR = DATA_DIR / "uploads"
PERSISTENT_STORAGE = DATA_DIR != INSTANCE_DIR
PULSE_PEER_URL = os.environ.get("PULSE_PEER_URL", "https://breathe-xozy.onrender.com").strip().rstrip("/")
PULSE_TIMEOUT_SECONDS = max(1, min(10, int(os.environ.get("PULSE_TIMEOUT_SECONDS", "4"))))
PULSE_ALLOWED_CALLBACK_HOSTS = {h.strip().lower() for h in os.environ.get("PULSE_ALLOWED_CALLBACK_HOSTS", "").split(",") if h.strip()}
ALLOWED_RESTORE_EXT = {"db", "sqlite", "sqlite3"}
PUBLIC_ROLES = ("Teacher", "Student", "Parent")
HIDDEN_ROLES = ("Admin", "ICT", "Finance", "Librarian")
QR_LOGIN_ROLES = {"Admin", "ICT", "Finance", "Teacher", "Librarian"}
QR_LOGIN_WORKSPACES = {"Teaching", "Driver", "Reception", "Guard", "Cook", "Other Staff"}
RECEPTION_WORKSPACE = "Reception"
ALL_PORTAL_ROLES = HIDDEN_ROLES + PUBLIC_ROLES
ADMIN_LOGIN_PATH = "/xtspolsjhulupjoppsup-lmkzcodup"
ADMIN_ROLES = {"Admin"}
ALL_ROLES = ALL_PORTAL_ROLES
SYSTEM_ROLE = "System"
INSTITUTION_TYPES = ("Kindergarten", "Primary School", "High School", "Secondary School", "TVET", "College", "University", "Mixed Institution")
DEFAULT_DEPARTMENTS = ("Communications", "Computer Studies")
_PORTAL_ROLE_COOKIE = "school_portal_role"
_AUTH_COOKIE = "school_auth_token"
_AUTH_COOKIE_MAX_AGE = 60 * 60 * 24 * 30
_AUTH_TICKET_COOKIE = "school_auth_ticket"
_PORTAL_CONTEXT_MAX_AGE = 60 * 60 * 12
_PORTAL_CONTEXT_SALT = "school-portal-context-v1"

app = Flask(__name__, instance_path=str(INSTANCE_DIR), instance_relative_config=True, static_folder=str(BASE_DIR / "static"), static_url_path="/static")
app.config.update(
    MAX_CONTENT_LENGTH=20 * 1024 * 1024,
    TEMPLATES_AUTO_RELOAD=True,
)

# -------------------------------------------------------------------
# Signed authentication session
# -------------------------------------------------------------------
def _load_or_create_secret() -> str:
    if os.environ.get("SECRET_KEY"):
        return os.environ["SECRET_KEY"]
    if SECRET_FILE.exists():
        value = SECRET_FILE.read_text(encoding="utf-8").strip()
        if value:
            return value
    legacy_secret = INSTANCE_DIR / "secret.key"
    if SECRET_FILE != legacy_secret and legacy_secret.exists():
        value = legacy_secret.read_text(encoding="utf-8").strip()
        if value:
            SECRET_FILE.write_text(value, encoding="utf-8")
            return value
    value = uuid.uuid4().hex + uuid.uuid4().hex
    try:
        fd = os.open(str(SECRET_FILE), os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o600)
        with os.fdopen(fd, "w", encoding="utf-8") as fh:
            fh.write(value)
        return value
    except FileExistsError:
        # Another Gunicorn worker created the shared secret at the same time.
        return SECRET_FILE.read_text(encoding="utf-8").strip()
    except OSError:
        SECRET_FILE.write_text(value, encoding="utf-8")
        return value

_secret_key = _load_or_create_secret()
app.config.update(
    SECRET_KEY=_secret_key,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=bool(os.environ.get("COOKIE_SECURE") == "1" or os.environ.get("RENDER")),
    SESSION_COOKIE_NAME="school_portal_session",
    SESSION_COOKIE_PATH="/",
    SESSION_REFRESH_EACH_REQUEST=True,
    PERMANENT_SESSION_LIFETIME=60 * 60 * 24 * 30,
)
session = flask_session

def flash(message: str, category: str = "message") -> None:
    flask_flash(message, category)


UPLOAD_DIR.mkdir(exist_ok=True)
DOC_DIR = UPLOAD_DIR / "documents"
DOC_DIR.mkdir(exist_ok=True)
QR_DIR = UPLOAD_DIR / "qr"
QR_DIR.mkdir(exist_ok=True)


# -------------------------
# Database helpers
# -------------------------
def _sqlite_integrity_ok(path: Path) -> bool:
    """Return True only when SQLite can open the database and verify it."""
    if not path.exists():
        return False
    try:
        with sqlite3.connect(path, timeout=10) as check_conn:
            result = check_conn.execute("PRAGMA integrity_check").fetchone()
            return bool(result and str(result[0]).strip().lower() == "ok")
    except (sqlite3.DatabaseError, OSError):
        return False


def _quarantine_corrupt_db(path: Path) -> None:
    """Preserve a damaged SQLite file without allowing it to stop startup."""
    if not path.exists():
        return
    stamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    target = path.with_name(f"{path.name}.corrupt-{stamp}")
    try:
        path.replace(target)
    except OSError:
        try:
            path.unlink(missing_ok=True)
        except OSError:
            pass
    # WAL/SHM sidecars belong to the damaged database and can poison a fresh copy.
    for suffix in ("-wal", "-shm"):
        try:
            path.with_name(path.name + suffix).unlink(missing_ok=True)
        except OSError:
            pass


def get_db() -> sqlite3.Connection:
    if "db" not in g:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        conn.execute("""CREATE TABLE IF NOT EXISTS auth_sessions (
            token_hash TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            expires_at TEXT NOT NULL,
            revoked INTEGER NOT NULL DEFAULT 0,
            last_seen_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_auth_sessions_user ON auth_sessions(user_id, revoked, expires_at)")
        conn.commit()
        g.db = conn
    return g.db


@app.teardown_appcontext
def close_db(_: Any) -> None:
    db = g.pop("db", None)
    if db is not None:
        db.close()


def migrate_legacy_student_store() -> None:
    """One-time migration from the old standalone students.db into school.db.

    The portal now has one authoritative database. If an older deployment still
    has students.db, merge any learners not already present by admission number,
    then remove the old file only after the school database transaction succeeds.
    """
    legacy = DATA_DIR / "students.db"
    if not legacy.exists():
        return
    try:
        if not _sqlite_integrity_ok(legacy):
            app.logger.warning("Ignoring invalid legacy students.db; school.db remains authoritative.")
            return
        with sqlite3.connect(legacy, timeout=30) as src:
            src.row_factory = sqlite3.Row
            rows = src.execute("SELECT * FROM students ORDER BY id").fetchall()
        db = get_db()
        db.execute("BEGIN")
        try:
            main_cols = table_columns(db, "students")
            for row in rows:
                data = {k: row[k] for k in row.keys()}
                admission = str(data.get("admission_no") or "").strip()
                name = str(data.get("full_name") or "").strip()
                grade = str(data.get("grade") or "").strip()
                if not admission or not name or not grade:
                    continue
                existing = db.execute("SELECT id FROM students WHERE admission_no=?", (admission,)).fetchone()
                if existing:
                    continue
                values = {
                    "admission_no": admission, "full_name": name, "grade": grade,
                    "guardian_name": data.get("guardian_name") or "",
                    "guardian_phone": data.get("guardian_phone") or "",
                    "guardian_email": data.get("guardian_email") or "",
                    "alt_guardian_name": data.get("alt_guardian_name") or "",
                    "alt_guardian_phone": data.get("alt_guardian_phone") or "",
                    "alt_guardian_email": data.get("alt_guardian_email") or "",
                    "student_phone": data.get("student_phone") or "", "student_email": data.get("student_email") or "",
                    "medical_condition": data.get("medical_condition") or "", "allergies": data.get("allergies") or "",
                    "special_info": data.get("special_info") or "", "notes": data.get("notes") or "",
                    "payment_status": data.get("payment_status") or "Pending", "balance": data.get("balance") or 0,
                    "active": 1 if data.get("active", 1) else 0, "age": data.get("age") or "",
                }
                cols=[c for c in values if c in main_cols]
                db.execute(f"INSERT INTO students({','.join(cols)}) VALUES({','.join('?' for _ in cols)})", tuple(values[c] for c in cols))
            db.commit()
        except Exception:
            db.rollback()
            raise
        for suffix in ("", "-wal", "-shm"):
            try:
                legacy.with_name(legacy.name + suffix).unlink(missing_ok=True)
            except OSError:
                pass
        app.logger.info("Legacy students.db migrated and removed; school.db is now the only learner database.")
    except Exception:
        app.logger.exception("Legacy student database migration failed; leaving students.db untouched for safety.")


def normalize_grade(value: str) -> str:
    value = re.sub(r"\s+", " ", (value or "").strip())
    if not value:
        return ""
    m = re.fullmatch(r"(?:grade\s*)?(\d+)", value, re.I)
    if m:
        return f"Grade {int(m.group(1))}"
    return value.title() if value.islower() else value


def next_admission_no() -> str:
    """Return the next collision-free admission number from school_settings."""
    settings = school_settings()
    prefix = str(settings["admission_prefix"] or "ADM-")
    suffix = str(settings["admission_suffix"] or "")
    rows = q("SELECT admission_no FROM students WHERE admission_no IS NOT NULL", one=False)
    pattern = re.compile(r"^" + re.escape(prefix) + r"(\d+)" + re.escape(suffix) + r"$", re.I)
    highest = 0
    for row in rows:
        value = str(row["admission_no"] or "").strip()
        match = pattern.match(value)
        if match:
            highest = max(highest, int(match.group(1)))
    candidate = highest + 1
    while True:
        admission = f"{prefix}{candidate:03d}{suffix}"
        if not q("SELECT id FROM students WHERE admission_no=?", (admission,), one=True):
            return admission
        candidate += 1


def auto_place_new_student(student_id: int, grade: str, actor_id: int) -> dict:
    """Enroll configured subjects for a grade and attach class/subject teachers.

    Matching is deliberately tolerant of legacy labels such as ``8`` / ``Grade 8``
    so old school configurations cannot strand a newly registered learner.
    """
    summary = {"class_teacher": None, "subjects": 0, "subject_teachers": 0}
    class_name = normalize_grade(grade)
    raw_grade = str(grade or "").strip()
    grade_variants = [class_name]
    if class_name.lower().startswith("grade "):
        grade_variants.append(class_name[6:].strip())
    if raw_grade and raw_grade not in grade_variants:
        grade_variants.append(raw_grade)

    class_teacher = None
    for variant in grade_variants:
        class_teacher = q("SELECT teacher_user_id,class_name FROM class_teacher_assignments WHERE lower(trim(class_name))=lower(trim(?)) ORDER BY id LIMIT 1", (variant,), one=True)
        if class_teacher:
            break
    if class_teacher:
        teacher = q("SELECT id,full_name FROM users WHERE id=? AND role='Teacher' AND active=1", (class_teacher['teacher_user_id'],), one=True)
        if teacher:
            execute("INSERT OR IGNORE INTO student_teacher_assignments(student_id,teacher_user_id,class_name,subject,scope,assigned_by,active) VALUES(?,?,?,'General','Class Teacher',?,1)", (student_id,teacher['id'],class_name,actor_id))
            summary['class_teacher'] = teacher['full_name']
    # Build the subject set from both compulsory-subject configuration and explicit
    # teacher assignments. Resolve against every grade variant so older records work.
    subjects = set()
    for variant in grade_variants:
        for row in q("SELECT subject FROM compulsory_subjects WHERE active=1 AND lower(trim(class_name))=lower(trim(?))", (variant,)):
            subject = str(row["subject"] or "").strip()
            if subject:
                subjects.add(subject)
        for row in q("SELECT subject FROM teacher_assignments WHERE active=1 AND lower(trim(class_name))=lower(trim(?))", (variant,)):
            subject = str(row["subject"] or "").strip()
            if subject:
                subjects.add(subject)

    for subject in sorted(subjects, key=str.casefold):
        # Missing catalog rows must not stop registration. Create the catalog entry
        # from the configured subject assignment and then enroll the learner.
        cat = q("SELECT id FROM subjects_catalog WHERE lower(trim(subject))=lower(trim(?) ) LIMIT 1", (subject,), one=True)
        if not cat:
            try:
                new_cat_id = execute("INSERT INTO subjects_catalog(subject,department,level_scope,description,active,created_by) VALUES(?,?,?,?,1,?)", (subject, "", class_name, "Auto-created from grade allocation", actor_id))
                cat = {"id": new_cat_id}
            except sqlite3.IntegrityError:
                cat = q("SELECT id FROM subjects_catalog WHERE lower(trim(subject))=lower(trim(?) ) LIMIT 1", (subject,), one=True)
        if cat:
            execute("INSERT INTO student_subjects(student_id,subject_id,status,selected_by) VALUES(?,?, 'Approved', ?) ON CONFLICT(student_id,subject_id) DO UPDATE SET status='Approved',selected_by=excluded.selected_by,updated_at=CURRENT_TIMESTAMP", (student_id,cat['id'],actor_id))
            summary['subjects'] += 1

        teacher_row = None
        for variant in grade_variants:
            teacher_row = q("SELECT teacher_user_id FROM teacher_assignments WHERE active=1 AND lower(trim(class_name))=lower(trim(?)) AND lower(trim(subject))=lower(trim(?)) ORDER BY id LIMIT 1", (variant,subject), one=True)
            if teacher_row:
                break
        if teacher_row:
            teacher = q("SELECT id,full_name FROM users WHERE id=? AND role='Teacher' AND active=1", (teacher_row['teacher_user_id'],), one=True)
            if teacher:
                execute("INSERT OR IGNORE INTO student_teacher_assignments(student_id,teacher_user_id,class_name,subject,scope,assigned_by,active) VALUES(?,?,?,?, 'Subject',?,1)", (student_id,teacher['id'],class_name,subject,actor_id))
                summary['subject_teachers'] += 1
    return summary


def backfill_student_allocations() -> None:
    """Bring existing learners into the canonical grade/class/subject allocation model."""
    actor = q("SELECT id FROM users WHERE active=1 AND role IN ('Admin','ICT') ORDER BY CASE role WHEN 'Admin' THEN 0 ELSE 1 END, id LIMIT 1", one=True)
    if not actor:
        return
    students = q("SELECT id,grade FROM students WHERE active=1 AND TRIM(COALESCE(grade,''))!=''")
    for st in students:
        try:
            auto_place_new_student(st['id'], normalize_grade(st['grade']), actor['id'])
        except sqlite3.DatabaseError:
            # A single malformed legacy relationship must never stop startup.
            continue


def table_columns(conn: sqlite3.Connection, table: str) -> set[str]:
    rows = conn.execute(f"PRAGMA table_info({table})").fetchall()
    return {row[1] for row in rows}


def ensure_column(conn: sqlite3.Connection, table: str, column_def: str) -> None:
    name = column_def.split()[0]
    if name not in table_columns(conn, table):
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column_def}")


def _init_db_once() -> None:
    INSTANCE_DIR.mkdir(exist_ok=True)
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    legacy_db = INSTANCE_DIR / "school.db"

    # Render may retain a corrupted SQLite database from a failed deploy or an
    # interrupted write. Never let that make Gunicorn crash before the app starts.
    # Preserve the damaged file, then recover from the known-good bundled DB when
    # available; otherwise SQLite will create a clean database below.
    if DB_PATH.exists() and not _sqlite_integrity_ok(DB_PATH):
        _quarantine_corrupt_db(DB_PATH)

    if PERSISTENT_STORAGE and not DB_PATH.exists() and legacy_db.exists() and _sqlite_integrity_ok(legacy_db):
        import shutil
        shutil.copy2(legacy_db, DB_PATH)

    with sqlite3.connect(DB_PATH, timeout=30) as conn:
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=FULL")
        conn.execute("PRAGMA busy_timeout=30000")
        conn.row_factory = sqlite3.Row
        conn.executescript(
            """
            PRAGMA foreign_keys = ON;

            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                full_name TEXT NOT NULL,
                username TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL CHECK(role IN ('Admin','ICT','Finance','Teacher','Student','Parent','Librarian','System')),
                student_id INTEGER,
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS school_settings (
                id INTEGER PRIMARY KEY CHECK (id = 1),
                school_name TEXT NOT NULL DEFAULT 'School',
                admission_prefix TEXT NOT NULL DEFAULT 'ADM-',
                admission_suffix TEXT NOT NULL DEFAULT '',
                student_name_prefix TEXT NOT NULL DEFAULT '',
                student_name_suffix TEXT NOT NULL DEFAULT '',
                currency_code TEXT NOT NULL DEFAULT 'KES',
                school_fee REAL NOT NULL DEFAULT 0
            );

            CREATE TABLE IF NOT EXISTS advertisements (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                body TEXT NOT NULL DEFAULT '',
                link_url TEXT NOT NULL DEFAULT '',
                start_date TEXT NOT NULL DEFAULT '',
                end_date TEXT NOT NULL DEFAULT '',
                priority INTEGER NOT NULL DEFAULT 0,
                active INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                admission_no TEXT NOT NULL UNIQUE,
                full_name TEXT NOT NULL,
                grade TEXT NOT NULL,
                guardian_name TEXT,
                guardian_phone TEXT,
                guardian_email TEXT,
                alt_guardian_name TEXT,
                alt_guardian_phone TEXT,
                alt_guardian_email TEXT,
                student_phone TEXT,
                student_email TEXT,
                medical_condition TEXT,
                allergies TEXT,
                special_info TEXT,
                notes TEXT,
                payment_status TEXT NOT NULL DEFAULT 'Pending' CHECK(payment_status IN ('Paid', 'Pending')),
                balance REAL NOT NULL DEFAULT 0,
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS payments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER NOT NULL,
                amount REAL NOT NULL,
                method TEXT NOT NULL,
                reference_no TEXT,
                recorded_by INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Posted' CHECK(status IN ('Posted', 'Reversed')),
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                reversed_at TEXT,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(recorded_by) REFERENCES users(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                actor_id INTEGER,
                actor_name TEXT NOT NULL,
                action TEXT NOT NULL,
                details TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(actor_id) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS login_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER NOT NULL,
                method TEXT NOT NULL DEFAULT 'Password', logged_in_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                ip_address TEXT, user_agent TEXT, latitude REAL, longitude REAL, accuracy REAL, location_label TEXT,
                FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_login_events_user_time ON login_events(user_id, logged_in_at);

            CREATE TABLE IF NOT EXISTS exam_batches (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                grade TEXT NOT NULL,
                term TEXT NOT NULL,
                submitted_by INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Submitted',
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(submitted_by) REFERENCES users(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS exam_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                batch_id INTEGER NOT NULL,
                grade TEXT NOT NULL,
                term TEXT NOT NULL,
                subject TEXT NOT NULL,
                student_id INTEGER NOT NULL,
                student_name TEXT NOT NULL,
                admission_no TEXT NOT NULL,
                mark REAL NOT NULL,
                max_mark REAL NOT NULL DEFAULT 100,
                submitted_by INTEGER NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(batch_id) REFERENCES exam_batches(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(submitted_by) REFERENCES users(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS assignments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                subject TEXT NOT NULL,
                grade TEXT NOT NULL,
                description TEXT,
                deadline TEXT,
                attachment_path TEXT,
                posted_by INTEGER NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                assignment_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                attachment_path TEXT,
                note TEXT,
                score REAL,
                feedback TEXT,
                submitted_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(assignment_id) REFERENCES assignments(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS portal_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                sender_role TEXT NOT NULL,
                sender_name TEXT NOT NULL,
                recipient_role TEXT NOT NULL,
                recipient_student_id INTEGER,
                body TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS portal_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_type TEXT NOT NULL CHECK(document_type IN ('Result', 'Exam Card')),
                student_id INTEGER NOT NULL,
                batch_id INTEGER,
                file_path TEXT,
                qr_token TEXT NOT NULL UNIQUE,
                status TEXT NOT NULL DEFAULT 'Valid' CHECK(status IN ('Valid', 'Revoked')),
                issued_by INTEGER,
                issued_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(batch_id) REFERENCES exam_batches(id) ON DELETE SET NULL,
                FOREIGN KEY(issued_by) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS elections (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                description TEXT,
                start_at TEXT,
                end_at TEXT,
                visible INTEGER NOT NULL DEFAULT 0,
                active INTEGER NOT NULL DEFAULT 0,
                created_by INTEGER,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS election_candidates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                election_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                position TEXT NOT NULL,
                manifesto TEXT,
                image_path TEXT,
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(election_id) REFERENCES elections(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS election_votes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                election_id INTEGER NOT NULL,
                candidate_id INTEGER NOT NULL,
                voter_user_id INTEGER NOT NULL,
                position TEXT NOT NULL DEFAULT 'General',
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(election_id, voter_user_id, position),
                FOREIGN KEY(election_id) REFERENCES elections(id) ON DELETE CASCADE,
                FOREIGN KEY(candidate_id) REFERENCES election_candidates(id) ON DELETE CASCADE,
                FOREIGN KEY(voter_user_id) REFERENCES users(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS library_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                category TEXT NOT NULL DEFAULT 'Book',
                author TEXT,
                item_code TEXT,
                quantity INTEGER NOT NULL DEFAULT 1,
                available_quantity INTEGER NOT NULL DEFAULT 1,
                location TEXT,
                resource_type TEXT NOT NULL DEFAULT 'Physical',
                file_path TEXT,
                external_url TEXT,
                description TEXT,
                active INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS library_loans (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                item_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                issued_by INTEGER NOT NULL,
                issued_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                due_date TEXT,
                returned_at TEXT,
                status TEXT NOT NULL DEFAULT 'Issued' CHECK(status IN ('Issued','Returned','Overdue')),
                FOREIGN KEY(item_id) REFERENCES library_items(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(issued_by) REFERENCES users(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS departments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                category TEXT NOT NULL DEFAULT 'Academic',
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS guardian_links (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                guardian_user_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                relationship TEXT NOT NULL DEFAULT 'Guardian',
                is_primary INTEGER NOT NULL DEFAULT 0,
                notes TEXT NOT NULL DEFAULT '',
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(guardian_user_id, student_id),
                FOREIGN KEY(guardian_user_id) REFERENCES users(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );
            """
        )

        ensure_column(conn, "school_settings", "school_name TEXT NOT NULL DEFAULT 'School'")
        ensure_column(conn, "school_settings", "admission_prefix TEXT NOT NULL DEFAULT 'ADM-'")
        ensure_column(conn, "school_settings", "admission_suffix TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "student_name_prefix TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "student_name_suffix TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "currency_code TEXT NOT NULL DEFAULT 'KES'")
        ensure_column(conn, "school_settings", "school_fee REAL NOT NULL DEFAULT 0")
        ensure_column(conn, "school_settings", "menu_order TEXT NOT NULL DEFAULT 'Home,Academics,Assignments,Messages,Profile'")
        ensure_column(conn, "school_settings", "background_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "logo_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "portal_subtitle TEXT NOT NULL DEFAULT 'School Portal System'")
        ensure_column(conn, "school_settings", "panel_color TEXT NOT NULL DEFAULT '#40414f'")
        ensure_column(conn, "school_settings", "background_color TEXT NOT NULL DEFAULT '#343541'")
        ensure_column(conn, "school_settings", "accent_color TEXT NOT NULL DEFAULT '#3457d5'")
        ensure_column(conn, "school_settings", "primary_color TEXT NOT NULL DEFAULT '#3457d5'")
        ensure_column(conn, "school_settings", "branding_label TEXT NOT NULL DEFAULT 'Branding'")
        ensure_column(conn, "school_settings", "finance_label TEXT NOT NULL DEFAULT 'Finance'")
        ensure_column(conn, "school_settings", "messages_label TEXT NOT NULL DEFAULT 'Messages'")
        ensure_column(conn, "school_settings", "results_label TEXT NOT NULL DEFAULT 'Results'")
        ensure_column(conn, "school_settings", "assignments_label TEXT NOT NULL DEFAULT 'Assignments'")
        ensure_column(conn, "school_settings", "home_label TEXT NOT NULL DEFAULT 'Home'")
        ensure_column(conn, "school_settings", "result_download_balance_limit REAL NOT NULL DEFAULT 0")
        ensure_column(conn, "school_settings", "result_release_mode TEXT NOT NULL DEFAULT 'Finance Approval'")
        ensure_column(conn, "school_settings", "exam_card_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "auth_required INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "school_settings", "prelogin_sections TEXT NOT NULL DEFAULT 'institution,history,achievements,owners,developer,company'")
        ensure_column(conn, "school_settings", "institution_owners TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "developer_name TEXT NOT NULL DEFAULT 'Toror Technology and Innovations Ltd.'")
        ensure_column(conn, "school_settings", "developer_about TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "company_name TEXT NOT NULL DEFAULT 'Toror Technology and Innovations Ltd.'")
        ensure_column(conn, "school_settings", "company_about TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "elections_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "library_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "institution_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "institution_history TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_performance TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_religion TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_affiliations TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_help TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_contact TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_image_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_image_2_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_image_3_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_image_1_position TEXT NOT NULL DEFAULT '50% 50%'")
        ensure_column(conn, "school_settings", "institution_image_2_position TEXT NOT NULL DEFAULT '50% 50%'")
        ensure_column(conn, "school_settings", "institution_image_3_position TEXT NOT NULL DEFAULT '50% 50%'")
        ensure_column(conn, "school_settings", "landing_background_color TEXT NOT NULL DEFAULT '#e7efff'")
        ensure_column(conn, "school_settings", "landing_panel_color TEXT NOT NULL DEFAULT '#f8fbff'")
        ensure_column(conn, "school_settings", "landing_text_color TEXT NOT NULL DEFAULT '#152033'")
        ensure_column(conn, "school_settings", "landing_accent_color TEXT NOT NULL DEFAULT '#2457d6'")
        ensure_column(conn, "school_settings", "landing_background_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "landing_font_family TEXT NOT NULL DEFAULT 'Inter'")
        ensure_column(conn, "school_settings", "landing_heading_font TEXT NOT NULL DEFAULT 'Manrope'")
        ensure_column(conn, "school_settings", "landing_content_width INTEGER NOT NULL DEFAULT 1240")
        ensure_column(conn, "school_settings", "landing_hero_layout TEXT NOT NULL DEFAULT 'split'")
        ensure_column(conn, "school_settings", "landing_role_columns INTEGER NOT NULL DEFAULT 3")
        ensure_column(conn, "school_settings", "online_class_provider TEXT NOT NULL DEFAULT 'https://meet.jit.si/'")
        ensure_column(conn, "school_settings", "parent_portal_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "welcome_animation_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "welcome_animation_name TEXT NOT NULL DEFAULT 'Toror Technology and Innovations Ltd.'")
        ensure_column(conn, "school_settings", "welcome_animation_duration_ms INTEGER NOT NULL DEFAULT 2200")
        ensure_column(conn, "school_settings", "welcome_animation_style TEXT NOT NULL DEFAULT 'clean'")
        ensure_column(conn, "school_settings", "landing_hero_title TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "landing_hero_text TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "landing_cta_primary TEXT NOT NULL DEFAULT 'Sign in to your workspace'")
        ensure_column(conn, "school_settings", "landing_cta_secondary TEXT NOT NULL DEFAULT 'View school information'")
        ensure_column(conn, "school_settings", "landing_announcement TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "landing_contact TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "landing_show_dates INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "landing_show_gallery INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "landing_show_roles INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "theme_preset TEXT NOT NULL DEFAULT 'classic'")

        ensure_column(conn, "students", "age TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "guardian_name TEXT")
        ensure_column(conn, "students", "guardian_phone TEXT")
        ensure_column(conn, "students", "guardian_email TEXT")
        ensure_column(conn, "students", "alt_guardian_name TEXT")
        ensure_column(conn, "students", "alt_guardian_phone TEXT")
        ensure_column(conn, "students", "alt_guardian_email TEXT")
        ensure_column(conn, "students", "student_phone TEXT")
        ensure_column(conn, "students", "student_email TEXT")
        ensure_column(conn, "students", "medical_condition TEXT")
        ensure_column(conn, "students", "allergies TEXT")
        ensure_column(conn, "students", "special_info TEXT")
        ensure_column(conn, "payments", "receipt_path TEXT")
        ensure_column(conn, "students", "notes TEXT")
        ensure_column(conn, "students", "date_of_birth TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "gender TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "id_reference TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "address TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "emergency_contact TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "blood_group TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "medical_notes TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "students", "accountability_notes TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "exam_batches", "finance_status TEXT NOT NULL DEFAULT 'Pending'")
        ensure_column(conn, "exam_batches", "finance_note TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "exam_batches", "approved_by INTEGER")
        ensure_column(conn, "exam_batches", "approved_at TEXT")
        if "election_votes" in [r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]:
            ensure_column(conn, "election_votes", "position TEXT NOT NULL DEFAULT 'General'")

        if conn.execute("SELECT COUNT(*) AS c FROM school_settings").fetchone()["c"] == 0:
            conn.execute(
                "INSERT INTO school_settings(id, school_name, admission_prefix, admission_suffix, student_name_prefix, student_name_suffix, currency_code, school_fee) VALUES (1, ?, ?, ?, ?, ?, ?, ?)",
                ("School Portal System", "ADM-", "", "", "", "KES", 0),
            )
        else:
            conn.execute("UPDATE school_settings SET school_name = 'School Portal System' WHERE id = 1 AND TRIM(school_name) IN ('', 'School', 'Legacy Portal')")

        # Keep the two baseline college/TVET departments available immediately.
        for dept_name in DEFAULT_DEPARTMENTS:
            conn.execute("INSERT OR IGNORE INTO departments(name, category) VALUES(?, 'Academic')", (dept_name,))

        # Extend the users role constraint to support the dedicated Librarian role.
        # Rebuild the user table and its direct FK dependants together. SQLite can
        # otherwise leave stale FK metadata pointing at a temporary migration table.
        cols = table_columns(conn, "users")
        if cols and "role" in cols:
            table_sql = conn.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name='users'").fetchone()
            sql_text = (table_sql[0] or "") if table_sql else ""
            if "'Librarian'" not in sql_text:
                conn.execute("PRAGMA foreign_keys=OFF")
                for table, tmp, create_sql, columns in [
                    ("payments", "payments_user_roles_fix", """CREATE TABLE {tmp} (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER NOT NULL, amount REAL NOT NULL,
                        method TEXT NOT NULL, reference_no TEXT, recorded_by INTEGER NOT NULL,
                        status TEXT NOT NULL DEFAULT 'Posted' CHECK(status IN ('Posted','Reversed')),
                        created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, reversed_at TEXT, receipt_path TEXT,
                        FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                        FOREIGN KEY(recorded_by) REFERENCES users_new(id) ON DELETE CASCADE
                    )""",
                     ["id","student_id","amount","method","reference_no","recorded_by","status","created_at","reversed_at","receipt_path"]),
                    ("audit_log", "audit_log_user_roles_fix", """CREATE TABLE {tmp} (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, actor_id INTEGER, actor_name TEXT NOT NULL,
                        action TEXT NOT NULL, details TEXT NOT NULL, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY(actor_id) REFERENCES users_new(id) ON DELETE SET NULL
                    )""",
                     ["id","actor_id","actor_name","action","details","created_at"]),
                ]:
                    conn.execute(f"DROP TABLE IF EXISTS {tmp}")
                conn.execute("ALTER TABLE users RENAME TO users_legacy_roles")
                conn.execute("""CREATE TABLE users_new (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, full_name TEXT NOT NULL, username TEXT NOT NULL UNIQUE,
                    password_hash TEXT NOT NULL,
                    role TEXT NOT NULL CHECK(role IN ('Admin','ICT','Finance','Teacher','Student','Parent','Librarian','System')),
                    student_id INTEGER, active INTEGER NOT NULL DEFAULT 1,
                    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE SET NULL
                )""")
                conn.execute("INSERT INTO users_new(id,full_name,username,password_hash,role,student_id,active,created_at) SELECT id,full_name,username,password_hash,role,student_id,active,created_at FROM users_legacy_roles")
                for table, tmp, create_sql, columns in [
                    ("payments", "payments_user_roles_fix", """CREATE TABLE {tmp} (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER NOT NULL, amount REAL NOT NULL,
                        method TEXT NOT NULL, reference_no TEXT, recorded_by INTEGER NOT NULL,
                        status TEXT NOT NULL DEFAULT 'Posted' CHECK(status IN ('Posted','Reversed')),
                        created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, reversed_at TEXT, receipt_path TEXT,
                        FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                        FOREIGN KEY(recorded_by) REFERENCES users_new(id) ON DELETE CASCADE
                    )""",
                     ["id","student_id","amount","method","reference_no","recorded_by","status","created_at","reversed_at","receipt_path"]),
                    ("audit_log", "audit_log_user_roles_fix", """CREATE TABLE {tmp} (
                        id INTEGER PRIMARY KEY AUTOINCREMENT, actor_id INTEGER, actor_name TEXT NOT NULL,
                        action TEXT NOT NULL, details TEXT NOT NULL, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                        FOREIGN KEY(actor_id) REFERENCES users_new(id) ON DELETE SET NULL
                    )""",
                     ["id","actor_id","actor_name","action","details","created_at"]),
                ]:
                    exists = conn.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone()
                    if exists:
                        conn.execute(create_sql.format(tmp=tmp))
                        conn.execute(f"INSERT INTO {tmp} ({','.join(columns)}) SELECT {','.join(columns)} FROM {table}")
                        conn.execute(f"DROP TABLE {table}")
                        conn.execute(f"ALTER TABLE {tmp} RENAME TO {table}")
                conn.execute("DROP TABLE users_legacy_roles")
                conn.execute("ALTER TABLE users_new RENAME TO users")
                conn.execute("PRAGMA foreign_keys=ON")

        # Authentication is account-based. A fresh installation starts with only a
        # non-login system account so existing demo payments/audit records remain valid.
        # The first real Administrator is created through /register.
        if "student_id" not in table_columns(conn, "users") or "active" not in table_columns(conn, "users"):
            # Migrate older users table while preserving existing rows.
            old_rows = conn.execute("SELECT id, full_name, username, password_hash, role, created_at FROM users").fetchall()
            conn.execute("ALTER TABLE users RENAME TO users_legacy")
            conn.execute("""CREATE TABLE users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                full_name TEXT NOT NULL,
                username TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL CHECK(role IN ('Admin','ICT','Finance','Teacher','Student','Parent','Librarian','System')),
                student_id INTEGER,
                active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE SET NULL
            )""")
            for r in old_rows:
                role = r[4] if r[4] in ALL_PORTAL_ROLES else SYSTEM_ROLE
                conn.execute("INSERT INTO users(id,full_name,username,password_hash,role,active,created_at) VALUES(?,?,?,?,?,?,?)",
                             (r[0],r[1],r[2],r[3],role,1,r[5] or datetime.utcnow().isoformat(timespec="seconds")))
            conn.execute("DROP TABLE users_legacy")
        ensure_column(conn, "users", "student_id INTEGER")
        ensure_column(conn, "users", "active INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "users", "title TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "department TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "phone TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "email TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "date_of_birth TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "gender TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "id_reference TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "address TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "emergency_contact TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "blood_group TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "title TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "department TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "leadership_role TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "leadership_level INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "users", "medical_notes TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "accountability_notes TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "profile_photo TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "archived_at TEXT")
        conn.execute("CREATE TABLE IF NOT EXISTS portal_contexts (token_id TEXT PRIMARY KEY, user_id INTEGER NOT NULL, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, revoked INTEGER NOT NULL DEFAULT 0, FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE)")
        ensure_column(conn, "school_settings", "auth_initialized INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "school_settings", "institution_type TEXT NOT NULL DEFAULT 'Secondary School'")
        ensure_column(conn, "school_settings", "learner_label TEXT NOT NULL DEFAULT 'Student'")
        ensure_column(conn, "school_settings", "staff_label TEXT NOT NULL DEFAULT 'Teacher'")
        ensure_column(conn, "school_settings", "academic_period_label TEXT NOT NULL DEFAULT 'Term'")
        ensure_column(conn, "school_settings", "class_label TEXT NOT NULL DEFAULT 'Class'")
        ensure_column(conn, "school_settings", "department_label TEXT NOT NULL DEFAULT 'Department'")
        ensure_column(conn, "school_settings", "theme_mode TEXT NOT NULL DEFAULT 'dark'")
        ensure_column(conn, "school_settings", "font_family TEXT NOT NULL DEFAULT 'Inter'")
        ensure_column(conn, "school_settings", "heading_font TEXT NOT NULL DEFAULT 'Inter'")
        ensure_column(conn, "school_settings", "radius_px INTEGER NOT NULL DEFAULT 12")
        ensure_column(conn, "school_settings", "sidebar_color TEXT NOT NULL DEFAULT '#40414f'")
        ensure_column(conn, "school_settings", "header_color TEXT NOT NULL DEFAULT '#40414f'")
        ensure_column(conn, "school_settings", "text_color TEXT NOT NULL DEFAULT '#ececf1'")
        ensure_column(conn, "school_settings", "muted_text_color TEXT NOT NULL DEFAULT '#b5bac7'")
        ensure_column(conn, "school_settings", "button_radius_px INTEGER NOT NULL DEFAULT 10")
        ensure_column(conn, "school_settings", "sidebar_style TEXT NOT NULL DEFAULT 'drawer'")
        ensure_column(conn, "school_settings", "custom_css TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "ai_enabled INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "school_settings", "ai_provider TEXT NOT NULL DEFAULT 'openai_responses'")
        ensure_column(conn, "school_settings", "ai_model TEXT NOT NULL DEFAULT 'gpt-5.6'")
        ensure_column(conn, "school_settings", "help_phone TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "help_email TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "footer_title TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "footer_text TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "footer_contact TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "footer_links TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "platform_credit_enabled INTEGER NOT NULL DEFAULT 1")
        conn.execute("UPDATE school_settings SET footer_title = school_name WHERE id=1 AND TRIM(COALESCE(footer_title,''))=''")
        ensure_column(conn, "users", "qr_access_token TEXT")
        ensure_column(conn, "users", "qr_login_enabled INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "users", "last_password_login_at TEXT")

        # V25 institution upgrades: learner fee composition, transport zones,
        # promotion workflow, generated learner email domain, and safer payroll bulk actions.
        for table, coldef in [
            ("students", "fee_assessed_total REAL NOT NULL DEFAULT 0"),
            ("students", "fee_override_enabled INTEGER NOT NULL DEFAULT 0"),
            ("students", "transport_zone TEXT NOT NULL DEFAULT ''"),
            ("students", "uses_school_bus INTEGER NOT NULL DEFAULT 0"),
            ("students", "meal_plan TEXT NOT NULL DEFAULT 'None'"),
            ("students", "transport_charge REAL NOT NULL DEFAULT 0"),
            ("students", "last_promotion_action TEXT NOT NULL DEFAULT ''"),
            ("students", "academic_year TEXT NOT NULL DEFAULT ''"),
            ("students", "promoted_from_grade TEXT NOT NULL DEFAULT ''"),
            ("users", "email TEXT NOT NULL DEFAULT ''"),
            ("school_settings", "student_email_domain TEXT NOT NULL DEFAULT 'school.ac.ke'"),
            ("school_settings", "public_location_title TEXT NOT NULL DEFAULT 'Visit us today at this location'"),
            ("school_settings", "public_address TEXT NOT NULL DEFAULT ''"),
            ("school_settings", "public_location_notes TEXT NOT NULL DEFAULT ''"),
            ("school_settings", "public_map_query TEXT NOT NULL DEFAULT ''"),
        ]:
            ensure_column(conn, table, coldef)
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS transport_rates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            zone_name TEXT NOT NULL UNIQUE,
            amount REAL NOT NULL DEFAULT 0,
            period TEXT NOT NULL DEFAULT 'Term 1',
            active INTEGER NOT NULL DEFAULT 1,
            created_by INTEGER,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE TABLE IF NOT EXISTS promotion_runs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            academic_year TEXT NOT NULL,
            effective_date TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'Draft',
            created_by INTEGER,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            completed_at TEXT,
            FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE TABLE IF NOT EXISTS promotion_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            run_id INTEGER NOT NULL,
            student_id INTEGER NOT NULL,
            from_grade TEXT NOT NULL,
            to_grade TEXT NOT NULL,
            decision TEXT NOT NULL DEFAULT 'Promote',
            notes TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(run_id) REFERENCES promotion_runs(id) ON DELETE CASCADE,
            FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_promotion_items_run ON promotion_items(run_id,student_id);
        """)
        # Seed a single generic transport zone only when the institution has none.
        if conn.execute("SELECT COUNT(*) FROM transport_rates").fetchone()[0] == 0:
            conn.execute("INSERT INTO transport_rates(zone_name,amount,period,active) VALUES('Local / Zone A',0,'Term 1',1)")
        conn.execute("UPDATE school_settings SET public_location_title=COALESCE(NULLIF(public_location_title,''),'Visit us today at this location') WHERE id=1")
        conn.execute("UPDATE school_settings SET public_address='', public_location_notes='', public_map_query='' WHERE id=1 AND (public_address LIKE '%Kimbo%' OR public_map_query LIKE '%Kimbo%')")
        # Reconcile legacy learner balances after migration/restore.
        for row in conn.execute("SELECT id FROM students").fetchall():
            sid=row[0]
            assessed=conn.execute("SELECT COALESCE(SUM(amount),0) FROM fee_charges WHERE student_id=? AND status='Posted'",(sid,)).fetchone()[0] or 0
            st=conn.execute("SELECT fee_assessed_total,fee_override_enabled FROM students WHERE id=?",(sid,)).fetchone()
            if not assessed and st and float(st[0] or 0)>0: assessed=float(st[0] or 0)
            if st and int(st[1] or 0): assessed=float(st[0] or 0)
            paid=conn.execute("SELECT COALESCE(SUM(amount),0) FROM payments WHERE student_id=? AND status='Posted'",(sid,)).fetchone()[0] or 0
            conn.execute("UPDATE students SET fee_assessed_total=?,balance=?,payment_status=? WHERE id=?",(assessed,float(assessed)-float(paid),'Paid' if float(assessed)-float(paid)<=0 else 'Pending',sid))
        conn.execute("CREATE TABLE IF NOT EXISTS school_calendar (id INTEGER PRIMARY KEY AUTOINCREMENT,title TEXT NOT NULL,start_date TEXT NOT NULL,end_date TEXT NOT NULL,kind TEXT NOT NULL DEFAULT 'School Day',school_day INTEGER NOT NULL DEFAULT 1,notes TEXT NOT NULL DEFAULT '',created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_school_calendar_range ON school_calendar(start_date,end_date,school_day)")
        ensure_column(conn, "users", "position_code TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "school_unit TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "school_location TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "users", "reception_enabled INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "users", "staff_code TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "class_level TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "subject TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "image_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "youtube_url TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "source_url TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "library_items", "source_name TEXT NOT NULL DEFAULT ''")

        conn.executescript("""
        CREATE TABLE IF NOT EXISTS system_help (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT NOT NULL DEFAULT 'Getting Started',
            content TEXT NOT NULL,
            role_scope TEXT NOT NULL DEFAULT 'All',
            sort_order INTEGER NOT NULL DEFAULT 100,
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS ai_usage_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            role TEXT,
            provider TEXT NOT NULL,
            model TEXT,
            prompt_preview TEXT,
            response_preview TEXT,
            status TEXT NOT NULL DEFAULT 'Success',
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE TABLE IF NOT EXISTS student_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id INTEGER NOT NULL,
            author_user_id INTEGER NOT NULL,
            category TEXT NOT NULL DEFAULT 'General',
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            visible_to_parent INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
            FOREIGN KEY(author_user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_student_records_student ON student_records(student_id, created_at);
        CREATE TABLE IF NOT EXISTS finance_ledger (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entry_type TEXT NOT NULL CHECK(entry_type IN ('Income','Expense','Payroll','Adjustment')),
            category TEXT NOT NULL,
            payee_user_id INTEGER,
            amount REAL NOT NULL CHECK(amount > 0),
            description TEXT NOT NULL,
            reference_no TEXT,
            status TEXT NOT NULL DEFAULT 'Posted' CHECK(status IN ('Posted','Reversed')),
            posted_by INTEGER NOT NULL,
            posted_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            reversed_by INTEGER,
            reversed_at TEXT,
            FOREIGN KEY(payee_user_id) REFERENCES users(id) ON DELETE SET NULL,
            FOREIGN KEY(posted_by) REFERENCES users(id) ON DELETE RESTRICT,
            FOREIGN KEY(reversed_by) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE TABLE IF NOT EXISTS finance_accounts (
            id INTEGER PRIMARY KEY CHECK(id=1),
            account_name TEXT NOT NULL DEFAULT 'Institution Operating Account',
            opening_balance REAL NOT NULL DEFAULT 0,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS backup_registry (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            backup_type TEXT NOT NULL,
            file_name TEXT NOT NULL,
            row_count INTEGER NOT NULL DEFAULT 0,
            created_by INTEGER,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE TABLE IF NOT EXISTS important_dates (id INTEGER PRIMARY KEY AUTOINCREMENT,title TEXT NOT NULL,event_date TEXT NOT NULL,event_time TEXT,location TEXT,description TEXT,visible INTEGER NOT NULL DEFAULT 1,landing_visible INTEGER NOT NULL DEFAULT 1,created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE INDEX IF NOT EXISTS idx_important_dates_date ON important_dates(event_date,visible,landing_visible);
        CREATE TABLE IF NOT EXISTS notifications (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER,title TEXT NOT NULL,body TEXT NOT NULL,link TEXT,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,read_at TEXT,priority TEXT NOT NULL DEFAULT 'Normal',FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE INDEX IF NOT EXISTS idx_notifications_user_read ON notifications(user_id,read_at,created_at);
        CREATE TABLE IF NOT EXISTS password_reset_tokens (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER NOT NULL,token_hash TEXT NOT NULL UNIQUE,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,expires_at TEXT NOT NULL,used_at TEXT,requested_ip TEXT,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE INDEX IF NOT EXISTS idx_password_reset_tokens_hash ON password_reset_tokens(token_hash);
        CREATE TABLE IF NOT EXISTS password_reset_requests (id INTEGER PRIMARY KEY AUTOINCREMENT,username TEXT NOT NULL,reason TEXT NOT NULL DEFAULT '',status TEXT NOT NULL DEFAULT 'Open',created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,resolved_at TEXT,resolved_by INTEGER,FOREIGN KEY(resolved_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE INDEX IF NOT EXISTS idx_password_reset_requests_status ON password_reset_requests(status,created_at);
        CREATE TABLE IF NOT EXISTS groups (id INTEGER PRIMARY KEY AUTOINCREMENT,name TEXT NOT NULL,group_type TEXT NOT NULL DEFAULT 'Academic',description TEXT,owner_user_id INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,active INTEGER NOT NULL DEFAULT 1,FOREIGN KEY(owner_user_id) REFERENCES users(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS group_members (id INTEGER PRIMARY KEY AUTOINCREMENT,group_id INTEGER NOT NULL,user_id INTEGER NOT NULL,student_id INTEGER,role TEXT NOT NULL DEFAULT 'Member',created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,UNIQUE(group_id,user_id),FOREIGN KEY(group_id) REFERENCES groups(id) ON DELETE CASCADE,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS group_posts (id INTEGER PRIMARY KEY AUTOINCREMENT,group_id INTEGER NOT NULL,user_id INTEGER NOT NULL,body TEXT NOT NULL,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(group_id) REFERENCES groups(id) ON DELETE CASCADE,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS class_sessions (id INTEGER PRIMARY KEY AUTOINCREMENT,teacher_user_id INTEGER NOT NULL,class_name TEXT NOT NULL,subject TEXT NOT NULL,title TEXT NOT NULL,starts_at TEXT NOT NULL,ends_at TEXT,room_name TEXT NOT NULL UNIQUE,provider_url TEXT NOT NULL,description TEXT,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,active INTEGER NOT NULL DEFAULT 1,FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE INDEX IF NOT EXISTS idx_class_sessions_start ON class_sessions(starts_at,active);
        CREATE TABLE IF NOT EXISTS leadership_assignments (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER NOT NULL,leadership_role TEXT NOT NULL,level INTEGER NOT NULL DEFAULT 1,department TEXT NOT NULL DEFAULT '',appointed_by INTEGER,active INTEGER NOT NULL DEFAULT 1,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(appointed_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE INDEX IF NOT EXISTS idx_library_class ON library_items(class_level, subject, active);
        CREATE INDEX IF NOT EXISTS idx_users_role_active ON users(role, active);
        CREATE INDEX IF NOT EXISTS idx_finance_status_date ON finance_ledger(status, posted_at);
        """)
        # These migrations are Python helpers and must run outside executescript().
        ensure_column(conn, "finance_ledger", "receipt_path TEXT NOT NULL DEFAULT ''")
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS attendance_events (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER NOT NULL,action TEXT NOT NULL CHECK(action IN ('IN','OUT')),method TEXT NOT NULL DEFAULT 'QR',office_token TEXT,event_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,source TEXT NOT NULL DEFAULT 'online',latitude REAL,longitude REAL,speed_kph REAL,device_note TEXT,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE INDEX IF NOT EXISTS idx_attendance_user_time ON attendance_events(user_id,event_at);
        CREATE TABLE IF NOT EXISTS attendance_days (attendance_date TEXT PRIMARY KEY,status TEXT NOT NULL DEFAULT 'Open' CHECK(status IN ('Open','Closed')),closed_at TEXT,closed_by INTEGER,FOREIGN KEY(closed_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS teacher_assignments (id INTEGER PRIMARY KEY AUTOINCREMENT,teacher_user_id INTEGER NOT NULL,class_name TEXT NOT NULL,subject TEXT NOT NULL,unit_code TEXT,active INTEGER NOT NULL DEFAULT 1,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS communication_messages (id INTEGER PRIMARY KEY AUTOINCREMENT,sender_user_id INTEGER NOT NULL,recipient_user_id INTEGER NOT NULL,body TEXT NOT NULL,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,read_at TEXT,FOREIGN KEY(sender_user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(recipient_user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS class_attendance (id INTEGER PRIMARY KEY AUTOINCREMENT,teacher_user_id INTEGER NOT NULL,student_id INTEGER NOT NULL,class_name TEXT NOT NULL,subject TEXT NOT NULL,attendance_date TEXT NOT NULL,status TEXT NOT NULL CHECK(status IN ('Present','Absent','Late','Excused')),note TEXT,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,UNIQUE(teacher_user_id,student_id,class_name,subject,attendance_date),FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS flashcard_decks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT NOT NULL DEFAULT '',
            subject TEXT NOT NULL DEFAULT 'General',
            class_name TEXT NOT NULL DEFAULT '',
            owner_user_id INTEGER NOT NULL,
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(owner_user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_flashcard_decks_owner ON flashcard_decks(owner_user_id, active);
        CREATE INDEX IF NOT EXISTS idx_flashcard_decks_scope ON flashcard_decks(class_name, subject, active);
        CREATE TABLE IF NOT EXISTS flashcards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            deck_id INTEGER NOT NULL,
            front TEXT NOT NULL,
            back TEXT NOT NULL,
            hint TEXT NOT NULL DEFAULT '',
            position INTEGER NOT NULL DEFAULT 0,
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(deck_id) REFERENCES flashcard_decks(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_flashcards_deck ON flashcards(deck_id, active, position);
        CREATE TABLE IF NOT EXISTS flashcard_progress (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            card_id INTEGER NOT NULL,
            rating INTEGER NOT NULL DEFAULT 0 CHECK(rating BETWEEN 0 AND 3),
            correct_count INTEGER NOT NULL DEFAULT 0,
            review_count INTEGER NOT NULL DEFAULT 0,
            next_review_at TEXT,
            last_reviewed_at TEXT,
            UNIQUE(user_id, card_id),
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY(card_id) REFERENCES flashcards(id) ON DELETE CASCADE
        );
        CREATE INDEX IF NOT EXISTS idx_flashcard_progress_user ON flashcard_progress(user_id, next_review_at);

        CREATE TABLE IF NOT EXISTS fee_structures (id INTEGER PRIMARY KEY AUTOINCREMENT,class_level TEXT NOT NULL,item_name TEXT NOT NULL,amount REAL NOT NULL CHECK(amount >= 0),period TEXT NOT NULL DEFAULT 'Term 1',active INTEGER NOT NULL DEFAULT 1,created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS fee_charges (id INTEGER PRIMARY KEY AUTOINCREMENT,student_id INTEGER NOT NULL,fee_structure_id INTEGER,amount REAL NOT NULL,description TEXT NOT NULL,status TEXT NOT NULL DEFAULT 'Posted',created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,FOREIGN KEY(fee_structure_id) REFERENCES fee_structures(id) ON DELETE SET NULL,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS payment_integrations (id INTEGER PRIMARY KEY CHECK(id=1),provider TEXT NOT NULL DEFAULT 'Manual',account_name TEXT NOT NULL DEFAULT '',collection_account TEXT NOT NULL DEFAULT '',callback_secret TEXT NOT NULL DEFAULT '',auto_match INTEGER NOT NULL DEFAULT 0,notes TEXT NOT NULL DEFAULT '',updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS external_payment_events (id INTEGER PRIMARY KEY AUTOINCREMENT,provider TEXT NOT NULL,external_reference TEXT NOT NULL UNIQUE,amount REAL NOT NULL,payer_name TEXT,payer_phone TEXT,admission_no TEXT,payload_json TEXT,status TEXT NOT NULL DEFAULT 'Received',matched_student_id INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,processed_at TEXT,FOREIGN KEY(matched_student_id) REFERENCES students(id) ON DELETE SET NULL);
        CREATE TABLE IF NOT EXISTS transport_trips (id INTEGER PRIMARY KEY AUTOINCREMENT,driver_user_id INTEGER NOT NULL,vehicle TEXT NOT NULL,route_name TEXT,status TEXT NOT NULL DEFAULT 'Active',started_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,ended_at TEXT,FOREIGN KEY(driver_user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS driver_locations (id INTEGER PRIMARY KEY AUTOINCREMENT,trip_id INTEGER,driver_user_id INTEGER NOT NULL,latitude REAL NOT NULL,longitude REAL NOT NULL,speed_kph REAL DEFAULT 0,accuracy REAL,recorded_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,source TEXT NOT NULL DEFAULT 'online',FOREIGN KEY(trip_id) REFERENCES transport_trips(id) ON DELETE SET NULL,FOREIGN KEY(driver_user_id) REFERENCES users(id) ON DELETE CASCADE);
        CREATE INDEX IF NOT EXISTS idx_driver_locations_latest ON driver_locations(driver_user_id,recorded_at);
        CREATE TABLE IF NOT EXISTS markbook_entries (id INTEGER PRIMARY KEY AUTOINCREMENT,teacher_user_id INTEGER NOT NULL,class_name TEXT NOT NULL,subject TEXT NOT NULL,student_id INTEGER NOT NULL,assessment TEXT NOT NULL,mark REAL NOT NULL,max_mark REAL NOT NULL DEFAULT 100,status TEXT NOT NULL DEFAULT 'Draft',created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE);
        CREATE TABLE IF NOT EXISTS attendance_qr_settings (id INTEGER PRIMARY KEY CHECK(id=1),office_name TEXT NOT NULL DEFAULT 'Main Office',token TEXT NOT NULL,updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS finance_closings (id INTEGER PRIMARY KEY AUTOINCREMENT,closing_date TEXT NOT NULL,submitted_by INTEGER NOT NULL,notes TEXT,status TEXT NOT NULL DEFAULT 'Submitted',submitted_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(submitted_by) REFERENCES users(id) ON DELETE RESTRICT);
        INSERT OR IGNORE INTO attendance_qr_settings(id,office_name,token) VALUES(1,'Main Office',lower(hex(randomblob(16))));
        CREATE TABLE IF NOT EXISTS reception_visits (
            id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
            person_type TEXT NOT NULL DEFAULT 'Visitor', full_name TEXT NOT NULL,
            phone TEXT NOT NULL DEFAULT '', gender TEXT NOT NULL DEFAULT '',
            reason TEXT NOT NULL DEFAULT '', position TEXT NOT NULL DEFAULT '',
            staff_code TEXT NOT NULL DEFAULT '', school_unit TEXT NOT NULL DEFAULT '',
            school_location TEXT NOT NULL DEFAULT '', device_token TEXT NOT NULL DEFAULT '',
            check_in TEXT, check_out TEXT, source TEXT NOT NULL DEFAULT 'online',
            method TEXT NOT NULL DEFAULT 'QR', latitude REAL, longitude REAL, accuracy REAL,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE SET NULL
        );
        CREATE INDEX IF NOT EXISTS idx_reception_open ON reception_visits(check_out,check_in);
        CREATE INDEX IF NOT EXISTS idx_reception_device ON reception_visits(device_token,created_at);
        INSERT OR IGNORE INTO payment_integrations(id,provider) VALUES(1,'Manual');
        """)
        ensure_column(conn, "teacher_assignments", "online_url TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "markbook_entries", "weight REAL NOT NULL DEFAULT 100")
        ensure_column(conn, "attendance_events", "location_label TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "attendance_events", "accuracy REAL")
        conn.execute("CREATE TABLE IF NOT EXISTS attendance_absence_requests (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER NOT NULL,absence_date TEXT NOT NULL,reason TEXT NOT NULL,status TEXT NOT NULL DEFAULT 'Pending' CHECK(status IN ('Pending','Approved','Denied')),requested_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,reviewed_by INTEGER,reviewed_at TEXT,review_note TEXT,UNIQUE(user_id,absence_date),FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(reviewed_by) REFERENCES users(id) ON DELETE SET NULL)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_absence_date ON attendance_absence_requests(absence_date,status)")
        conn.execute("CREATE TABLE IF NOT EXISTS staff_timetable (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER NOT NULL,day_of_week INTEGER NOT NULL CHECK(day_of_week BETWEEN 0 AND 6),start_time TEXT NOT NULL,end_time TEXT NOT NULL,title TEXT NOT NULL,location TEXT NOT NULL DEFAULT '',notes TEXT NOT NULL DEFAULT '',active INTEGER NOT NULL DEFAULT 1,created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_staff_timetable_user_day ON staff_timetable(user_id,day_of_week,start_time)")
        conn.execute("CREATE TABLE IF NOT EXISTS staff_reminders (id INTEGER PRIMARY KEY AUTOINCREMENT,user_id INTEGER,role_scope TEXT NOT NULL DEFAULT 'All',title TEXT NOT NULL,due_at TEXT NOT NULL,notes TEXT NOT NULL DEFAULT '',priority TEXT NOT NULL DEFAULT 'Normal',completed INTEGER NOT NULL DEFAULT 0,created_by INTEGER,created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE,FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_staff_reminders_due ON staff_reminders(due_at,completed)")
        conn.execute("""CREATE TABLE IF NOT EXISTS class_teacher_assignments (
            id INTEGER PRIMARY KEY AUTOINCREMENT, class_name TEXT NOT NULL UNIQUE, teacher_user_id INTEGER NOT NULL,
            assigned_by INTEGER, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY(assigned_by) REFERENCES users(id) ON DELETE SET NULL
        )""")
        conn.execute("""CREATE TABLE IF NOT EXISTS scheme_of_work (
            id INTEGER PRIMARY KEY AUTOINCREMENT, teacher_user_id INTEGER NOT NULL, class_name TEXT NOT NULL,
            subject TEXT NOT NULL, term TEXT NOT NULL DEFAULT 'Current Term', week_no INTEGER NOT NULL DEFAULT 1,
            topic TEXT NOT NULL DEFAULT '', objectives TEXT NOT NULL DEFAULT '', activities TEXT NOT NULL DEFAULT '',
            resources TEXT NOT NULL DEFAULT '', assessment TEXT NOT NULL DEFAULT '', status TEXT NOT NULL DEFAULT 'Planned',
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_scheme_teacher ON scheme_of_work(teacher_user_id,class_name,subject,term,week_no)")
        conn.execute("CREATE TABLE IF NOT EXISTS compulsory_subjects (id INTEGER PRIMARY KEY AUTOINCREMENT, class_name TEXT NOT NULL, subject TEXT NOT NULL, unit_name TEXT NOT NULL DEFAULT '', active INTEGER NOT NULL DEFAULT 1, created_by INTEGER, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, UNIQUE(class_name,subject), FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL)")
        conn.execute("""CREATE TABLE IF NOT EXISTS subjects_catalog (
            id INTEGER PRIMARY KEY AUTOINCREMENT, subject TEXT NOT NULL UNIQUE, department TEXT NOT NULL DEFAULT '',
            level_scope TEXT NOT NULL DEFAULT 'All', description TEXT NOT NULL DEFAULT '', active INTEGER NOT NULL DEFAULT 1,
            created_by INTEGER, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL
        )""")
        conn.execute("""CREATE TABLE IF NOT EXISTS student_subjects (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER NOT NULL, subject_id INTEGER NOT NULL,
            status TEXT NOT NULL DEFAULT 'Approved' CHECK(status IN ('Approved','Pending','Dropped')), selected_by INTEGER,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(student_id,subject_id), FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
            FOREIGN KEY(subject_id) REFERENCES subjects_catalog(id) ON DELETE CASCADE, FOREIGN KEY(selected_by) REFERENCES users(id) ON DELETE SET NULL
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_student_subjects_student ON student_subjects(student_id,status)")
        conn.execute("""CREATE TABLE IF NOT EXISTS student_departments (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER NOT NULL, department_id INTEGER NOT NULL,
            status TEXT NOT NULL DEFAULT 'Approved' CHECK(status IN ('Approved','Pending','Dropped')), selected_by INTEGER,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(student_id,department_id), FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
            FOREIGN KEY(department_id) REFERENCES departments(id) ON DELETE CASCADE, FOREIGN KEY(selected_by) REFERENCES users(id) ON DELETE SET NULL
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_student_departments_student ON student_departments(student_id,status)")
        conn.execute("""CREATE TABLE IF NOT EXISTS student_teacher_assignments (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER NOT NULL, teacher_user_id INTEGER NOT NULL,
            class_name TEXT NOT NULL DEFAULT '', subject TEXT NOT NULL DEFAULT '', scope TEXT NOT NULL DEFAULT 'Subject',
            assigned_by INTEGER, active INTEGER NOT NULL DEFAULT 1, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(student_id,teacher_user_id,class_name,subject),
            FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
            FOREIGN KEY(teacher_user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY(assigned_by) REFERENCES users(id) ON DELETE SET NULL
        )""")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_student_teacher_student ON student_teacher_assignments(student_id,active)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_student_teacher_teacher ON student_teacher_assignments(teacher_user_id,class_name,subject,active)")
        conn.execute("""CREATE TABLE IF NOT EXISTS department_leadership (
            department_id INTEGER PRIMARY KEY, dean_user_id INTEGER, deputy_user_id INTEGER, updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(department_id) REFERENCES departments(id) ON DELETE CASCADE,
            FOREIGN KEY(dean_user_id) REFERENCES users(id) ON DELETE SET NULL,
            FOREIGN KEY(deputy_user_id) REFERENCES users(id) ON DELETE SET NULL
        )""")
        ensure_column(conn, "transport_trips", "vehicle_type TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "transport_trips", "number_plate TEXT NOT NULL DEFAULT ''")
        conn.execute("""CREATE TABLE IF NOT EXISTS student_face_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER UNIQUE, user_id INTEGER UNIQUE, image_path TEXT NOT NULL DEFAULT '',
            descriptor_json TEXT NOT NULL DEFAULT '', active INTEGER NOT NULL DEFAULT 1, enrolled_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE, FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        )""")
        ensure_column(conn, "reception_visits", "id_number TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "reception_visits", "visit_reason_code TEXT NOT NULL DEFAULT ''")
        conn.execute("CREATE TABLE IF NOT EXISTS demo_seed_meta (id INTEGER PRIMARY KEY CHECK(id=1), version INTEGER NOT NULL DEFAULT 0, seeded_at TEXT)")
        ensure_column(conn, "assignments", "allowed_types TEXT NOT NULL DEFAULT 'pdf,doc,docx,xls,xlsx,ppt,pptx,csv,txt,png,jpg,jpeg,webp,zip'")
        ensure_column(conn, "assignments", "max_submissions INTEGER NOT NULL DEFAULT 2")
        ensure_column(conn, "assignments", "allow_any_file INTEGER NOT NULL DEFAULT 0")
        ensure_column(conn, "class_sessions", "audience_mode TEXT NOT NULL DEFAULT 'Class'")
        ensure_column(conn, "class_sessions", "recording_path TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "class_sessions", "recording_url TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "class_sessions", "saved_at TEXT")
        ensure_column(conn, "class_sessions", "library_item_id INTEGER")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_compulsory_subjects_class ON compulsory_subjects(class_name,active)")
        ensure_column(conn, "users", "workspace_type TEXT NOT NULL DEFAULT 'Teaching'")
        ensure_column(conn, "school_settings", "offline_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "backup_reminder_time TEXT NOT NULL DEFAULT '16:00'")
        ensure_column(conn, "school_settings", "backup_auto_time TEXT NOT NULL DEFAULT '16:30'")
        ensure_column(conn, "school_settings", "tracking_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "scanner_enabled INTEGER NOT NULL DEFAULT 1")
        ensure_column(conn, "school_settings", "institution_portal_guide TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_admin_guide TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_ict_guide TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_finance_guide TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_driver_guide_en TEXT NOT NULL DEFAULT ''")
        ensure_column(conn, "school_settings", "institution_driver_guide_sw TEXT NOT NULL DEFAULT ''")
        conn.execute("""CREATE TABLE IF NOT EXISTS theme_snapshots (id INTEGER PRIMARY KEY AUTOINCREMENT, snapshot_type TEXT NOT NULL, settings_json TEXT NOT NULL, created_by INTEGER, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, FOREIGN KEY(created_by) REFERENCES users(id) ON DELETE SET NULL)""")
        conn.execute("INSERT OR IGNORE INTO finance_accounts(id, account_name, opening_balance) VALUES(1, 'Institution Operating Account', 0)")
        conn.execute("UPDATE users SET qr_access_token=lower(hex(randomblob(16))) WHERE role!='System' AND (qr_access_token IS NULL OR qr_access_token='')")
        conn.execute("UPDATE users SET qr_login_enabled=0 WHERE qr_login_enabled IS NULL")
        if conn.execute("SELECT COUNT(*) FROM system_help").fetchone()[0] == 0:
            help_rows=[
                ('Getting started','Getting Started','Use the navigation to open your workspace. Administrators manage people, institution settings, security, backups and permissions.','All',10),
                ('Managing records','Staff & Learners','Use Staff for workforce accounts and Academics for learner and parent records. Role controls the security boundary; title and department describe institutional responsibility.','Admin',20),
                ('Learner records','Learners','Learner profiles contain identity, class/grade, guardian links, contact details, accountability information, academic and finance references.','All',30),
                ('Finance posting','Finance','Finance officers can post institutional income, expenses and payroll. Posted transactions are immutable; only an Administrator may reverse one.','Finance,Admin',40),
                ('Library resources','Library','Library staff can catalogue books, past papers, notes, images, uploaded documents, YouTube resources and external websites and assign resources to classes or subjects.','Teacher,Student,Parent,Librarian,Admin,ICT',50),
                ('Backup and recovery','Administration','A full JSON backup contains database records, system settings and available uploaded assets. Restore it through Administration > Backup & Recovery. Keep an additional SQLite backup for fast rollback.','Admin',60),
                ('AI assistant','AI','AI features require an OpenAI API key on the server environment. The assistant uses the configured provider/model and records a small audit preview without storing API keys in the database.','All',70),
            ]
            conn.executemany('INSERT INTO system_help(title,category,content,role_scope,sort_order) VALUES(?,?,?,?,?)',help_rows)

        guide_seed = [
            ("portal-user-guide", "Portal Guide", "A professional orientation to the institution portal: sign in through the appropriate role, use the workspace navigation to reach academic, communication, attendance and support services, review notifications and important dates routinely, protect your credentials, and use the institution help centre when a process requires assistance. Administrative, ICT and Finance procedures are intentionally kept within their restricted workspaces and are not published as public navigation links.", "All", 5),
            ("admin-guide", "Administrator Guide", "Administrators manage institutional structure, accounts, permissions, security controls, backups, public-facing information and system-wide visual settings. Use Staff for workforce account lifecycle management and Academics for learner and parent records; use institution and branding controls for public content; review audit and backup facilities regularly; and use direct account access only for legitimate support or oversight. Restricted administrator paths are deliberately omitted from public guides.", "Admin", 15),
            ("ict-guide", "ICT Guide", "ICT personnel maintain technical operations, user support, attendance technology, device access, public landing-page presentation and institution-wide workspace settings. Use the ICT control deck for technical configuration and support tasks. Avoid exposing administrator-only security controls or finance processes to ordinary portal users.", "ICT", 25),
            ("finance-guide", "Finance Guide", "Finance personnel manage approved fee structures, payments, ledger entries, institutional income, expenditure, payroll and daily handovers. Post records against the correct category and reference, preserve supporting documentation, and escalate reversals or exceptional corrections to an Administrator. Financial controls are intentionally concealed from public and teaching workspaces.", "Finance", 35),
            ("driver-guide", "Driver Guide", "English: Start a trip with the assigned vehicle and route, keep the trip active only while operating the institution journey, share location only while on duty when enabled, use attendance check-in/out when required, review notifications and important dates, and stop the trip when the journey ends.\n\nKiswahili: Anzisha safari kwa gari na njia uliyopewa, acha safari ikiwa hai tu unapokuwa unaendesha safari ya taasisi, tumia ushiriki wa eneo ukiwa kazini unapowashwa, tumia kuingia/kutoka kazini inapohitajika, angalia taarifa na tarehe muhimu, kisha simamisha safari safari inapokamilika.", "Driver", 45),
        ]
        for slug,title,content,scope,sort_order in guide_seed:
            if not conn.execute("SELECT 1 FROM system_help WHERE title=? LIMIT 1", (title,)).fetchone():
                conn.execute('INSERT INTO system_help(title,category,content,role_scope,sort_order) VALUES(?,?,?,?,?)',(title,"Portal Guidance",content,scope,sort_order))

        # Convert legacy seeded accounts to an inert System account on this build.
        auth_row = conn.execute("SELECT auth_initialized FROM school_settings WHERE id=1").fetchone()
        auth_ready = bool(auth_row and auth_row["auth_initialized"])
        admin_count = conn.execute("SELECT COUNT(*) AS c FROM users WHERE role='Admin' AND active=1").fetchone()["c"]
        if admin_count == 0 and not auth_ready:
            # New installations start with the non-Administrator login page disabled.
            conn.execute("UPDATE school_settings SET auth_required=0 WHERE id=1")
            # Preserve FK-backed sample data by using a single disabled system identity.
            system = conn.execute("SELECT id FROM users WHERE role='System' LIMIT 1").fetchone()
            if not system:
                conn.execute("INSERT INTO users(full_name,username,password_hash,role,active) VALUES(?,?,?,?,0)",
                             ("Portal System", "__system__", generate_password_hash(uuid.uuid4().hex), SYSTEM_ROLE))
                system_id = conn.execute("SELECT id FROM users WHERE username='__system__'").fetchone()["id"]
            else:
                system_id = system["id"]
            # Repoint legacy references to system account, then remove legacy accounts.
            for table, col in [("payments","recorded_by"),("audit_log","actor_id"),("exam_batches","submitted_by"),("exam_results","submitted_by"),("assignments","posted_by")]:
                try:
                    conn.execute(f"UPDATE {table} SET {col}=? WHERE {col} IS NOT NULL", (system_id,))
                except sqlite3.OperationalError:
                    pass
            conn.execute("DELETE FROM users WHERE role!='System'")
            conn.execute("UPDATE school_settings SET auth_initialized=0 WHERE id=1")
        elif admin_count > 0:
            conn.execute("UPDATE school_settings SET auth_initialized=1 WHERE id=1")

        # Repair an earlier authentication migration that accidentally left
        # payments/audit_log foreign keys pointing at users_old_auth. Preserve data.
        def repair_user_foreign_keys(table: str, create_sql: str, columns: list[str]) -> None:
            tbl = conn.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone()
            sql_text = (tbl[0] or "") if tbl else ""
            if "users_old_auth" not in sql_text:
                return
            conn.execute("PRAGMA foreign_keys=OFF")
            tmp = f"{table}_userfk_fix"
            conn.execute(f"DROP TABLE IF EXISTS {tmp}")
            conn.execute(create_sql.format(tmp=tmp))
            conn.execute(f"INSERT INTO {tmp} ({','.join(columns)}) SELECT {','.join(columns)} FROM {table}")
            conn.execute(f"DROP TABLE {table}")
            conn.execute(f"ALTER TABLE {tmp} RENAME TO {table}")
            conn.execute("PRAGMA foreign_keys=ON")

        repair_user_foreign_keys(
            "payments",
            """CREATE TABLE {tmp} (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER NOT NULL,
                amount REAL NOT NULL,
                method TEXT NOT NULL,
                reference_no TEXT,
                recorded_by INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Posted' CHECK(status IN ('Posted','Reversed')),
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                reversed_at TEXT,
                receipt_path TEXT,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(recorded_by) REFERENCES users(id) ON DELETE CASCADE
            )""",
            ["id","student_id","amount","method","reference_no","recorded_by","status","created_at","reversed_at","receipt_path"],
        )
        repair_user_foreign_keys(
            "audit_log",
            """CREATE TABLE {tmp} (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                actor_id INTEGER,
                actor_name TEXT NOT NULL,
                action TEXT NOT NULL,
                details TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(actor_id) REFERENCES users(id) ON DELETE SET NULL
            )""",
            ["id","actor_id","actor_name","action","details","created_at"],
        )


        # Remove stale SQLite triggers/objects left by older authentication migrations.
        # Some older builds created references to temporary users_roles_legacy/users_old_auth tables.
        stale_rows = conn.execute("SELECT type, name, sql FROM sqlite_master WHERE sql IS NOT NULL").fetchall()
        for obj_type, obj_name, obj_sql in stale_rows:
            sql_text = (obj_sql or "")
            if "users_roles_legacy" in sql_text or "users_old_auth" in sql_text or "users_legacy" in sql_text:
                if obj_type == "trigger":
                    conn.execute(f"DROP TRIGGER IF EXISTS [{obj_name}]")

        # Provision the production Administrator from environment variables when supplied.
        # This keeps credentials out of source control and makes Render deployment deterministic.
        env_admin_username = os.environ.get("ADMIN_USERNAME", "").strip()
        env_admin_password = os.environ.get("ADMIN_PASSWORD", "")
        env_admin_name = os.environ.get("ADMIN_NAME", "").strip()
        if env_admin_username and env_admin_password and env_admin_name:
            existing = conn.execute(
                "SELECT id, role FROM users WHERE username=? LIMIT 1",
                (env_admin_username,),
            ).fetchone()
            password_hash = generate_password_hash(env_admin_password)
            if existing:
                if existing["role"] != "Admin":
                    raise RuntimeError(
                        f"ADMIN_USERNAME '{env_admin_username}' already belongs to a non-Administrator account."
                    )
                conn.execute(
                    "UPDATE users SET full_name=?, password_hash=?, active=1 WHERE id=?",
                    (env_admin_name, password_hash, existing["id"]),
                )
            else:
                conn.execute(
                    "INSERT INTO users(full_name, username, password_hash, role, active) VALUES(?,?,?,?,1)",
                    (env_admin_name, env_admin_username, password_hash, "Admin"),
                )
            conn.execute("UPDATE school_settings SET auth_initialized=1, auth_required=1 WHERE id=1")

        # Optional coherent demonstration institution. Set SEED_DEMO_DATA=0 in production to disable.
        if os.environ.get("SEED_DEMO_DATA", "0") != "0":
            meta = conn.execute("SELECT version FROM demo_seed_meta WHERE id=1").fetchone()
            if not meta or int(meta["version"] or 0) < 2:
                demo_teacher_pw = generate_password_hash(os.environ.get("DEMO_TEACHER_PASSWORD", "DemoTeacher@123"))
                demo_student_pw = generate_password_hash(os.environ.get("DEMO_STUDENT_PASSWORD", "DemoStudent@123"))
                demo_admin_pw = generate_password_hash(os.environ.get("DEMO_ADMIN_PASSWORD", "DemoAdmin@123"))
                # Ensure an administrator exists for the demo controls. Existing admin credentials are never overwritten.
                admin = conn.execute("SELECT id FROM users WHERE role='Admin' AND active=1 LIMIT 1").fetchone()
                if not admin:
                    conn.execute("INSERT INTO users(full_name,username,password_hash,role,active) VALUES(?,?,?,?,1)",("Demo Administrator","demo.admin",demo_admin_pw,"Admin"))
                    admin = conn.execute("SELECT id FROM users WHERE username='demo.admin'").fetchone()
                admin_id=admin["id"]
                def ensure_user(full_name, username, pw, role, student_id=None, title='', dept='', workspace='Teaching'):
                    row=conn.execute("SELECT id FROM users WHERE username=? LIMIT 1",(username,)).fetchone()
                    if row:
                        return row["id"]
                    conn.execute("INSERT INTO users(full_name,username,password_hash,role,student_id,active,title,department,workspace_type,school_unit,school_location,position_code,staff_code,reception_enabled) VALUES(?,?,?,?,?,1,?,?,?,?,?,?,?,0)",(full_name,username,pw,role,student_id,title,dept,workspace,"Demo Institution","Juja, Kiambu, Kenya",role[:3].upper(),role[:3].upper()))
                    return conn.execute("SELECT id FROM users WHERE username=?",(username,)).fetchone()["id"]
                # Learners by level/class, deliberately realistic enough for a presentation.
                demo_students=[
                    ("PP1-001","Amara Wanjiku","PP1","Peter Wanjiku","0711000001",18000,"Pending"),
                    ("PP1-002","Brian Mwangi","PP1","Mary Mwangi","0711000002",0,"Paid"),
                    ("G4-001","Amani Kariuki","Grade 4","Jane Kariuki","0711000003",8500,"Pending"),
                    ("G4-002","Brian Otieno","Grade 4","Otieno Family","0711000004",0,"Paid"),
                    ("G4-003","Cynthia Njeri","Grade 4","Njeri Family","0711000005",4200,"Pending"),
                    ("G4-004","David Kamau","Grade 4","Kamau Family","0711000006",0,"Paid"),
                    ("G4-005","Elijah Maina","Grade 4","Maina Family","0711000007",12600,"Pending"),
                    ("G4-006","Faith Atieno","Grade 4","Atieno Family","0711000008",0,"Paid"),
                    ("G4-007","Grace Wambui","Grade 4","Wambui Family","0711000009",3500,"Pending"),
                    ("G5-001","Hassan Ali","Grade 5","Ali Family","0711000010",0,"Paid"),
                    ("G5-002","Ivy Chebet","Grade 5","Chebet Family","0711000011",6000,"Pending"),
                    ("G5-003","Joel Kiptoo","Grade 5","Kiptoo Family","0711000012",0,"Paid"),
                    ("G5-004","Lydia Akinyi","Grade 5","Akinyi Family","0711000013",20000,"Pending"),
                ]
                for adm,name,grade,guardian,phone,balance,status in demo_students:
                    conn.execute("INSERT OR IGNORE INTO students(admission_no,full_name,grade,guardian_name,guardian_phone,balance,payment_status,active) VALUES(?,?,?,?,?,?,?,1)",(adm,name,grade,guardian,phone,balance,status))
                student_rows={r["admission_no"]:r for r in conn.execute("SELECT * FROM students WHERE admission_no IN (%s)" % ','.join('?'*len(demo_students)),[x[0] for x in demo_students])}
                teacher_id=ensure_user("James Mwangi","demo.teacher",demo_teacher_pw,"Teacher",None,"Mathematics & ICT Teacher","Academics","Teaching")
                student_id=student_rows["G4-001"]["id"]
                ensure_user("Amani Kariuki","demo.student",demo_student_pw,"Student",student_id,"","","Student")
                conn.execute("UPDATE users SET workspace_type='Teaching', school_unit='Demo Institution', school_location='Juja, Kiambu, Kenya' WHERE id=?",(teacher_id,))
                for adm,name,grade,guardian,phone,balance,status in demo_students[1:]:
                    sid=student_rows[adm]["id"]
                    uname='demo.'+adm.lower().replace('-','')
                    if not conn.execute("SELECT 1 FROM users WHERE username=?",(uname,)).fetchone():
                        ensure_user(name,uname,demo_student_pw,"Student",sid,'','','Student')
                # Keep the demo student usable through both the named demo credentials and the same bootstrap convention used for normal learners.
                for _adm, _name, _grade, _guardian, _phone, _balance, _status in demo_students:
                    _sid=student_rows[_adm]["id"]
                    _u=conn.execute("SELECT id,username FROM users WHERE student_id=? AND role='Student' AND active=1 LIMIT 1",(_sid,)).fetchone()
                    if _u and _u["username"] != "demo.student":
                        conn.execute("UPDATE users SET password_hash=? WHERE id=?",(generate_password_hash(_adm),_u["id"]))
                # Teaching load and class-teacher authority.
                for cls,subj,unit in [("Grade 4","Mathematics","MAT-G4"),("Grade 4","Integrated Science","SCI-G4"),("Grade 5","Mathematics","MAT-G5")]:
                    conn.execute("INSERT OR IGNORE INTO teacher_assignments(teacher_user_id,class_name,subject,unit_code,active) VALUES(?,?,?,?,1)",(teacher_id,cls,subj,unit))
                conn.execute("INSERT INTO class_teacher_assignments(class_name,teacher_user_id,assigned_by) VALUES('Grade 4',?,?) ON CONFLICT(class_name) DO UPDATE SET teacher_user_id=excluded.teacher_user_id,assigned_by=excluded.assigned_by",(teacher_id,admin_id))
                conn.execute("INSERT OR IGNORE INTO compulsory_subjects(class_name,subject,unit_name,created_by) VALUES('Grade 4','Mathematics','Core Mathematics',?)",(admin_id,))
                conn.execute("INSERT OR IGNORE INTO compulsory_subjects(class_name,subject,unit_name,created_by) VALUES('Grade 5','Mathematics','Core Mathematics',?)",(admin_id,))
                for dept_name,dept_cat in [('Sciences','Academic'),('Languages','Academic'),('Humanities','Academic'),('ICT','Academic'),('Arts','Academic')]:
                    conn.execute("INSERT OR IGNORE INTO departments(name,category) VALUES(?,?)",(dept_name,dept_cat))
                demo_subjects=[('Mathematics','Sciences'),('English','Languages'),('Integrated Science','Sciences'),('Social Studies','Humanities'),('Computer Studies','ICT'),('Creative Arts','Arts')]
                for subj,dept in demo_subjects:
                    conn.execute("INSERT OR IGNORE INTO subjects_catalog(subject,department,level_scope,created_by) VALUES(?,?,?,?)",(subj,dept,'All',admin_id))
                cat={r['subject']:r['id'] for r in conn.execute("SELECT id,subject FROM subjects_catalog WHERE subject IN (?,?,?,?,?,?)",tuple(x[0] for x in demo_subjects))}
                for adm,name,grade,guardian,phone,balance,status in demo_students:
                    sid=student_rows[adm]['id']
                    subs=['Mathematics','English','Integrated Science','Social Studies','Computer Studies','Creative Arts'] if grade=='Grade 4' else ['Mathematics','English','Integrated Science','Social Studies']
                    for subj in subs:
                        conn.execute("INSERT OR IGNORE INTO student_subjects(student_id,subject_id,status,selected_by) VALUES(?,?,'Approved',?)",(sid,cat[subj],admin_id))
                # Sample weighted marks for visible ordering.
                for idx,r in enumerate(conn.execute("SELECT id FROM students WHERE grade='Grade 4' AND active=1 ORDER BY full_name"),1):
                    for assessment,mark,max_mark,weight in [("CAT 1",55+((idx*7)%36),100,20),("CAT 2",48+((idx*9)%46),100,20),("Term Exam",50+((idx*11)%48),100,60)]:
                        exists=conn.execute("SELECT 1 FROM markbook_entries WHERE teacher_user_id=? AND student_id=? AND subject='Mathematics' AND assessment=?",(teacher_id,r["id"],assessment)).fetchone()
                        if not exists: conn.execute("INSERT INTO markbook_entries(teacher_user_id,class_name,subject,student_id,assessment,mark,max_mark,status,weight) VALUES(?,?,?,?,?,?,?,?,?)",(teacher_id,"Grade 4","Mathematics",r["id"],assessment,mark,max_mark,"Submitted",weight))
                # One real assignment, submission history and a scheduled lesson.
                assignment=conn.execute("SELECT id FROM assignments WHERE title='Fractions Practice & Reflection' AND grade='Grade 4' LIMIT 1").fetchone()
                if not assignment:
                    aid=conn.execute("INSERT INTO assignments(title,subject,grade,description,deadline,posted_by,allowed_types,max_submissions,allow_any_file) VALUES(?,?,?,?,?,?,?,?,0)",("Fractions Practice & Reflection","Mathematics","Grade 4","Complete the worked examples, then submit your reflection. Two attempts are permitted before the deadline.",(datetime.now()+timedelta(days=3)).strftime('%Y-%m-%dT%H:%M'),teacher_id,"pdf,doc,docx,xls,xlsx,ppt,pptx,csv,txt,png,jpg,jpeg,webp",2)).lastrowid
                    for n,note in [(1,"First attempt: worked answers attached."),(2,"Second attempt: corrected question 6 and 8.")]:
                        if n==1:
                            conn.execute("INSERT INTO submissions(assignment_id,student_id,attachment_path,note) VALUES(?,?,?,?)",(aid,student_id,"uploads/demo-fractions-attempt1.pdf",note))
                        else:
                            conn.execute("INSERT INTO submissions(assignment_id,student_id,attachment_path,note,score,feedback) VALUES(?,?,?,?,?,?)",(aid,student_id,"uploads/demo-fractions-attempt2.pdf",note,92,"Excellent correction and clear working."))
                session=conn.execute("SELECT id FROM class_sessions WHERE title='Grade 4 Live Mathematics Clinic' LIMIT 1").fetchone()
                if not session:
                    start=(datetime.now()+timedelta(hours=3)).replace(second=0,microsecond=0)
                    end=start+timedelta(minutes=50)
                    room=f"Demo-Institution-Grade-4-{uuid.uuid4().hex[:8]}"
                    conn.execute("INSERT INTO class_sessions(teacher_user_id,class_name,subject,title,starts_at,ends_at,room_name,provider_url,description,audience_mode) VALUES(?,?,?,?,?,?,?,?,?,?)",(teacher_id,"Grade 4","Mathematics","Grade 4 Live Mathematics Clinic",start.strftime('%Y-%m-%d %H:%M'),end.strftime('%Y-%m-%d %H:%M'),room,"https://meet.jit.si/","Live worked examples, Q&A and revision clinic.","Compulsory"))
                    sess_id=conn.execute("SELECT id FROM class_sessions WHERE room_name=?",(room,)).fetchone()["id"]
                    ids=[r["id"] for r in conn.execute("SELECT u.id FROM users u JOIN students s ON s.id=u.student_id WHERE u.role='Student' AND u.active=1 AND lower(s.grade)=lower('Grade 4')")]
                    for uid in ids:
                        conn.execute("INSERT INTO notifications(user_id,title,body,link,priority) VALUES(?,?,?,?,?)",(uid,"Upcoming live Mathematics class",f"Grade 4 Live Mathematics Clinic starts {start.strftime('%A, %d %b at %H:%M')}. Join from your Student Dashboard.",f"/online-class/{sess_id}","High"))
                conn.execute("INSERT OR REPLACE INTO demo_seed_meta(id,version,seeded_at) VALUES(1,2,CURRENT_TIMESTAMP)")
                conn.execute("UPDATE school_settings SET school_name=COALESCE(NULLIF(school_name,''),'Prime Demonstration Institution'), school_fee=CASE WHEN school_fee=0 THEN 20000 ELSE school_fee END WHERE id=1")
                conn.execute("UPDATE school_settings SET auth_initialized=1, auth_required=1 WHERE id=1")
        # Force SQLite to materialize/validate the final schema after migration cleanup.
        conn.execute("PRAGMA foreign_key_check")
        conn.commit()


def init_db() -> None:
    """Initialize SQLite, with a hard recovery path for malformed Render databases."""
    try:
        _init_db_once()
        return
    except sqlite3.DatabaseError as exc:
        # Render can retain a damaged /var/data/school.db.  A PRAGMA integrity
        # check is not sufficient protection because corruption can surface only
        # when a schema statement touches a damaged page/index.  Therefore any
        # SQLite DatabaseError during startup triggers a full, deterministic
        # recovery from the bundled known-good seed database.
        recovery_message = str(exc)

    legacy_db = INSTANCE_DIR / "school.db"
    if not legacy_db.exists() or not _sqlite_integrity_ok(legacy_db):
        # Do not silently continue with a known-bad seed database.  Re-raise a
        # clear startup error so the deployment failure points to the package
        # itself rather than the persistent Render disk.
        raise sqlite3.DatabaseError(
            f"Prime database startup failed and bundled recovery database is unavailable. Original error: {recovery_message}"
        )

    # Preserve the broken persistent database for later forensic recovery, then
    # replace it with the known-good bundled database.  This is intentionally
    # destructive to the active persistent DB path only; the damaged file is
    # retained with a .corrupt-* name first.
    try:
        _quarantine_corrupt_db(DB_PATH)
    except Exception:
        # If rename is unavailable but the file can be removed, continue.  The
        # bundled seed is the authoritative recovery source.
        try:
            DB_PATH.unlink(missing_ok=True)
        except OSError:
            pass
        for suffix in ("-wal", "-shm"):
            try:
                DB_PATH.with_name(DB_PATH.name + suffix).unlink(missing_ok=True)
            except OSError:
                pass

    import shutil
    shutil.copy2(legacy_db, DB_PATH)

    # Never carry WAL/SHM files from the corrupted database into the replacement.
    for suffix in ("-wal", "-shm"):
        try:
            DB_PATH.with_name(DB_PATH.name + suffix).unlink(missing_ok=True)
        except OSError:
            pass

    # The retry is deliberately uncaught: if the bundled seed is somehow broken,
    # we want Render to report that real problem instead of looping indefinitely.
    _init_db_once()

def q(sql: str, params: Iterable[Any] = (), one: bool = False):
    cur = get_db().execute(sql, tuple(params))
    rows = cur.fetchall()
    cur.close()
    if one:
        return rows[0] if rows else None
    return rows


def execute(sql: str, params: Iterable[Any] = ()) -> int:
    db = get_db()
    cur = db.execute(sql, tuple(params))
    db.commit()
    return cur.lastrowid



def recalculate_student_balance(student_id: int) -> dict:
    """Rebuild a learner's account from assessed charges and posted payments.
    Manual fee override wins when enabled; otherwise active fee charges are authoritative.
    """
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True)
    if not student:
        return {"assessed":0.0,"paid":0.0,"balance":0.0,"status":"Pending"}
    if int(student["fee_override_enabled"] or 0):
        assessed=float(student["fee_assessed_total"] or 0)
    else:
        assessed=float(q("SELECT COALESCE(SUM(amount),0) AS n FROM fee_charges WHERE student_id=? AND status='Posted'",(student_id,),one=True)["n"] or 0)
        # Preserve legacy manually entered balances until the first structured charge is applied.
        if assessed == 0 and float(student["fee_assessed_total"] or 0) > 0:
            assessed=float(student["fee_assessed_total"] or 0)
    paid=float(q("SELECT COALESCE(SUM(amount),0) AS n FROM payments WHERE student_id=? AND status='Posted'",(student_id,),one=True)["n"] or 0)
    balance=assessed-paid
    status="Paid" if balance <= 0 else "Pending"
    execute("UPDATE students SET fee_assessed_total=?, balance=?, payment_status=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",(assessed,balance,status,student_id))
    return {"assessed":assessed,"paid":paid,"balance":balance,"status":status}

def student_email_for(admission_no: str, existing: str = "") -> str:
    existing=(existing or "").strip().lower()
    if existing:
        return existing
    domain=(school_settings()["student_email_domain"] or "school.ac.ke").strip().lstrip("@").lower()
    local=re.sub(r"[^a-z0-9]+",".",(admission_no or "student").lower()).strip(".") or "student"
    return f"{local}@{domain}"

def next_grade_for(grade: str) -> str:
    raw=(grade or "").strip().lower().replace("_"," ").replace("-"," ")
    mapping={
        "playgroup":"PP1","play group":"PP1","pp1":"PP2","pp2":"Grade 1",
        "grade 1":"Grade 2","grade 2":"Grade 3","grade 3":"Grade 4","grade 4":"Grade 5",
        "grade 5":"Grade 6","grade 6":"Grade 7","grade 7":"Grade 8","grade 8":"Grade 9",
        "grade 9":"Grade 10","grade 10":"Grade 11","grade 11":"Grade 12"
    }
    return mapping.get(raw, "")

def return_to_referrer(fallback_endpoint: str):
    ref=request.referrer
    if ref and ref.startswith(request.host_url.rstrip("/") + "/"):
        return ref
    return url_for(fallback_endpoint)

def school_settings():
    """Return the single school-wide portal settings row, creating it when needed."""
    row = q("SELECT * FROM school_settings WHERE id = 1", one=True)
    if row:
        return row
    execute(
        """
        INSERT INTO school_settings(
            id, school_name, admission_prefix, admission_suffix,
            student_name_prefix, student_name_suffix, currency_code, school_fee,
            auth_required, prelogin_sections, developer_name, company_name
        ) VALUES (1, 'School Portal System', 'ADM-', '', '', '', 'KES', 0, 0, 'institution,history,achievements,owners,developer,company', 'Toror Technology and Innovations Ltd.', 'Toror Technology and Innovations Ltd.')
        """
    )
    return q("SELECT * FROM school_settings WHERE id = 1", one=True)


def audit(actor_id: int | None, actor_name: str, action: str, details: str) -> None:
    execute(
        "INSERT INTO audit_log(actor_id, actor_name, action, details) VALUES (?, ?, ?, ?)",
        (actor_id, actor_name, action, details),
    )


def record_login_event(user, method='Password', latitude=None, longitude=None, accuracy=None):
    ip=request.headers.get('X-Forwarded-For', request.remote_addr or '')
    if ',' in ip: ip=ip.split(',')[0].strip()
    ua=(request.headers.get('User-Agent') or '')[:500]
    return execute("INSERT INTO login_events(user_id,method,ip_address,user_agent,latitude,longitude,accuracy) VALUES(?,?,?,?,?,?,?)",(user['id'],method,ip[:120],ua,latitude,longitude,accuracy))


# -------------------------
# Local role session helpers
# -------------------------
def _auth_serializer():
    return URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="school-auth-v3")

def _auth_token_for(user_id: int) -> str:
    return _auth_serializer().dumps({"uid": int(user_id)})

def _user_from_auth_token(token: str):
    try:
        data = _auth_serializer().loads(token, max_age=_AUTH_COOKIE_MAX_AGE)
        uid = int(data.get("uid", 0))
    except (BadSignature, SignatureExpired, ValueError, TypeError, AttributeError):
        return None
    if not uid:
        return None
    return q("SELECT id, full_name, username, role, student_id, active, title, department, leadership_role, leadership_level, workspace_type, school_unit, school_location, position_code, staff_code, reception_enabled, qr_access_token, profile_photo FROM users WHERE id = ? AND active = 1 AND role != 'System'", (uid,), one=True)

def _portal_context_serializer():
    return URLSafeTimedSerializer(app.config["SECRET_KEY"], salt=_PORTAL_CONTEXT_SALT)

def _ensure_portal_context_table():
    execute("CREATE TABLE IF NOT EXISTS portal_contexts (token_id TEXT PRIMARY KEY, user_id INTEGER NOT NULL, created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP, revoked INTEGER NOT NULL DEFAULT 0, FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE)")

def _portal_context_for(user_id: int) -> str:
    _ensure_portal_context_table()
    token_id = uuid.uuid4().hex
    execute("INSERT INTO portal_contexts(token_id,user_id) VALUES(?,?)", (token_id, int(user_id)))
    return _portal_context_serializer().dumps({"uid": int(user_id), "tid": token_id})

def _user_from_portal_context(token: str):
    try:
        data = _portal_context_serializer().loads(token, max_age=_PORTAL_CONTEXT_MAX_AGE)
        uid = int(data.get("uid", 0)); tid = str(data.get("tid", ""))
    except (BadSignature, SignatureExpired, ValueError, TypeError, AttributeError):
        return None
    if not uid or not tid:
        return None
    _ensure_portal_context_table()
    valid = q("SELECT token_id FROM portal_contexts WHERE token_id=? AND user_id=? AND revoked=0", (tid, uid), one=True)
    if not valid:
        return None
    return q("SELECT id, full_name, username, role, student_id, active, title, department, leadership_role, leadership_level, workspace_type, school_unit, school_location, position_code, staff_code, reception_enabled, qr_access_token, profile_photo FROM users WHERE id=? AND active=1 AND role!='System'", (uid,), one=True)

def _portal_context_id(token: str):
    try:
        data = _portal_context_serializer().loads(token, max_age=_PORTAL_CONTEXT_MAX_AGE)
        return str(data.get("tid", "")) or None
    except Exception:
        return None

def _auth_ticket_hash(token: str) -> str:
    return hashlib.sha256((token or "").encode("utf-8")).hexdigest()

def _issue_auth_ticket(user_id: int) -> str:
    raw = secrets.token_urlsafe(48)
    token_hash = _auth_ticket_hash(raw)
    expires = datetime.utcnow() + timedelta(seconds=_AUTH_COOKIE_MAX_AGE)
    execute("INSERT INTO auth_sessions(token_hash,user_id,expires_at,revoked) VALUES(?,?,?,0)", (token_hash, int(user_id), expires.isoformat(timespec="seconds")))
    return raw

def _user_from_auth_ticket(token: str):
    if not token:
        return None
    row = q("""SELECT u.id, u.full_name, u.username, u.role, u.student_id, u.active, u.title, u.department,
                     u.leadership_role, u.leadership_level, u.workspace_type, u.school_unit, u.school_location,
                     u.position_code, u.staff_code, u.reception_enabled, u.qr_access_token, u.profile_photo,
                     a.expires_at
              FROM auth_sessions a JOIN users u ON u.id=a.user_id
              WHERE a.token_hash=? AND a.revoked=0 AND u.active=1 AND u.role!='System'
                AND datetime(a.expires_at) > datetime('now') LIMIT 1""", (_auth_ticket_hash(token),), one=True)
    if row:
        execute("UPDATE auth_sessions SET last_seen_at=CURRENT_TIMESTAMP WHERE token_hash=?", (_auth_ticket_hash(token),))
    return row

def _revoke_auth_ticket(token: str) -> None:
    if token:
        execute("UPDATE auth_sessions SET revoked=1 WHERE token_hash=?", (_auth_ticket_hash(token),))


@app.before_request
def load_current_user() -> None:
    g.user = None
    context_token = request.args.get("portal_context") or request.form.get("portal_context")
    if context_token:
        contextual = _user_from_portal_context(context_token)
        if contextual:
            g.user = contextual
            g.portal_context = context_token
            session.permanent = True
            # Administrator account-access links intentionally run in a portal
            # context without replacing the Admin session. This keeps the
            # original Admin dashboard usable in another tab while the selected
            # account is being viewed. Normal user login contexts still become
            # the active session as before.
            if not session.get("admin_impersonation"):
                session["user_id"] = contextual["id"]
                session["active_portal_role"] = contextual["role"]
            return
    # Server-backed auth ticket is the primary identity source. The signed
    # Flask session remains a convenience layer, but losing it can never log
    # the person out as long as their account is still active.
    ticket = request.cookies.get(_AUTH_TICKET_COOKIE, "")
    recovered = _user_from_auth_ticket(ticket)
    if recovered:
        g.user = recovered
        session.permanent = True
        session["user_id"] = recovered["id"]
        session["active_portal_role"] = recovered["role"]
        g.portal_context = None
        return
    user_id = session.get("user_id")
    if user_id:
        g.user = q("SELECT id, full_name, username, role, student_id, active, title, department, leadership_role, leadership_level, workspace_type, school_unit, school_location, position_code, staff_code, reception_enabled, qr_access_token, profile_photo FROM users WHERE id = ? AND active = 1 AND role != 'System'", (user_id,), one=True)
        if g.user:
            session.permanent = True
            g.portal_context = None
            return
    token = request.cookies.get(_AUTH_COOKIE, "")
    if token:
        recovered = _user_from_auth_token(token)
        if recovered:
            g.user = recovered
            session.permanent = True
            session["user_id"] = recovered["id"]
            g.portal_context = None

@app.after_request
def persist_auth_cookie(response):
    user = getattr(g, "user", None)
    if user and not getattr(g, "logging_out", False):
        # Keep a durable server-backed identity ticket independent of Flask
        # session mutations performed by unrelated Admin operations.
        ticket = request.cookies.get(_AUTH_TICKET_COOKIE, "")
        if not _user_from_auth_ticket(ticket):
            ticket = _issue_auth_ticket(user["id"])
        response.set_cookie(_AUTH_TICKET_COOKIE, ticket, max_age=_AUTH_COOKIE_MAX_AGE, httponly=True, secure=app.config.get("SESSION_COOKIE_SECURE", False), samesite="Lax", path="/")
        # Do not replace the administrator's persistent auth cookie while a
        # direct-access portal context is being viewed. The Admin session remains
        # the source of truth for the original command-centre tab.
        if not (getattr(g, "portal_context", None) and flask_session.get("admin_impersonation")):
            response.set_cookie(_AUTH_COOKIE, _auth_token_for(user["id"]), max_age=_AUTH_COOKIE_MAX_AGE, httponly=True, secure=app.config.get("SESSION_COOKIE_SECURE", False), samesite="Lax", path="/")
        if request.method == "GET" and response.content_type and response.content_type.startswith("text/html"):
            token = getattr(g, "portal_context", None)
            if token:
                try:
                    body = response.get_data(as_text=True)
                    marker = "</body>"
                    script = """<script>(function(){const t=%r;try{sessionStorage.setItem('prime_portal_context',t)}catch(e){};function apply(){document.querySelectorAll('a[href]').forEach(function(a){try{const u=new URL(a.href,location.href);if(u.origin===location.origin&&!u.searchParams.has('portal_context')&&!u.pathname.startsWith('/static/')){u.searchParams.set('portal_context',t);a.href=u.toString()}}catch(e){}});document.querySelectorAll('form[action]').forEach(function(f){if(!f.querySelector('input[name="portal_context"]')){const i=document.createElement('input');i.type='hidden';i.name='portal_context';i.value=t;f.appendChild(i)}})}apply();new MutationObserver(apply).observe(document.documentElement,{subtree:true,childList:true});})();</script>""" % token
                    if marker in body:
                        response.set_data(body.replace(marker, script+marker, 1))
                except Exception:
                    pass
    token = getattr(g, "portal_context", None)
    if token and not getattr(g, "logging_out", False) and response.status_code in {301,302,303,307,308}:
        location = response.headers.get("Location", "")
        if location:
            try:
                parsed = urllib.parse.urlsplit(location)
                if not parsed.netloc or parsed.netloc == request.host:
                    params = urllib.parse.parse_qs(parsed.query, keep_blank_values=True)
                    params.setdefault("portal_context", [token])
                    query = urllib.parse.urlencode(params, doseq=True)
                    response.headers["Location"] = urllib.parse.urlunsplit((parsed.scheme, parsed.netloc, parsed.path, query, parsed.fragment))
            except Exception:
                pass
    # Institution-wide notification bell + role shortcuts. Kept out of the public landing page.
    if getattr(g, "user", None) and request.path != "/" and not request.path.startswith("/static/") and response.content_type and response.content_type.startswith("text/html"):
        try:
            body=response.get_data(as_text=True)
            if 'id="prime-global-tools"' not in body and "</body>" in body:
                shell="""<button id="prime-mobile-nav" class="prime-mobile-nav" type="button" aria-label="Open navigation" aria-expanded="false" title="Open navigation"><span class="icon-bars" aria-hidden="true"><i></i><i></i><i></i></span></button><div id="prime-mobile-menu" class="prime-mobile-menu"><a href="/dashboard">Dashboard</a><a href="/calendar">School calendar</a><a href="/notifications">Notifications</a><a href="/system-help">System help</a><a href="/logout">Logout</a></div><div id="prime-global-tools" class="prime-global-tools"><a class="prime-search-link" href="/system-search" aria-label="Search system" title="Search system"><svg class="ui-icon" viewBox="0 0 24 24" aria-hidden="true"><circle cx="11" cy="11" r="6.5"></circle><path d="m16 16 4.5 4.5"></path></svg></a><a class="prime-bell" href="/notifications" aria-label="Notifications" title="Notifications"><svg class="ui-icon" viewBox="0 0 24 24" aria-hidden="true"><path d="M18 9a6 6 0 0 0-12 0c0 7-3 7-3 9h18c0-2-3-2-3-9"></path><path d="M10 21h4"></path></svg><b id="prime-notification-count" class="prime-count hidden"></b></a><button type="button" class="prime-shortcuts-btn" aria-label="Open shortcuts" onclick="document.getElementById('prime-shortcuts').classList.toggle('open')">☰</button><div id="prime-shortcuts" class="prime-shortcuts"><strong>Quick access</strong><a href="/calendar">Calendar</a><a href="/notifications">Notifications</a><a href="/online-classes">Live classes</a><a href="/groups">Groups</a><a href="/leadership">Leadership</a></div></div><button id="prime-mobile-text" class="prime-mobile-text" type="button" aria-label="Adjust text size" title="Adjust text size">Aa</button><div id="prime-text-sheet" class="prime-text-sheet" role="dialog" aria-modal="true" aria-label="Text size settings"><div class="prime-text-sheet-card"><div><strong>Text size</strong><span class="muted">Adjust this device only.</span></div><div class="prime-text-choices"><button type="button" data-prime-text="normal">Normal</button><button type="button" data-prime-text="large">Large</button><button type="button" data-prime-text="xlarge">Extra large</button></div><button type="button" class="btn btn-ghost btn-block" id="prime-text-close">Done</button></div></div><style>.prime-global-tools{position:fixed;right:18px;top:16px;z-index:5000;display:flex;gap:8px;align-items:flex-start;font-family:system-ui,sans-serif}.prime-search-link,.prime-bell,.prime-shortcuts-btn{width:42px;height:42px;border-radius:50%;display:grid;place-items:center;text-decoration:none;border:1px solid color-mix(in srgb,var(--primary-blue,#3457d5) 35%,transparent);background:var(--panel,#fff);color:var(--primary-text,#152033);box-shadow:0 8px 30px rgba(0,0,0,.18);cursor:pointer}.prime-search-link,.prime-bell{font-size:18px}.prime-bell span{color:inherit;font-size:18px;line-height:1}.prime-count{position:absolute;right:45px;top:-3px;min-width:17px;height:17px;padding:0 4px;border-radius:999px;background:#dc143c;color:#fff;font:700 10px/17px system-ui;text-align:center}.prime-count.dot{width:8px;min-width:8px;height:8px;padding:0;line-height:8px;right:47px}.prime-count.hidden{display:none}.prime-shortcuts{display:none;position:absolute;right:0;top:48px;min-width:190px;padding:10px;border-radius:14px;background:var(--panel,#fff);border:1px solid color-mix(in srgb,var(--primary-blue,#3457d5) 20%,transparent);box-shadow:0 18px 40px rgba(0,0,0,.22)}.prime-shortcuts.open{display:grid;gap:5px}.prime-shortcuts strong{padding:5px 8px}.prime-shortcuts a{padding:8px 10px;border-radius:9px;color:inherit;text-decoration:none}.prime-shortcuts a:hover{background:rgba(127,127,127,.12)}.prime-mobile-nav{display:none}.prime-mobile-nav.open{display:grid}.prime-mobile-text,.prime-text-sheet{display:none}body.auth-body .prime-global-tools,body.auth-body .prime-mobile-nav,body.auth-body .prime-mobile-menu,body.auth-body .prime-mobile-text{display:none}@media(max-width:820px){.prime-shortcuts-btn{display:none!important}.prime-mobile-nav{display:grid;place-items:center;position:fixed;left:12px;top:12px;width:46px;height:46px;border-radius:12px;border:1px solid var(--text-border,var(--border));background:var(--panel,#fff);color:var(--primary-text,#152033);box-shadow:0 10px 28px rgba(0,0,0,.20);font-size:22px;cursor:pointer;z-index:5001}.prime-global-tools{right:12px;top:12px}.prime-mobile-text{display:grid;place-items:center;position:fixed;right:64px;top:12px;width:46px;height:46px;border-radius:12px;border:1px solid var(--text-border,var(--border));background:var(--panel,#fff);color:var(--primary-text,#152033);box-shadow:0 10px 28px rgba(0,0,0,.20);font-size:15px;font-weight:900;cursor:pointer;z-index:5001}.prime-text-sheet{position:fixed;inset:0;background:rgba(0,0,0,.48);z-index:6000;align-items:flex-end;justify-content:center;padding:14px}.prime-text-sheet.open{display:flex}.prime-text-sheet-card{width:min(460px,100%);border:1px solid var(--text-border,var(--border));border-radius:20px 20px 14px 14px;background:var(--panel);box-shadow:0 -18px 50px rgba(0,0,0,.26);padding:18px;display:grid;gap:16px}.prime-text-sheet-card>div:first-child{display:grid;gap:4px}.prime-text-choices{display:grid;grid-template-columns:repeat(3,1fr);gap:8px}.prime-text-choices button{border:1px solid var(--text-border,var(--border));background:var(--panel-3);color:var(--primary-text);border-radius:12px;padding:12px 8px;font-weight:800;cursor:pointer}.prime-text-choices button.active{border-color:var(--primary-blue);box-shadow:0 0 0 3px rgba(16,163,127,.14)}}</style><script>(function(){var m=document.getElementById('prime-mobile-nav');if(m){if(document.getElementById('sidebarToggle')){m.remove();}else{m.addEventListener('click',function(e){e.stopPropagation();var hasSidebar=!!document.querySelector('.sidebar');document.body.classList.toggle(hasSidebar?'mobile-nav-open':'prime-smart-menu-open');});document.addEventListener('click',function(e){if(!m.contains(e.target)){document.body.classList.remove('mobile-nav-open','prime-smart-menu-open');}});}}var tb=document.getElementById('prime-mobile-text'),sheet=document.getElementById('prime-text-sheet'),close=document.getElementById('prime-text-close'),choices=document.querySelectorAll('[data-prime-text]');var saved=localStorage.getItem('prime_text_size')||'large';if(saved==='normal'||saved==='large'||saved==='xlarge')document.documentElement.dataset.primeText=saved;if(tb&&sheet){tb.addEventListener('click',function(){sheet.classList.add('open');});sheet.addEventListener('click',function(e){if(e.target===sheet)sheet.classList.remove('open');});close&&close.addEventListener('click',function(){sheet.classList.remove('open');});choices.forEach(function(b){b.classList.toggle('active',b.dataset.primeText===saved);b.addEventListener('click',function(){saved=b.dataset.primeText;localStorage.setItem('prime_text_size',saved);document.documentElement.dataset.primeText=saved;choices.forEach(function(x){x.classList.toggle('active',x.dataset.primeText===saved);});});});}fetch('/api/notifications').then(r=>r.json()).then(d=>{var n=document.getElementById('prime-notification-count');if(!n)return;var c=Number(d.count||0);if(c<=0){n.classList.add('hidden');return;}n.classList.remove('hidden');if(c>5){n.textContent='';n.classList.add('dot');}else{n.textContent=String(c);n.classList.remove('dot');}}).catch(function(){});})();</script>"""
                response.set_data(body.replace("</body>",shell+"</body>",1))
        except Exception:
            pass
    response.headers.setdefault("Permissions-Policy", "camera=(self), microphone=(self), geolocation=(self), display-capture=(self)")
    return response


def current_user():
    return getattr(g, "user", None)


def login_required(view: Callable):
    @wraps(view)
    def wrapper(*args, **kwargs):
        # Only an uninitialized presentation/demo installation may auto-seat a role.
        # Once the school has initialized authentication, every protected route
        # requires an actual authenticated account.
        if not current_user() and not auth_initialized() and not auth_required():
            _ensure_demo_identity()
        if not current_user():
            return redirect(url_for("login", role=request.args.get("role", "")))
        return view(*args, **kwargs)
    return wrapper


def role_required(*roles: str):
    def decorator(view: Callable):
        @wraps(view)
        def wrapper(*args, **kwargs):
            user = current_user()
            if not user:
                return redirect(url_for("login", role=(roles[0] if roles else "")))
            # Admin and ICT are deliberately separate workspaces. A role-specific
            # dashboard URL can never switch identities. When an authenticated user
            # reaches the other dashboard URL, send them back to their own dashboard
            # instead of producing a dead-end Access denied page.
            if view.__name__ in {"admin_dashboard", "ict_dashboard", "admin_entry", "ict_entry"}:
                expected = "Admin" if view.__name__.startswith("admin") else "ICT"
                if user["role"] != expected:
                    if user["role"] in ALL_PORTAL_ROLES:
                        return redirect(role_target(user["role"]))
                    abort(403)
            elif user["role"] not in roles:
                if user["role"] == "Teacher":
                    return redirect(url_for("teacher_dashboard"))
                abort(403)
            if user["role"] == "Teacher" and "Teacher" in roles and workspace_type_for_user(user) != "Teaching":
                return redirect(url_for("teacher_dashboard"))
            return view(*args, **kwargs)
        return wrapper
    return decorator


def auth_initialized() -> bool:
    row = q("SELECT auth_initialized FROM school_settings WHERE id=1", one=True)
    return bool(row and row["auth_initialized"])


# DEMO_AUTH_BYPASS is intentionally enabled for presentation/local testing.
# Set it to False before commercial deployment to restore the normal login gate.
DEMO_AUTH_BYPASS = False

def auth_required() -> bool:
    if DEMO_AUTH_BYPASS:
        return False
    row = q("SELECT auth_required, auth_initialized FROM school_settings WHERE id=1", one=True)
    return bool(row and (row["auth_required"] or row["auth_initialized"]))

def _demo_role_for_request() -> str:
    path = request.path.lower()
    endpoint = (request.endpoint or "").lower()
    rules = [
        (("/admin", "admin_"), "Admin"),
        (("/ict", "ict_"), "ICT"),
        (("/finance", "finance_"), "Finance"),
        (("/teacher", "teacher_", "/dashboard"), "Teacher"),
        (("/student", "student_"), "Student"),
        (("/parent", "parent_"), "Parent"),
        (("/librarian", "librarian_", "/library"), "Librarian"),
        (("/driver", "driver_"), "Driver"),
        (("/workforce", "workforce_"), "Other Staff"),
    ]
    for needles, role in rules:
        if any(n in path or n in endpoint for n in needles):
            return role
    return (session.get("active_portal_role") or "Admin")

def _ensure_demo_identity() -> None:
    if not DEMO_AUTH_BYPASS or current_user():
        return
    desired = _demo_role_for_request()
    if desired == "Parent" and not parent_portal_enabled():
        desired = "Admin"
    candidates = q("SELECT id FROM users WHERE role=? AND active=1 AND role!='System' ORDER BY id", (desired,))
    if not candidates and desired in {"Driver", "Guard", "Cook", "Other Staff"}:
        candidates = q("SELECT id FROM users WHERE active=1 AND role NOT IN ('System','Student','Parent') ORDER BY id")
    if not candidates:
        candidates = q("SELECT id FROM users WHERE role='Admin' AND active=1 ORDER BY id")
    if candidates:
        uid = candidates[0]["id"]
        session.permanent = True
        session["user_id"] = uid
        session["active_portal_role"] = desired
        g.user = q("SELECT id, full_name, username, role, student_id, active, title, department, leadership_role, leadership_level, workspace_type, school_unit, school_location, position_code, staff_code, reception_enabled, qr_access_token, profile_photo FROM users WHERE id=? AND active=1 AND role!='System'", (uid,), one=True)


def role_target(role: str) -> str:
    return {
        "Admin": url_for("admin_dashboard"),
        "ICT": url_for("ict_dashboard"),
        "Finance": url_for("finance_dashboard"),
        "Teacher": url_for("teacher_dashboard"),
        "Student": url_for("student_dashboard"),
        "Parent": url_for("parent_dashboard"),
        "Librarian": url_for("librarian_dashboard"),
    }[role]


def workspace_type_for_user(user) -> str:
    if not user: return "Teaching"
    try:
        row=q("SELECT workspace_type FROM users WHERE id=?",(user["id"],),one=True)
        wt=(row["workspace_type"] if row else "") or ""
    except Exception:
        wt=""
    return wt or ("Teaching" if user["role"]=="Teacher" else user["role"])

def is_reception_user(user) -> bool:
    if not user:
        return False
    return bool(user["reception_enabled"]) or user["role"] in {"Admin", "ICT"}

def staff_code_for(user_role: str, workspace_type: str) -> str:
    prefix={"Teacher":"TCH","Driver":"DRV","Reception":"REC","Guard":"SEC","Cook":"CAT","Finance":"FIN","ICT":"ICT","Librarian":"LIB","Admin":"ADM"}.get(workspace_type if workspace_type in {"Driver","Reception","Guard","Cook"} else user_role,"STF")
    row=q("SELECT COUNT(*) AS c FROM users WHERE staff_code LIKE ?",(prefix+"-%",),one=True)
    return f"{prefix}-{int(row['c'] or 0)+1:03d}"

def resolve_staff_token(raw_token: str):
    token=(raw_token or '').strip()
    if token.startswith('STAFF|'):
        parts=token.split('|',5); token=parts[1] if len(parts)>1 else ''
    m=re.search(r'/qr/([A-Za-z0-9_-]+)',token)
    if m: token=m.group(1)
    return q("SELECT * FROM users WHERE qr_access_token=? AND active=1 AND role!='System'",(token,),one=True) if token else None

def reception_admin_ids():
    return [r['id'] for r in q("SELECT id FROM users WHERE active=1 AND role IN ('Admin','ICT')")]

def attendance_admin_ids():
    """Attendance notifications are private to active Admin accounts."""
    return [r['id'] for r in q("SELECT id FROM users WHERE active=1 AND role='Admin'")]

def record_reception_scan(action, token='', device_token='', full_name='', phone='', gender='', reason='', source='online', method='QR', latitude=None, longitude=None, accuracy=None, event_at=None, school_unit='', school_location='', id_number=''):
    action=str(action or '').upper()
    if action not in {'IN','OUT'}: raise ValueError('Action must be IN or OUT')
    user=resolve_staff_token(token) if token else None
    now=event_at or datetime.utcnow().isoformat(timespec='seconds')
    if user:
        name=user['full_name']; phone=user['phone'] or ''; gender=user['gender'] or ''; position=user['title'] or user['role']; code=user['position_code'] or user['staff_code'] or ''; unit=user['school_unit'] or school_settings()['school_name']; loc=user['school_location'] or school_settings()['institution_affiliations'] or ''
        open_visit=q("SELECT * FROM reception_visits WHERE user_id=? AND check_out IS NULL ORDER BY id DESC LIMIT 1",(user['id'],),one=True)
        if action=='IN':
            vid=execute("INSERT INTO reception_visits(user_id,person_type,full_name,phone,gender,reason,position,staff_code,school_unit,school_location,device_token,check_in,source,method,latitude,longitude,accuracy,id_number) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(user['id'],'Staff',name,phone,gender,reason,position,code,unit,loc,device_token,now,source,method,latitude,longitude,accuracy,id_number))
        elif open_visit:
            vid=open_visit['id']; execute("UPDATE reception_visits SET check_out=?,source=?,method=?,latitude=?,longitude=?,accuracy=? WHERE id=?",(now,source,method,latitude,longitude,accuracy,vid))
        else:
            vid=execute("INSERT INTO reception_visits(user_id,person_type,full_name,phone,gender,reason,position,staff_code,school_unit,school_location,device_token,check_out,source,method,latitude,longitude,accuracy,id_number) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(user['id'],'Staff',name,phone,gender,reason,position,code,unit,loc,device_token,now,source,method,latitude,longitude,accuracy,id_number))
        execute("INSERT INTO attendance_events(user_id,action,method,event_at,source,latitude,longitude,device_note) VALUES(?,?,?,?,?,?,?,?)",(user['id'],action,method,now,source,latitude,longitude,(device_token or '')[:120]))
        notify_users(reception_admin_ids(),f'Reception: {name} checked {"IN" if action=="IN" else "OUT"}',f'{name} ({position}) {"arrived at" if action=="IN" else "left"} reception at {now}.',url_for('reception_dashboard'))
        return {'ok':True,'visit_id':vid,'name':name,'position':position,'staff_code':code,'message':f'{name} checked {"in" if action=="IN" else "out"} at {now}.','registered':True}
    if not device_token: device_token=uuid.uuid4().hex
    unit=school_unit or school_settings()['school_name']; loc=school_location or school_settings()['institution_affiliations'] or ''
    existing=q("SELECT * FROM reception_visits WHERE device_token=? AND check_out IS NULL ORDER BY id DESC LIMIT 1",(device_token,),one=True)
    if action=='OUT' and existing:
        execute("UPDATE reception_visits SET check_out=?,source=?,method=?,latitude=?,longitude=?,accuracy=? WHERE id=?",(now,source,method,latitude,longitude,accuracy,existing['id']))
        notify_users(reception_admin_ids(),'Reception: unregistered person checked OUT',f'{existing["full_name"]} left reception at {now}.',url_for('reception_dashboard'))
        return {'ok':True,'visit_id':existing['id'],'name':existing['full_name'],'message':f'{existing["full_name"]} checked out at {now}.','registered':False}
    if action=='OUT' and not existing:
        vid=execute("INSERT INTO reception_visits(person_type,full_name,phone,gender,reason,school_unit,school_location,device_token,check_out,source,method,latitude,longitude,accuracy,id_number) VALUES('Anonymous',?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(full_name or 'Unregistered person',phone,gender,reason,unit,loc,device_token,now,source,method,latitude,longitude,accuracy,id_number))
        return {'ok':True,'visit_id':vid,'name':full_name or 'Unregistered person','message':f'Unregistered checkout captured at {now}.','registered':False}
    vid=execute("INSERT INTO reception_visits(person_type,full_name,phone,gender,reason,school_unit,school_location,device_token,check_in,source,method,latitude,longitude,accuracy,id_number) VALUES('Anonymous',?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(full_name or 'Unregistered person',phone,gender,reason,unit,loc,device_token,now,source,method,latitude,longitude,accuracy,id_number))
    notify_users(reception_admin_ids(),'Reception: unregistered person checked IN',f'{full_name or "Unregistered visitor"} arrived at reception at {now}.',url_for('reception_dashboard'))
    return {'ok':True,'visit_id':vid,'name':full_name or 'Unregistered person','message':f'{full_name or "Unregistered person"} checked in at {now}.','registered':False}

def specialized_dashboard_for(user) -> str:
    # One account -> one operational destination. Student linkage always wins
    # over stale legacy role/workspace data.
    role=(user.get("role") if hasattr(user, "get") else user["role"]) or ""
    try:
        if user.get("student_id"):
            linked = q("SELECT id FROM students WHERE id=? AND active=1", (user["student_id"],), one=True)
            if linked:
                role = "Student"
    except Exception:
        pass
    if role in {"Admin","ICT","Finance","Teacher","Student","Parent","Librarian"}:
        return role_target(role)
    wt=workspace_type_for_user(user)
    if wt=="Driver": return url_for("driver_dashboard")
    if wt in {"Guard","Cook","Other Staff"}: return url_for("workforce_dashboard", kind=wt)
    # Unknown/legacy staff is sent to the generic dashboard dispatcher rather
    # than generating a Not Found page.
    return url_for("dashboard")

def enter_role_without_login(role: str):
    if role == "Parent" and not parent_portal_enabled():
        flash("Parent accounts are disabled for this institution mode.", "warning")
        return redirect(url_for("index"))
    # Passwordless mode is intended only as a controlled demo/single-user setup.
    # Once a role has multiple active accounts, require an actual identity so one
    # user's dashboard cannot silently become another user's account.
    if auth_required():
        return redirect(url_for("login", role=role))
    users = q(
        "SELECT id FROM users WHERE role=? AND active=1 AND role!='System' ORDER BY id",
        (role,),
    )
    if not users:
        # In presentation mode there may be no real account yet. Use the closest
        # permitted account so the workspace can still be demonstrated.
        _ensure_demo_identity()
        target = role_target(role) if role in ALL_PORTAL_ROLES else url_for("dashboard")
        return redirect(target)
    session.clear()
    session.permanent = True
    session["user_id"] = users[0]["id"]
    session["active_portal_role"] = role
    user = q("SELECT * FROM users WHERE id=?", (users[0]["id"],), one=True)
    return redirect(specialized_dashboard_for(user))


def selected_role_from_request(default=""):
    role = (request.args.get("role") or request.form.get("role") or default).strip()
    return role if role in ALL_PORTAL_ROLES else default




def important_dates(limit=20, landing=False):
    if landing:
        return q("SELECT * FROM important_dates WHERE visible=1 AND landing_visible=1 ORDER BY event_date,event_time,id LIMIT ?", (limit,))
    return q("SELECT * FROM important_dates WHERE visible=1 ORDER BY event_date,event_time,id LIMIT ?", (limit,))


def school_day_status(day=None):
    """Return the institution's operational day state without ever blocking login/attendance."""
    day = day or datetime.now().date()
    iso = day.isoformat()
    if day.weekday() >= 5:
        return {"is_school_day": False, "label": "Weekend", "kind": "Weekend"}
    rows = q("SELECT * FROM school_calendar WHERE start_date<=? AND end_date>=? ORDER BY id DESC", (iso, iso))
    if any(int(r["school_day"] or 0) == 0 for r in rows):
        r = next((r for r in rows if int(r["school_day"] or 0) == 0), rows[0])
        return {"is_school_day": False, "label": r["title"], "kind": r["kind"]}
    openings = q("SELECT MIN(start_date) AS d FROM school_calendar WHERE kind='Opening Date'")
    closings = q("SELECT MAX(end_date) AS d FROM school_calendar WHERE kind='Closing Date'")
    opening = openings[0]["d"] if openings and openings[0]["d"] else None
    closing = closings[0]["d"] if closings and closings[0]["d"] else None
    if opening and iso < opening:
        return {"is_school_day": False, "label": "Before school opening", "kind": "School Closure"}
    if closing and iso > closing:
        return {"is_school_day": False, "label": "After school closing", "kind": "School Closure"}
    if rows:
        r = next((r for r in rows if int(r["school_day"] or 0) == 1), rows[0])
        return {"is_school_day": True, "label": r["title"], "kind": r["kind"]}
    return {"is_school_day": True, "label": "Regular school day", "kind": "School Day"}

def qr_login_allowed(user):
    return bool(user and user["active"] and user["role"] != "System" and user["role"] in QR_LOGIN_ROLES and ((user["workspace_type"] or "") in QR_LOGIN_WORKSPACES or user["role"] in QR_LOGIN_ROLES) and int(user["qr_login_enabled"] or 0) == 1)

def notify_user(user_id, title, body, link='', priority='Normal'):
    if not user_id:
        return
    execute("INSERT INTO notifications(user_id,title,body,link,priority) VALUES(?,?,?,?,?)", (int(user_id), title[:160], body[:5000], link[:500] if link else None, priority[:20]))

def notify_users(user_ids, title, body, link='', priority='Normal'):
    for uid in set(int(x) for x in user_ids if x):
        notify_user(uid,title,body,link,priority)

def notification_count(user_id):
    row=q("SELECT COUNT(*) AS n FROM notifications WHERE user_id=? AND read_at IS NULL",(user_id,),one=True) if user_id else None
    return int(row['n']) if row else 0

def role_shortcuts(user):
    if not user: return []
    common=[('Calendar','/calendar','Important dates'),('Notifications','/notifications','Updates & alerts')]
    role='Reception' if is_reception_user(user) and user['role'] not in {'Admin','ICT'} else user['role']
    extras={
        'Teacher':[('Live classes','/online-classes','Host & schedule classes'),('My groups','/groups','Groups, discussions & meetings'),('My leadership','/leadership','Department / HOD responsibilities')],
        'Student':[('Student shortcuts','/student-dashboard#shortcuts','Learning shortcuts'),('Live classes','/online-classes','Join scheduled classes'),('My groups','/groups','Groups & discussions'),('Student leadership','/leadership','Leadership workspace')],
        'Parent':[('Family dates','/calendar','School calendar'),('Teacher communication','/communication','Talk to teachers'),('Learning updates','/notifications','Family notifications')],
        'ICT':[('Attendance','/admin/attendance','People & attendance'),('Timetable','/staff/timetable','Staff schedules'),('Reminders','/staff/reminders','Work reminders'),('Content & dates','/calendar','Publish dates'),('Groups','/groups','Manage groups'),('Leadership','/leadership','Set structure')],
        'Admin':[('Attendance','/admin/attendance','People & attendance'),('Timetable','/staff/timetable','Staff schedules'),('Reminders','/staff/reminders','Work reminders'),('Content & dates','/calendar','Publish dates'),('Groups','/groups','Manage groups'),('Leadership','/leadership','Manage structure')],
    }
    return common+extras.get(role,[('Groups','/groups','Groups & activities')])

def _hex_rgb(value, fallback=(52,53,65)):
    value=str(value or '').strip().lstrip('#')
    if len(value)==3:
        value=''.join(ch*2 for ch in value)
    if len(value)!=6 or not re.fullmatch(r'[0-9a-fA-F]{6}', value):
        return fallback
    return tuple(int(value[i:i+2],16) for i in (0,2,4))


def _rel_luminance(rgb):
    vals=[]
    for c in rgb:
        x=c/255.0
        vals.append(x/12.92 if x<=0.03928 else ((x+0.055)/1.055)**2.4)
    return 0.2126*vals[0]+0.7152*vals[1]+0.0722*vals[2]


def _contrast_ratio(a,b):
    la=_rel_luminance(_hex_rgb(a)); lb=_rel_luminance(_hex_rgb(b))
    hi=max(la,lb); lo=min(la,lb)
    return (hi+0.05)/(lo+0.05)


def _best_text(background, requested, minimum=4.5):
    requested=str(requested or '').strip() or '#ececf1'
    if _contrast_ratio(background, requested) >= minimum:
        return requested
    return '#101828' if _contrast_ratio(background, '#101828') >= _contrast_ratio(background, '#ffffff') else '#ffffff'


def _rgba(hex_value, alpha):
    r,g,b=_hex_rgb(hex_value)
    return f'rgba({r},{g},{b},{alpha})'


def theme_style(settings=None) -> str:
    settings = settings or school_settings()
    def esc(v):
        return str(v).replace('<','').replace('>','').replace('"','').replace(';','')
    bg=esc(settings['background_color'] or '#343541')
    panel=esc(settings['panel_color'] or bg)
    sidebar=esc(settings['sidebar_color'] or panel)
    header=esc(settings['header_color'] or panel)
    requested_text=esc(settings['text_color'] or '#ececf1')
    requested_muted=esc(settings['muted_text_color'] or '#b5bac7')
    # User-selected colours remain authoritative, but unreadable combinations are automatically corrected.
    body_text=_best_text(bg, requested_text)
    panel_text=_best_text(panel, body_text)
    sidebar_text=_best_text(sidebar, body_text)
    header_text=_best_text(header, body_text)
    muted=_best_text(bg, requested_muted, minimum=3.0)
    input_bg=_best_text(panel, bg, minimum=1.2) if _contrast_ratio(panel,bg)<1.12 else bg
    primary_button_text=_best_text(settings['primary_color'] or '#3457d5', '#ffffff')
    heading_font=esc(settings['heading_font'] or settings['font_family'] or 'Inter')
    font_family=esc(settings['font_family'] or 'Inter')
    bg_path=str(settings.get('background_path','') or '').strip() if hasattr(settings,'get') else ''
    bg_image = f"body.app-body{{background-image:linear-gradient(180deg,rgba(15,23,42,.18),rgba(15,23,42,.24)),url('/{bg_path}');background-size:cover;background-position:center;background-attachment:fixed;}}" if bg_path else ''
    css=(
        f":root{{--bg:{bg};--panel:{panel};--panel-3:{input_bg};"
        f"--primary-blue:{esc(settings['primary_color'] or '#3457d5')};--deep-accent-blue:{esc(settings['accent_color'] or '#3457d5')};"
        f"--primary-text:{body_text};--muted-text:{muted};--panel-text:{panel_text};--sidebar-text:{sidebar_text};--header-text:{header_text};"
        f"--text-soft:{_rgba(body_text,.06)};--text-border:{_rgba(body_text,.14)};--text-hover:{_rgba(body_text,.10)};"
        f"--input-text:{_best_text(input_bg,body_text)};--primary-button-text:{primary_button_text};"
        f"--font:'{font_family}',Inter,system-ui,sans-serif;--heading-font:'{heading_font}',{font_family},Inter,sans-serif;"
        f"--radius:{int(settings['radius_px'] or 12)}px;--radius-sm:{int(settings['button_radius_px'] or 10)}px;--sidebar-bg:{sidebar};--header-bg:{header};}}"
    )
    extra=settings['custom_css'] or ''
    if len(extra)>12000 or re.search(r'@import|javascript:|expression\s*\(', extra, re.I):
        extra=''
    return css+extra


def theme_preset_style(settings=None) -> str:
    settings=settings or school_settings()
    preset=str(settings["theme_preset"] or "classic").lower() if "theme_preset" in settings.keys() else "classic"
    presets={
        "classic":"",
        "christmas":":root{--theme-glow:rgba(220,38,38,.18)} body.app-body:before{content:'';position:fixed;inset:0;pointer-events:none;background:radial-gradient(circle at 12% 0%,rgba(22,163,74,.16),transparent 28%),radial-gradient(circle at 88% 0%,rgba(220,38,38,.16),transparent 28%);z-index:-1} body.app-body .topbar{border-top-color:rgba(220,38,38,.5);box-shadow:0 14px 34px rgba(220,38,38,.12)} body.app-body .panel{border-color:rgba(255,255,255,.10)}",
        "easter":":root{--theme-glow:rgba(168,85,247,.16)} body.app-body:before{content:'';position:fixed;inset:0;pointer-events:none;background:radial-gradient(circle at 15% 0%,rgba(250,204,21,.16),transparent 25%),radial-gradient(circle at 82% 0%,rgba(168,85,247,.16),transparent 25%);z-index:-1} body.app-body .topbar{border-top-color:rgba(168,85,247,.45);box-shadow:0 14px 34px rgba(168,85,247,.12)}",
        "madaraka":":root{--theme-glow:rgba(220,38,38,.16)} body.app-body:before{content:'';position:fixed;inset:0;pointer-events:none;background:linear-gradient(120deg,rgba(0,0,0,.10),rgba(220,38,38,.10),rgba(22,163,74,.10));z-index:-1} body.app-body .topbar{border-top-color:rgba(22,163,74,.55);box-shadow:0 14px 34px rgba(22,163,74,.10)}",
        "school-pride":":root{--theme-glow:rgba(37,99,235,.18)} body.app-body .topbar{border-top-color:var(--primary-blue);box-shadow:0 14px 34px rgba(37,99,235,.18)}"
    }
    return presets.get(preset,presets["classic"])

def active_advertisements(limit=4):
    today=datetime.utcnow().strftime("%Y-%m-%d")
    return q("SELECT a.*,u.full_name AS poster FROM advertisements a LEFT JOIN users u ON u.id=a.created_by WHERE a.active=1 AND (a.start_date='' OR a.start_date<=?) AND (a.end_date='' OR a.end_date>=?) ORDER BY a.priority DESC,a.created_at DESC LIMIT ?",(today,today,limit))

def landing_style(settings=None) -> str:
    settings=settings or school_settings()
    ff=str(settings['landing_font_family'] or 'Inter').replace('<','').replace('>','').replace(';','').replace('"','')
    hf=str(settings['landing_heading_font'] or ff).replace('<','').replace('>','').replace(';','').replace('"','')
    width=max(900,min(1600,int(settings['landing_content_width'] or 1240)))
    cols=max(1,min(3,int(settings['landing_role_columns'] or 3)))
    hero=str(settings['landing_hero_layout'] or 'split')
    hero_css='grid-template-columns:minmax(0,1.55fr) minmax(260px,.65fr);' if hero=='split' else 'grid-template-columns:1fr;'
    return f".landing-shell{{width:min({width}px,calc(100% - 48px));}} .landing-hero{{{hero_css}}} .role-grid{{grid-template-columns:repeat({cols},minmax(0,1fr));}} .landing-font-scope{{font-family:'{ff}',Inter,system-ui,sans-serif;}} .landing-heading-scope{{font-family:'{hf}',{ff},Inter,system-ui,sans-serif;}}"


def current_landing_url() -> str:
    return url_for('index', _external=True)

def portal_qr_data_uri() -> str:
    payload=url_for("index", _external=True)
    qr=qrcode.QRCode(version=2,box_size=5,border=2); qr.add_data(payload); qr.make(fit=True)
    buf=io.BytesIO(); qr.make_image().save(buf,format="PNG"); return "data:image/png;base64,"+base64.b64encode(buf.getvalue()).decode("ascii")

@app.context_processor
def auth_template_context():
    settings=school_settings()
    return {
        "current_user": current_user(), "school_settings": settings, "portal_title": settings["school_name"], "theme_color": settings["primary_color"], "all_roles": ALL_PORTAL_ROLES, "public_roles": PUBLIC_ROLES,
        "theme_style": theme_style(settings), "theme_preset_style": theme_preset_style(settings), "landing_style": landing_style(settings), "portal_landing_url": current_landing_url(),
        "active_adverts": active_advertisements(),
        "welcome_animation": bool(settings["welcome_animation_enabled"]), "welcome_animation_name": settings["welcome_animation_name"], "welcome_animation_duration_ms": int(settings["welcome_animation_duration_ms"] or 2200), "welcome_animation_style": settings["welcome_animation_style"] if "welcome_animation_style" in settings.keys() else "clean",
        "important_dates": important_dates(12, landing=request.path == '/'),
        "school_day": school_day_status(),
        "notification_count": notification_count(current_user()['id']) if current_user() else 0,
        "role_shortcuts": role_shortcuts(current_user()),
        "institution_type": settings['institution_type'], "learner_label": settings['learner_label'],
        "learner_plural": "Pupils" if settings['learner_label'].lower().startswith('pupil') else "Students",
        "staff_label": settings['staff_label'],
        "staff_plural": "Teachers / Lecturers" if 'Lecturer' in (settings['staff_label'] or '') else "Employees",
        "help_enabled": True,
        "parent_portal_enabled": parent_portal_enabled(),
        "portal_qr_data_uri": portal_qr_data_uri(),
        "footer_settings": {
            "title": settings["footer_title"] or settings["school_name"],
            "text": settings["footer_text"] or "Institution portal for learning, communication, finance and school services.",
            "contact": settings["footer_contact"] or settings["help_email"] or settings["help_phone"],
            "links": settings["footer_links"] or "Home,Library,Help,AI Assistant",
            "platform_credit": bool(settings["platform_credit_enabled"]),
        },
        "current_year": datetime.utcnow().year,
    }


def workspace_for(role: str) -> str:
    return {
        "Admin": "Admin Command Centre",
        "ICT": "ICT Control Panel",
        "Finance": "Finance Workspace",
        "Teacher": "Teacher Dashboard",
        "Student": "Student Dashboard",
        "Parent": "Parent Dashboard",
        "Librarian": "Library Workspace",
        "Staff": "Staff Workspace",
    }.get(role, "School Portal System")


def navigation_items(role: str, settings):
    order=[x.strip() for x in (settings["menu_order"] or "").split(",") if x.strip()]
    labels={
        "Home": settings["home_label"],
        "Assignments": settings["assignments_label"],
        "Results": settings["results_label"],
        "Messages": settings["messages_label"],
        "Finance": settings["finance_label"],
        "Branding": settings["branding_label"],
        "Submissions": "Submissions",
        "Online classes": "Online classes",
        "My children": "My children",
        "Results & fees": "Results & fees",
        "Teacher communication": "Teacher communication",
        "Payments": "Payments",
        "Elections": "Elections",
        "Library": "Library",
        "Institution": "About the institution",
        "Theme": "Theme",
        "Navigation order": "Navigation order",
        "Members": "Members",
        "Flashcards": "Flashcards",
    }
    # Admin is the reference ordering for management workspaces. ICT mirrors that
    # order for shared capabilities but receives only the management features it
    # is actually cleared to use. ICT-only presentation controls are then added
    # in a stable block rather than scrambling the common order.
    management_order=["Home","Finance","Flashcards","Elections","Library","Institution","Members"]
    role_allowed={
        "Teacher": ["Home","Assignments","Submissions","Flashcards","Online classes","Elections","Library","Institution"],
        "Student": ["Home","Assignments","Results","Flashcards","Online classes","Elections","Library","Institution"],
        "Parent": ["Home","My children","Results & fees","Teacher communication","Library","Institution"],
        "Finance": ["Home","Finance","Payments","Library","Institution"],
        "ICT": ["Home","Elections","Library","Institution","Members","Branding","Theme","Navigation order"],
        "Librarian": ["Home","Library","Institution"],
        "Admin": ["Home","Finance","Flashcards","Elections","Library","Institution","Members"],
    }
    allowed=role_allowed.get(role, ["Home","Institution"])
    if role in {"Admin","ICT"}:
        allowed=[k for k in management_order if k in allowed] + [k for k in ("Branding","Theme","Navigation order") if k in allowed]
    if not int(settings["elections_enabled"] or 0) and role not in {"Admin","ICT"}:
        allowed=[x for x in allowed if x!="Elections"]
    if not int(settings["library_enabled"] or 0) and role not in {"Admin","ICT","Librarian"}:
        allowed=[x for x in allowed if x!="Library"]
    anchor_map={"Home":"home","Assignments":"assignments","Submissions":"submissions","Flashcards":"flashcards","Online classes":"classes","Results":"results","My children":"children","Results & fees":"results","Teacher communication":"messages","Finance":"finance","Payments":"payments","Branding":"branding","Theme":"theme","Navigation order":"navigation","Elections":"elections","Library":"library","Institution":"institution","Members":"users"}
    result=[]
    for key in order:
        if key in allowed and key not in [r["key"] for r in result]:
            result.append({"key":key,"label":labels.get(key,key),"anchor":anchor_map[key]})
    for key in allowed:
        if key not in [r["key"] for r in result]:
            result.append({"key":key,"label":labels.get(key,key),"anchor":anchor_map[key]})
    return result


def portal_student(student_id=None):
    user=current_user()
    if user and user["role"] in {"Student","Parent"}:
        linked_id=user["student_id"]
        if not linked_id:
            return None
        # Never allow a URL/query parameter to switch a Student/Parent to another pupil.
        return q("SELECT * FROM students WHERE id=? AND active=1", (linked_id,), one=True)
    if student_id:
        return q("SELECT * FROM students WHERE id=? AND active=1", (student_id,), one=True)
    return q("SELECT * FROM students WHERE active=1 ORDER BY id LIMIT 1", one=True)


def parent_portal_enabled() -> bool:
    row = q("SELECT parent_portal_enabled FROM school_settings WHERE id=1", one=True)
    if row is not None:
        return bool(row["parent_portal_enabled"])
    return school_settings()["institution_type"] in {"Primary School", "Secondary School", "Mixed Institution"}

def teacher_or_admin() -> bool:
    return bool(current_user() and current_user()["role"] in {"Teacher", "Admin"})

def parent_children(user):
    """Return every learner explicitly linked to the authenticated Parent.

    Explicit guardian_links are authoritative; legacy contact matching is only
    a fallback for older records that predate the multi-child guardian links.
    """
    if not user or user["role"] != "Parent":
        return []
    linked = q("""SELECT s.*
                 FROM guardian_links gl
                 JOIN students s ON s.id=gl.student_id
                 WHERE gl.guardian_user_id=? AND gl.active=1 AND s.active=1
                 ORDER BY s.grade,s.full_name""", (user["id"],))
    if linked:
        return linked
    if not user["student_id"]:
        return []
    child = q("SELECT * FROM students WHERE id=? AND active=1", (user["student_id"],), one=True)
    if not child:
        return []
    filters=[]; params=[]
    for column in ("guardian_phone", "guardian_email", "alt_guardian_phone", "alt_guardian_email"):
        value=child[column]
        if value:
            filters.append(f"{column}=?"); params.append(value)
    if not filters and child["guardian_name"]:
        filters.append("guardian_name=?"); params.append(child["guardian_name"])
    if not filters:
        return [child]
    return q("SELECT * FROM students WHERE active=1 AND ("+" OR ".join(filters)+") ORDER BY full_name", params)


def can_access_student(student_id: int, *, write: bool=False) -> bool:
    user=current_user()
    if not user:
        return False
    role=user["role"]
    if role in {"Admin", "ICT", "Finance", "Librarian"}:
        return not write or role in {"Admin", "ICT"}
    if role == "Student":
        return bool(user["student_id"] == student_id and not write)
    if role == "Parent":
        return any(row["id"] == student_id for row in parent_children(user)) and not write
    return False


def assignment_rows(grade=None):
    if grade:
        return q("""SELECT a.*, u.full_name AS teacher_name, COUNT(s.id) AS submissions
                   FROM assignments a LEFT JOIN users u ON u.id=a.posted_by
                   LEFT JOIN submissions s ON s.assignment_id=a.id
                   WHERE a.grade=? GROUP BY a.id ORDER BY a.created_at DESC, a.id DESC""", (grade,))
    return q("""SELECT a.*, u.full_name AS teacher_name, COUNT(s.id) AS submissions
               FROM assignments a LEFT JOIN users u ON u.id=a.posted_by
               LEFT JOIN submissions s ON s.assignment_id=a.id
               GROUP BY a.id ORDER BY a.created_at DESC, a.id DESC""")


def make_qr_token(document_type: str, student_id: int) -> str:
    return f"{document_type.lower().replace(' ', '-')}-{student_id}-{uuid.uuid4().hex}"

def create_qr_file(token: str) -> Path:
    qr=qrcode.QRCode(version=1, box_size=8, border=3)
    qr.add_data(url_for("verify_document", token=token, _external=True)); qr.make(fit=True)
    path=QR_DIR/f"{token}.png"; qr.make_image().save(path); return path

def result_release_info(student_id: int):
    settings=school_settings(); limit=float(settings["result_download_balance_limit"] or 0)
    student=q("SELECT balance FROM students WHERE id=?",(student_id,),one=True)
    balance=float(student["balance"] or 0) if student else 0
    batches=q("SELECT b.*, COUNT(r.id) AS entries FROM exam_batches b JOIN exam_results r ON r.batch_id=b.id WHERE r.student_id=? GROUP BY b.id ORDER BY b.created_at DESC,b.id DESC",(student_id,))
    return [{"batch":b,"eligible":b["finance_status"]=="Approved" and balance<=limit,"limit":limit,"balance":balance} for b in batches]

def generate_result_pdf(student,batch,results,doc,qr_path):
    buf=io.BytesIO(); c=canvas.Canvas(buf,pagesize=A4); w,h=A4; settings=school_settings()
    c.setTitle(f"Result - {student['full_name']}"); c.setFont("Helvetica-Bold",16); c.drawString(44,h-52,settings["school_name"])
    c.setFont("Helvetica",10); c.drawString(44,h-70,"Official Student Result")
    c.drawString(44,h-94,f"Student: {student['full_name']}"); c.drawString(44,h-110,f"Admission No: {student['admission_no']}   Grade: {student['grade']}"); c.drawString(44,h-126,f"Term: {batch['term']}   Finance status: {batch['finance_status']}")
    y=h-160; c.setFont("Helvetica-Bold",10); c.drawString(44,y,"Subject"); c.drawString(250,y,"Mark"); c.drawString(330,y,"Out of"); y-=18; c.setFont("Helvetica",10)
    for r in results:
        c.drawString(44,y,str(r['subject'])); c.drawString(250,y,f"{r['mark']:.1f}"); c.drawString(330,y,f"{r['max_mark']:.1f}"); y-=16
        if y<120: c.showPage(); y=h-60
    c.drawImage(str(qr_path),w-145,60,width=90,height=90,preserveAspectRatio=True,mask='auto'); c.setFont("Helvetica",7); c.drawString(w-148,50,"Scan to verify authenticity"); c.drawString(44,44,f"Verification token: {doc['qr_token']}"); c.drawString(44,32,"Issued by School Portal System"); c.save(); buf.seek(0); return buf

def embed_qr_in_document(source_path: Path, qr_path: Path, token: str) -> str:
    """Create a verified copy with the QR visible on the document when possible."""
    suffix=source_path.suffix.lower()
    if suffix in {".png",".jpg",".jpeg",".webp"}:
        img=Image.open(source_path).convert("RGB")
        qr=Image.open(qr_path).convert("RGB").resize((220,220))
        pad=20; canvas_img=Image.new("RGB",(img.width,max(img.height,260)),"white"); canvas_img.paste(img,(0,0));
        x=max(0,canvas_img.width-qr.width-pad); canvas_img.paste(qr,(x,20)); out=DOC_DIR/f"verified-{token}.jpg"; canvas_img.save(out,quality=92); return "uploads/documents/"+out.name
    if suffix==".pdf":
        reader=PdfReader(str(source_path)); writer=PdfWriter()
        overlay=io.BytesIO(); c=canvas.Canvas(overlay,pagesize=A4); w,h=A4; c.drawImage(str(qr_path),w-150,h-150,width=90,height=90,mask='auto'); c.setFont("Helvetica",7); c.drawString(w-154,h-158,"Scan to verify authenticity"); c.save(); overlay.seek(0); ov=PdfReader(overlay).pages[0]
        for i,page in enumerate(reader.pages):
            if i==0: page.merge_page(ov)
            writer.add_page(page)
        out=DOC_DIR/f"verified-{token}.pdf"; writer.write(str(out)); return "uploads/documents/"+out.name
    return ""

def generate_exam_card_pdf(student,batch_name,doc,qr_path):
    buf=io.BytesIO(); c=canvas.Canvas(buf,pagesize=A4); w,h=A4; settings=school_settings(); c.setTitle(f"Exam Card - {student['full_name']}")
    c.setFont("Helvetica-Bold",18); c.drawString(44,h-58,settings["school_name"]); c.setFont("Helvetica-Bold",14); c.drawString(44,h-84,"EXAMINATION CARD")
    c.setFont("Helvetica",11); y=h-130
    photo_path=(student['profile_photo'] if 'profile_photo' in student.keys() else '') if student else ''
    if photo_path:
        try:
            target=(BASE_DIR/photo_path).resolve() if photo_path.startswith('uploads/') else (UPLOAD_DIR/photo_path).resolve()
            if target.exists() and target.is_file():
                c.drawImage(str(target),44,h-255,width=88,height=110,preserveAspectRatio=True,mask='auto')
        except Exception:
            pass
    for line in [f"Student: {student['full_name']}",f"Admission No: {student['admission_no']}",f"Grade: {student['grade']}",f"Examination: {batch_name}",f"Fee balance: {settings['currency_code']} {float(student['balance'] or 0):,.0f}","Status: APPROVED"]:
        c.drawString(150 if photo_path else 60,y,line); y-=22
    c.drawString(60,y-8,"Present this card when requested by the school."); c.drawImage(str(qr_path),w-170,95,width=100,height=100,preserveAspectRatio=True,mask='auto'); c.setFont("Helvetica",7); c.drawString(w-174,84,"Scan to verify authenticity"); c.drawString(44,44,f"Verification token: {doc['qr_token']}"); c.save(); buf.seek(0); return buf

# -------------------------
# Context / templates
# -------------------------
@app.context_processor
def inject_globals():
    settings = school_settings()
    return {
        "now_year": datetime.now().year,
        "current_user": current_user(),
        "workspace_for": workspace_for,
        "school_settings": settings,
        "admin_login_path": ADMIN_LOGIN_PATH,
        "public_roles": PUBLIC_ROLES,
        "hidden_roles": HIDDEN_ROLES,
        "theme_color": settings["background_color"],
        "theme_accent": settings["primary_color"],
        "elections_enabled": bool(settings["elections_enabled"]),
        "library_enabled": bool(settings["library_enabled"]),
        "institution_types": INSTITUTION_TYPES,
    }


# -------------------------
# Routes
# -------------------------
@app.route("/offline")
def offline():
    settings = school_settings()
    return render_template("offline.html", portal_title=settings["school_name"], school_settings=settings, theme_color=settings["primary_color"])


@app.route("/pulse_receiver", methods=["POST", "GET"])
def pulse_receiver():
    payload = request.get_json(silent=True) or {}
    request_id = str(payload.get("request_id") or request.headers.get("X-Pulse-ID") or uuid.uuid4().hex)
    source = str(payload.get("source") or payload.get("sender") or request.headers.get("X-Pulse-Source") or "unknown")[:200]
    reply_payload = {
        "type": "school_portal_pulse",
        "event": "received",
        "request_id": request_id,
        "received_from": source,
        "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "reply": True,
    }
    requested_reply = str(payload.get("reply_to") or "").strip()
    configured_peer_host = (urllib.parse.urlparse(PULSE_PEER_URL).hostname or "").lower()
    allowed_hosts = set(PULSE_ALLOWED_CALLBACK_HOSTS)
    if configured_peer_host:
        allowed_hosts.add(configured_peer_host)
    reply_url = PULSE_PEER_URL
    if requested_reply:
        parsed_reply = urllib.parse.urlparse(requested_reply)
        if parsed_reply.scheme == "https" and parsed_reply.netloc and (parsed_reply.hostname or "").lower() in allowed_hosts:
            reply_url = requested_reply.rstrip("/")

    def post_json(target: str) -> tuple[bool, int]:
        parsed = urllib.parse.urlparse(target)
        if parsed.scheme != "https" or not parsed.netloc:
            return False, 0
        body = json.dumps(reply_payload, separators=(",", ":")).encode("utf-8")
        req = urllib.request.Request(target, data=body, method="POST", headers={"Content-Type":"application/json","Accept":"application/json","User-Agent":"SchoolPortal-Pulse/1.0"})
        try:
            with urllib.request.urlopen(req, timeout=PULSE_TIMEOUT_SECONDS) as response:
                return True, int(getattr(response, "status", 200))
        except urllib.error.HTTPError as exc:
            return False, int(exc.code)
        except (urllib.error.URLError, TimeoutError, OSError, ValueError):
            return False, 0

    def reply_worker():
        ok, status = post_json(reply_url)
        # The configured peer was supplied as a bare origin in the request.
        # If that origin does not expose POST, retry its conventional pulse route.
        parsed = urllib.parse.urlparse(reply_url)
        if not ok and status == 404 and parsed.path in ("", "/"):
            post_json(reply_url.rstrip("/") + "/pulse_receiver")
    threading.Thread(target=reply_worker, daemon=True, name="pulse-peer-reply").start()
    threading.Thread(target=reply_worker, daemon=True, name="pulse-peer-reply").start()
    return jsonify({"ok": True, "received": True, "reply_scheduled": True, "request_id": request_id, "service": "school-portal", "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z"}), 200


@app.route("/pulse", methods=["POST", "GET"])
def pulse():
    return pulse_receiver()


@app.route("/health")
def health():
    db_ok = True
    try:
        q("SELECT 1", one=True)
    except Exception:
        db_ok = False
    return jsonify({"ok": db_ok, "database": "ok" if db_ok else "error", "persistent_storage": PERSISTENT_STORAGE, "pulse_peer": PULSE_PEER_URL, "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z"}), (200 if db_ok else 503)


@app.route("/")
def index():
    settings = school_settings()
    return render_template("login.html", portal_title=settings["school_name"], school_settings=settings, theme_color=settings["primary_color"], setup_required=not auth_initialized())


def _hash_reset_token(token: str) -> str:
    return hashlib.sha256(token.encode("utf-8")).hexdigest()


def _create_password_reset_token(user_id: int, minutes: int = 20) -> str:
    raw = secrets.token_urlsafe(40)
    token_hash = _hash_reset_token(raw)
    expires = (datetime.utcnow() + timedelta(minutes=minutes)).strftime("%Y-%m-%d %H:%M:%S")
    execute("UPDATE password_reset_tokens SET used_at=CURRENT_TIMESTAMP WHERE user_id=? AND used_at IS NULL", (user_id,))
    execute("INSERT INTO password_reset_tokens(user_id,token_hash,expires_at,requested_ip) VALUES(?,?,?,?)", (user_id, token_hash, expires, request.remote_addr or ""))
    return raw


def _reset_user_from_token(raw_token: str):
    token_hash = _hash_reset_token(raw_token or "")
    row = q("""
        SELECT prt.id, prt.user_id, prt.expires_at, prt.used_at,
               u.full_name, u.username, u.role, u.active
        FROM password_reset_tokens prt
        JOIN users u ON u.id=prt.user_id
        WHERE prt.token_hash=?
        LIMIT 1
    """, (token_hash,), one=True)
    if not row or row["used_at"] or not row["active"]:
        return None
    try:
        expiry = datetime.strptime(row["expires_at"], "%Y-%m-%d %H:%M:%S")
        if datetime.utcnow() > expiry:
            return None
    except Exception:
        return None
    return row


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    settings = school_settings()
    message = None
    message_type = "success"
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        email = request.form.get("email", "").strip().lower()
        generic = (
            "If the account details match a registered account, a secure password-reset "
            "link has been sent to the registered email address. The link is valid for 30 minutes. "
            "If no email is registered or email delivery is unavailable, Admin / ICT will receive a recovery request."
        )
        row = q("SELECT id,full_name,username,role,email,active FROM users WHERE lower(username)=? AND role!='System' LIMIT 1", (username,), one=True) if username else None

        # Basic anti-abuse throttling by IP and account. We still return the same public message.
        recent_ip = q("SELECT COUNT(*) AS n FROM password_reset_tokens WHERE requested_ip=? AND created_at >= datetime('now','-10 minutes')", (request.remote_addr or "",), one=True)
        recent_user = q("SELECT COUNT(*) AS n FROM password_reset_tokens WHERE user_id=? AND created_at >= datetime('now','-10 minutes')", (row["id"],), one=True) if row else {"n": 0}
        throttled = int((recent_ip or {"n": 0})["n"] or 0) >= 5 or int((recent_user or {"n": 0})["n"] or 0) >= 3

        if row and row["active"] and email and row["email"] and email == row["email"].strip().lower() and not throttled:
            raw = _create_password_reset_token(row["id"], minutes=30)
            reset_link = url_for("password_reset", token=raw, _external=True)
            sent = _send_password_reset_email(row, reset_link)
            if sent:
                message = generic
                audit(row["id"], row["full_name"], "Password Reset Request", "Self-service password reset link emailed to the account's registered email address.")
            else:
                # Email delivery is intentionally required for self-service reset.
                # A recovery request is created so the institution can still help the user.
                execute("INSERT INTO password_reset_requests(username,reason) VALUES(?,?)", (row["username"], "Self-service reset matched the registered email, but outbound email delivery is not configured or failed."))
                recipients = q("SELECT id FROM users WHERE role IN ('Admin','ICT') AND active=1")
                for admin_row in recipients:
                    notify_user(admin_row["id"], "Password reset assistance requested", f"A password reset was requested for username '{row['username']}', but email delivery was unavailable.", url_for("password_reset_requests"), "High")
                message_type = "warning"
                message = "We could not send the reset email right now. Admin / ICT has received a recovery request for this account."
        else:
            # Never disclose whether a username/email pair exists. Create an assistance
            # request only when a plausible account exists and the request is not obviously abusive.
            if row and not throttled:
                execute("INSERT INTO password_reset_requests(username,reason) VALUES(?,?)", (row["username"], "Self-service password reset could not be completed. The username/email pair was not verified or no recovery email is registered."))
                recipients = q("SELECT id FROM users WHERE role IN ('Admin','ICT') AND active=1")
                for admin_row in recipients:
                    notify_user(admin_row["id"], "Password reset assistance requested", f"A user requested password reset help for username '{row['username']}'. Open the recovery inbox to review.", url_for("password_reset_requests"), "High")
            message_type = "success"
            message = generic
    return render_template("forgot_password.html", settings=settings, message=message, message_type=message_type)


def _send_password_reset_email(user, reset_link: str) -> bool:
    host = os.environ.get("MAIL_SERVER", "").strip()
    username = os.environ.get("MAIL_USERNAME", "").strip()
    password = os.environ.get("MAIL_PASSWORD", "")
    if not host or not username or not password or not user["email"]:
        return False
    try:
        port = int(os.environ.get("MAIL_PORT", "587"))
    except Exception:
        port = 587
    use_tls = os.environ.get("MAIL_USE_TLS", "1").strip().lower() not in {"0", "false", "no", "off"}
    sender = os.environ.get("MAIL_FROM", "").strip() or username
    school_name = school_settings()["school_name"] or "School Portal System"
    msg = EmailMessage()
    msg["Subject"] = f"{school_name} password reset"
    msg["From"] = sender
    msg["To"] = user["email"].strip()
    msg.set_content(
        f"Hello {user['full_name']},\n\n"
        f"A password reset was requested for your {school_name} account ({user['username']}).\n\n"
        f"Use this secure link within 30 minutes:\n{reset_link}\n\n"
        "The link can be used once. If you did not request this, you can ignore this email.\n\n"
        f"{school_name}\n"
    )
    try:
        context = ssl.create_default_context()
        if use_tls:
            with smtplib.SMTP(host, port, timeout=15) as smtp:
                smtp.ehlo(); smtp.starttls(context=context); smtp.ehlo(); smtp.login(username, password); smtp.send_message(msg)
        else:
            with smtplib.SMTP_SSL(host, port, timeout=15, context=context) as smtp:
                smtp.login(username, password); smtp.send_message(msg)
        return True
    except Exception as exc:
        app.logger.warning("Password reset email failed: %s", exc)
        return False


@app.route("/password-reset/<token>", methods=["GET", "POST"])
def password_reset(token):
    row = _reset_user_from_token(token)
    if not row:
        return render_template("reset_password.html", settings=school_settings(), invalid=True, token="", user=None, error="This reset link is invalid, expired, or has already been used.", completed=False), 400
    if request.method == "POST":
        password = request.form.get("password", "")
        confirm = request.form.get("confirm_password", "")
        if len(password) < 8:
            return render_template("reset_password.html", settings=school_settings(), invalid=False, token=token, user=row, error="Use a new password with at least 8 characters.", completed=False)
        if password != confirm:
            return render_template("reset_password.html", settings=school_settings(), invalid=False, token=token, user=row, error="The two new passwords do not match.", completed=False)
        execute("UPDATE users SET password_hash=? WHERE id=?", (generate_password_hash(password), row["user_id"]))
        execute("UPDATE password_reset_tokens SET used_at=CURRENT_TIMESTAMP WHERE id=?", (row["id"],))
        audit(row["user_id"], row["full_name"], "Password Reset Completed", "Account password changed through the verified self-service reset link.")
        notify_user(row["user_id"], "Password reset complete", "Your password was successfully changed. If you did not make this change, contact Admin / ICT immediately.", "/login")
        return render_template("reset_password.html", settings=school_settings(), invalid=False, token="", user=row, error=None, completed=True)
    return render_template("reset_password.html", settings=school_settings(), invalid=False, token=token, user=row, error=None, completed=False)


@app.route("/admin/password-reset-requests")
@login_required
@role_required("Admin", "ICT")
def password_reset_requests():
    rows = q("""
        SELECT pr.*, u.full_name AS resolved_by_name
        FROM password_reset_requests pr
        LEFT JOIN users u ON u.id=pr.resolved_by
        ORDER BY CASE WHEN pr.status='Open' THEN 0 ELSE 1 END, pr.created_at DESC, pr.id DESC
        LIMIT 200
    """)
    return render_template("password_reset_requests.html", settings=school_settings(), requests=rows, actor_name=current_user()["full_name"], role=current_user()["role"])


@app.route("/admin/password-reset-requests/<int:request_id>/close", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def close_password_reset_request(request_id: int):
    execute("UPDATE password_reset_requests SET status='Closed',resolved_at=CURRENT_TIMESTAMP,resolved_by=? WHERE id=?", (current_user()["id"], request_id))
    flash("Password reset assistance request closed.", "success")
    return redirect(url_for("password_reset_requests"))


@app.route("/admin/password-reset-requests/<int:request_id>/notify", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def notify_reset_requester(request_id: int):
    item = q("SELECT * FROM password_reset_requests WHERE id=?", (request_id,), one=True)
    if not item:
        abort(404)
    user = q("SELECT id,full_name FROM users WHERE lower(username)=? LIMIT 1", (item["username"].lower(),), one=True)
    if user:
        notify_user(user["id"], "Your password reset request is being handled", "Admin / ICT has received your password reset request. Please follow the institution's sign-in guidance or contact the office directly if you need further assistance.", "/")
    flash("The account was notified where a matching account exists.", "success")
    return redirect(url_for("password_reset_requests"))


@app.route("/admin/password-reset-requests/<int:request_id>/generate", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def generate_reset_for_request(request_id: int):
    item = q("SELECT * FROM password_reset_requests WHERE id=?", (request_id,), one=True)
    if not item:
        abort(404)
    user = q("SELECT id,full_name,username,role,active FROM users WHERE lower(username)=? AND role!='System' LIMIT 1", (item["username"].lower(),), one=True)
    if not user or not user["active"]:
        flash("No active account matches this recovery request.", "warning")
        return redirect(url_for("password_reset_requests"))
    raw = _create_password_reset_token(user["id"], minutes=20)
    reset_link = url_for("password_reset", token=raw, _external=True)
    actor=current_user()
    audit(actor["id"], actor["full_name"], "Password Reset Link Generated", f"Generated a supervised 20-minute reset link for {user['username']}.")
    flash(f"Supervised reset link for {user['username']} (valid 20 minutes): {reset_link}", "success")
    return redirect(url_for("password_reset_requests"))


@app.route("/login", methods=["GET", "POST"])
def login():
    settings = school_settings()
    role = selected_role_from_request()
    if role == "Parent" and not parent_portal_enabled():
        flash("Parent / Guardian portal is disabled for this institution mode.", "warning")
        return redirect(url_for("index"))
    # Keep the legacy single-account/passwordless portal shortcut on GET, but
    # NEVER bypass an explicit credential POST. Accounts created by Admin/ICT
    # must always be able to authenticate with their own username/password.
    if request.method == "GET" and role and role in ALL_PORTAL_ROLES:
        # All portal roles use explicit credentials. This prevents one logged-in role
        # from silently taking over another role workspace in the same browser.
        pass
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        password = request.form.get("password", "")
        role = selected_role_from_request().strip()
        # Identity is established by username + password. The role selected on the
        # login screen is only a convenience hint and must never lock out a valid
        # account when its stored role is authoritative. After credentials are
        # verified, the account's own role determines the dashboard and permissions.
        user = q("SELECT * FROM users WHERE lower(username)=? AND active=1", (username,), one=True)
        # Student-friendly login: exact full name + admission number is the bootstrap
        # credential. Once the student changes their password, the admission number
        # no longer authenticates the account.
        admission_login_ok = False
        parent_name_login_ok = False
        if not user:
            student_candidates = q("""
                SELECT u.* FROM users u
                JOIN students s ON s.id=u.student_id
                WHERE u.active=1 AND u.role='Student' AND lower(trim(s.full_name))=?
                ORDER BY u.id DESC
            """, (username,),)
            for candidate in student_candidates:
                student = q("SELECT admission_no FROM students WHERE id=? AND active=1", (candidate["student_id"],), one=True)
                if student and str(password).strip().lower() == str(student["admission_no"] or "").strip().lower():
                    user = candidate
                    try:
                        admission_login_ok = check_password_hash(candidate["password_hash"], password)
                    except Exception:
                        admission_login_ok = False
                    # Older student accounts may have a different initial hash.
                    # A matching admission number is still accepted once as the
                    # bootstrap credential, then the hash is normalized to it.
                    if not admission_login_ok:
                        admission_login_ok = True
                        execute("UPDATE users SET password_hash=?, username=?, full_name=? WHERE id=?", (generate_password_hash(student["admission_no"]), student["admission_no"], student["full_name"], candidate["id"]))
                    break
        if not user:
            email_candidate=q("SELECT * FROM users WHERE active=1 AND lower(trim(email))=? ORDER BY id DESC LIMIT 1", ((username or '').strip().lower(),), one=True)
            if email_candidate:
                user=email_candidate
                if user['role']=='Student':
                    st=q("SELECT admission_no FROM students WHERE id=? AND active=1",(user['student_id'],),one=True)
                    if st and str(password).strip().lower()==str(st['admission_no'] or '').strip().lower():
                        admission_login_ok=True
        if not user:
            parent_candidates = q("SELECT * FROM users WHERE active=1 AND role='Parent' AND lower(trim(full_name))=? ORDER BY id DESC", (username,))
            for candidate in parent_candidates:
                try:
                    if check_password_hash(candidate["password_hash"], password):
                        user = candidate
                        parent_name_login_ok = True
                        break
                except Exception:
                    continue

        elif user["role"] == "Student":
            student = q("SELECT admission_no FROM students WHERE id=? AND active=1", (user["student_id"],), one=True)
            if student and str(password).strip().lower() == str(student["admission_no"] or "").strip().lower():
                try:
                    admission_login_ok = check_password_hash(user["password_hash"], password)
                except Exception:
                    admission_login_ok = False
        try:
            password_ok = bool(user) and check_password_hash(user["password_hash"], password)
        except Exception:
            password_ok = False
        if not admission_login_ok and not parent_name_login_ok and not password_ok:
            if user and user["role"] == "Student":
                flash("Invalid student name/admission number or password.", "danger")
                return render_template("login.html", portal_title=settings["school_name"], school_settings=settings, theme_color=settings["primary_color"], login_role=role, error="Invalid student name/admission number or password.", success=("Password reset successfully. You can now sign in with your new password." if request.args.get("reset")=="success" else None), setup_required=not auth_initialized())
            flash("Invalid username or password.", "danger")
            return render_template("login.html", portal_title=settings["school_name"], school_settings=settings, theme_color=settings["primary_color"], login_role=role, error="Invalid username or password.", success=("Password reset successfully. You can now sign in with your new password." if request.args.get("reset")=="success" else None), setup_required=not auth_initialized())
        # A student-linked account is always a Student account. This repairs any
        # legacy/migrated record that accidentally retained a staff role and prevents
        # student credentials from ever entering the Teacher workspace.
        if user and user["student_id"]:
            execute("UPDATE users SET role='Student', workspace_type='Student' WHERE id=?", (user["id"],))
            user = q("SELECT * FROM users WHERE id=? AND active=1", (user["id"],), one=True)
        # A selected portal role can never override the account's stored role.
        # This is especially important for Admin vs ICT separation.
        role = user["role"]
        # Normal password login always goes directly to the person's dashboard.
        # Staff QR is only an optional convenience after the first successful password login.
        if user["role"] not in {"Student", "Parent", "System"} and "qr_login_enabled" in user.keys():
            execute("UPDATE users SET qr_login_enabled=1,last_password_login_at=CURRENT_TIMESTAMP,qr_access_token=COALESCE(NULLIF(qr_access_token,''),lower(hex(randomblob(16)))) WHERE id=?", (user["id"],))
            user=q("SELECT * FROM users WHERE id=?",(user["id"],),one=True)
        login_id=record_login_event(user,'Password')
        session.clear(); session.permanent=True
        session["user_id"]=user["id"]; session["active_portal_role"]=user["role"]; session["login_event_id"]=login_id; session["login_location_pending"]=1
        session["auth_ticket"] = _issue_auth_ticket(user["id"])
        return redirect(specialized_dashboard_for(user))
    return render_template("login.html", portal_title=settings["school_name"], school_settings=settings, theme_color=settings["primary_color"], login_role=role, success=("Password reset successfully. You can now sign in with your new password." if request.args.get("reset")=="success" else None), setup_required=not auth_initialized())


@app.route("/register", methods=["GET", "POST"])
def register_admin():
    if auth_initialized():
        return redirect(url_for("login"))
    settings = school_settings()
    if request.method == "POST":
        full_name=request.form.get("full_name","").strip(); username=request.form.get("username","").strip(); password=request.form.get("password","")
        confirm=request.form.get("confirm_password","")
        if not full_name or not username or len(password) < 4 or password != confirm:
            flash("Enter a name, unique username, and matching password of at least 4 characters.", "danger")
            return render_template("register.html", school_settings=settings, portal_title=settings["school_name"])
        try:
            execute("INSERT INTO users(full_name,username,password_hash,role,active) VALUES(?,?,?,?,1)", (full_name,username,generate_password_hash(password),"Admin"))
            execute("UPDATE school_settings SET auth_initialized=1 WHERE id=1")
        except sqlite3.IntegrityError:
            flash("That username is already in use.", "danger")
            return render_template("register.html", school_settings=settings, portal_title=settings["school_name"])
        flash("Administrator account created. Please log in.", "success")
        return redirect(url_for("login", role="Admin"))
    return render_template("register.html", school_settings=settings, portal_title=settings["school_name"])


def enter_role(role: str):
    """Enter a role portal without ever guessing or swapping identities.

    Public role launchers (/teacher and /student) are login entry points when no
    account is authenticated. Once logged in, the account's stored role is the
    only authority and the launcher redirects to that account's own dashboard.
    """
    if role not in ALL_PORTAL_ROLES:
        abort(404)
    existing=current_user()
    if not existing:
        # Never fabricate a user here. The old implementation could fall into
        # presentation/demo selection and make one role appear as another.
        return redirect(url_for("login", role=role))
    if existing["role"] != role:
        if existing["role"] in ALL_PORTAL_ROLES:
            return redirect(role_target(existing["role"]))
        abort(403)
    return redirect(role_target(role))


@app.route("/admin")
def admin_entry():
    abort(404)

@app.route(ADMIN_LOGIN_PATH)
def admin_hidden_entry():
    return enter_role("Admin")
@app.route("/ict")
def ict_entry(): return enter_role("ICT")
@app.route("/finance")
def finance_entry(): return enter_role("Finance")
@app.route("/teacher")
def teacher_entry(): return enter_role("Teacher")
@app.route("/student")
def student_entry(): return enter_role("Student")
@app.route("/parent")
def parent_entry(): return enter_role("Parent")
@app.route("/librarian")
def librarian_entry(): return enter_role("Librarian")


@app.route("/communication")
@login_required
def communication_center():
    user=current_user(); users=q("SELECT id,full_name,username,role,workspace_type FROM users WHERE active=1 AND id!=? AND role!='System' ORDER BY full_name",(user['id'],)); inbox=q("SELECT m.*,u.full_name AS sender_name,u.role AS sender_role FROM communication_messages m JOIN users u ON u.id=m.sender_user_id WHERE m.recipient_user_id=? ORDER BY m.created_at DESC LIMIT 80",(user['id'],)); sent=q("SELECT m.*,u.full_name AS recipient_name,u.role AS recipient_role FROM communication_messages m JOIN users u ON u.id=m.recipient_user_id WHERE m.sender_user_id=? ORDER BY m.created_at DESC LIMIT 40",(user['id'],)); return render_template('communication.html',settings=school_settings(),actor_name=user['full_name'],role=user['role'],users=users,inbox=inbox,sent=sent)

@app.route("/communication/send",methods=['POST'])
@login_required
def communication_send():
    rid=request.form.get('recipient_user_id',type=int); body=request.form.get('body','').strip(); target=q('SELECT id,full_name FROM users WHERE id=? AND active=1',(rid,),one=True)
    if not target or not body: flash('Choose a recipient and enter a message.','danger'); return redirect(url_for('communication_center'))
    message_id=execute('INSERT INTO communication_messages(sender_user_id,recipient_user_id,body) VALUES(?,?,?)',(current_user()['id'],rid,body)); notify_user(rid, f'New message from {current_user()["full_name"]}', body[:180], url_for('communication_center'), 'High'); audit(current_user()['id'],current_user()['full_name'],'Internal Message',f'Message sent to {target["full_name"]}.'); flash('Message sent.','success'); return redirect(url_for('communication_center'))

@app.route("/communication/read/<int:message_id>",methods=['POST'])
@login_required
def communication_read(message_id:int):
    execute('UPDATE communication_messages SET read_at=CURRENT_TIMESTAMP WHERE id=? AND recipient_user_id=?',(message_id,current_user()['id'])); return redirect(url_for('communication_center'))

@app.route("/teacher/class-attendance",methods=['GET','POST'])
@login_required
@role_required('Teacher')
def teacher_class_attendance():
    user=current_user()
    class_rows=q("""SELECT DISTINCT class_name FROM (
        SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?
        UNION ALL
        SELECT class_name FROM student_teacher_assignments WHERE teacher_user_id=? AND active=1 AND class_name!=''
        UNION ALL
        SELECT s.grade AS class_name FROM student_teacher_assignments sta JOIN students s ON s.id=sta.student_id WHERE sta.teacher_user_id=? AND sta.active=1
    ) WHERE TRIM(COALESCE(class_name,''))!='' ORDER BY class_name""",(user['id'],user['id'],user['id']))
    assigned=[r['class_name'] for r in class_rows]
    subject_rows=q("""SELECT DISTINCT subject FROM (
        SELECT subject FROM teacher_assignments WHERE teacher_user_id=? AND active=1
        UNION ALL SELECT subject FROM student_teacher_assignments WHERE teacher_user_id=? AND active=1
    ) WHERE TRIM(COALESCE(subject,''))!='' ORDER BY subject""",(user['id'],user['id']))
    subject_options=[r['subject'] for r in subject_rows]
    cls=request.values.get('class_name','').strip()
    if cls and cls not in assigned:
        flash('That class register is available from your Admin-assigned learners and classes.','warning')
        return redirect(url_for('teacher_dashboard'))
    subject=request.values.get('subject','').strip() or (subject_options[0] if subject_options else 'General')
    date=request.values.get('attendance_date','').strip() or datetime.utcnow().strftime('%Y-%m-%d')
    students=[]
    if cls:
        # Class teachers get the full class plus any individually allocated
        # learners in that class. Subject teachers do not receive this route.
        students=q("""SELECT DISTINCT s.id,s.full_name,s.admission_no
                     FROM students s
                     LEFT JOIN class_teacher_assignments cta ON cta.class_name=s.grade AND cta.teacher_user_id=?
                     LEFT JOIN student_teacher_assignments sta ON sta.student_id=s.id AND sta.teacher_user_id=? AND sta.active=1 AND (sta.class_name='' OR sta.class_name=s.grade)
                     WHERE s.active=1 AND s.grade=? AND (cta.id IS NOT NULL OR sta.id IS NOT NULL)
                     ORDER BY s.full_name""",(user['id'],user['id'],cls))
    existing={r['student_id']:r for r in q('SELECT * FROM class_attendance WHERE teacher_user_id=? AND class_name=? AND subject=? AND attendance_date=?',(user['id'],cls,subject,date))} if cls else {}
    if request.method=='POST':
        saved=0
        for st in students:
            status=request.form.get(f"attendance_{st['id']}",'').strip()
            note=request.form.get(f"note_{st['id']}",'').strip()
            if status not in {'Present','Absent','Late','Excused'}:
                execute("DELETE FROM class_attendance WHERE teacher_user_id=? AND student_id=? AND class_name=? AND subject=? AND attendance_date=?",(user['id'],st['id'],cls,subject,date)); continue
            execute("INSERT INTO class_attendance(teacher_user_id,student_id,class_name,subject,attendance_date,status,note) VALUES(?,?,?,?,?,?,?) ON CONFLICT(teacher_user_id,student_id,class_name,subject,attendance_date) DO UPDATE SET status=excluded.status,note=excluded.note",(user['id'],st['id'],cls,subject,date,status,note))
            # An absence reason is a permanent student record. Keep one
            # attendance-reason record per learner/day/subject and update it
            # when the teacher changes the reason.
            if status == 'Absent' and note:
                title=f'Attendance absence — {date}'
                existing_record=q("SELECT id FROM student_records WHERE student_id=? AND author_user_id=? AND category='Attendance' AND title=? ORDER BY id DESC LIMIT 1",(st['id'],user['id'],title),one=True)
                if existing_record:
                    execute("UPDATE student_records SET content=?,updated_at=CURRENT_TIMESTAMP WHERE id=?",(note,existing_record['id']))
                else:
                    execute("INSERT INTO student_records(student_id,author_user_id,category,title,content,visible_to_parent) VALUES(?,?,?,?,?,?)",(st['id'],user['id'],'Attendance',title,note,0))
            saved+=1
        flash(f'{saved} class attendance records saved.','success')
        return redirect(url_for('teacher_class_attendance',class_name=cls,subject=subject,attendance_date=date))
    return render_template('teacher_class_attendance.html',settings=school_settings(),actor_name=user['full_name'],assigned_classes=assigned,classes=assigned,subject_options=subject_options,students=students,selected_class=cls,subject=subject,attendance_date=date,existing=existing,role=user['role'])

@app.route("/finance/external/<int:event_id>/match",methods=['POST'])
@login_required
@role_required('Finance','Admin')
def finance_match_external(event_id:int):
    event=q("SELECT * FROM external_payment_events WHERE id=? AND status!='Matched'",(event_id,),one=True); sid=request.form.get('student_id',type=int); student=q('SELECT * FROM students WHERE id=? AND active=1',(sid,),one=True) if sid else None
    if not event or not student: flash('Payment event or student was not found.','danger'); return redirect(url_for('finance_dashboard'))
    poster=current_user()['id']; pid=execute("INSERT INTO payments(student_id,amount,method,reference_no,recorded_by,status) VALUES(?,?,?,?,?,'Posted')",(sid,event['amount'],event['provider'],event['external_reference'],poster)); recalculate_student_balance(sid); new_balance=float(q("SELECT balance FROM students WHERE id=?",(sid,),one=True)['balance'] or 0); notify_teachers_of_payment(q("SELECT * FROM students WHERE id=?",(sid,),one=True),event['amount'],new_balance); execute("UPDATE external_payment_events SET status='Matched',matched_student_id=?,processed_at=CURRENT_TIMESTAMP WHERE id=?",(sid,event_id)); audit(current_user()['id'],current_user()['full_name'],'Match External Payment',f'External payment {event["external_reference"]} matched to {student["admission_no"]}.'); flash('External payment matched and posted.','success'); return redirect(url_for('finance_dashboard'))

@app.route("/allocate", methods=["GET","POST"])
@login_required
@role_required("Admin","ICT")
def allocate_canonical():
    """Canonical short path for the People → Allocate workspace."""
    return admin_student_allocation()

@app.route("/admin/student-allocation", methods=["GET","POST"])
@login_required
@role_required("Admin","ICT")
def admin_student_allocation():
    user=current_user()
    if request.method=="POST":
        action=(request.form.get("action") or "allocation").strip()
        if action=="leadership":
            department_id=request.form.get("department_id",type=int)
            dean_id=request.form.get("dean_user_id",type=int) or None
            deputy_id=request.form.get("deputy_user_id",type=int) or None
            dept=q("SELECT id,name FROM departments WHERE id=? AND active=1",(department_id,),one=True) if department_id else None
            for uid in [dean_id,deputy_id]:
                if uid and not q("SELECT 1 FROM users WHERE id=? AND active=1 AND role IN ('Teacher','Admin','ICT')",(uid,),one=True):
                    if dean_id==uid: dean_id=None
                    if deputy_id==uid: deputy_id=None
            if dept:
                execute("INSERT INTO department_leadership(department_id,dean_user_id,deputy_user_id) VALUES(?,?,?) ON CONFLICT(department_id) DO UPDATE SET dean_user_id=excluded.dean_user_id,deputy_user_id=excluded.deputy_user_id,updated_at=CURRENT_TIMESTAMP",(department_id,dean_id,deputy_id))
                audit(user["id"],user["full_name"],"Department Leadership",f"{dept['name']} leadership updated.")
                flash("Department leadership saved.","success")
            else:
                flash("Choose a valid department.","danger")
            return redirect(url_for("admin_student_allocation"))
        if action=="teacher_profile":
            teacher_id=request.form.get("teacher_user_id",type=int)
            teacher=q("SELECT id,full_name,role FROM users WHERE id=? AND active=1 AND role='Teacher'",(teacher_id,),one=True) if teacher_id else None
            department=(request.form.get("teacher_department") or "").strip()
            title=(request.form.get("teacher_title") or "").strip()
            leadership_role=(request.form.get("leadership_role") or "").strip()
            if teacher and department and q("SELECT 1 FROM departments WHERE active=1 AND name=?",(department,),one=True):
                execute("UPDATE users SET department=?,title=?,leadership_role=?,leadership_level=? WHERE id=?",(department,title,leadership_role,1 if leadership_role in {'Dean','Deputy','Deputy Principal','HOD','Head of Department'} else 0,teacher_id))
                audit(user["id"],user["full_name"],"Teacher Placement",f"{teacher['full_name']} placed in {department} as {leadership_role or title or 'Teacher'}.")
                flash("Teacher department and role saved.","success")
            else:
                flash("Choose a valid teacher and department.","danger")
            return redirect(url_for("admin_student_allocation"))

        teacher_id=request.form.get("teacher_user_id",type=int)
        scope=(request.form.get("scope") or "Subject").strip()
        subject=(request.form.get("subject") or "General").strip() or "General"
        selected_ids=[]
        for raw in request.form.getlist("student_ids"):
            try: selected_ids.append(int(raw))
            except Exception: pass
        selected_ids=list(dict.fromkeys(selected_ids))
        grades=[]
        for raw in request.form.getlist("grade_names"):
            value=(raw or "").strip()
            if value: grades.append(value)
        if grades:
            marks=','.join('?'*len(grades))
            rows=q(f"SELECT id FROM students WHERE active=1 AND grade IN ({marks})",tuple(grades))
            selected_ids.extend(int(r['id']) for r in rows)
        selected_ids=list(dict.fromkeys(selected_ids))
        teacher=q("SELECT id,full_name FROM users WHERE id=? AND active=1 AND role='Teacher'",(teacher_id,),one=True) if teacher_id else None
        if not teacher:
            flash("Choose a valid teacher.","danger")
            return redirect(url_for("admin_student_allocation"))
        students=q(f"SELECT id,full_name,grade FROM students WHERE active=1 AND id IN ({','.join('?'*len(selected_ids))})",tuple(selected_ids)) if selected_ids else []
        if not students:
            flash("Select at least one student or one grade/class.","danger")
            return redirect(url_for("admin_student_allocation"))
        class_name=(request.form.get("class_name") or "").strip()
        saved=0
        for student in students:
            this_class=class_name or student['grade']
            if scope=="Class Teacher":
                execute("INSERT INTO class_teacher_assignments(class_name,teacher_user_id,assigned_by) VALUES(?,?,?) ON CONFLICT(class_name) DO UPDATE SET teacher_user_id=excluded.teacher_user_id,assigned_by=excluded.assigned_by",(this_class,teacher_id,user["id"]))
            else:
                execute("INSERT OR IGNORE INTO student_teacher_assignments(student_id,teacher_user_id,class_name,subject,scope,assigned_by) VALUES(?,?,?,?,?,?)",(student['id'],teacher_id,this_class,subject,scope,user['id']))
            saved+=1
        audit(user["id"],user["full_name"],"Student Teacher Allocation",f"{saved} learners allocated to {teacher['full_name']} ({subject} / {scope}).")
        flash(f"{saved} learner{'s' if saved != 1 else ''} allocated to {teacher['full_name']}.","success")
        return redirect(url_for("admin_student_allocation"))

    students=q("SELECT id,full_name,admission_no,grade FROM students WHERE active=1 ORDER BY grade,full_name")
    teachers=q("SELECT id,full_name,title,department,leadership_role FROM users WHERE active=1 AND role='Teacher' ORDER BY full_name")
    departments=q("SELECT id,name,category FROM departments WHERE active=1 ORDER BY name")
    grades=q("SELECT DISTINCT grade FROM students WHERE active=1 AND TRIM(COALESCE(grade,''))!='' ORDER BY grade")
    subjects=q("SELECT subject,department FROM subjects_catalog WHERE active=1 ORDER BY department,subject")
    allocations=q("""SELECT sta.*,s.full_name AS student_name,s.admission_no,t.full_name AS teacher_name,t.department
                     FROM student_teacher_assignments sta JOIN students s ON s.id=sta.student_id JOIN users t ON t.id=sta.teacher_user_id
                     WHERE sta.active=1 ORDER BY s.grade,s.full_name,t.full_name,sta.subject""")
    class_teachers=q("""SELECT cta.class_name,u.full_name AS teacher_name,u.department FROM class_teacher_assignments cta JOIN users u ON u.id=cta.teacher_user_id ORDER BY cta.class_name""")
    class_teacher_map={r["class_name"]: r["teacher_name"] for r in class_teachers}
    leadership=q("""SELECT d.id,d.name,du.full_name AS dean_name,pu.full_name AS deputy_name
                    FROM departments d LEFT JOIN department_leadership dl ON dl.department_id=d.id
                    LEFT JOIN users du ON du.id=dl.dean_user_id LEFT JOIN users pu ON pu.id=dl.deputy_user_id
                    WHERE d.active=1 ORDER BY d.name""")
    return render_template("admin_student_allocation.html",settings=school_settings(),actor_name=user["full_name"],role=user["role"],students=students,teachers=teachers,departments=departments,grades=grades,subjects=subjects,allocations=allocations,class_teachers=class_teachers,class_teacher_map=class_teacher_map,leadership=leadership)

@app.route("/admin/class-teachers", methods=["GET","POST"])
@login_required
@role_required("Admin","ICT")
def admin_class_teachers():
    if request.method == "POST":
        class_name=request.form.get("class_name","").strip(); teacher_id=request.form.get("teacher_user_id",type=int)
        teacher=q("SELECT id,full_name FROM users WHERE id=? AND active=1 AND role='Teacher'",(teacher_id,),one=True) if teacher_id else None
        if not class_name or not teacher:
            flash("Choose a class and an active Teacher.","danger")
        else:
            actor_id=current_user()['id']
            execute("INSERT INTO class_teacher_assignments(class_name,teacher_user_id,assigned_by) VALUES(?,?,?) ON CONFLICT(class_name) DO UPDATE SET teacher_user_id=excluded.teacher_user_id,assigned_by=excluded.assigned_by",(class_name,teacher_id,actor_id))
            for st in q("SELECT id FROM students WHERE active=1 AND lower(grade)=lower(?)", (class_name,)):
                execute("INSERT OR IGNORE INTO student_teacher_assignments(student_id,teacher_user_id,class_name,subject,scope,assigned_by,active) VALUES(?,?,?,'General','Class Teacher',?,1)",(st['id'],teacher_id,class_name,actor_id))
                auto_place_new_student(st['id'], class_name, actor_id)
            audit(actor_id,current_user()['full_name'],'Class Teacher Assignment',f'{class_name} assigned to {teacher["full_name"]}; existing learners were routed automatically.')
            flash("Class teacher assignment saved.","success")
        return redirect(url_for('admin_class_teachers'))
    classes=q("SELECT DISTINCT grade AS class_name FROM students WHERE active=1 AND grade!='' ORDER BY grade")
    teachers=q("SELECT id,full_name,title,department FROM users WHERE active=1 AND role='Teacher' ORDER BY full_name")
    assignments=q("SELECT a.*,u.full_name AS teacher_name,u.title,u.department FROM class_teacher_assignments a JOIN users u ON u.id=a.teacher_user_id ORDER BY a.class_name")
    return render_template('admin_class_teachers.html',settings=school_settings(),classes=classes,teachers=teachers,assignments=assignments)


# -------------------------------------------------------------------
# Flashcards: a small, canonical learning service shared by Teacher/Admin/Student
# -------------------------------------------------------------------
def _flashcard_role_scope(user):
    return user and user["role"] in {"Teacher", "Admin", "ICT", "Student"}

def _teacher_has_student(user_id, student_id):
    return bool(q("""SELECT 1 FROM student_teacher_assignments WHERE teacher_user_id=? AND student_id=? AND active=1
                   UNION SELECT 1 FROM class_teacher_assignments c JOIN students s ON s.grade=c.class_name WHERE c.teacher_user_id=? AND s.id=?""",
                  (user_id, student_id, user_id, student_id), one=True))

def _flashcard_deck_visible(deck, user):
    if not deck or not user:
        return False
    if user["role"] in {"Admin", "ICT"} or deck["owner_user_id"] == user["id"]:
        return True
    if user["role"] == "Student":
        student_id=user["student_id"]
        if not student_id: return False
        student=q("SELECT grade FROM students WHERE id=? AND active=1", (student_id,), one=True)
        return bool(student and (not deck["class_name"] or deck["class_name"].lower() == (student["grade"] or "").lower()))
    return False

@app.route("/flashcards", methods=["GET","POST"])
@login_required
def flashcards_workspace():
    user=current_user()
    if not _flashcard_role_scope(user): abort(403)
    if request.method == "POST":
        action=(request.form.get("action") or "").strip()
        if action == "create_deck" and user["role"] in {"Teacher","Admin","ICT"}:
            title=(request.form.get("title") or "").strip()
            subject=(request.form.get("subject") or "General").strip() or "General"
            class_name=(request.form.get("class_name") or "").strip()
            description=(request.form.get("description") or "").strip()
            if not title:
                flash("Deck title is required.","danger")
            else:
                execute("INSERT INTO flashcard_decks(title,description,subject,class_name,owner_user_id) VALUES(?,?,?,?,?)",(title,description,subject,class_name,user["id"]))
                flash("Flashcard deck created.","success")
            return redirect(url_for("flashcards_workspace"))
        if action == "create_card" and user["role"] in {"Teacher","Admin","ICT"}:
            deck_id=request.form.get("deck_id",type=int)
            deck=q("SELECT * FROM flashcard_decks WHERE id=? AND active=1",(deck_id,),one=True) if deck_id else None
            if not deck or not _flashcard_deck_visible(deck,user): abort(403)
            front=(request.form.get("front") or "").strip(); back=(request.form.get("back") or "").strip(); hint=(request.form.get("hint") or "").strip()
            if not front or not back: flash("Both the question and answer are required.","danger")
            else:
                pos=q("SELECT COALESCE(MAX(position),0)+1 AS n FROM flashcards WHERE deck_id=?",(deck_id,),one=True)["n"]
                execute("INSERT INTO flashcards(deck_id,front,back,hint,position) VALUES(?,?,?,?,?)",(deck_id,front,back,hint,pos))
                execute("UPDATE flashcard_decks SET updated_at=CURRENT_TIMESTAMP WHERE id=?",(deck_id,))
                flash("Flashcard added.","success")
            return redirect(url_for("flashcards_workspace",deck_id=deck_id))
        if action == "delete_deck" and user["role"] in {"Teacher","Admin","ICT"}:
            deck_id=request.form.get("deck_id",type=int); deck=q("SELECT * FROM flashcard_decks WHERE id=? AND active=1",(deck_id,),one=True)
            if deck and _flashcard_deck_visible(deck,user): execute("UPDATE flashcard_decks SET active=0,updated_at=CURRENT_TIMESTAMP WHERE id=?",(deck_id,)); flash("Deck archived.","success")
            return redirect(url_for("flashcards_workspace"))
    decks=q("""SELECT d.*,u.full_name AS owner_name,COUNT(f.id) AS card_count
                FROM flashcard_decks d JOIN users u ON u.id=d.owner_user_id
                LEFT JOIN flashcards f ON f.deck_id=d.id AND f.active=1
                WHERE d.active=1 GROUP BY d.id ORDER BY d.updated_at DESC,d.id DESC""")
    visible=[d for d in decks if _flashcard_deck_visible(d,user)]
    selected_id=request.args.get("deck_id",type=int) or (visible[0]["id"] if visible else None)
    selected=q("SELECT * FROM flashcard_decks WHERE id=? AND active=1",(selected_id,),one=True) if selected_id else None
    if selected and not _flashcard_deck_visible(selected,user): selected=None
    cards=q("SELECT * FROM flashcards WHERE deck_id=? AND active=1 ORDER BY position,id",(selected["id"],)) if selected else []
    progress={r["card_id"]:r for r in q("SELECT * FROM flashcard_progress WHERE user_id=?",(user["id"],))}
    return render_template("flashcards.html",settings=school_settings(),role=user["role"],actor_name=user["full_name"],decks=visible,selected=selected,cards=cards,progress=progress,today=datetime.utcnow().strftime("%Y-%m-%d"))

@app.route("/flashcards/<int:deck_id>/review", methods=["GET","POST"])
@login_required
def flashcards_review(deck_id:int):
    user=current_user(); deck=q("SELECT * FROM flashcard_decks WHERE id=? AND active=1",(deck_id,),one=True)
    if not deck or not _flashcard_deck_visible(deck,user): abort(404)
    cards=q("SELECT * FROM flashcards WHERE deck_id=? AND active=1 ORDER BY position,id",(deck_id,))
    if not cards: return redirect(url_for("flashcards_workspace",deck_id=deck_id))
    idx=max(0,min(request.values.get("index",0,type=int),len(cards)-1))
    card=cards[idx]
    if request.method=="POST":
        rating=request.form.get("rating",type=int)
        if rating not in {0,1,2,3}: abort(400)
        gap={0:0,1:1,2:3,3:7}[rating]
        execute("""INSERT INTO flashcard_progress(user_id,card_id,rating,correct_count,review_count,next_review_at,last_reviewed_at)
                   VALUES(?,?,?,?,?,datetime('now',?),CURRENT_TIMESTAMP)
                   ON CONFLICT(user_id,card_id) DO UPDATE SET rating=excluded.rating,correct_count=flashcard_progress.correct_count+CASE WHEN excluded.rating>=2 THEN 1 ELSE 0 END,review_count=flashcard_progress.review_count+1,next_review_at=excluded.next_review_at,last_reviewed_at=CURRENT_TIMESTAMP""",
                (user["id"],card["id"],rating,1 if rating>=2 else 0,1,f"+{gap} day"))
        if idx+1 < len(cards): return redirect(url_for("flashcards_review",deck_id=deck_id,index=idx+1))
        return redirect(url_for("flashcards_workspace",deck_id=deck_id))
    return render_template("flashcard_review.html",settings=school_settings(),role=user["role"],actor_name=user["full_name"],deck=deck,card=card,index=idx,total=len(cards))

@app.route("/admin/subjects", methods=["GET","POST"])
@login_required
@role_required("Admin","ICT")
def admin_subjects():
    actor=current_user()
    if request.method=='POST':
        action=request.form.get('action','').strip()
        if action == 'add_subject':
            subject=request.form.get('subject','').strip(); dept=request.form.get('department','').strip(); scope=request.form.get('level_scope','All').strip() or 'All'
            if not subject:
                flash('Subject name is required.','danger')
            else:
                existing=q("SELECT id FROM subjects_catalog WHERE lower(trim(subject))=lower(trim(?)) LIMIT 1",(subject,),one=True)
                if existing:
                    execute("UPDATE subjects_catalog SET department=?,level_scope=?,description=?,active=1,created_by=? WHERE id=?",(dept,scope,request.form.get('description','').strip(),actor['id'],existing['id']))
                    flash('Subject restored/updated and made available.','success')
                else:
                    execute("INSERT INTO subjects_catalog(subject,department,level_scope,description,created_by,active) VALUES(?,?,?,?,?,1)",(subject,dept,scope,request.form.get('description','').strip(),actor['id']))
                    flash('Subject published for registration and teaching assignments.','success')
        elif action == 'add_department':
            name=request.form.get('department_name','').strip(); category=request.form.get('department_category','Academic').strip() or 'Academic'
            if not name:
                flash('Department name is required.','danger')
            else:
                existing=q("SELECT id FROM departments WHERE lower(trim(name))=lower(trim(?)) LIMIT 1",(name,),one=True)
                if existing:
                    execute("UPDATE departments SET category=?,active=1 WHERE id=?",(category,existing['id']))
                    flash('Department restored/updated and made available.','success')
                else:
                    execute("INSERT INTO departments(name,category,active) VALUES(?,?,1)",(name,category))
                    flash('Department added.','success')
        elif action == 'deactivate_subjects':
            ids=[]
            for raw in request.form.getlist('subject_ids'):
                try: ids.append(int(raw))
                except (TypeError,ValueError): pass
            if ids:
                placeholders=','.join('?'*len(ids))
                execute(f"UPDATE subjects_catalog SET active=0 WHERE id IN ({placeholders})",tuple(ids))
                flash(f'{len(ids)} subject(s) removed from new selections. Existing academic records were preserved.','success')
        elif action == 'deactivate_departments':
            ids=[]
            for raw in request.form.getlist('department_ids'):
                try: ids.append(int(raw))
                except (TypeError,ValueError): pass
            if ids:
                placeholders=','.join('?'*len(ids))
                execute(f"UPDATE departments SET active=0 WHERE id IN ({placeholders})",tuple(ids))
                flash(f'{len(ids)} department(s) removed from new selections. Existing assignments were preserved.','success')
        elif action == 'restore_subjects':
            ids=[]
            for raw in request.form.getlist('subject_ids'):
                try: ids.append(int(raw))
                except (TypeError,ValueError): pass
            if ids:
                placeholders=','.join('?'*len(ids))
                execute(f"UPDATE subjects_catalog SET active=1 WHERE id IN ({placeholders})",tuple(ids))
                flash(f'{len(ids)} subject(s) restored.','success')
        elif action == 'restore_departments':
            ids=[]
            for raw in request.form.getlist('department_ids'):
                try: ids.append(int(raw))
                except (TypeError,ValueError): pass
            if ids:
                placeholders=','.join('?'*len(ids))
                execute(f"UPDATE departments SET active=1 WHERE id IN ({placeholders})",tuple(ids))
                flash(f'{len(ids)} department(s) restored.','success')
        return redirect(url_for('admin_subjects'))

    rows=q("SELECT s.*,COUNT(ss.id) AS learner_count FROM subjects_catalog s LEFT JOIN student_subjects ss ON ss.subject_id=s.id AND ss.status!='Dropped' GROUP BY s.id ORDER BY s.active DESC,s.department,s.subject")
    departments=q("SELECT d.*,COUNT(DISTINCT u.id) AS staff_count,COUNT(DISTINCT sd.student_id) AS learner_count FROM departments d LEFT JOIN users u ON u.department=d.name AND u.active=1 LEFT JOIN student_departments sd ON sd.department_id=d.id AND sd.status!='Dropped' GROUP BY d.id ORDER BY d.active DESC,d.name")
    return render_template('admin_subjects.html',settings=school_settings(),rows=rows,departments=departments,actor_name=actor['full_name'],role=actor['role'])

@app.route("/teacher/roster")
@login_required
@role_required("Teacher","Admin","ICT")
def teacher_roster():
    user=current_user(); assignments=q("SELECT * FROM teacher_assignments WHERE active=1" + (" AND teacher_user_id=?" if user['role']=='Teacher' else "") + " ORDER BY class_name,subject",(user['id'],) if user['role']=='Teacher' else ())
    selected_class=request.args.get('class_name','').strip(); selected_subject=request.args.get('subject','').strip()
    if user['role']=='Teacher' and selected_class and not q("SELECT 1 FROM teacher_assignments WHERE teacher_user_id=? AND class_name=? AND subject=? AND active=1",(user['id'],selected_class,selected_subject),one=True): abort(403)
    students=[]
    if selected_class and selected_subject:
        students=q("SELECT s.id,s.full_name,s.admission_no,s.grade,ss.status FROM students s JOIN student_subjects ss ON ss.student_id=s.id JOIN subjects_catalog sc ON sc.id=ss.subject_id WHERE s.active=1 AND lower(s.grade)=lower(?) AND lower(sc.subject)=lower(?) AND ss.status='Approved' ORDER BY s.full_name",(selected_class,selected_subject))
    return render_template('teacher_roster.html',settings=school_settings(),assignments=assignments,students=students,selected_class=selected_class,selected_subject=selected_subject,actor_name=user['full_name'],role=user['role'])

@app.route("/performance")
@login_required
def performance_view():
    user=current_user(); leadership=(user['leadership_role'] or '').lower()
    if user['role'] not in {'Admin','ICT','Teacher'} and leadership not in {'dean','hod','deputy','deputy principal','deputy head','principal'}: abort(403)
    classes=sorted({r['grade'] for r in q("SELECT DISTINCT grade FROM students WHERE active=1")})
    selected=request.args.get('class_name','').strip() or (classes[0] if classes else '')
    selected_subject=request.args.get('subject','').strip()
    senior = user['role'] in {'Admin','ICT'} or leadership in {'dean','hod','deputy','deputy principal','deputy head','principal'}
    if user['role']=='Teacher' and not senior:
        assigned_classes=[r['class_name'] for r in q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?",(user['id'],))]
        is_class_teacher=selected in assigned_classes
        if not is_class_teacher:
            allowed=[r['subject'] for r in q("SELECT subject FROM teacher_assignments WHERE teacher_user_id=? AND class_name=? AND active=1 ORDER BY subject",(user['id'],selected))]
            if selected_subject not in allowed:
                selected_subject=allowed[0] if allowed else ''
            if not allowed:
                selected_subject='' 
        else:
            selected_subject=selected_subject if selected_subject else ''
    subjects=q("SELECT DISTINCT subject FROM markbook_entries WHERE class_name=? ORDER BY subject",(selected,)) if selected else []
    if user['role']=='Teacher' and not senior and selected not in assigned_classes:
        subjects=[{'subject':x['subject']} for x in q("SELECT DISTINCT subject FROM teacher_assignments WHERE teacher_user_id=? AND class_name=? AND active=1 ORDER BY subject",(user['id'],selected))]
    rows=markbook_class_summary(selected,selected_subject) if selected else []
    return render_template('performance.html',settings=school_settings(),classes=classes,selected=selected,selected_subject=selected_subject,rows=rows,subjects=subjects,actor_name=user['full_name'],role=user['role'],leadership=user['leadership_role'],senior=senior)

@app.route("/reception")
@app.route("/reception/")
@login_required
def reception_dashboard():
    if not is_reception_user(current_user()): abort(403)
    settings=school_settings(); open_visits=q("SELECT * FROM reception_visits WHERE check_in IS NOT NULL AND check_out IS NULL ORDER BY check_in ASC,id ASC"); recent=q("SELECT * FROM reception_visits ORDER BY id DESC LIMIT 120"); staff=q("SELECT * FROM users WHERE active=1 AND role NOT IN ('Student','Parent','System','Admin') ORDER BY full_name"); students=q("SELECT id,full_name,admission_no,grade FROM students WHERE active=1 ORDER BY grade,full_name LIMIT 500"); subjects=q("SELECT * FROM subjects_catalog WHERE active=1 ORDER BY department,subject")
    me=current_user(); token=me['qr_access_token'] or uuid.uuid4().hex
    if not me['qr_access_token']:
        execute("UPDATE users SET qr_access_token=? WHERE id=?",(token,me['id']))
    office=q("SELECT token FROM attendance_qr_settings WHERE id=1",one=True); payload='ATTEND:'+office['token'] if office else ''
    code=qrcode.QRCode(version=3,box_size=9,border=3); code.add_data(payload); code.make(fit=True); buf=io.BytesIO(); code.make_image().save(buf,format='PNG'); self_qr_data='data:image/png;base64,'+base64.b64encode(buf.getvalue()).decode('ascii')
    return render_template('reception_dashboard.html',settings=settings,actor_name=me['full_name'],role=me['role'],open_visits=open_visits,recent=recent,staff=staff,school_unit=settings['school_name'],school_location=settings['institution_affiliations'] or '',self_qr_data=self_qr_data,self_qr_code=me['position_code'] or me['staff_code'] or '',students=students,subjects=subjects,settings_departments=q("SELECT id,name,category FROM departments WHERE active=1 ORDER BY name"))

@app.route("/reception/scan",methods=["POST"])
@login_required
def reception_scan():
    if not is_reception_user(current_user()): abort(403)
    data=request.get_json(silent=True) or request.form.to_dict()
    try:
        result=record_reception_scan(data.get('action'),data.get('token',''),data.get('device_token',''),data.get('full_name',''),data.get('phone',''),data.get('gender',''),data.get('reason',''),data.get('source','online'),data.get('method','QR'),float(data['latitude']) if data.get('latitude') not in (None,'') else None,float(data['longitude']) if data.get('longitude') not in (None,'') else None,float(data['accuracy']) if data.get('accuracy') not in (None,'') else None,data.get('event_at') or None,data.get('school_unit',''),data.get('school_location',''))
        return jsonify(result)
    except ValueError as exc: return jsonify({'ok':False,'message':str(exc)}),400
    except Exception:
        app.logger.exception('Reception scan failed'); return jsonify({'ok':False,'message':'The reception record could not be saved.'}),500

@app.route("/reception/sync",methods=["POST"])
@login_required
def reception_sync():
    if not is_reception_user(current_user()): abort(403)
    payload=request.get_json(silent=True) or {}; events=payload.get('events') if isinstance(payload,dict) else []; saved=0
    for item in events if isinstance(events,list) else []:
        try:
            r=record_reception_scan(item.get('action'),item.get('token',''),item.get('device_token',''),item.get('full_name',''),item.get('phone',''),item.get('gender',''),item.get('reason',''),'offline-sync',item.get('method','QR'),item.get('latitude'),item.get('longitude'),item.get('accuracy'),item.get('event_at'),item.get('school_unit',''),item.get('school_location','')); saved += 1 if r.get('ok') else 0
        except Exception: continue
    return jsonify({'ok':True,'saved':saved})

@app.route("/reception/visitor",methods=["POST"])
@login_required
def reception_register_visitor():
    if not is_reception_user(current_user()): abort(403)
    name=request.form.get('full_name','').strip(); reason=request.form.get('reason','').strip(); id_number=request.form.get('id_number','').strip()
    if not name or not reason: flash('Visitor name and reason for visit are required.','danger'); return redirect(url_for('reception_dashboard'))
    r=record_reception_scan('IN',device_token=request.form.get('device_token','').strip() or uuid.uuid4().hex,full_name=name,phone=request.form.get('phone','').strip(),gender=request.form.get('gender','').strip(),reason=reason,method='Reception desk',school_unit=request.form.get('school_unit',''),school_location=request.form.get('school_location',''),id_number=id_number)
    if r.get('ok') and r.get('visit_id'):
        execute("UPDATE reception_visits SET person_type='Visitor',id_number=? WHERE id=?",(id_number,r['visit_id']))
    flash(r['message'],'success'); return redirect(url_for('reception_dashboard'))

@app.route("/reception/visitors/search")
@login_required
def reception_visitor_search():
    if not is_reception_user(current_user()): abort(403)
    term=request.args.get('q','').strip()
    if not term: return jsonify([])
    rows=q("SELECT full_name,id_number,phone,MAX(check_in) AS last_visit,COUNT(*) AS visits FROM reception_visits WHERE person_type='Visitor' AND (full_name LIKE ? OR id_number LIKE ? OR phone LIKE ?) GROUP BY full_name,id_number,phone ORDER BY last_visit DESC LIMIT 20",(f'%{term}%',f'%{term}%',f'%{term}%'))
    return jsonify([dict(r) for r in rows])

@app.route("/reception/visit/<int:visit_id>/checkout",methods=["POST"])
@login_required
def reception_checkout_visit(visit_id):
    if not is_reception_user(current_user()): abort(403)
    visit=q("SELECT * FROM reception_visits WHERE id=?",(visit_id,),one=True)
    if not visit: abort(404)
    now=datetime.utcnow().isoformat(timespec='seconds'); execute("UPDATE reception_visits SET check_out=?,source='online',method='Reception desk' WHERE id=?",(now,visit_id))
    if visit['user_id']: execute("INSERT INTO attendance_events(user_id,action,method,event_at,source,device_note) VALUES(?,?,?,?,?,?)",(visit['user_id'],'OUT','Reception desk',now,'online',visit['device_token']))
    notify_users(reception_admin_ids(),'Reception: person checked OUT',f"{visit['full_name']} left reception at {now}.",url_for('reception_dashboard'))
    flash(f"{visit['full_name']} checked out.",'success'); return redirect(url_for('reception_dashboard'))

@app.route("/reception/staff/register",methods=["POST"])
@login_required
def reception_register_staff():
    actor=current_user()
    if not is_reception_user(actor): abort(403)
    role=request.form.get('role','Teacher'); role=role if role in {'Teacher','Librarian'} else 'Teacher'; workspace=request.form.get('workspace_type','Teaching'); workspace=workspace if workspace in {'Teaching','Driver','Reception','Guard','Cook','Other Staff'} else 'Teaching'
    name=request.form.get('full_name','').strip(); username=request.form.get('username','').strip().lower(); password=request.form.get('password','').strip()
    if not name or not username or len(password)<4: flash('Staff name, username and a temporary password are required.','danger'); return redirect(url_for('reception_dashboard'))
    if q("SELECT id FROM users WHERE lower(username)=?",(username,),one=True): flash('That username is already in use.','danger'); return redirect(url_for('reception_dashboard'))
    code=staff_code_for(role,workspace); unit=request.form.get('school_unit','').strip() or school_settings()['school_name']; loc=request.form.get('school_location','').strip(); uid=execute("INSERT INTO users(full_name,username,password_hash,role,active,title,department,phone,gender,workspace_type,school_unit,school_location,reception_enabled,position_code,staff_code) VALUES(?,?,?,?,1,?,?,?,?,?,?,?,?,?,?,?)",(name,username,generate_password_hash(password),role,request.form.get('title','').strip(),request.form.get('department','').strip(),request.form.get('phone','').strip(),request.form.get('gender','').strip(),workspace,unit,loc,1 if workspace=='Reception' else 0,code,code)); execute("UPDATE users SET qr_access_token=? WHERE id=?",(uuid.uuid4().hex,uid)); audit(actor['id'],actor['full_name'],'Reception staff registration',f'{name} registered with staff code {code}.'); flash(f'{name} registered. Staff code: {code}.','success'); return redirect(url_for('reception_dashboard'))

@app.route("/reception/face/enrol",methods=["POST"])
@login_required
def reception_face_enrol():
    if not is_reception_user(current_user()): abort(403)
    student_id=request.form.get('student_id',type=int); user_id=request.form.get('user_id',type=int); descriptor=request.form.get('descriptor_json','').strip()
    if not descriptor or (not student_id and not user_id): return jsonify({'ok':False,'message':'A person and a face descriptor are required.'}),400
    if student_id and not q("SELECT id FROM students WHERE id=?",(student_id,),one=True): return jsonify({'ok':False,'message':'Student not found.'}),404
    if user_id and not q("SELECT id FROM users WHERE id=? AND active=1",(user_id,),one=True): return jsonify({'ok':False,'message':'User not found.'}),404
    if student_id:
        execute("INSERT INTO student_face_profiles(student_id,user_id,image_path,descriptor_json,active) VALUES(?,?,?,?,1) ON CONFLICT(student_id) DO UPDATE SET user_id=excluded.user_id,descriptor_json=excluded.descriptor_json,active=1,enrolled_at=CURRENT_TIMESTAMP",(student_id,user_id or None,'',descriptor))
    else:
        execute("INSERT INTO student_face_profiles(user_id,image_path,descriptor_json,active) VALUES(?,?,?,1) ON CONFLICT(user_id) DO UPDATE SET descriptor_json=excluded.descriptor_json,active=1,enrolled_at=CURRENT_TIMESTAMP",(user_id,'',descriptor))
    return jsonify({'ok':True,'message':'Face profile enrolled.'})

@app.route("/reception/face/verify",methods=["POST"])
@login_required
def reception_face_verify():
    if not is_reception_user(current_user()): abort(403)
    raw=request.form.get('descriptor_json','').strip(); action=(request.form.get('action') or 'IN').upper()
    if not raw or action not in {'IN','OUT'}: return jsonify({'ok':False,'message':'Face data and IN/OUT action are required.'}),400
    try: probe=[float(x) for x in __import__('json').loads(raw)]
    except Exception: return jsonify({'ok':False,'message':'Invalid face descriptor.'}),400
    profiles=q("SELECT fp.*,s.full_name AS student_name,s.id AS sid,u.id AS uid,u.full_name AS user_name,u.role FROM student_face_profiles fp LEFT JOIN students s ON s.id=fp.student_id LEFT JOIN users u ON u.id=fp.user_id WHERE fp.active=1")
    best=None; best_dist=999.0
    for row in profiles:
        try: base=[float(x) for x in __import__('json').loads(row['descriptor_json'] or '[]')]
        except Exception: continue
        if len(base)!=len(probe): continue
        dist=(sum((a-b)*(a-b) for a,b in zip(base,probe))/len(probe))**0.5
        if dist<best_dist: best_dist=dist; best=row
    if not best or best_dist>0.58: return jsonify({'ok':False,'message':'Face not recognized. Use the institution QR or try again.','distance':round(best_dist,4) if best else None}),403
    matched_user=q("SELECT * FROM users WHERE id=? AND active=1",(best['uid'],),one=True) if best['uid'] else None
    if not matched_user and best['sid']:
        matched_user=q("SELECT * FROM users WHERE student_id=? AND active=1 AND role='Student' ORDER BY id LIMIT 1",(best['sid'],),one=True)
    if not matched_user: return jsonify({'ok':False,'message':'The recognized person has no active system account.'}),403
    now=datetime.utcnow().isoformat(timespec='seconds')
    execute("INSERT INTO attendance_events(user_id,action,method,event_at,source,device_note) VALUES(?,?,?,?,?,?)",(matched_user['id'],action, 'Face recognition', now,'online',f'face-distance={best_dist:.4f}'))
    return jsonify({'ok':True,'name':matched_user['full_name'],'action':action,'event_at':now,'message':f'{matched_user["full_name"]} checked {"in" if action=="IN" else "out"} successfully.'})

@app.route("/reception/face/profiles")
@login_required
def reception_face_profiles():
    if not is_reception_user(current_user()): abort(403)
    rows=q("SELECT fp.id,fp.student_id,fp.user_id,fp.descriptor_json,s.full_name AS student_name,u.full_name AS user_name,u.role FROM student_face_profiles fp LEFT JOIN students s ON s.id=fp.student_id LEFT JOIN users u ON u.id=fp.user_id WHERE fp.active=1 ORDER BY COALESCE(s.full_name,u.full_name)")
    return jsonify([dict(r) for r in rows])

@app.route("/reception/export")
@login_required
def reception_export():
    if not is_reception_user(current_user()): abort(403)
    rows=q("SELECT * FROM reception_visits ORDER BY id DESC"); out=io.StringIO(); w=csv.writer(out); w.writerow(['Name','Type','Position','Staff code','Phone','Gender','Reason','School','Location','Time in','Time out','Source','Method','Device token'])
    for r in rows: w.writerow([r['full_name'],r['person_type'],r['position'],r['staff_code'],r['phone'],r['gender'],r['reason'],r['school_unit'],r['school_location'],r['check_in'] or '',r['check_out'] or '',r['source'],r['method'],r['device_token']])
    b=io.BytesIO(out.getvalue().encode('utf-8-sig')); b.seek(0); return send_file(b,mimetype='text/csv',as_attachment=True,download_name=f"{secure_filename(school_settings()['school_name'])}-reception-register.csv")

KENYA_TZ_OFFSET = timedelta(hours=3)

def _utc_now_naive():
    # The application stores timestamps in UTC as naive ISO strings.
    return datetime.utcnow()

def _local_now_naive():
    return datetime.utcnow() + KENYA_TZ_OFFSET

def _local_iso(dt: datetime | None):
    if not dt:
        return ''
    return (dt + KENYA_TZ_OFFSET).strftime('%Y-%m-%d %H:%M:%S')

def _parse_stored_event(raw):
    if not raw:
        return None
    txt=str(raw).replace('Z','').strip()
    try:
        return datetime.fromisoformat(txt)
    except Exception:
        try:
            return datetime.strptime(txt[:19], '%Y-%m-%d %H:%M:%S')
        except Exception:
            return None

def attendance_day_bounds_utc(day):
    day=attendance_date_from_value(day)
    local_start=datetime.strptime(day,'%Y-%m-%d')
    local_end=local_start+timedelta(days=1)
    return (local_start-KENYA_TZ_OFFSET).strftime('%Y-%m-%d %H:%M:%S'), (local_end-KENYA_TZ_OFFSET).strftime('%Y-%m-%d %H:%M:%S')

def _reverse_geocode(latitude, longitude):
    """Turn the exact captured coordinates into a human-readable place without changing the coordinates."""
    try:
        if latitude is None or longitude is None:
            return ''
        lat=float(latitude); lon=float(longitude)
        url='https://nominatim.openstreetmap.org/reverse?format=jsonv2&lat='+urllib.parse.quote(str(lat))+'&lon='+urllib.parse.quote(str(lon))+'&zoom=18&addressdetails=1'
        req=urllib.request.Request(url,headers={'User-Agent':'Prime-School-Portal/1.0 (attendance location lookup)'})
        with urllib.request.urlopen(req,timeout=3) as resp:
            data=json.loads(resp.read().decode('utf-8'))
        addr=data.get('address') or {}
        parts=[]
        for key in ('house_number','road','neighbourhood','suburb','village','town','city','municipality','county','state'):
            val=(addr.get(key) or '').strip()
            if val and val not in parts:
                parts.append(val)
        label=', '.join(parts[:7])
        return label or (data.get('display_name') or '')
    except Exception:
        return ''

def _attendance_event_location(latitude, longitude, location_label=''):
    label=''
    # Never borrow or trust a label from another scan/device. Resolve this event's exact
    # coordinates on the server whenever coordinates exist; only use the submitted label
    # as a last-resort fallback when the reverse-geocoder is unavailable.
    if latitude is not None and longitude is not None:
        label=_reverse_geocode(latitude,longitude)
        if not label:
            label=(location_label or '').strip()
    return label

def _expected_school_days(start_date, end_date):
    days=[]
    cur=start_date
    rules=q("SELECT start_date,end_date,school_day FROM school_calendar WHERE end_date>=? AND start_date<=? ORDER BY start_date",(start_date,end_date))
    while cur<=end_date:
        ds=cur.isoformat(); school_day=1
        for r in rules:
            if r['start_date']<=ds<=r['end_date']:
                school_day=int(r['school_day'])
        if cur.weekday()<5 and school_day:
            days.append(ds)
        cur+=timedelta(days=1)
    return days

def _attendance_range(value=None):
    value=(value or 'day').strip().lower()
    today=_local_now_naive().date()
    spans={'day':1,'week':7,'month':30,'3m':90,'6m':180,'8m':240,'year':365}
    if value not in spans: value='day'
    start=today-timedelta(days=spans[value]-1)
    return value,start,today

def _decorate_attendance_rows(rows):
    out=[]
    for row in rows:
        d=dict(row)
        for key in ('check_in_at','check_out_at','event_at'):
            if d.get(key):
                dt=_parse_stored_event(d[key]); d[key+'_local']=_local_iso(dt) if dt else d[key]
        out.append(d)
    return out

def attendance_date_from_value(value=None):
    raw=(value or '').strip()
    if raw:
        try:
            parsed=datetime.fromisoformat(raw.replace('Z','+00:00'))
            if parsed.tzinfo is not None:
                parsed=parsed.astimezone(__import__('datetime').timezone.utc).replace(tzinfo=None) + KENYA_TZ_OFFSET
            elif 'T' in raw or (' ' in raw and len(raw)>=19):
                parsed=parsed + KENYA_TZ_OFFSET
            return parsed.date().isoformat()
        except Exception:
            try: return datetime.strptime(raw[:10], '%Y-%m-%d').date().isoformat()
            except Exception: pass
    return _local_now_naive().date().isoformat()

def attendance_day_is_closed(day):
    row=q("SELECT status FROM attendance_days WHERE attendance_date=?", (attendance_date_from_value(day),), one=True)
    return bool(row and row['status']=='Closed')

def set_attendance_day(day,status,actor_id=None):
    day=attendance_date_from_value(day)
    if status=='Closed':
        execute("INSERT INTO attendance_days(attendance_date,status,closed_at,closed_by) VALUES(?, 'Closed', CURRENT_TIMESTAMP, ?) ON CONFLICT(attendance_date) DO UPDATE SET status='Closed',closed_at=CURRENT_TIMESTAMP,closed_by=excluded.closed_by",(day,actor_id))
    else:
        execute("INSERT INTO attendance_days(attendance_date,status,closed_at,closed_by) VALUES(?, 'Open', NULL, NULL) ON CONFLICT(attendance_date) DO UPDATE SET status='Open',closed_at=NULL,closed_by=NULL",(day,))

def next_attendance_action(user_id):
    today=_local_now_naive().date().isoformat(); start_utc,end_utc=attendance_day_bounds_utc(today)
    row=q("SELECT action,event_at FROM attendance_events WHERE user_id=? AND event_at>=? AND event_at<? ORDER BY event_at DESC,id DESC LIMIT 1",(user_id,start_utc,end_utc),one=True)
    return 'OUT' if row and row['action']=='IN' else 'IN'

def _payload_float(payload,key):
    raw=payload.get(key) if hasattr(payload,'get') else None
    try: return float(raw) if raw not in (None,'') else None
    except (TypeError,ValueError): return None

def record_account_attendance(user,action,event_at=None,source='online',method='QR',latitude=None,longitude=None,accuracy=None,device_note='',location_label=''):
    action=str(action or '').upper()
    if action not in {'IN','OUT'}: return {'ok':False,'message':'Invalid attendance action.'}
    stamp=event_at or datetime.utcnow().isoformat(timespec='seconds')
    if event_at and 'T' in event_at:
        try:
            stamp=(datetime.fromisoformat(event_at)-KENYA_TZ_OFFSET).isoformat(timespec='seconds')
        except ValueError:
            pass
    if attendance_day_is_closed(stamp): return {'ok':False,'message':f'Attendance for {attendance_date_from_value(stamp)} is closed by the school.','closed':True}
    location_label=_attendance_event_location(latitude,longitude,location_label)
    execute("INSERT INTO attendance_events(user_id,action,method,office_token,event_at,source,latitude,longitude,accuracy,speed_kph,device_note,location_label) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",(user['id'],action,method,'',stamp,source,latitude,longitude,accuracy,None,device_note,location_label))
    position=user['title'] or user['role'] or 'Staff'
    notify_users(attendance_admin_ids(),f'Attendance: {user["full_name"]} checked {"IN" if action=="IN" else "OUT"}',f'{user["full_name"]} ({position}) checked {"in" if action=="IN" else "out"} at {_local_iso(_parse_stored_event(stamp)) or stamp}. Location: {location_label or "Exact coordinates captured; address lookup unavailable"}.',url_for('admin_attendance'))
    return {'ok':True,'message':f'{user["full_name"]} checked {"in" if action=="IN" else "out"}.','action':action,'event_at':stamp,'location_label':location_label,'dashboard':specialized_dashboard_for(user)}

@app.route("/attendance")
@login_required
def attendance_center():
    user=current_user(); settings=school_settings(); events=q("SELECT * FROM attendance_events WHERE user_id=? ORDER BY event_at DESC,id DESC LIMIT 30",(user['id'],)); office=q("SELECT * FROM attendance_qr_settings WHERE id=1",one=True); qr_data=''
    if office:
        img=qrcode.make(f"ATTEND:{office['token']}"); buf=io.BytesIO(); img.save(buf,format='PNG'); qr_data='data:image/png;base64,'+base64.b64encode(buf.getvalue()).decode('ascii')
    return render_template('attendance_center.html',settings=settings,actor_name=user['full_name'],role=user['role'],workspace_type=workspace_type_for_user(user),events=events,office=office,qr_data=qr_data)

@app.route("/attendance/office-qr")
@login_required
def attendance_office_qr():
    if not is_reception_user(current_user()): abort(403)
    office=q("SELECT * FROM attendance_qr_settings WHERE id=1",one=True)
    if not office: abort(404)
    img=qrcode.make(f"ATTEND:{office['token']}"); buf=io.BytesIO(); img.save(buf,format='PNG'); buf.seek(0)
    return send_file(buf,mimetype='image/png',download_name='institution-attendance-qr.png',as_attachment=False)

@app.route("/attendance/rotate",methods=['POST'])
@login_required
@role_required("Admin","ICT")
def attendance_rotate():
    execute("UPDATE attendance_qr_settings SET token=?,updated_at=CURRENT_TIMESTAMP WHERE id=1",(uuid.uuid4().hex,)); audit(current_user()['id'],current_user()['full_name'],'Attendance QR Rotated','Institution attendance QR token rotated.'); flash('Attendance QR rotated. Previous QR is now invalid.','success'); return redirect(request.referrer or url_for('attendance'))

@app.route("/attendance/record",methods=['POST'])
@login_required
def attendance_record():
    token=(request.form.get('token') or '').strip(); action=(request.form.get('action') or '').upper(); event_at=(request.form.get('event_at') or '').strip() or None; office=q("SELECT token FROM attendance_qr_settings WHERE id=1",one=True)
    if action not in {'IN','OUT'} or not office or token!=office['token']: return jsonify({'ok':False,'message':'Invalid institution attendance QR.'}),400
    stamp=event_at or datetime.utcnow().isoformat(timespec='seconds')
    if attendance_day_is_closed(stamp): return jsonify({'ok':False,'message':f'Attendance for {attendance_date_from_value(stamp)} is closed by the school.'}),409
    lat=request.form.get('latitude',type=float); lon=request.form.get('longitude',type=float); accuracy=request.form.get('accuracy',type=float)
    resolved_location=_attendance_event_location(lat,lon,request.form.get('location_label','').strip())
    execute("INSERT INTO attendance_events(user_id,action,method,office_token,event_at,source,latitude,longitude,accuracy,speed_kph,device_note,location_label) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",(current_user()['id'],action,'QR',token,stamp,request.form.get('source','online'),lat,lon,accuracy,request.form.get('speed_kph',type=float),request.form.get('device_note',''),resolved_location))
    notify_users(attendance_admin_ids(),f'Attendance: {current_user()["full_name"]} checked {"IN" if action=="IN" else "OUT"}',f'{current_user()["full_name"]} checked {"in" if action=="IN" else "out"} at {_local_iso(_parse_stored_event(stamp)) or stamp}. Location: {resolved_location or "Exact coordinates captured; address lookup unavailable"}.',url_for('admin_attendance'))
    return jsonify({'ok':True,'message':f'Checked {"in" if action=="IN" else "out"}.','event_at':stamp,'location_label':resolved_location})

@app.route("/attendance/sync",methods=['POST'])
@login_required
def attendance_sync():
    payload=request.get_json(silent=True) or {}; events=payload.get('events') if isinstance(payload,dict) else None; office=q("SELECT token FROM attendance_qr_settings WHERE id=1",one=True)
    if not isinstance(events,list) or not office: return jsonify({'ok':False,'message':'Invalid offline queue.'}),400
    saved=0
    for item in events[:100]:
        if item.get('token')!=office['token'] or str(item.get('action','')).upper() not in {'IN','OUT'}: continue
        if attendance_day_is_closed(item.get('event_at')): continue
        execute("INSERT INTO attendance_events(user_id,action,method,office_token,event_at,source,latitude,longitude,accuracy,speed_kph,device_note,location_label) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",(current_user()['id'],str(item['action']).upper(),'QR',office['token'],item.get('event_at') or None,'offline-sync',item.get('latitude'),item.get('longitude'),item.get('accuracy'),item.get('speed_kph'),item.get('device_note',''),_attendance_event_location(item.get('latitude'),item.get('longitude'),item.get('location_label','')))); saved+=1
    return jsonify({'ok':True,'saved':saved})


SYSTEM_SEARCH_INDEX = [
    {"title":"Home / QR attendance", "description":"Scan or manually record staff check-in and check-out", "url":"/admin-dashboard", "roles":{"Admin","ICT"}, "keywords":"home qr attendance check in check out login offline manual"},
    {"title":"View all employees", "description":"Staff directory, edit, add, delete and attendance history", "url":"/admin/employees", "roles":{"Admin","ICT"}, "keywords":"employees staff people users edit delete add attendance"},
    {"title":"View all learners", "description":"Student directory and learner records", "url":"/admin/learners", "roles":{"Admin","ICT"}, "keywords":"students learners pupils people access directory"},
    {"title":"Attendance", "description":"Day, week, month, 3 months, 6 months, 8 months and year reports", "url":"/admin/attendance", "roles":{"Admin","ICT"}, "keywords":"attendance qr checkin checkout signed in signed out missed absence reports"},
    {"title":"Staff timetable", "description":"Schedules for teachers and relevant staff", "url":"/staff/timetable", "roles":{"Admin","ICT","Teacher","Finance","Driver","Reception","Librarian"}, "keywords":"timetable schedule staff teacher class time"},
    {"title":"Reminders", "description":"Operational reminders and due work", "url":"/staff/reminders", "roles":{"Admin","ICT","Teacher","Finance","Driver","Reception","Librarian"}, "keywords":"reminders tasks alerts due"},
    {"title":"School settings", "description":"School identity, admissions, fees and institution configuration", "url":"/admin-dashboard#settings-panel", "roles":{"Admin","ICT"}, "keywords":"settings school setting configuration fees admissions school name institution"},
    {"title":"Branding & logo", "description":"Logo, background images, fonts, colors and portal appearance", "url":"/admin-dashboard#settings-panel", "roles":{"Admin","ICT"}, "keywords":"logo upload background image branding colors theme font appearance"},
    {"title":"Analytics & summary", "description":"Performance charts, school totals and breakdowns", "url":"/admin-dashboard#analytics-panel", "roles":{"Admin","ICT"}, "keywords":"analytics charts graph graphs summary metrics breakdown performance"},
    {"title":"Finance", "description":"Income, expenses and ledger", "url":"/finance", "roles":{"Admin","Finance"}, "keywords":"finance income expenses payments ledger money"},
    {"title":"Examinations", "description":"Compile grades and school totals", "url":"/admin-dashboard#exam-panel", "roles":{"Admin","ICT"}, "keywords":"exams examinations grades marks totals"},
    {"title":"Subjects", "description":"Subject catalogue and compulsory subjects", "url":"/admin/subjects", "roles":{"Admin","ICT"}, "keywords":"subjects curriculum catalogue compulsory"},
    {"title":"Student & teacher allocation", "description":"Allocate learners to subject teachers, class teachers and departments", "url":"/admin/student-allocation", "roles":{"Admin","ICT"}, "keywords":"students teachers allocation cohort class teacher department dean deputy"},
    {"title":"Class teacher assignments", "description":"Assign teaching staff to classes", "url":"/admin/class-teachers", "roles":{"Admin","ICT"}, "keywords":"class teacher assignments teachers"},
    {"title":"Reception desk", "description":"Reception and operational front desk", "url":"/reception", "roles":{"Admin","ICT","Reception"}, "keywords":"reception desk visitors"},
    {"title":"Backup & restore", "description":"Full system backup and recovery", "url":"/admin/dashboard#backup-panel", "roles":{"Admin"}, "keywords":"backup restore recovery export"},
    {"title":"Exports", "description":"Students, employees, payments and audit exports", "url":"/admin/dashboard#exports-panel", "roles":{"Admin","ICT"}, "keywords":"export csv audit employees students payments"},
    {"title":"System help", "description":"Guides and operational help", "url":"/system-help", "roles":set(ALL_PORTAL_ROLES), "keywords":"help guide support"},
    {"title":"AI assistant", "description":"System assistant", "url":"/ai-assistant", "roles":set(ALL_PORTAL_ROLES), "keywords":"ai assistant help search"},
]

def _search_terms(text):
    return re.findall(r'[a-z0-9]+', (text or '').lower())

@app.route('/system-search')
@login_required
def system_search():
    user=current_user(); role=user['role']; term=(request.args.get('q') or '').strip()
    terms=set(_search_terms(term)); results=[]
    if term:
        for item in SYSTEM_SEARCH_INDEX:
            if role not in item['roles']:
                continue
            hay=' '.join([item['title'],item['description'],item['keywords']]).lower()
            hay_terms=set(_search_terms(hay))
            score=0
            if term.lower() in item['title'].lower(): score+=8
            if term.lower() in item['description'].lower(): score+=4
            score += len(terms & hay_terms)*2
            if score:
                results.append((score,item))
        # Search live people too, but only within the user's allowed scope.
        if role in {'Admin','ICT'}:
            for u in q("SELECT id,full_name,username,role,COALESCE(title,'') AS title,COALESCE(department,'') AS department FROM users WHERE active=1 AND role NOT IN ('System') ORDER BY full_name LIMIT 500"):
                hay=' '.join([u['full_name'] or '',u['username'] or '',u['role'] or '',u['title'] or '',u['department'] or '']).lower()
                score=(8 if term.lower() in (u['full_name'] or '').lower() else 0)+(4 if term.lower() in hay else 0)+len(terms & set(_search_terms(hay)))
                if score:
                    results.append((score,{"title":u['full_name'],"description":f"{u['role']} · {u['title'] or u['department'] or 'Staff account'}","url":f"/users/{u['id']}/edit","kind":"User"}))
            for st in q("SELECT id,full_name,admission_no,grade FROM students WHERE active=1 ORDER BY full_name LIMIT 1000"):
                hay=' '.join([st['full_name'] or '',st['admission_no'] or '',st['grade'] or '']).lower()
                score=(8 if term.lower() in (st['full_name'] or '').lower() else 0)+(4 if term.lower() in hay else 0)+len(terms & set(_search_terms(hay)))
                if score:
                    results.append((score,{"title":st['full_name'],"description":f"Student · {st['admission_no'] or 'No admission number'} · {st['grade'] or ''}","url":f"/students/{st['id']}","kind":"Student"}))
    results=[item for _,item in sorted(results,key=lambda x:(-x[0],x[1].get('title','')))][:40]
    return render_template('system_search.html',settings=school_settings(),actor_name=user['full_name'],role=role,term=term,results=results)


@app.route('/admin/attendance/qr')
@login_required
@role_required('Admin','ICT')
def admin_attendance_qr_image():
    office=q("SELECT * FROM attendance_qr_settings WHERE id=1",one=True)
    if not office: abort(404)
    img=qrcode.make(f"ATTEND:{office['token']}"); buf=io.BytesIO(); img.save(buf,format='PNG'); buf.seek(0)
    return send_file(buf,mimetype='image/png',download_name='institution-attendance-qr.png',as_attachment=False)

@app.route('/admin/attendance/live')
@login_required
@role_required('Admin','ICT')
def admin_attendance_live():
    local_today=_local_now_naive().date().isoformat()
    today_start,today_end=attendance_day_bounds_utc(local_today)
    rows=q("""
        SELECT u.id,u.full_name,u.role,COALESCE(NULLIF(u.title,''),u.role) AS title,
               (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS sign_in_at,
               (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='OUT' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at DESC,a.id DESC LIMIT 1) AS sign_out_at,
               (SELECT a.location_label FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.location_label!='' ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS location,
               (SELECT a.latitude FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS latitude,
               (SELECT a.longitude FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS longitude,
               (SELECT a.accuracy FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS accuracy
        FROM users u
        WHERE u.active=1 AND u.role NOT IN ('Student','Parent','System')
        ORDER BY u.full_name
    """,(today_start,today_end,today_start,today_end,today_start,today_end,today_start,today_end,today_start,today_end))
    out=[]
    for r in rows:
        item=dict(r)
        item['sign_in_local']=_local_iso(_parse_stored_event(item['sign_in_at'])) if item.get('sign_in_at') else None
        item['sign_out_local']=_local_iso(_parse_stored_event(item['sign_out_at'])) if item.get('sign_out_at') else None
        # Role is always read from the staff account itself; never infer it from the scanner/admin account.
        out.append(item)
    return jsonify({'ok':True,'date':local_today,'rows':out})

@app.route('/admin/attendance/manual', methods=['POST'])
@login_required
@role_required('Admin','ICT')
def admin_manual_attendance():
    user_id=request.form.get('user_id',type=int)
    action=(request.form.get('action') or '').upper().strip()
    event_at=(request.form.get('event_at') or '').strip() or None
    location_label=(request.form.get('location_label') or '').strip()[:500]
    note=(request.form.get('device_note') or 'Manual attendance entry').strip()[:500]
    target=q("SELECT * FROM users WHERE id=? AND active=1 AND role NOT IN ('Student','Parent','System')",(user_id,),one=True)
    if not target or action not in {'IN','OUT'}:
        flash('Choose a valid employee and check-in/check-out action.','danger')
        return redirect(url_for('admin_dashboard'))
    stamp=event_at or datetime.utcnow().isoformat(timespec='seconds')
    if attendance_day_is_closed(stamp):
        flash(f'Attendance for {attendance_date_from_value(stamp)} is closed.','danger')
        return redirect(url_for('admin_dashboard'))
    result=record_account_attendance(target,action,stamp,'manual','Manual',None,None,None,note,location_label)
    if result.get('ok'):
        notify_users(attendance_admin_ids(),f'Attendance: {target["full_name"]} checked {"IN" if action=="IN" else "OUT"}',f'Manual attendance recorded for {target["full_name"]}. Location: {location_label or "Not supplied"}.',url_for('admin_attendance'))
        flash(f'{target["full_name"]} marked {"in" if action=="IN" else "out"}.','success')
    else:
        flash(result.get('message','Attendance could not be recorded.'),'danger')
    return redirect(url_for('admin_dashboard'))

@app.route("/admin/attendance", methods=["GET","POST"])
@login_required
@role_required('Admin','ICT')
def admin_attendance():
    selected_date=attendance_date_from_value(request.values.get('date'))
    range_key=(request.values.get('range') or 'day').strip().lower()
    if request.method=='POST':
        selected_date=attendance_date_from_value(request.form.get('attendance_date') or selected_date)
        action=request.form.get('day_action','').strip().lower()
        if action in {'close','open'}:
            set_attendance_day(selected_date,'Closed' if action=='close' else 'Open',current_user()['id'])
            audit(current_user()['id'],current_user()['full_name'],f'Attendance Day {action.title()}',f'Attendance day {selected_date} set to {"Closed" if action=="close" else "Open"}.')
            flash(f'Attendance for {selected_date} is now {"closed" if action=="close" else "open"}.','success')
        return redirect(url_for('admin_attendance',date=selected_date,range=range_key))

    range_key,start_date,end_date=_attendance_range(range_key)
    # Keep the explicit day selector tied to the selected day while range buttons expand the reporting window.
    if request.values.get('date'):
        start_date=end_date=datetime.strptime(selected_date,'%Y-%m-%d').date()
        range_key='day'
    start_utc,end_utc=attendance_day_bounds_utc(start_date.isoformat())[0], attendance_day_bounds_utc(end_date.isoformat())[1]
    day_start_utc,day_end_utc=attendance_day_bounds_utc(selected_date)

    employees=q("SELECT * FROM users WHERE active=1 AND role NOT IN ('Student','Parent','System') ORDER BY full_name")
    rows=q("""
        SELECT u.id,u.full_name,u.username,u.role,COALESCE(u.title,u.role) AS title,u.department,u.school_unit,u.staff_code,u.active,
               (SELECT MIN(a.event_at) FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<?) AS check_in_at,
               (SELECT MAX(a.event_at) FROM attendance_events a WHERE a.user_id=u.id AND a.action='OUT' AND a.event_at>=? AND a.event_at<?) AS check_out_at,
               (SELECT a.location_label FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.location_label!='' ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS location_label,
               (SELECT a.latitude FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.latitude IS NOT NULL ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS latitude,
               (SELECT a.longitude FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.longitude IS NOT NULL ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS longitude,
               (SELECT a.accuracy FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.accuracy IS NOT NULL ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS accuracy
        FROM users u WHERE u.active=1 AND u.role NOT IN ('Student','Parent','System') ORDER BY u.full_name
    """,(start_utc,end_utc,start_utc,end_utc,start_utc,end_utc,start_utc,end_utc,start_utc,end_utc,start_utc,end_utc))
    rows=_decorate_attendance_rows(rows)
    expected_days=_expected_school_days(start_date,end_date)
    expected_count=len(expected_days)
    # Per-person missed days are computed from the same exact expected school-day set; approved reasons are separated.
    for row in rows:
        rid=row['id']; present={r['day'] for r in q("SELECT DISTINCT substr(a.event_at,1,10) AS day FROM attendance_events a WHERE a.user_id=? AND a.action='IN' AND a.event_at>=? AND a.event_at<?",(rid,start_utc,end_utc))}
        # event_at is UTC, so convert each event day into Kenya local date for range reporting.
        local_present=set()
        evs=q("SELECT event_at FROM attendance_events WHERE user_id=? AND action='IN' AND event_at>=? AND event_at<?",(rid,start_utc,end_utc))
        for ev in evs:
            dt=_parse_stored_event(ev['event_at'])
            if dt: local_present.add((dt+KENYA_TZ_OFFSET).date().isoformat())
        row['expected_days']=expected_count
        row['present_days']=len(local_present)
        row['missed_days']=max(0,expected_count-len(local_present))
        row['approved_missed']=q("SELECT COUNT(*) AS n FROM attendance_absence_requests WHERE user_id=? AND absence_date>=? AND absence_date<=? AND status='Approved'",(rid,start_date.isoformat(),end_date.isoformat()),one=True)['n']
        row['absence_reasons']=q("SELECT * FROM attendance_absence_requests WHERE user_id=? AND absence_date>=? AND absence_date<=? ORDER BY absence_date DESC",(rid,start_date.isoformat(),end_date.isoformat()))

    events=q("SELECT a.*,u.full_name,u.username,u.role,COALESCE(u.title,u.role) AS position,u.department,u.school_unit FROM attendance_events a JOIN users u ON u.id=a.user_id WHERE a.event_at>=? AND a.event_at<? ORDER BY a.event_at DESC,a.id DESC",(start_utc,end_utc))
    events=_decorate_attendance_rows(events)
    day_row=q("SELECT * FROM attendance_days WHERE attendance_date=?",(selected_date,),one=True)
    summary=q("""
      SELECT COUNT(*) AS total_events,COUNT(DISTINCT user_id) AS people,
             SUM(CASE WHEN action='IN' THEN 1 ELSE 0 END) AS checkins,
             SUM(CASE WHEN action='OUT' THEN 1 ELSE 0 END) AS checkouts,
             COUNT(DISTINCT CASE WHEN latitude IS NOT NULL AND longitude IS NOT NULL THEN user_id END) AS located_people
      FROM attendance_events WHERE event_at>=? AND event_at<?
    """,(start_utc,end_utc),one=True)
    summary=dict(summary)
    summary['expected_staff']=len(employees)
    summary['missed_staff_days']=sum(int(r['missed_days']) for r in rows)
    summary['excused_days']=sum(int(r['approved_missed']) for r in rows)
    summary['complete_checkout']=sum(1 for r in rows if r['check_in_at'] and r['check_out_at'])
    reminders=q("SELECT sr.*,u.full_name AS assigned_name FROM staff_reminders sr LEFT JOIN users u ON u.id=sr.user_id WHERE sr.completed=0 AND (sr.user_id IS NULL OR sr.user_id IN (SELECT id FROM users WHERE role IN ('Admin','ICT'))) ORDER BY sr.due_at LIMIT 12")
    timetable=q("SELECT st.*,u.full_name FROM staff_timetable st JOIN users u ON u.id=st.user_id WHERE st.active=1 AND u.active=1 AND u.role IN ('Admin','ICT','Teacher','Finance','Librarian','Driver','Reception') ORDER BY st.day_of_week,st.start_time,u.full_name")
    return render_template('admin_attendance.html',settings=school_settings(),events=events,office=q("SELECT * FROM attendance_qr_settings WHERE id=1",one=True),actor_name=current_user()['full_name'],role=current_user()['role'],selected_date=selected_date,day_status=(day_row['status'] if day_row else 'Open'),summary=summary,latest_by_user=rows,range_key=range_key,range_start=start_date.isoformat(),range_end=end_date.isoformat(),timetable=timetable,reminders=reminders)

@app.route('/admin/attendance/employee/<int:user_id>')
@login_required
@role_required('Admin','ICT')
def admin_attendance_employee(user_id):
    user=q("SELECT * FROM users WHERE id=? AND role NOT IN ('Student','Parent','System')",(user_id,),one=True)
    if not user: abort(404)
    events=q("SELECT * FROM attendance_events WHERE user_id=? ORDER BY event_at DESC,id DESC LIMIT 500",(user_id,))
    events=_decorate_attendance_rows(events)
    reasons=q("SELECT * FROM attendance_absence_requests WHERE user_id=? ORDER BY absence_date DESC,id DESC LIMIT 200",(user_id,))
    return render_template('attendance_employee.html',settings=school_settings(),user=user,events=events,reasons=reasons,role=current_user()['role'],actor_name=current_user()['full_name'])

@app.route('/attendance/absence',methods=['POST'])
@login_required
def attendance_absence():
    user=current_user(); day=attendance_date_from_value(request.form.get('absence_date')); reason=(request.form.get('reason') or '').strip()[:2000]
    if not reason: flash('Please give a reason for the absence.','danger'); return redirect(request.referrer or url_for('attendance_center'))
    execute("INSERT INTO attendance_absence_requests(user_id,absence_date,reason) VALUES(?,?,?) ON CONFLICT(user_id,absence_date) DO UPDATE SET reason=excluded.reason,status='Pending',requested_at=CURRENT_TIMESTAMP,reviewed_by=NULL,reviewed_at=NULL,review_note=NULL",(user['id'],day,reason))
    admin_ids=[r['id'] for r in q("SELECT id FROM users WHERE active=1 AND role IN ('Admin','ICT')")]
    notify_users(admin_ids,f'Attendance reason submitted: {user["full_name"]}',f'{user["full_name"]} submitted a reason for {day}: {reason}',url_for('admin_attendance'))
    flash('Attendance reason submitted for review.','success')
    return redirect(request.referrer or url_for('attendance_center'))

@app.route('/admin/attendance/absence/<int:request_id>/review',methods=['POST'])
@login_required
@role_required('Admin','ICT')
def review_absence(request_id):
    status=request.form.get('status','Pending').strip().title()
    if status not in {'Pending','Approved','Denied'}: status='Pending'
    note=(request.form.get('review_note') or '').strip()[:1000]
    req=q("SELECT * FROM attendance_absence_requests WHERE id=?",(request_id,),one=True)
    if not req: abort(404)
    execute("UPDATE attendance_absence_requests SET status=?,reviewed_by=?,reviewed_at=CURRENT_TIMESTAMP,review_note=? WHERE id=?",(status,current_user()['id'],note,request_id))
    notify_users([req['user_id']],f'Attendance reason {status.lower()}',f'Your attendance reason for {req["absence_date"]} was marked {status.lower()}. {note}'.strip(),url_for('attendance_center'))
    return redirect(request.referrer or url_for('admin_attendance'))

@app.route('/staff/timetable',methods=['GET','POST'])
@login_required
def staff_timetable():
    user=current_user()
    if request.method=='POST' and user['role'] not in {'Admin','ICT'}:
        abort(403)
    if request.method=='POST':
        uid=request.form.get('user_id',type=int) or user['id']; day=max(0,min(6,request.form.get('day_of_week',type=int) or 0)); start=(request.form.get('start_time') or '').strip(); end=(request.form.get('end_time') or '').strip(); title=(request.form.get('title') or '').strip(); loc=(request.form.get('location') or '').strip(); notes=(request.form.get('notes') or '').strip()
        if not title or not start or not end: flash('Staff, title, start and end time are required.','danger'); return redirect(url_for('staff_timetable'))
        execute("INSERT INTO staff_timetable(user_id,day_of_week,start_time,end_time,title,location,notes,created_by) VALUES(?,?,?,?,?,?,?,?)",(uid,day,start,end,title,loc,notes,user['id']))
        flash('Timetable entry saved.','success')
        return redirect(url_for('staff_timetable'))
    target=request.args.get('user_id',type=int) if user['role'] in {'Admin','ICT'} else user['id']
    if target:
        entries=q("SELECT st.*,u.full_name,u.role,u.title AS user_title FROM staff_timetable st JOIN users u ON u.id=st.user_id WHERE st.active=1 AND st.user_id=? ORDER BY st.day_of_week,st.start_time,u.full_name",(target,))
    else:
        entries=q("SELECT st.*,u.full_name,u.role,u.title AS user_title FROM staff_timetable st JOIN users u ON u.id=st.user_id WHERE st.active=1 AND u.active=1 ORDER BY st.day_of_week,st.start_time,u.full_name")
    staff=q("SELECT id,full_name,role,title FROM users WHERE active=1 AND role NOT IN ('Student','Parent','System') ORDER BY full_name") if user['role'] in {'Admin','ICT'} else []
    return render_template('staff_timetable.html',settings=school_settings(),entries=entries,staff=staff,target_user=target,actor_name=user['full_name'],role=user['role'])

@app.route('/staff/reminders',methods=['GET','POST'])
@login_required
def staff_reminders():
    user=current_user()
    if request.method=='POST' and user['role'] not in {'Admin','ICT'}: abort(403)
    if request.method=='POST':
        uid=request.form.get('user_id',type=int); scope=request.form.get('role_scope','All').strip() or 'All'; title=(request.form.get('title') or '').strip(); due=(request.form.get('due_at') or '').strip(); notes=(request.form.get('notes') or '').strip(); priority=request.form.get('priority','Normal').strip()
        if not title or not due: flash('Reminder title and due time are required.','danger'); return redirect(url_for('staff_reminders'))
        execute("INSERT INTO staff_reminders(user_id,role_scope,title,due_at,notes,priority,created_by) VALUES(?,?,?,?,?,?,?)",(uid,scope,title,due,notes,priority,user['id']))
        recipient_ids=[uid] if uid else [r['id'] for r in q("SELECT id FROM users WHERE active=1 AND role NOT IN ('System','Student','Parent')" if scope=='Staff' else "SELECT id FROM users WHERE active=1 AND role NOT IN ('System','Student','Parent')")]
        if uid: recipient_ids=[uid]
        notify_users(recipient_ids,'New reminder',f'{title} — {due}',url_for('staff_reminders'), 'High' if priority=='High' else 'Normal')
        flash('Reminder created.','success'); return redirect(url_for('staff_reminders'))
    if user['role'] in {'Admin','ICT'}:
        rows=q("SELECT sr.*,u.full_name AS assigned_name FROM staff_reminders sr LEFT JOIN users u ON u.id=sr.user_id WHERE sr.completed=0 ORDER BY sr.due_at,sr.id")
        staff=q("SELECT id,full_name,role FROM users WHERE active=1 AND role NOT IN ('Student','Parent','System') ORDER BY full_name")
    else:
        rows=q("SELECT sr.*,u.full_name AS assigned_name FROM staff_reminders sr LEFT JOIN users u ON u.id=sr.user_id WHERE sr.completed=0 AND (sr.user_id=? OR (sr.user_id IS NULL AND (sr.role_scope='All' OR sr.role_scope=?))) ORDER BY sr.due_at",(user['id'],user['role']))
        staff=[]
    return render_template('staff_reminders.html',settings=school_settings(),rows=rows,staff=staff,actor_name=user['full_name'],role=user['role'])

@app.route('/staff/reminders/<int:reminder_id>/complete',methods=['POST'])
@login_required
def complete_staff_reminder(reminder_id):
    user=current_user(); row=q("SELECT * FROM staff_reminders WHERE id=?",(reminder_id,),one=True)
    if not row: abort(404)
    if user['role'] not in {'Admin','ICT'} and row['user_id'] not in {None,user['id']}: abort(403)
    execute("UPDATE staff_reminders SET completed=1 WHERE id=?",(reminder_id,))
    return redirect(request.referrer or url_for('staff_reminders'))

@app.route("/driver")
@app.route("/driver-dashboard")
@login_required
def driver_dashboard():
    user=current_user()
    if workspace_type_for_user(user) != 'Driver' and user['role'] not in {'Driver','Admin','ICT'}:
        if user['role']=='Teacher': return redirect(url_for('teacher_dashboard'))
        abort(403)
    trip=q("SELECT * FROM transport_trips WHERE driver_user_id=? AND status='Active' ORDER BY id DESC LIMIT 1",(user['id'],),one=True)
    return render_template('driver_dashboard.html',settings=school_settings(),actor_name=user['full_name'],trip=trip,notification_count=notification_count(user['id']))

@app.route("/driver/trip/start",methods=['POST'])
@login_required
def driver_trip_start():
    user=current_user()
    if workspace_type_for_user(user)!='Driver' and user['role'] not in {'Admin','ICT'}: abort(403)
    existing=q("SELECT id FROM transport_trips WHERE driver_user_id=? AND status='Active'",(user['id'],),one=True)
    if existing:
        flash('A trip is already active.','warning'); return redirect(url_for('driver_dashboard'))
    vehicle_type=request.form.get('vehicle_type','').strip(); number_plate=request.form.get('number_plate','').strip().upper(); route=request.form.get('route_name','').strip()
    if not vehicle_type or not number_plate:
        flash('Vehicle type and number plate are required.','danger'); return redirect(url_for('driver_dashboard'))
    vehicle=f'{vehicle_type} · {number_plate}'
    execute("INSERT INTO transport_trips(driver_user_id,vehicle,vehicle_type,number_plate,route_name,status) VALUES(?,?,?,?,?,'Active')",(user['id'],vehicle,vehicle_type,number_plate,route))
    audit(user['id'],user['full_name'],'Driver Trip Started',f'{vehicle} · {route or "Route not specified"}.')
    flash('Trip started. Live tracking is ready.','success'); return redirect(url_for('driver_dashboard'))

@app.route("/driver/trip/stop",methods=['POST'])
@login_required
def driver_trip_stop():
    user=current_user()
    if workspace_type_for_user(user)!='Driver' and user['role'] not in {'Admin','ICT'}: abort(403)
    trip=q("SELECT id,vehicle,route_name FROM transport_trips WHERE driver_user_id=? AND status='Active' ORDER BY id DESC LIMIT 1",(user['id'],),one=True)
    if trip:
        execute("UPDATE transport_trips SET status='Completed',ended_at=CURRENT_TIMESTAMP WHERE id=?",(trip['id'],))
        audit(user['id'],user['full_name'],'Driver Trip Completed',f'{trip["vehicle"]} · {trip["route_name"] or "Route not specified"}.')
        flash('Trip completed and tracking stopped.','success')
    return redirect(url_for('driver_dashboard'))

@app.route("/workforce/<kind>")
@login_required
def workforce_dashboard(kind):
    user=current_user(); allowed={'Guard','Cook','Other Staff'}
    if user['role'] not in allowed and workspace_type_for_user(user) not in allowed and user['role'] not in {'Admin','ICT'}:
        if user['role']=='Teacher': return redirect(url_for('teacher_dashboard'))
        abort(403)
    return render_template('role_dashboard.html',school_settings=school_settings(),settings=school_settings(),role=kind,workspace=f'{kind} workspace',actor_name=user['full_name'],nav_items=navigation_items(kind,school_settings()),active_adverts=q("SELECT * FROM advertisements WHERE active=1 ORDER BY priority DESC,created_at DESC"),library_items=q("SELECT * FROM library_items WHERE active=1 ORDER BY category,title LIMIT 30"),notification_count=notification_count(user['id']),grades=[r['grade'] for r in q("SELECT DISTINCT grade FROM students ORDER BY grade")],students=[],assignments=[],submissions=[],results=[],teacher_online_url='/online-classes',student=None)

@app.route("/driver/location",methods=['POST'])
@login_required
def driver_location():
    if workspace_type_for_user(current_user())!='Driver' and current_user()['role'] not in {'Admin','ICT'}: abort(403)
    data=request.get_json(silent=True) or {}; trip=q("SELECT id FROM transport_trips WHERE driver_user_id=? AND status='Active' ORDER BY id DESC LIMIT 1",(current_user()['id'],),one=True)
    try: lat=float(data.get('latitude')); lon=float(data.get('longitude'))
    except (TypeError,ValueError): return jsonify({'ok':False}),400
    execute("INSERT INTO driver_locations(trip_id,driver_user_id,latitude,longitude,speed_kph,accuracy,source) VALUES(?,?,?,?,?,?,?)",(trip['id'] if trip else None,current_user()['id'],lat,lon,float(data.get('speed_kph') or 0),float(data.get('accuracy') or 0),'online')); return jsonify({'ok':True})

@app.route("/admin/transport")
@login_required
@role_required('Admin','ICT')
def admin_transport():
    rows=q("SELECT u.id,u.full_name,u.position_code,u.school_unit,u.school_location,t.vehicle,t.vehicle_type,t.number_plate,t.route_name,t.status,t.started_at,t.ended_at,l.latitude,l.longitude,l.speed_kph,l.accuracy,l.recorded_at,l.source FROM users u LEFT JOIN transport_trips t ON t.driver_user_id=u.id AND t.status='Active' LEFT JOIN driver_locations l ON l.id=(SELECT x.id FROM driver_locations x WHERE x.driver_user_id=u.id ORDER BY x.recorded_at DESC,x.id DESC LIMIT 1) WHERE u.active=1 AND u.workspace_type='Driver' ORDER BY u.full_name")
    return render_template('admin_transport.html',settings=school_settings(),rows=rows,actor_name=current_user()['full_name'],role=current_user()['role'])

@app.route("/finance/fee-structure",methods=['POST'])
@login_required
@role_required('Finance','Admin')
def finance_fee_structure():
    cls=request.form.get('class_level','').strip(); item=request.form.get('item_name','').strip()
    try: amount=float(request.form.get('amount','0') or 0)
    except ValueError: amount=-1
    period=request.form.get('period','Term 1').strip() or 'Term 1'
    if not cls or not item or amount<0: flash('Enter a class, fee item and valid amount.','danger'); return redirect(url_for('finance_dashboard'))
    execute("INSERT INTO fee_structures(class_level,item_name,amount,period,created_by) VALUES(?,?,?,?,?)",(cls,item,amount,period,current_user()['id'])); flash('Fee structure saved.','success'); return redirect(url_for('finance_dashboard'))

@app.route("/finance/apply-fee",methods=['POST'])
@login_required
@role_required('Finance','Admin')
def finance_apply_fee():
    sid=request.form.get('student_id',type=int); fid=request.form.get('fee_structure_id',type=int); student=q('SELECT * FROM students WHERE id=? AND active=1',(sid,),one=True); fee=q('SELECT * FROM fee_structures WHERE id=? AND active=1',(fid,),one=True)
    if not student or not fee: flash('Student or fee structure not found.','danger'); return redirect(url_for('finance_dashboard'))
    execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(sid,fid,float(fee['amount']),fee['item_name'],current_user()['id'])); recalculate_student_balance(sid); flash('Fee charge applied and student balance recalculated.','success'); return redirect(request.referrer or url_for('finance_dashboard'))

@app.route("/finance/integration",methods=['POST'])
@login_required
@role_required('Admin','Finance')
def finance_integration():
    provider=request.form.get('provider','Manual').strip() or 'Manual'; account_name=request.form.get('account_name','').strip(); collection=request.form.get('collection_account','').strip(); secret=request.form.get('callback_secret','').strip(); auto=1 if request.form.get('auto_match')=='1' else 0
    execute("UPDATE payment_integrations SET provider=?,account_name=?,collection_account=?,callback_secret=?,auto_match=?,notes=?,updated_at=CURRENT_TIMESTAMP WHERE id=1",(provider,account_name,collection,secret,auto,request.form.get('notes','').strip())); flash('Payment intake settings saved.','success'); return redirect(url_for('finance_dashboard'))

@app.route("/api/payments/webhook/<provider>",methods=['POST'])
def payments_webhook(provider):
    cfg=q('SELECT * FROM payment_integrations WHERE id=1',one=True)
    if not cfg or not cfg['callback_secret'] or request.headers.get('X-Webhook-Secret','')!=cfg['callback_secret']: return jsonify({'ok':False,'message':'Unauthorized webhook.'}),401
    data=request.get_json(silent=True) or {}; ref=str(data.get('reference') or data.get('transaction_id') or data.get('trans_id') or '').strip()
    try: amount=float(data.get('amount') or 0)
    except (TypeError,ValueError): amount=0
    admission=str(data.get('admission_no') or data.get('account') or '').strip(); phone=str(data.get('phone') or data.get('payer_phone') or '').strip(); name=str(data.get('payer_name') or data.get('name') or '').strip()
    if not ref or amount<=0: return jsonify({'ok':False,'message':'reference and amount required.'}),400
    try: eid=execute("INSERT INTO external_payment_events(provider,external_reference,amount,payer_name,payer_phone,admission_no,payload_json) VALUES(?,?,?,?,?,?,?)",(provider,ref,amount,name,phone,admission,json.dumps(data)))
    except sqlite3.IntegrityError: return jsonify({'ok':True,'message':'Already received.'})
    student=q('SELECT * FROM students WHERE admission_no=?',(admission,),one=True) if admission else None
    if not student and phone: student=q('SELECT * FROM students WHERE student_phone=? OR guardian_phone=? ORDER BY id LIMIT 1',(phone,phone),one=True)
    if cfg['auto_match'] and student:
        poster=q("SELECT id FROM users WHERE role='Admin' AND active=1 ORDER BY id LIMIT 1",one=True); poster_id=poster['id'] if poster else None
        if poster_id:
            pid=execute("INSERT INTO payments(student_id,amount,method,reference_no,recorded_by,status) VALUES(?,?,?,?,?,'Posted')",(student['id'],amount,provider,ref,poster_id)); recalculate_student_balance(student['id']); new_balance=float(q("SELECT balance FROM students WHERE id=?",(student['id'],),one=True)['balance'] or 0); execute("UPDATE external_payment_events SET status='Matched',matched_student_id=?,processed_at=CURRENT_TIMESTAMP WHERE id=?",(student['id'],eid)); return jsonify({'ok':True,'matched':True,'payment_id':pid})
    return jsonify({'ok':True,'matched':False,'event_id':eid})


@app.route('/finance/transport-rate', methods=['POST'])
@login_required
@role_required('Admin','Finance')
def finance_transport_rate():
    zone=(request.form.get('zone_name') or '').strip();
    try: amount=max(0,float(request.form.get('amount','0') or 0))
    except ValueError: amount=-1
    period=(request.form.get('period') or 'Term 1').strip() or 'Term 1'
    if not zone or amount<0:
        flash('Enter a transport zone and valid charge.','danger'); return redirect(request.referrer or url_for('finance_dashboard'))
    conn=get_db(); conn.execute('INSERT INTO transport_rates(zone_name,amount,period,active,created_by) VALUES(?,?,?,?,?) ON CONFLICT(zone_name) DO UPDATE SET amount=excluded.amount,period=excluded.period,active=1,created_by=excluded.created_by',(zone,amount,period,current_user()['id'],)); conn.commit()
    flash('Transport charge saved. New learners in that zone can be charged automatically.','success'); return redirect(request.referrer or url_for('finance_dashboard'))

@app.route('/admin/promotions', methods=['GET'])
@login_required
@role_required('Admin')
def admin_promotions():
    students=q("SELECT id,full_name,admission_no,grade,balance,payment_status,active FROM students WHERE active=1 ORDER BY grade,full_name")
    runs=q("SELECT r.*,COUNT(i.id) AS items_count FROM promotion_runs r LEFT JOIN promotion_items i ON i.run_id=r.id GROUP BY r.id ORDER BY r.id DESC LIMIT 20")
    return render_template('admin_promotions.html',settings=school_settings(),students=students,runs=runs,actor_name=current_user()['full_name'],role='Admin',current_year=datetime.utcnow().year,today=datetime.utcnow().strftime('%Y-%m-%d'))

@app.route('/admin/promotions/apply', methods=['POST'])
@login_required
@role_required('Admin')
def admin_promotions_apply():
    ids=[int(x) for x in request.form.getlist('student_ids') if str(x).isdigit()]
    decision=request.form.get('decision','Promote').strip().title()
    year=(request.form.get('academic_year') or datetime.utcnow().strftime('%Y')).strip()
    effective=(request.form.get('effective_date') or datetime.utcnow().strftime('%Y-%m-%d')).strip()
    if not ids: flash('Select at least one learner.','warning'); return redirect(url_for('admin_promotions'))
    run_id=execute("INSERT INTO promotion_runs(academic_year,effective_date,status,created_by) VALUES(?,?,?,?)",(year,effective,'Completed',current_user()['id']))
    changed=0
    for sid in ids:
        st=q('SELECT * FROM students WHERE id=? AND active=1',(sid,),one=True)
        if not st: continue
        target=next_grade_for(st['grade']) if decision=='Promote' else st['grade']
        item_decision=decision
        if decision=='Promote' and not target:
            target='Graduated'; st_target='Graduated'; item_decision='Graduate'
        elif decision=='Graduate': target='Graduated'; st_target='Graduated'
        elif decision=='Repeat': target=st['grade']; st_target=st['grade']
        elif decision=='Transfer': target='Transferred'; st_target='Transferred'
        else: st_target=target
        execute('INSERT INTO promotion_items(run_id,student_id,from_grade,to_grade,decision,notes) VALUES(?,?,?,?,?,?)',(run_id,sid,st['grade'],target,item_decision,(request.form.get('notes') or '').strip()))
        execute("UPDATE students SET promoted_from_grade=?, academic_year=?, grade=?, last_promotion_action=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",(st['grade'],year,st_target,item_decision,sid))
        changed+=1
    flash(f'Promotion cycle completed for {changed} learner(s).','success'); return redirect(url_for('admin_promotions'))

@app.route('/admin/bulk-students', methods=['POST'])
@login_required
@role_required('Admin')
def admin_bulk_students():
    ids=[int(x) for x in request.form.getlist('student_ids') if str(x).isdigit()]
    action=(request.form.get('action') or '').strip().lower()
    if not ids: flash('Select at least one learner first.','warning'); return redirect(request.referrer or url_for('admin_dashboard'))
    if action=='archive':
        for sid in ids: execute('UPDATE students SET active=0,updated_at=CURRENT_TIMESTAMP WHERE id=?',(sid,))
        flash(f'{len(ids)} learner(s) archived.','success')
    elif action=='promote':
        for sid in ids:
            st=q('SELECT grade FROM students WHERE id=? AND active=1',(sid,),one=True); target=next_grade_for(st['grade']) if st else ''
            if target: execute("UPDATE students SET promoted_from_grade=grade,grade=?,last_promotion_action='Promote',updated_at=CURRENT_TIMESTAMP WHERE id=?",(target,sid))
        flash(f'{len(ids)} learner(s) moved to their next mapped class.','success')
    else: flash('Unknown bulk learner action.','danger')
    return redirect(request.referrer or url_for('admin_dashboard'))

@app.route('/finance/payroll/bulk', methods=['POST'])
@login_required
@role_required('Admin','Finance')
def finance_payroll_bulk():
    ids=[int(x) for x in request.form.getlist('user_ids') if str(x).isdigit()]
    try: amount=float(request.form.get('amount','0') or 0)
    except ValueError: amount=0
    if not ids or amount<=0:
        flash('Select staff and enter a valid common salary amount.','danger'); return redirect(request.referrer or url_for('finance_dashboard'))
    for uid in ids:
        staff=q("SELECT id FROM users WHERE id=? AND active=1 AND role NOT IN ('Student','Parent','System')",(uid,),one=True)
        if staff:
            execute("INSERT INTO finance_ledger(entry_type,category,payee_user_id,amount,description,reference_no,posted_by) VALUES('Payroll','Salary / Bulk Payroll',?,?,?,?,?)",(uid,amount,'Common salary batch payment','BULK-'+datetime.utcnow().strftime('%Y%m%d'),current_user()['id']))
    flash(f'Bulk payroll posted for {len(ids)} staff member(s).','success'); return redirect(request.referrer or url_for('finance_dashboard'))

@app.route("/coming-soon/<feature>")
def coming_soon(feature: str):
    return redirect(url_for("index"))


@app.route("/api/login-location", methods=["POST"])
@login_required
def login_location():
    payload=request.get_json(silent=True) or {}
    try:
        lat=float(payload.get('latitude')); lon=float(payload.get('longitude')); acc=float(payload.get('accuracy')) if payload.get('accuracy') is not None else None
        if not (-90<=lat<=90 and -180<=lon<=180): raise ValueError
    except (TypeError,ValueError):
        return jsonify({'ok':False}),400
    event_id=session.get('login_event_id')
    if event_id:
        execute("UPDATE login_events SET latitude=?,longitude=?,accuracy=? WHERE id=? AND user_id=?",(lat,lon,acc,event_id,current_user()['id']))
    else:
        event_id=record_login_event(current_user(),'Location update',lat,lon,acc); session['login_event_id']=event_id
    session.pop('login_location_pending',None)
    return jsonify({'ok':True})


@app.route("/logout")
def logout():
    g.logging_out = True
    context_token=request.args.get("portal_context")
    tid=_portal_context_id(context_token) if context_token else None
    if tid:
        _ensure_portal_context_table()
        execute("UPDATE portal_contexts SET revoked=1 WHERE token_id=?", (tid,))
    _revoke_auth_ticket(request.cookies.get(_AUTH_TICKET_COOKIE, ""))
    session.clear()
    response=redirect(url_for("index"))
    response.delete_cookie(_AUTH_TICKET_COOKIE, path="/")
    response.delete_cookie(_AUTH_COOKIE, path="/")
    return response

@app.route("/favicon.ico")
def favicon():
    return send_file(BASE_DIR / "static" / "icons" / "favicon.ico", mimetype="image/vnd.microsoft.icon", conditional=True, max_age=86400)


@app.route("/sw.js")
def service_worker():
    return send_file(BASE_DIR / "static" / "sw.js", mimetype="application/javascript", conditional=True, max_age=0)


@app.route("/dashboard")
@login_required
def dashboard():
    # Canonical dashboard entry: always dispatch to the authenticated account's
    # own workspace. This route must never be a Teacher-only dashboard because
    # global navigation uses /dashboard.
    return redirect(specialized_dashboard_for(current_user()))


# Legacy teacher implementation retained below for compatibility with existing
# links/bookmarks; /dashboard itself is now only the role dispatcher.

@app.route("/admin/dashboard")
@login_required
@role_required("Admin")
def admin_dashboard():
    settings = school_settings()
    students = q("SELECT * FROM students ORDER BY created_at DESC, id DESC LIMIT 20")
    payments = q(
        """
        SELECT p.*, s.full_name AS student_name, s.admission_no, u.full_name AS recorded_by_name
        FROM payments p
        JOIN students s ON s.id = p.student_id
        JOIN users u ON u.id = p.recorded_by
        ORDER BY p.created_at DESC, p.id DESC
        LIMIT 20
        """
    )
    users = q("""SELECT u.*, COALESCE(u.title, '') AS title, COALESCE(u.department, '') AS department,
                      CASE WHEN u.active=1 THEN 'Active' ELSE 'Archived' END AS access_state
               FROM users u WHERE u.role!='System' ORDER BY u.active DESC, u.created_at DESC""")
    audits = q("SELECT * FROM audit_log ORDER BY created_at DESC, id DESC LIMIT 20")
    total_students = q("SELECT COUNT(*) AS c FROM students", one=True)["c"]
    active_students = q("SELECT COUNT(*) AS c FROM students WHERE active = 1", one=True)["c"]
    paid_students = q("SELECT COUNT(*) AS c FROM students WHERE payment_status = 'Paid'", one=True)["c"]
    pending_students = q("SELECT COUNT(*) AS c FROM students WHERE payment_status = 'Pending'", one=True)["c"]
    partial_students = q("SELECT COUNT(*) AS c FROM students WHERE active=1 AND balance > 0 AND fee_assessed_total > 0 AND balance < fee_assessed_total", one=True)["c"]
    unpaid_students = q("SELECT COUNT(*) AS c FROM students WHERE active=1 AND balance > 0 AND (fee_assessed_total <= 0 OR balance >= fee_assessed_total)", one=True)["c"]
    payment_paid_rows=q("SELECT admission_no,full_name,grade,balance,payment_status,fee_assessed_total FROM students WHERE active=1 AND balance<=0 ORDER BY grade,full_name")
    payment_partial_rows=q("SELECT admission_no,full_name,grade,balance,payment_status,fee_assessed_total FROM students WHERE active=1 AND balance>0 AND fee_assessed_total>0 AND balance < fee_assessed_total ORDER BY grade,full_name")
    payment_unpaid_rows=q("SELECT admission_no,full_name,grade,balance,payment_status,fee_assessed_total FROM students WHERE active=1 AND balance>0 AND (fee_assessed_total<=0 OR balance>=fee_assessed_total) ORDER BY grade,full_name")
    grade_people=q("SELECT grade,COUNT(*) AS c,SUM(CASE WHEN balance=0 THEN 1 ELSE 0 END) AS paid,SUM(CASE WHEN balance>0 THEN 1 ELSE 0 END) AS owing FROM students WHERE active=1 GROUP BY grade ORDER BY grade")
    staff_breakdown=q("SELECT role,COUNT(*) AS c FROM users WHERE active=1 AND role!='System' GROUP BY role ORDER BY role")
    total_income = q("SELECT COALESCE(SUM(amount), 0) AS total FROM payments WHERE status = 'Posted'", one=True)["total"]
    total_balance = q("SELECT COALESCE(SUM(balance), 0) AS total FROM students", one=True)["total"]
    avg_balance = q("SELECT COALESCE(AVG(balance), 0) AS total FROM students", one=True)["total"]
    categories = {
        "students": q("SELECT grade, COUNT(*) AS c FROM students GROUP BY grade ORDER BY grade"),
        "employees": q("SELECT role, COUNT(*) AS c FROM users GROUP BY role ORDER BY role"),
        "payments": q("SELECT method, COUNT(*) AS c FROM payments WHERE status = 'Posted' GROUP BY method ORDER BY method"),
    }

    exam_grade_summary = q(
        """
        SELECT grade,
               COUNT(DISTINCT student_id) AS pupils,
               COUNT(DISTINCT subject) AS subjects,
               COUNT(DISTINCT batch_id) AS batches,
               ROUND(AVG(mark), 1) AS mean_score,
               ROUND(MAX(mark), 1) AS highest_score,
               ROUND(MIN(mark), 1) AS lowest_score
        FROM exam_results
        GROUP BY grade
        ORDER BY grade
        """
    )
    exam_term_summary = q(
        """
        SELECT grade,
               term,
               COUNT(DISTINCT student_id) AS pupils,
               COUNT(DISTINCT subject) AS subjects,
               ROUND(AVG(mark), 1) AS mean_score,
               ROUND(MAX(mark), 1) AS highest_score,
               ROUND(MIN(mark), 1) AS lowest_score
        FROM exam_results
        GROUP BY grade, term
        ORDER BY grade, term
        """
    )
    exam_subject_summary = q(
        """
        SELECT subject,
               COUNT(*) AS entries,
               ROUND(AVG(mark), 1) AS mean_score,
               ROUND(MAX(mark), 1) AS highest_score,
               ROUND(MIN(mark), 1) AS lowest_score
        FROM exam_results
        GROUP BY subject
        ORDER BY mean_score DESC, subject
        """
    )
    recent_exam_batches = q(
        """
        SELECT b.*, u.full_name AS submitted_by_name, COUNT(r.id) AS entries, ROUND(AVG(r.mark), 1) AS mean_mark
        FROM exam_batches b
        LEFT JOIN exam_results r ON r.batch_id = b.id
        JOIN users u ON u.id = b.submitted_by
        GROUP BY b.id
        ORDER BY b.created_at DESC, b.id DESC
        LIMIT 12
        """
    )
    exam_school_summary = {
        "batches": q("SELECT COUNT(*) AS c FROM exam_batches", one=True)["c"],
        "results": q("SELECT COUNT(*) AS c FROM exam_results", one=True)["c"],
        "mean": q("SELECT COALESCE(AVG(mark), 0) AS c FROM exam_results", one=True)["c"],
        "best": q("SELECT ROUND(MAX(mark), 1) AS c FROM exam_results", one=True)["c"],
    }

    elections=q("SELECT * FROM elections ORDER BY created_at DESC")
    election_candidates={e["id"]:q("SELECT * FROM election_candidates WHERE election_id=? ORDER BY position,name",(e["id"],)) for e in elections}
    library_items=q("SELECT * FROM library_items ORDER BY active DESC,category,title")
    local_today=(datetime.utcnow()+KENYA_TZ_OFFSET).date().isoformat()
    today_start,today_end=attendance_day_bounds_utc(local_today)
    today_attendance=q("""
        SELECT u.id,u.full_name,u.role,COALESCE(u.title,u.role) AS title,
               (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS sign_in_at,
               (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='OUT' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at DESC,a.id DESC LIMIT 1) AS sign_out_at,
               (SELECT a.location_label FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.location_label!='' ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS location
        FROM users u
        WHERE u.active=1 AND u.role NOT IN ('Student','Parent','System')
        ORDER BY u.full_name
    """,(today_start,today_end,today_start,today_end,today_start,today_end))
    today_attendance=[dict(r) for r in today_attendance]
    for r in today_attendance:
        r['sign_in_local']=_local_iso(_parse_stored_event(r['sign_in_at'])) if r.get('sign_in_at') else None
        r['sign_out_local']=_local_iso(_parse_stored_event(r['sign_out_at'])) if r.get('sign_out_at') else None
    analytics_max_employee=max([int(r['c'] or 0) for r in categories['employees']] or [1])
    analytics_max_grade=max([int(r['c'] or 0) for r in categories['students']] or [1])
    analytics_max_payment=max([int(r['c'] or 0) for r in categories['payments']] or [1])

    return render_template(
        "admin_dashboard.html",
        workspace=workspace_for("Admin"),
        settings=settings,
        students=students,
        payments=payments,
        partial_students=partial_students, unpaid_students=unpaid_students, payment_paid_rows=payment_paid_rows, payment_partial_rows=payment_partial_rows, payment_unpaid_rows=payment_unpaid_rows, grade_people=grade_people, staff_breakdown=staff_breakdown,
        users=users,
        audits=audits,
        summary={
            "total_students": total_students,
            "active_students": active_students,
            "paid_students": paid_students,
            "pending_students": pending_students,
            "total_income": total_income,
            "total_balance": total_balance,
            "avg_balance": avg_balance,
        },
        categories=categories,
        exam_grade_summary=exam_grade_summary,
        exam_term_summary=exam_term_summary,
        exam_subject_summary=exam_subject_summary,
        recent_exam_batches=recent_exam_batches,
        exam_school_summary=exam_school_summary,
        elections=elections,
        election_candidates=election_candidates,
        library_items=library_items,
        departments=q("SELECT id, name, category, active FROM departments WHERE active=1 ORDER BY name"),
        subjects_catalog=q("SELECT id, subject, department, level_scope FROM subjects_catalog WHERE active=1 ORDER BY department,subject"),
        guardian_links=q("""SELECT gl.*, gu.full_name AS guardian_name, gu.username AS guardian_username, st.full_name AS student_name, st.admission_no
                            FROM guardian_links gl JOIN users gu ON gu.id=gl.guardian_user_id
                            JOIN students st ON st.id=gl.student_id ORDER BY gl.created_at DESC"""),
        archived_users=q("SELECT id, full_name, username, role, title, department, archived_at FROM users WHERE active=0 AND role!='System' ORDER BY archived_at DESC, full_name"),
        finance_closings=q("SELECT c.*,u.full_name AS submitted_name FROM finance_closings c JOIN users u ON u.id=c.submitted_by ORDER BY c.submitted_at DESC,c.id DESC LIMIT 20"),
        today_attendance=today_attendance,
        local_today=local_today,
        analytics_max_employee=analytics_max_employee,
        analytics_max_grade=analytics_max_grade,
        analytics_max_payment=analytics_max_payment,
        all_roles=ALL_PORTAL_ROLES,
        onboarding_students=q("SELECT id,full_name,admission_no,grade FROM students WHERE active=1 ORDER BY grade,full_name"),
    )


@app.route("/teacher/dashboard")
@login_required
def teacher_dashboard():
    if current_user()["role"] not in {"Teacher", "Admin"}:
        return redirect(specialized_dashboard_for(current_user()))
    settings=school_settings(); user=current_user()
    assignments=q("SELECT * FROM teacher_assignments WHERE teacher_user_id=? AND active=1 ORDER BY class_name,subject",(user["id"],)) if user["role"]=="Teacher" else q("SELECT a.*,u.full_name AS teacher_name FROM teacher_assignments a JOIN users u ON u.id=a.teacher_user_id WHERE a.active=1 ORDER BY a.class_name,a.subject")
    # A teacher can receive learners directly from Admin even when the
    # teacher has not created a separate teaching assignment. Include those
    # allocated class names in the workspace so the learners actually appear.
    allocated_class_rows=q("SELECT DISTINCT COALESCE(NULLIF(class_name,''), (SELECT grade FROM students WHERE id=student_id)) AS class_name FROM student_teacher_assignments WHERE teacher_user_id=? AND active=1 ORDER BY class_name",(user["id"],)) if user["role"]=="Teacher" else []
    allocated_classes={r["class_name"] for r in allocated_class_rows if r["class_name"]}
    classes=sorted({r["class_name"] for r in assignments} | {r["class_name"] for r in q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?",(user["id"],))} | allocated_classes)
    assigned_class_rows=q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=? ORDER BY class_name",(user['id'],)) if user['role']=='Teacher' else q("SELECT class_name FROM class_teacher_assignments ORDER BY class_name")
    class_teacher_classes=[r['class_name'] for r in assigned_class_rows]
    students=[]
    if user["role"]=="Teacher" and classes:
        placeholders=','.join('?'*len(classes))
        students=q(f"""SELECT DISTINCT s.id,s.full_name,s.admission_no,s.grade,s.balance,s.fee_assessed_total,s.payment_status,
            CASE WHEN s.balance<=0 THEN 'Paid' WHEN s.fee_assessed_total>0 AND s.balance<s.fee_assessed_total THEN 'Partial' ELSE 'Unpaid' END AS payment_bucket,
            (SELECT COUNT(*) FROM student_subjects ss WHERE ss.student_id=s.id AND ss.status='Approved') AS subject_count
            FROM students s
            LEFT JOIN student_teacher_assignments sta ON sta.student_id=s.id AND sta.teacher_user_id=? AND sta.active=1
            LEFT JOIN class_teacher_assignments cta ON cta.class_name=s.grade AND cta.teacher_user_id=?
            WHERE s.active=1 AND (s.grade IN ({placeholders}) OR sta.id IS NOT NULL OR cta.id IS NOT NULL)
            ORDER BY s.grade,s.full_name""",(user["id"],user["id"],*classes))
    elif user["role"] in {"Admin","ICT"}:
        students=q("SELECT s.id,s.full_name,s.admission_no,s.grade,s.balance,s.fee_assessed_total,s.payment_status,CASE WHEN s.balance<=0 THEN 'Paid' WHEN s.fee_assessed_total>0 AND s.balance<s.fee_assessed_total THEN 'Partial' ELSE 'Unpaid' END AS payment_bucket,COUNT(ss.id) AS subject_count FROM students s LEFT JOIN student_subjects ss ON ss.student_id=s.id AND ss.status='Approved' WHERE s.active=1 GROUP BY s.id ORDER BY s.grade,s.full_name")
    latest_marks=q("SELECT m.*,s.full_name FROM markbook_entries m JOIN students s ON s.id=m.student_id WHERE m.teacher_user_id=? ORDER BY m.created_at DESC LIMIT 80",(user['id'],))
    events=q("SELECT * FROM attendance_events WHERE user_id=? ORDER BY event_at DESC,id DESC LIMIT 20",(user["id"],)) if user["role"]=="Teacher" else []
    upcoming=q("SELECT * FROM class_sessions WHERE teacher_user_id=? AND active=1 AND (starts_at>=datetime('now') OR ends_at>=datetime('now')) ORDER BY starts_at LIMIT 8",(user['id'],)) if user['role']=='Teacher' else []
    schemes=q("SELECT * FROM scheme_of_work WHERE teacher_user_id=? ORDER BY updated_at DESC,id DESC LIMIT 8",(user['id'],)) if user['role']=='Teacher' else []
    summaries={}
    for cls in classes[:12]:
        if cls in class_teacher_classes:
            summaries[cls]=markbook_class_summary(cls)
        else:
            teacher_subject=q("SELECT subject FROM teacher_assignments WHERE teacher_user_id=? AND class_name=? AND active=1 ORDER BY subject LIMIT 1",(user['id'],cls),one=True) if user['role']=='Teacher' else None
            summaries[cls]=markbook_class_summary(cls,teacher_subject['subject']) if teacher_subject else []
    return render_template("teacher_dashboard_pro.html",settings=settings,school_settings=settings,actor_name=user["full_name"],role=user["role"],assignments=assignments,classes=classes,students=students,latest_marks=latest_marks,events=events,workspace_type=workspace_type_for_user(user),upcoming=upcoming,schemes=schemes,class_teacher_classes=class_teacher_classes,mark_summaries=summaries,nav_items=navigation_items("Teacher",settings))

@app.route("/teacher/scheme-of-work", methods=["GET", "POST"])
@login_required
@role_required("Teacher", "Admin")
def teacher_scheme_of_work():
    user=current_user(); settings=school_settings()
    if user["role"] == "Teacher":
        assignments=q("SELECT class_name,subject,unit_code FROM teacher_assignments WHERE teacher_user_id=? AND active=1 ORDER BY class_name,subject", (user["id"],))
        # Never dead-end a teacher simply because the requested class/subject has
        # not yet been formalized in an assignment. Use the supplied values and
        # allow the teacher to save a plan; the planner itself remains private to
        # that teacher.
    else:
        assignments=q("SELECT class_name,subject,unit_code FROM teacher_assignments WHERE active=1 ORDER BY class_name,subject")

    selected_class=(request.args.get("class_name") or request.form.get("class_name") or "").strip()
    selected_subject=(request.args.get("subject") or request.form.get("subject") or "").strip()
    if request.method == "POST":
        term=(request.form.get("term") or "Current Term").strip() or "Current Term"
        try: week_no=max(1,int(request.form.get("week_no") or 1))
        except ValueError: week_no=1
        topic=(request.form.get("topic") or "").strip()
        if not selected_class or not selected_subject or not topic:
            flash("Class, subject and topic are required.", "danger")
            return redirect(url_for("teacher_scheme_of_work", class_name=selected_class, subject=selected_subject))
        owner_id=user["id"] if user["role"] == "Teacher" else int(request.form.get("teacher_user_id") or user["id"])
        existing=q("SELECT id FROM scheme_of_work WHERE teacher_user_id=? AND class_name=? AND subject=? AND term=? AND week_no=? LIMIT 1", (owner_id,selected_class,selected_subject,term,week_no), one=True)
        values=(topic,request.form.get("objectives","").strip(),request.form.get("activities","").strip(),request.form.get("resources","").strip(),request.form.get("assessment","").strip(),request.form.get("status","Planned").strip() or "Planned")
        if existing:
            execute("UPDATE scheme_of_work SET topic=?,objectives=?,activities=?,resources=?,assessment=?,status=?,updated_at=CURRENT_TIMESTAMP WHERE id=?", (*values,existing["id"]))
        else:
            execute("INSERT INTO scheme_of_work(teacher_user_id,class_name,subject,term,week_no,topic,objectives,activities,resources,assessment,status) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (owner_id,selected_class,selected_subject,term,week_no,*values))
        audit(user["id"],user["full_name"],"Scheme of Work",f"{selected_class} / {selected_subject} / Week {week_no}")
        flash("Scheme of Work week saved.", "success")
        return redirect(url_for("teacher_scheme_of_work", class_name=selected_class, subject=selected_subject))

    if user["role"] == "Teacher":
        rows=q("SELECT * FROM scheme_of_work WHERE teacher_user_id=? AND (?='' OR class_name=?) AND (?='' OR subject=?) ORDER BY term,week_no,id", (user["id"],selected_class,selected_class,selected_subject,selected_subject))
    else:
        rows=q("SELECT s.*,u.full_name AS teacher_name FROM scheme_of_work s JOIN users u ON u.id=s.teacher_user_id WHERE (?='' OR s.class_name=?) AND (?='' OR s.subject=?) ORDER BY s.term,s.class_name,s.subject,s.week_no,s.id", (selected_class,selected_class,selected_subject,selected_subject))
    return render_template("scheme_of_work.html", settings=settings, assignments=assignments, rows=rows, selected_class=selected_class, selected_subject=selected_subject, theme_style="", theme_preset_style="")


@app.route("/api/teacher/fee-status")
@login_required
@role_required("Teacher", "Admin")
def teacher_fee_status_api():
    user=current_user()
    if user["role"]=="Admin":
        rows=q("SELECT id,balance,fee_assessed_total FROM students WHERE active=1")
    else:
        class_rows=q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?",(user["id"],))
        taught_rows=q("SELECT DISTINCT class_name FROM teacher_assignments WHERE teacher_user_id=? AND active=1",(user["id"],))
        allocated_rows=q("SELECT DISTINCT student_id FROM student_teacher_assignments WHERE teacher_user_id=? AND active=1",(user["id"],))
        classes={r["class_name"] for r in class_rows if r["class_name"]}|{r["class_name"] for r in taught_rows if r["class_name"]}
        ids={int(r["student_id"]) for r in allocated_rows if r["student_id"]}
        clauses=[]; params=[]
        if classes:
            ph=','.join('?'*len(classes)); clauses.append(f"grade IN ({ph})"); params.extend(sorted(classes))
        if ids:
            ph=','.join('?'*len(ids)); clauses.append(f"id IN ({ph})"); params.extend(sorted(ids))
        rows=q(f"SELECT id,balance,fee_assessed_total FROM students WHERE active=1 AND ({' OR '.join(clauses) or '0'})",tuple(params))
    items=[]
    for r in rows:
        balance=float(r["balance"] or 0); assessed=float(r["fee_assessed_total"] or 0)
        items.append({"id":r["id"],"balance":balance,"bucket":"Paid" if balance<=0 else ("Partial" if assessed>0 and balance<assessed else "Unpaid")})
    return jsonify({"ok":True,"items":items})

@app.route("/student/dashboard")
@login_required
def student_dashboard():
    user=current_user()
    if user["role"] != "Student" and not DEMO_AUTH_BYPASS:
        return redirect(specialized_dashboard_for(user))
    student=portal_student(request.args.get("student_id", type=int))
    if user["role"]=="Student":
        student=q("SELECT * FROM students WHERE id=? AND active=1", (user["student_id"],), one=True) if user["student_id"] else None
    elif not student and user["student_id"]:
        student=q("SELECT * FROM students WHERE id=? AND active=1", (user["student_id"],), one=True)
    if not student: abort(404)
    # Optional learner panels should never take down the whole portal. A partially
    # migrated school database can still open the Student dashboard while a single
    # secondary feature is repaired.
    assignments=[]; submissions=[]; results=[]; result_releases={}; messages=[]; elections=[]; election_candidates={}; voted_positions=set(); library_items=[]; student_subjects=[]; online_classes=[]; mark_rows=[]; my_rank=None
    try: assignments=assignment_rows(student["grade"])
    except Exception: assignments=[]
    try: submissions=q("SELECT * FROM submissions WHERE student_id=? ORDER BY submitted_at DESC", (student["id"],))
    except Exception: submissions=[]
    try: results=q("SELECT subject, term, mark, max_mark FROM exam_results WHERE student_id=? ORDER BY term DESC, subject", (student["id"],))
    except Exception: results=[]
    try: result_releases=result_release_info(student["id"])
    except Exception: result_releases={}
    try: messages=q("SELECT * FROM portal_messages WHERE recipient_student_id=? OR (recipient_role='Student' AND recipient_student_id IS NULL) ORDER BY created_at DESC LIMIT 30", (student["id"],))
    except Exception: messages=[]
    try:
        elections=q("SELECT * FROM elections WHERE visible=1 ORDER BY created_at DESC") if school_settings()["elections_enabled"] else []
        election_candidates={e["id"]:q("SELECT * FROM election_candidates WHERE election_id=? AND active=1 ORDER BY position,name",(e["id"],)) for e in elections}
        voted_positions={(r["election_id"], r["position"]) for r in q("SELECT election_id, position FROM election_votes WHERE voter_user_id=?",(current_user()["id"],))}
    except Exception:
        elections=[]; election_candidates={}; voted_positions=set()
    try: library_items=q("SELECT * FROM library_items WHERE active=1 ORDER BY category,title LIMIT 80") if school_settings()["library_enabled"] else []
    except Exception: library_items=[]
    try: student_subjects=q("SELECT sc.id,sc.subject,sc.department,ss.status FROM student_subjects ss JOIN subjects_catalog sc ON sc.id=ss.subject_id WHERE ss.student_id=? AND ss.status!='Dropped' ORDER BY sc.department,sc.subject",(student['id'],))
    except Exception: student_subjects=[]
    try: online_classes=q("SELECT cs.*,u.full_name AS teacher_name FROM class_sessions cs JOIN users u ON u.id=cs.teacher_user_id WHERE cs.active=1 AND lower(cs.class_name)=lower(?) AND (cs.starts_at>=datetime('now','-1 day') OR cs.ends_at>=datetime('now','-1 day') OR cs.ends_at IS NULL) ORDER BY cs.starts_at",(student["grade"],))
    except Exception: online_classes=[]
    try: mark_rows=markbook_class_summary(student['grade']); my_rank=next((r for r in mark_rows if r['student_id']==student['id']),None)
    except Exception: mark_rows=[]; my_rank=None
    settings=school_settings(); nav_items=navigation_items("Student",settings)
    return render_template("student_dashboard.html",school_settings=settings,settings=settings,role="Student",workspace=workspace_for("Student"),student=student,assignments=assignments,submissions=submissions,results=results,result_releases=result_releases,messages=messages,elections=elections,election_candidates=election_candidates,voted_positions=voted_positions,library_items=library_items,online_classes=online_classes,student_subjects=student_subjects,actor_name=student["full_name"],nav_items=nav_items,my_rank=my_rank,mark_rows=mark_rows)

@app.route("/parent-dashboard")
@login_required
def parent_dashboard():
    if not parent_portal_enabled(): abort(404)
    if current_user()["role"] != "Parent" and not DEMO_AUTH_BYPASS:
        abort(403)
    children=parent_children(current_user()) if current_user()["role"] == "Parent" else []
    if not children and DEMO_AUTH_BYPASS:
        demo_child=q("SELECT s.*, COALESCE((SELECT u.full_name FROM users u JOIN guardian_links gl ON gl.guardian_user_id=u.id WHERE gl.student_id=s.id LIMIT 1), '') AS guardian_name FROM students s WHERE s.active=1 ORDER BY s.grade,s.full_name LIMIT 1", one=True)
        children=[demo_child] if demo_child else []
    requested=request.args.get("child_id", type=int)
    child=next((row for row in children if row["id"] == requested), None) if requested else (children[0] if children else None)
    if not child:
        return render_template("parent_dashboard.html", settings=school_settings(), actor_name=current_user()["full_name"], child=None, children=[], assignments=[], results=[], result_releases=[], submissions=[], messages=[], library_items=[], online_classes=[], nav_items=navigation_items("Parent", school_settings()))
    assignments=assignment_rows(child["grade"])
    results=q("SELECT subject, term, mark, max_mark FROM exam_results WHERE student_id=? ORDER BY term DESC, subject", (child["id"],))
    result_releases=result_release_info(child["id"])
    submissions=q("""SELECT s.*, a.title, a.subject FROM submissions s JOIN assignments a ON a.id=s.assignment_id WHERE s.student_id=? ORDER BY s.submitted_at DESC""", (child["id"],))
    messages=q("SELECT * FROM portal_messages WHERE recipient_student_id=? ORDER BY created_at DESC LIMIT 30", (child["id"],))
    library_items=q("SELECT * FROM library_items WHERE active=1 ORDER BY category,title LIMIT 80") if school_settings()["library_enabled"] else []
    online_classes=q("SELECT a.*,u.full_name AS teacher_name FROM teacher_assignments a JOIN users u ON u.id=a.teacher_user_id WHERE a.active=1 AND lower(a.class_name)=lower(?) ORDER BY a.subject,u.full_name", (child["grade"],))
    settings=school_settings(); nav_items=navigation_items("Parent", settings)
    return render_template("parent_dashboard.html", role="Parent", workspace=workspace_for("Parent"), child=child, children=children, assignments=assignments, results=results, result_releases=result_releases, submissions=submissions, messages=messages, library_items=library_items, online_classes=online_classes, actor_name=child["guardian_name"] or "Parent", nav_items=nav_items)


@app.route("/library")
@login_required
def library():
    if not school_settings()["library_enabled"] and current_user()["role"] not in {"Admin","ICT","Librarian"}: abort(404)
    class_level=request.args.get("class","").strip(); subject=request.args.get("subject","").strip()
    where=["active=1"]; params=[]
    role=current_user()["role"]
    if role=="Student":
        st=portal_student()
        class_level=(st["grade"] if st else class_level) or ""
        where.append("(class_level=? OR class_level='')"); params.append(class_level)
    elif role=="Parent":
        children=parent_children(current_user())
        child=children[0] if children else None
        class_level=(child["grade"] if child else class_level) or ""
        where.append("(class_level=? OR class_level='')"); params.append(class_level)
    elif class_level:
        where.append("(class_level=? OR class_level='')"); params.append(class_level)
    if subject: where.append("(subject=? OR subject='')"); params.append(subject)
    items=q(f"SELECT * FROM library_items WHERE {' AND '.join(where)} ORDER BY class_level,subject,category,title",params)
    loans=q("""SELECT l.*, i.title, s.full_name AS student_name, s.admission_no FROM library_loans l JOIN library_items i ON i.id=l.item_id JOIN students s ON s.id=l.student_id WHERE l.status='Issued' ORDER BY l.due_date, l.issued_at""")
    classes=[r["grade"] for r in q("SELECT DISTINCT grade FROM students WHERE active=1 ORDER BY grade")]
    subjects=[r["subject"] for r in q("SELECT DISTINCT subject FROM library_items WHERE active=1 AND subject!='' ORDER BY subject")]
    return render_template("library.html", items=items, loans=loans, settings=school_settings(), actor_name=current_user()["full_name"], selected_class=class_level, selected_subject=subject, library_classes=classes, library_subjects=subjects)

@app.route("/librarian-dashboard")
@login_required
@role_required("Librarian","Admin","ICT")
def librarian_dashboard():
    items=q("SELECT * FROM library_items ORDER BY active DESC, category, title")
    loans=q("""SELECT l.*, i.title, s.full_name AS student_name, s.admission_no FROM library_loans l JOIN library_items i ON i.id=l.item_id JOIN students s ON s.id=l.student_id ORDER BY CASE WHEN l.status='Issued' THEN 0 ELSE 1 END, l.due_date DESC""")
    students=q("SELECT * FROM students WHERE active=1 ORDER BY grade,full_name")
    return render_template("library_manager.html", role=current_user()["role"], settings=school_settings(), items=items, loans=loans, students=students, actor_name=current_user()["full_name"])

@app.route("/institution")
@login_required
def institution():
    settings=school_settings()
    return render_template("institution.html", settings=settings, actor_name=current_user()["full_name"], role=current_user()["role"])

@app.route("/institution/save", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def institution_save():
    values={k:request.form.get(k,"").strip() for k in ["institution_history","institution_performance","institution_religion","institution_affiliations","institution_help","institution_contact","institution_portal_guide","institution_admin_guide","institution_ict_guide","institution_finance_guide","institution_driver_guide_en","institution_driver_guide_sw"]}
    image=request.files.get("institution_image"); image_path=school_settings()["institution_image_path"] or ""
    if image and image.filename:
        dest=UPLOAD_DIR/"institution"; dest.mkdir(exist_ok=True); fname=secure_filename(image.filename); out=dest/f"{uuid.uuid4().hex}-{fname}"; image.save(out); image_path="uploads/institution/"+out.name
    execute("UPDATE school_settings SET institution_history=?, institution_performance=?, institution_religion=?, institution_affiliations=?, institution_help=?, institution_contact=?, institution_portal_guide=?, institution_admin_guide=?, institution_ict_guide=?, institution_finance_guide=?, institution_driver_guide_en=?, institution_driver_guide_sw=?, institution_image_path=?, institution_enabled=1 WHERE id=1",(values["institution_history"],values["institution_performance"],values["institution_religion"],values["institution_affiliations"],values["institution_help"],values["institution_contact"],values["institution_portal_guide"],values["institution_admin_guide"],values["institution_ict_guide"],values["institution_finance_guide"],values["institution_driver_guide_en"],values["institution_driver_guide_sw"],image_path))
    flash("Institution information updated.","success"); return redirect(url_for("institution"))

@app.route("/admin/theme/preset", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def save_theme_preset():
    preset=(request.form.get("theme_preset") or "classic").strip().lower()
    if preset not in {"classic","christmas","easter","madaraka","school-pride"}: preset="classic"
    _save_theme_snapshot("workspace", current_user()["id"])
    execute("UPDATE school_settings SET theme_preset=? WHERE id=1",(preset,))
    audit(current_user()["id"],current_user()["full_name"],"Theme Preset",f"Workspace theme set to {preset}.")
    flash(f"Workspace theme set to {preset.replace('-', ' ').title()}.","success")
    return redirect(url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("admin_dashboard"))

@app.route("/advertisements/create", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def create_advertisement():
    title=request.form.get("title","").strip()[:160]
    body=request.form.get("body","").strip()[:5000]
    link=request.form.get("link_url","").strip()[:600]
    start=request.form.get("start_date","").strip()[:10]
    end=request.form.get("end_date","").strip()[:10]
    try: priority=max(0,min(100,int(request.form.get("priority",0))))
    except Exception: priority=0
    if not title or not body:
        flash("An advert needs a title and message.","danger")
        return redirect(request.referrer or url_for("ict_dashboard"))
    execute("INSERT INTO advertisements(title,body,link_url,start_date,end_date,priority,active,created_by) VALUES(?,?,?,?,?,?,1,?)",(title,body,link,start,end,priority,current_user()["id"]))
    audit(current_user()["id"],current_user()["full_name"],"Create Advert",f"Advert '{title}' published.")
    flash("Advert published to the school portal.","success")
    return redirect(request.referrer or url_for("ict_dashboard"))

@app.route("/advertisements/<int:advert_id>/toggle", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def toggle_advertisement(advert_id:int):
    row=q("SELECT * FROM advertisements WHERE id=?",(advert_id,),one=True)
    if not row: abort(404)
    execute("UPDATE advertisements SET active=? WHERE id=?",(0 if row["active"] else 1,advert_id))
    flash("Advert status updated.","success")
    return redirect(request.referrer or url_for("ict_dashboard"))

@app.route("/ict/features", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def ict_features():
    elections=1 if request.form.get("elections_enabled") else 0; library=1 if request.form.get("library_enabled") else 0
    execute("UPDATE school_settings SET elections_enabled=?, library_enabled=? WHERE id=1",(elections,library)); flash("Module visibility updated.","success"); return redirect(url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("admin_dashboard"))

@app.route("/finance-dashboard")
@login_required
@role_required("Finance", "Admin")
def finance_dashboard():
    settings=school_settings()
    payments=q("""SELECT p.*, s.full_name AS student_name, s.admission_no FROM payments p JOIN students s ON s.id=p.student_id ORDER BY p.created_at DESC LIMIT 50""")
    total_income=q("SELECT COALESCE(SUM(amount),0) AS n FROM payments WHERE status='Posted'", one=True)["n"]
    balance=q("SELECT COALESCE(SUM(balance),0) AS n FROM students", one=True)["n"]
    students=q("SELECT * FROM students WHERE active=1 ORDER BY grade, full_name")
    batches=q("""SELECT b.*,u.full_name AS submitted_by_name,COUNT(r.id) AS entries,ROUND(AVG(r.mark),1) AS mean_mark FROM exam_batches b JOIN users u ON u.id=b.submitted_by LEFT JOIN exam_results r ON r.batch_id=b.id GROUP BY b.id ORDER BY b.created_at DESC,b.id DESC LIMIT 30""")
    documents=q("""SELECT d.*,s.full_name AS student_name,s.admission_no FROM portal_documents d JOIN students s ON s.id=d.student_id ORDER BY d.issued_at DESC LIMIT 30""")
    ledger=q("""SELECT f.*,u.full_name AS poster,pu.full_name AS payee_name FROM finance_ledger f JOIN users u ON u.id=f.posted_by LEFT JOIN users pu ON pu.id=f.payee_user_id ORDER BY f.posted_at DESC,f.id DESC LIMIT 60""")
    ledger_income=q("SELECT COALESCE(SUM(amount),0) AS n FROM finance_ledger WHERE status='Posted' AND entry_type='Income'",one=True)["n"]
    ledger_expense=q("SELECT COALESCE(SUM(amount),0) AS n FROM finance_ledger WHERE status='Posted' AND entry_type IN ('Expense','Payroll')",one=True)["n"]
    account=q("SELECT * FROM finance_accounts WHERE id=1",one=True)
    system_balance=float(account["opening_balance"] if account else 0)+float(total_income or 0)+float(ledger_income or 0)-float(ledger_expense or 0)
    employees=q("SELECT id,full_name,username,role,title,department FROM users WHERE active=1 AND role NOT IN ('System','Student','Parent') ORDER BY full_name")
    fee_structures=q("SELECT * FROM fee_structures WHERE active=1 ORDER BY class_level,item_name")
    fee_charges=q("SELECT c.*,s.full_name,s.admission_no FROM fee_charges c JOIN students s ON s.id=c.student_id ORDER BY c.created_at DESC LIMIT 80")
    payint=q("SELECT * FROM payment_integrations WHERE id=1",one=True)
    external_events=q("SELECT e.*,s.full_name AS matched_name FROM external_payment_events e LEFT JOIN students s ON s.id=e.matched_student_id ORDER BY e.created_at DESC LIMIT 50")
    daily_usage=q("SELECT date(posted_at) day,entry_type,category,COALESCE(SUM(amount),0) amount,COUNT(*) count FROM finance_ledger WHERE status='Posted' GROUP BY date(posted_at),entry_type,category ORDER BY day DESC LIMIT 60")
    closings=q("SELECT c.*,u.full_name AS submitted_name FROM finance_closings c JOIN users u ON u.id=c.submitted_by ORDER BY c.id DESC LIMIT 30")
    nav_items=navigation_items("Finance", settings)
    return render_template("finance_dashboard.html", role="Finance", workspace=workspace_for("Finance"), settings=settings, payments=payments, total_income=total_income, total_balance=balance, students=students, batches=batches, documents=documents, actor_name=current_user()["full_name"], nav_items=nav_items, ledger=ledger, ledger_income=ledger_income, ledger_expense=ledger_expense, system_balance=system_balance, employees=employees, fee_structures=fee_structures, fee_charges=fee_charges, payint=payint, external_events=external_events, daily_usage=daily_usage, closings=closings)


@app.route("/ict-dashboard")
@login_required
@role_required("ICT")
def ict_dashboard():
    settings=school_settings()
    nav_items=navigation_items("ICT", settings)
    elections=q("SELECT * FROM elections ORDER BY created_at DESC")
    election_candidates={e["id"]:q("SELECT * FROM election_candidates WHERE election_id=? ORDER BY position,name",(e["id"],)) for e in elections}
    library_items=q("SELECT * FROM library_items ORDER BY active DESC,category,title")
    students=q("SELECT * FROM students WHERE active=1 ORDER BY grade,full_name")
    users=q("SELECT id,full_name,username,role,student_id,active FROM users WHERE role!='System' ORDER BY role,full_name")
    return render_template("ict_dashboard.html", role="ICT", workspace=workspace_for("ICT"), settings=settings, actor_name=current_user()["full_name"], nav_items=nav_items, onboarding_students=q("SELECT id, full_name, admission_no FROM students ORDER BY full_name"), elections=elections, election_candidates=election_candidates, library_items=library_items, students=students, users=users)


def _theme_snapshot_payload(settings):
    keys=["school_name","portal_subtitle","landing_hero_title","landing_hero_text","landing_cta_primary","landing_cta_secondary","landing_announcement","landing_contact","landing_show_dates","landing_show_gallery","landing_show_roles","primary_color","accent_color","background_color","panel_color","sidebar_color","header_color","text_color","muted_text_color","font_family","heading_font","radius_px","button_radius_px","theme_mode","sidebar_style","menu_order","home_label","assignments_label","results_label","messages_label","finance_label","branding_label","custom_css","footer_title","footer_text","footer_contact","footer_links","platform_credit_enabled","landing_background_color","landing_panel_color","landing_text_color","landing_accent_color","landing_font_family","landing_heading_font","landing_content_width","landing_hero_layout","landing_role_columns","landing_background_path","institution_image_path","institution_image_2_path","institution_image_3_path","institution_image_1_position","institution_image_2_position","institution_image_3_position","welcome_animation_enabled","welcome_animation_name","welcome_animation_duration_ms","welcome_animation_style"]
    return {k: settings[k] for k in keys if k in settings.keys()}

def _save_theme_snapshot(snapshot_type, actor_id):
    settings=q("SELECT * FROM school_settings WHERE id=1",one=True)
    if not settings: return
    execute("INSERT INTO theme_snapshots(snapshot_type,settings_json,created_by) VALUES(?,?,?)",(snapshot_type,json.dumps(_theme_snapshot_payload(settings),ensure_ascii=False),actor_id))
    # Keep a compact history; the latest three restore points are enough for practical rollback.
    ids=q("SELECT id FROM theme_snapshots WHERE snapshot_type=? ORDER BY id DESC",(snapshot_type,))
    for row in ids[3:]: execute("DELETE FROM theme_snapshots WHERE id=?",(row["id"],))

@app.route("/ict/settings", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def ict_settings():
    _save_theme_snapshot("workspace", current_user()["id"])
    school_name=request.form.get("school_name", "School").strip() or "School"
    portal_subtitle=request.form.get("portal_subtitle", "School Portal System").strip() or "School Portal System"
    primary=request.form.get("primary_color", "#3457d5").strip() or "#3457d5"
    accent=request.form.get("accent_color", "#3457d5").strip() or "#3457d5"
    bg=request.form.get("background_color", "#343541").strip() or "#343541"
    panel=request.form.get("panel_color", "#40414f").strip() or "#40414f"
    sidebar=request.form.get("sidebar_color", panel).strip() or panel
    header=request.form.get("header_color", panel).strip() or panel
    text_color=request.form.get("text_color", "#ececf1").strip() or "#ececf1"
    muted=request.form.get("muted_text_color", "#b5bac7").strip() or "#b5bac7"
    footer_title=request.form.get("footer_title", "").strip()
    footer_text=request.form.get("footer_text", "").strip()
    footer_contact=request.form.get("footer_contact", "").strip()
    footer_links=request.form.get("footer_links", "").strip()
    platform_credit_enabled=1 if request.form.get("platform_credit_enabled") in {"1","on","true","yes"} else 0
    font_family=request.form.get("font_family", "Inter").strip() or "Inter"
    heading_font=request.form.get("heading_font", font_family).strip() or font_family
    try: radius=max(4,min(28,int(request.form.get("radius_px", "12"))))
    except ValueError: radius=12
    try: button_radius=max(4,min(28,int(request.form.get("button_radius_px", "10"))))
    except ValueError: button_radius=10
    theme_mode=request.form.get("theme_mode", "dark").strip().lower() if request.form.get("theme_mode") else "dark"
    if theme_mode not in {"dark","light","system"}: theme_mode="dark"
    sidebar_style=request.form.get("sidebar_style", "drawer").strip().lower()
    if sidebar_style not in {"drawer","left","top","hover"}: sidebar_style="drawer"
    menu_order=request.form.get("menu_order", "Home,Assignments,Submissions,Online classes").strip() or "Home,Assignments,Submissions,Online classes"
    labels={k: request.form.get(k, defaults).strip() or defaults for k,defaults in [("home_label","Home"),("assignments_label","Assignments"),("results_label","Results"),("messages_label","Messages"),("finance_label","Finance"),("branding_label","Branding")]}
    custom_css=request.form.get("custom_css", "").strip()[:12000]
    if re.search(r'@import|javascript:|expression\s*\(', custom_css, re.I): custom_css=''
    execute("""UPDATE school_settings SET school_name=?, portal_subtitle=?, primary_color=?, accent_color=?, background_color=?, panel_color=?, sidebar_color=?, header_color=?, text_color=?, muted_text_color=?, font_family=?, heading_font=?, radius_px=?, button_radius_px=?, theme_mode=?, sidebar_style=?, menu_order=?, home_label=?, assignments_label=?, results_label=?, messages_label=?, finance_label=?, branding_label=?, custom_css=?, footer_title=?, footer_text=?, footer_contact=?, footer_links=?, platform_credit_enabled=? WHERE id=1""", (school_name,portal_subtitle,primary,accent,bg,panel,sidebar,header,text_color,muted,font_family,heading_font,radius,button_radius,theme_mode,sidebar_style,menu_order,labels["home_label"],labels["assignments_label"],labels["results_label"],labels["messages_label"],labels["finance_label"],labels["branding_label"],custom_css,footer_title,footer_text,footer_contact,footer_links,platform_credit_enabled))
    audit(current_user()["id"], current_user()["full_name"], "Portal Theme Update", f"Institution-wide interface theme updated for {school_name}.")
    flash("The institution-wide interface has been redesigned and saved.", "success")
    return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))


@app.route("/ict/landing-branding", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def ict_landing_branding():
    _save_theme_snapshot("landing", current_user()["id"])
    # Public landing page branding is intentionally separate from the in-system workspace theme.
    vals={
        "landing_background_color": request.form.get("landing_background_color", "#e7efff").strip(),
        "landing_panel_color": request.form.get("landing_panel_color", "#f8fbff").strip(),
        "landing_text_color": request.form.get("landing_text_color", "#152033").strip(),
        "landing_accent_color": request.form.get("landing_accent_color", "#2457d6").strip(),
        "landing_font_family": request.form.get("landing_font_family", "Inter").strip() or "Inter",
        "landing_heading_font": request.form.get("landing_heading_font", "Manrope").strip() or "Manrope",
        "landing_content_width": max(900,min(1600,int(request.form.get("landing_content_width", "1240") or 1240))),
        "landing_hero_layout": request.form.get("landing_hero_layout", "split").strip().lower() if request.form.get("landing_hero_layout") else "split",
        "landing_role_columns": max(1,min(3,int(request.form.get("landing_role_columns", "3") or 3))),
        "welcome_animation_enabled": 1 if request.form.get("welcome_animation_enabled") in {"1","on","true","yes"} else 0,
        "welcome_animation_name": request.form.get("welcome_animation_name", "Toror Technology and Innovations Ltd.").strip()[:120] or "Toror Technology and Innovations Ltd.",
        "welcome_animation_duration_ms": max(1200,min(5000,int(request.form.get("welcome_animation_duration_ms", "2200") or 2200))),
        "welcome_animation_style": request.form.get("welcome_animation_style", "clean").strip().lower(),
    }
    if vals["landing_hero_layout"] not in {"split","stacked"}: vals["landing_hero_layout"]="split"
    positions=[request.form.get(f"institution_image_{i}_position","50% 50%").strip()[:40] for i in (1,2,3)]
    paths=[]
    landing_file=request.files.get("landing_background")
    landing_path=school_settings()["landing_background_path"] or ""
    if landing_file and landing_file.filename:
        ext=landing_file.filename.rsplit('.',1)[-1].lower() if '.' in landing_file.filename else ''
        if ext not in {"png","jpg","jpeg","webp"}:
            flash("Landing background must be PNG, JPG, JPEG or WEBP.","danger"); return redirect((url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("admin_dashboard"))+"#branding")
        folder=UPLOAD_DIR/"institution"; folder.mkdir(exist_ok=True)
        out=folder/f"landing-bg-{uuid.uuid4().hex[:10]}.{ext}"; landing_file.save(out); landing_path="uploads/institution/"+out.name
    for i in (1,2,3):
        file=request.files.get(f"institution_image_{i}")
        existing=school_settings()[f"institution_image_{i if i>1 else 1}_path"] if i>1 else school_settings()["institution_image_path"]
        path=existing or ""
        if file and file.filename:
            ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
            if ext not in {"png","jpg","jpeg","webp"}:
                flash(f"History image {i} must be PNG, JPG, JPEG or WEBP.","danger"); return redirect((url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("admin_dashboard"))+"#branding")
            folder=UPLOAD_DIR/"institution"; folder.mkdir(exist_ok=True)
            out=folder/f"history-{i}-{uuid.uuid4().hex[:10]}.{ext}"; file.save(out); path="uploads/institution/"+out.name
        paths.append(path)
    if vals["welcome_animation_style"] not in {"clean","zoom","wave","particles","split"}: vals["welcome_animation_style"]="clean"
    execute("""UPDATE school_settings SET landing_background_color=?,landing_panel_color=?,landing_text_color=?,landing_accent_color=?,landing_font_family=?,landing_heading_font=?,landing_content_width=?,landing_hero_layout=?,landing_role_columns=?,landing_background_path=?,institution_image_path=?,institution_image_2_path=?,institution_image_3_path=?,institution_image_1_position=?,institution_image_2_position=?,institution_image_3_position=?,welcome_animation_enabled=?,welcome_animation_name=?,welcome_animation_duration_ms=?,welcome_animation_style=? WHERE id=1""",(vals["landing_background_color"],vals["landing_panel_color"],vals["landing_text_color"],vals["landing_accent_color"],vals["landing_font_family"],vals["landing_heading_font"],vals["landing_content_width"],vals["landing_hero_layout"],vals["landing_role_columns"],landing_path,paths[0],paths[1],paths[2],positions[0],positions[1],positions[2],vals["welcome_animation_enabled"],vals["welcome_animation_name"],vals["welcome_animation_duration_ms"],vals["welcome_animation_style"]))
    audit(current_user()["id"],current_user()["full_name"],"Landing Page Branding Update","Public landing branding, history visuals and welcome animation updated.")
    flash("Landing-page branding and institution history visuals saved separately from the logged-in system theme.","success")
    return redirect((url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("admin_dashboard"))+"#branding")

@app.route("/ict/logo", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def ict_logo():
    file=request.files.get("logo")
    if not file or not file.filename: return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))
    ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
    if ext not in {"png","jpg","jpeg","webp","svg"}:
        flash("Logo must be PNG, JPG, JPEG, WEBP or SVG.", "danger")
        return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))
    name="school-logo."+ext
    path=UPLOAD_DIR/name; file.save(path)
    execute("UPDATE school_settings SET logo_path=? WHERE id=1", ("uploads/"+name,))
    flash("School logo updated.", "success")
    return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))


@app.route("/ict/background", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def ict_background():
    file=request.files.get("background")
    if not file or not file.filename: return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))
    ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
    if ext not in {"png","jpg","jpeg","webp"}:
        flash("Background must be PNG, JPG, JPEG or WEBP.", "danger")
        return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))
    name="school-background."+ext
    file.save(UPLOAD_DIR/name)
    execute("UPDATE school_settings SET background_path=? WHERE id=1", ("uploads/"+name,))
    flash("Portal background updated.", "success")
    return redirect(url_for("ict_dashboard" if current_user()["role"]=="ICT" else "admin_dashboard"))


@app.route("/calendar")
@login_required
def calendar_view():
    return render_template("calendar.html", settings=school_settings(), dates=important_dates(100), calendar_rules=q("SELECT * FROM school_calendar ORDER BY start_date,end_date,id"), day_status=school_day_status(), actor_name=current_user()["full_name"], role=current_user()["role"])

@app.route("/calendar/create", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def calendar_create():
    title=request.form.get("title","").strip(); date=request.form.get("event_date","").strip()
    kind=request.form.get("kind","Important Date").strip() or "Important Date"
    if not title or not date:
        flash("Event title and date are required.","danger"); return redirect(url_for("calendar_view"))
    execute("INSERT INTO important_dates(title,event_date,event_time,location,description,visible,landing_visible,created_by) VALUES(?,?,?,?,?,?,?,?)",(title,date,request.form.get("event_time",""),request.form.get("location",""),request.form.get("description","")+ (f" [{kind}]" if kind else ""),1,1,current_user()["id"]))
    ids=[r["id"] for r in q("SELECT id FROM users WHERE active=1 AND role!='System'")]
    notify_users(ids,"Important date added",f"{title} — {date}",url_for("calendar_view"))
    flash("Important date published to the calendar and landing page.","success"); return redirect(url_for("calendar_view"))

@app.route("/calendar/rule/create", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def calendar_rule_create():
    title=request.form.get("title","").strip(); start=request.form.get("start_date","").strip(); end=request.form.get("end_date","").strip() or start; kind=request.form.get("kind","School Day").strip() or "School Day"
    try:
        sd=datetime.strptime(start,"%Y-%m-%d").date(); ed=datetime.strptime(end,"%Y-%m-%d").date()
    except ValueError:
        flash("Enter valid calendar dates.","danger"); return redirect(url_for("calendar_view"))
    if ed < sd or not title:
        flash("The title and a valid date range are required.","danger"); return redirect(url_for("calendar_view"))
    closed_kinds={"Mid-term Break","National Holiday","Public Holiday","School Holiday","School Closed"}
    school_day=0 if kind in closed_kinds else 1
    execute("INSERT INTO school_calendar(title,start_date,end_date,kind,school_day,notes,created_by) VALUES(?,?,?,?,?,?,?)",(title,sd.isoformat(),ed.isoformat(),kind,school_day,request.form.get("notes","").strip(),current_user()["id"]))
    flash(f"{kind} calendar rule saved.","success"); return redirect(url_for("calendar_view"))

@app.route("/calendar/rule/<int:rule_id>/delete", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def calendar_rule_delete(rule_id):
    execute("DELETE FROM school_calendar WHERE id=?",(rule_id,)); flash("School calendar rule removed.","success"); return redirect(url_for("calendar_view"))

@app.route("/calendar/<int:event_id>/delete", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def calendar_delete(event_id):
    execute("UPDATE important_dates SET visible=0, landing_visible=0 WHERE id=?",(event_id,)); flash("Calendar event removed.","success"); return redirect(url_for("calendar_view"))

def _notification_recipients_for(actor, audience, selected_ids=None, class_name=""):
    audience=(audience or "").strip().lower()
    selected_ids=selected_ids or []
    ids=[]
    if audience == "all":
        ids=[r["id"] for r in q("SELECT id FROM users WHERE active=1 AND role!='System'")]
    elif audience == "staff":
        ids=[r["id"] for r in q("SELECT id FROM users WHERE active=1 AND role IN ('Admin','ICT','Finance','Teacher','Librarian','Reception','Driver')")]
    elif audience == "students":
        ids=[r["id"] for r in q("SELECT id FROM users WHERE active=1 AND role='Student'")]
    elif audience == "parents":
        ids=[r["id"] for r in q("SELECT id FROM users WHERE active=1 AND role='Parent'")]
    elif audience == "class":
        if actor["role"] == "Teacher":
            allowed=[r["class_name"] for r in q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?",(actor["id"],))]
            taught=[r["class_name"] for r in q("SELECT DISTINCT class_name FROM teacher_assignments WHERE teacher_user_id=? AND active=1",(actor["id"],))]
            allowed=sorted(set(allowed+taught))
            if class_name and class_name not in allowed:
                abort(403)
        rows=q("SELECT u.id FROM users u JOIN students s ON s.id=u.student_id WHERE u.active=1 AND u.role='Student' AND lower(s.grade)=lower(?)",(class_name,))
        ids=[r["id"] for r in rows]
    elif audience == "selected":
        clean=[]
        for raw in selected_ids:
            try: clean.append(int(raw))
            except Exception: pass
        if clean:
            marks=','.join('?' for _ in clean)
            ids=[r["id"] for r in q(f"SELECT id FROM users WHERE active=1 AND role!='System' AND id IN ({marks})", clean)]
    elif audience == "self":
        ids=[actor["id"]]
    else:
        ids=[]
    return sorted(set(ids))

@app.route("/notifications/send", methods=["POST"])
@login_required
@role_required("Admin","ICT","Teacher")
def notifications_send():
    actor=current_user()
    audience=request.form.get("audience","selected")
    title=request.form.get("title","").strip()[:160]
    body=request.form.get("body","").strip()[:5000]
    link=request.form.get("link","").strip()[:500]
    priority=request.form.get("priority","Normal").strip() or "Normal"
    class_name=request.form.get("class_name","").strip()
    selected=request.form.getlist("user_ids")
    if not title or not body:
        flash("Enter a notification title and message.","danger")
        return redirect(url_for("notifications_view"))
    ids=_notification_recipients_for(actor,audience,selected,class_name)
    if not ids:
        flash("No recipients matched that audience.","warning")
        return redirect(url_for("notifications_view"))
    # A Teacher can notify learners only through their assigned/taught class or a
    # deliberate selected-user action. Direct selected users are constrained to Students.
    if actor["role"]=="Teacher":
        allowed_teacher_ids=set(_notification_recipients_for(actor,"students",[],""))
        ids=[uid for uid in ids if uid in allowed_teacher_ids or uid==actor["id"]]
        if audience not in {"class","selected","self"} and audience not in {"students"}:
            abort(403)
    notify_users(ids,title,body,link,priority)
    audit(actor["id"],actor["full_name"],"Notification sent",f"{title} · {len(ids)} recipient(s) · {audience}")
    flash(f"Notification sent to {len(ids)} recipient(s).","success")
    return redirect(url_for("notifications_view"))

@app.route("/notifications")
@login_required
def notifications_view():
    user=current_user()
    rows=q("SELECT * FROM notifications WHERE user_id=? ORDER BY created_at DESC,id DESC LIMIT 100",(user["id"],))
    recipients=[]
    if user["role"] in {"Admin","ICT","Teacher"}:
        recipients=q("SELECT id,full_name,username,role,student_id FROM users WHERE active=1 AND id!=? AND role!='System' ORDER BY full_name",(user["id"],))
    classes=[]
    if user["role"]=="Teacher":
        own=[r["class_name"] for r in q("SELECT class_name FROM teacher_assignments WHERE teacher_user_id=? AND active=1",(user["id"],))]
        own+= [r["class_name"] for r in q("SELECT class_name FROM class_teacher_assignments WHERE teacher_user_id=?",(user["id"],))]
        classes=sorted(set(own))
    elif user["role"] in {"Admin","ICT"}:
        classes=[r["grade"] for r in q("SELECT DISTINCT grade FROM students WHERE active=1 AND TRIM(COALESCE(grade,''))!='' ORDER BY grade")]
    return render_template("notifications.html",settings=school_settings(),notifications=rows,actor_name=user["full_name"],role=user["role"],notification_recipients=recipients,notification_classes=classes)

@app.route("/notifications/read", methods=["POST"])
@login_required
def notifications_read():
    nid=request.form.get("id",type=int)
    if nid: execute("UPDATE notifications SET read_at=CURRENT_TIMESTAMP WHERE id=? AND user_id=?",(nid,current_user()["id"]))
    else: execute("UPDATE notifications SET read_at=CURRENT_TIMESTAMP WHERE user_id=? AND read_at IS NULL",(current_user()["id"],))
    return redirect(request.referrer or url_for("notifications_view"))

@app.route("/api/notifications")
@login_required
def notifications_api():
    rows=q("SELECT id,title,body,link,created_at,priority FROM notifications WHERE user_id=? AND read_at IS NULL ORDER BY created_at DESC,id DESC LIMIT 8",(current_user()["id"],))
    return jsonify({"count":notification_count(current_user()["id"]),"items":[dict(r) for r in rows]})

@app.route("/online-classes")
@login_required
def online_classes():
    user=current_user()
    sessions=q("SELECT cs.*,u.full_name AS teacher_name FROM class_sessions cs JOIN users u ON u.id=cs.teacher_user_id WHERE cs.active=1 AND (cs.starts_at >= datetime('now','-1 day') OR cs.ends_at IS NULL OR cs.ends_at >= datetime('now','-1 day')) ORDER BY cs.starts_at",())
    if user["role"]=="Teacher": sessions=[r for r in sessions if r["teacher_user_id"]==user["id"]] + [r for r in sessions if r["teacher_user_id"]!=user["id"]]
    elif user["role"]=="Student" and user["student_id"]:
        st=q("SELECT grade FROM students WHERE id=?",(user["student_id"],),one=True); grade=st["grade"] if st else ""; sessions=q("SELECT cs.*,u.full_name AS teacher_name FROM class_sessions cs JOIN users u ON u.id=cs.teacher_user_id WHERE cs.active=1 AND lower(cs.class_name)=lower(?) ORDER BY cs.starts_at",(grade,))
    return render_template("online_classes.html",settings=school_settings(),sessions=sessions,assignments=(q("SELECT * FROM teacher_assignments WHERE teacher_user_id=? AND active=1 ORDER BY class_name,subject",(user["id"],)) if user["role"]=="Teacher" else []),actor_name=user["full_name"],role=user["role"],today=datetime.now().strftime('%Y-%m-%d'))

@app.route("/online-class/<int:session_id>")
@login_required
def live_classroom(session_id:int):
    user=current_user(); sess=q("SELECT cs.*,u.full_name AS teacher_name FROM class_sessions cs JOIN users u ON u.id=cs.teacher_user_id WHERE cs.id=? AND cs.active=1",(session_id,),one=True)
    if not sess: abort(404)
    allowed=False; mode='student'
    if user['role']=='Teacher':
        allowed=True; mode='teacher' if user['id']==sess['teacher_user_id'] else 'teacher_observer'
    elif user['role']=='Admin': allowed=True; mode='moderator'
    elif user['role']=='Student':
        st=q("SELECT grade FROM students WHERE id=? AND active=1",(user['student_id'],),one=True) if user['student_id'] else None
        allowed=bool(st and str(st['grade']).lower()==str(sess['class_name']).lower())
    if not allowed:
        if user['role']=='Teacher': return redirect(url_for('online_classes'))
        abort(403)
    roster=q("SELECT s.id,s.full_name,s.admission_no FROM students s WHERE s.active=1 AND lower(s.grade)=lower(?) ORDER BY s.full_name",(sess['class_name'],))
    provider=sess['provider_url'] or 'https://meet.jit.si/'
    provider=provider if provider.endswith('/') else provider+'/'
    return render_template('live_classroom.html',settings=school_settings(),session=sess,roster=roster,actor_name=user['full_name'],role=user['role'],mode=mode,jitsi_domain=urllib.parse.urlparse(provider).netloc or 'meet.jit.si')

@app.route("/teacher/online-class/create", methods=["POST"])
@login_required
@role_required("Teacher")
def create_online_class():
    user=current_user(); title=request.form.get("title","").strip() or "Live Class"; cls=request.form.get("class_name","").strip(); subject=request.form.get("subject","").strip(); starts=request.form.get("starts_at","").strip(); ends=request.form.get("ends_at","").strip(); desc=request.form.get("description","").strip(); audience=request.form.get("audience_mode","Class").strip() or "Class"
    if not cls or not subject or not starts: flash("Class, subject and start time are required.","danger"); return redirect(url_for("online_classes"))
    # A teacher may create a cohort/room even before students have been allocated.
    # Existing assignments still determine the normal class population; an empty
    # class is a valid preparation state rather than an access-denied dead end.
    compulsory=q("SELECT id FROM compulsory_subjects WHERE active=1 AND lower(class_name)=lower(?) AND lower(subject)=lower(?)",(cls,subject),one=True)
    if compulsory: audience="Compulsory"
    room=f"{school_settings()['school_name'].replace(' ','-')}-{cls.replace(' ','-')}-{uuid.uuid4().hex[:8]}"
    provider="https://meet.jit.si/"; url=provider+urllib.parse.quote(room)
    sid=execute("INSERT INTO class_sessions(teacher_user_id,class_name,subject,title,starts_at,ends_at,room_name,provider_url,description,audience_mode) VALUES(?,?,?,?,?,?,?,?,?,?)",(user['id'],cls,subject,title,starts,ends,room,provider,desc,audience))
    ids=[r["id"] for r in q("SELECT u.id FROM users u JOIN students s ON s.id=u.student_id WHERE u.active=1 AND u.role='Student' AND lower(s.grade)=lower(?)",(cls,))]
    when=starts.replace('T',' ')
    notify_users(ids,"Class scheduled" if audience!='Compulsory' else "Compulsory class scheduled",f"{title} — {subject} at {when}. Open your Student Dashboard to join when it starts.",url_for("online_classroom",session_id=sid),"High" if audience=='Compulsory' else "Normal")
    flash("Live class scheduled. Students have been notified, and the lesson is now on their classroom schedule.","success"); return redirect(url_for("online_classes"))

@app.route("/online-class/<int:session_id>/save", methods=["POST"])
@login_required
def save_online_class(session_id:int):
    user=current_user(); sess=q("SELECT * FROM class_sessions WHERE id=? AND active=1",(session_id,),one=True)
    if not sess: abort(404)
    if user["role"] not in {"Admin","ICT"} and not (user["role"]=="Teacher" and user["id"]==sess["teacher_user_id"]): abort(403)
    notes=request.form.get("archive_notes","").strip()[:4000]
    recording_url=request.form.get("recording_url","").strip()[:1000]
    uploaded=request.files.get("recording")
    recording_path=sess["recording_path"] or ""
    if uploaded and uploaded.filename:
        ext=uploaded.filename.rsplit('.',1)[-1].lower() if '.' in uploaded.filename else ''
        if ext not in {"mp4","webm","mov","m4v","mp3","m4a","wav","ogg"}:
            flash("Recording must be MP4, WEBM, MOV, M4V, MP3, M4A, WAV or OGG.","danger"); return redirect(url_for("live_classroom",session_id=session_id))
        folder=UPLOAD_DIR/"online_classes"; folder.mkdir(exist_ok=True)
        out=folder/f"class-{session_id}-{uuid.uuid4().hex[:10]}.{ext}"; uploaded.save(out); recording_path="uploads/online_classes/"+out.name
    title=f"{sess['title']} — {sess['subject']} ({sess['class_name']})"
    description=(sess["description"] or "").strip()
    combined=(description+(("\n\n"+notes) if notes else "")).strip()
    library_id=None
    if request.form.get("publish_library") in {"1","on","true","yes"}:
        existing=q("SELECT id FROM library_items WHERE active=1 AND title=? LIMIT 1",(title,),one=True)
        if existing:
            library_id=existing["id"]
            execute("UPDATE library_items SET description=?,class_level=?,subject=?,resource_type=?,file_path=?,external_url=? WHERE id=?",(combined,sess["class_name"],sess["subject"],"Digital",recording_path,recording_url,library_id))
        else:
            library_id=execute("INSERT INTO library_items(title,category,author,quantity,available_quantity,resource_type,file_path,external_url,description,created_by,class_level,subject) VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",(title,"Recorded Class",sess["teacher_name"] if "teacher_name" in sess.keys() else user["full_name"],1,1,"Digital",recording_path,recording_url,combined,user["id"],sess["class_name"],sess["subject"]))
    execute("UPDATE class_sessions SET recording_path=?,recording_url=?,saved_at=CURRENT_TIMESTAMP,library_item_id=? WHERE id=?",(recording_path,recording_url,library_id,session_id))
    student_ids=[r["id"] for r in q("SELECT u.id FROM users u JOIN students s ON s.id=u.student_id WHERE u.active=1 AND u.role='Student' AND lower(s.grade)=lower(?)",(sess["class_name"],))]
    resource_link=url_for("online_classroom",session_id=session_id)
    if recording_path or recording_url:
        resource_link=url_for("download_online_class",session_id=session_id)
    notify_users(student_ids,"Recorded class available",f"{sess['title']} — {sess['subject']} is now available for revision. Open the class resource from your notifications or Library.",resource_link,"Normal")
    audit(user["id"],user["full_name"],"Online Class Archive",f"Saved learning resource for online class {sess['title']}.")
    flash("Class saved as a reusable learning resource" + (" and added to the Library." if library_id else "." ) + " Students have been notified.","success")
    return redirect(url_for("live_classroom",session_id=session_id))

@app.route("/online-class/<int:session_id>/download")
@login_required
def download_online_class(session_id:int):
    user=current_user()
    sess=q("SELECT * FROM class_sessions WHERE id=? AND active=1",(session_id,),one=True)
    if not sess: abort(404)
    allowed=False
    if user["role"] in {"Admin","ICT"}: allowed=True
    elif user["role"]=="Teacher" and user["id"]==sess["teacher_user_id"]: allowed=True
    elif user["role"]=="Student" and user["student_id"]:
        st=q("SELECT grade FROM students WHERE id=? AND active=1",(user["student_id"],),one=True)
        allowed=bool(st and str(st["grade"]).lower()==str(sess["class_name"]).lower())
    if not allowed: abort(403)
    if sess["recording_path"]:
        target=(BASE_DIR / sess["recording_path"]).resolve()
        root=UPLOAD_DIR.resolve()
        if root not in target.parents or not target.exists() or not target.is_file(): abort(404)
        return send_file(target,as_attachment=True,download_name=Path(target).name)
    if sess["recording_url"]:
        return redirect(sess["recording_url"])
    abort(404)

@app.route("/groups")
@login_required
def groups_view():
    user=current_user(); groups=q("SELECT g.*,u.full_name AS owner_name,(SELECT COUNT(*) FROM group_members gm WHERE gm.group_id=g.id) AS member_count FROM groups g LEFT JOIN users u ON u.id=g.owner_user_id WHERE g.active=1 ORDER BY g.created_at DESC")
    if user["role"]=="Student": groups=[r for r in groups if q("SELECT 1 FROM group_members WHERE group_id=? AND user_id=?",(r["id"],user["id"]),one=True)]
    mine=[r for r in groups if r["owner_user_id"]==user["id"]]
    students=q("SELECT u.id,u.full_name,u.student_id,s.grade,s.admission_no FROM users u LEFT JOIN students s ON s.id=u.student_id WHERE u.role='Student' AND u.active=1 ORDER BY u.full_name") if user["role"] in {"Teacher","Admin","ICT"} else []
    return render_template("groups.html",settings=school_settings(),groups=groups,mine=mine,students=students,actor_name=user["full_name"],role=user["role"])

@app.route("/groups/create", methods=["POST"])
@login_required
@role_required("Teacher","Admin","ICT")
def group_create():
    name=request.form.get("name","").strip(); desc=request.form.get("description","").strip(); typ=request.form.get("group_type","Academic").strip() or "Academic"
    if not name: flash("Group name is required.","danger"); return redirect(url_for("groups_view"))
    gid=execute("INSERT INTO groups(name,group_type,description,owner_user_id) VALUES(?,?,?,?)",(name,typ,desc,current_user()["id"]))
    flash("Group created. Add learners below.","success"); return redirect(url_for("groups_view"))

@app.route("/groups/<int:group_id>/members", methods=["POST"])
@login_required
@role_required("Teacher","Admin","ICT")
def group_add_members(group_id):
    group=q("SELECT * FROM groups WHERE id=?",(group_id,),one=True)
    if not group: abort(404)
    if current_user()["role"]=="Teacher" and group["owner_user_id"]!=current_user()["id"]: abort(403)
    ids=request.form.getlist("user_ids")
    for uid in ids:
        u=q("SELECT id,student_id FROM users WHERE id=? AND role='Student' AND active=1",(int(uid),),one=True)
        if u: execute("INSERT OR IGNORE INTO group_members(group_id,user_id,student_id,role) VALUES(?,?,?,'Member')",(group_id,u["id"],u["student_id"]))
    flash("Group membership updated.","success"); return redirect(url_for("groups_view"))

@app.route("/groups/<int:group_id>/post", methods=["POST"])
@login_required
def group_post(group_id):
    member=q("SELECT 1 FROM group_members WHERE group_id=? AND user_id=?",(group_id,current_user()["id"]),one=True); group=q("SELECT * FROM groups WHERE id=? AND active=1",(group_id,),one=True)
    if not group or (current_user()["role"] not in {"Admin","ICT","Teacher"} and not member): abort(403)
    body=request.form.get("body","").strip()
    if body: execute("INSERT INTO group_posts(group_id,user_id,body) VALUES(?,?,?)",(group_id,current_user()["id"],body))
    return redirect(url_for("groups_view"))

@app.route("/student-leadership")
@login_required
@role_required("Student")
def student_leadership():
    user=current_user()
    if not user["leadership_role"]:
        flash("No student leadership assignment is attached to this account.","warning")
        return redirect(url_for("student_dashboard"))
    peers=q("SELECT u.full_name,u.leadership_role,u.department,u.leadership_level FROM users u WHERE u.active=1 AND u.leadership_role!='' ORDER BY u.leadership_level DESC,u.department,u.full_name")
    return render_template("student_leadership.html",settings=school_settings(),role="Student Leader",actor_name=user["full_name"],user=user,peers=peers)

@app.route("/leadership")
@login_required
def leadership_view():
    rows=q("SELECT u.id,u.full_name,u.role,u.title,u.department,u.leadership_role,u.leadership_level FROM users u WHERE u.active=1 AND u.leadership_role!='' ORDER BY u.leadership_level DESC,u.department,u.full_name")
    return render_template("leadership.html",settings=school_settings(),rows=rows,role=current_user()["role"],actor_name=current_user()["full_name"],users=q("SELECT id,full_name,role FROM users WHERE active=1 AND role!='System' ORDER BY full_name") if current_user()["role"] in {"Admin","ICT"} else [])

@app.route("/leadership/assign", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def leadership_assign():
    uid=request.form.get("user_id",type=int); role_name=request.form.get("leadership_role","").strip(); level=request.form.get("level",type=int) or 1; dept=request.form.get("department","").strip()
    u=q("SELECT * FROM users WHERE id=?",(uid,),one=True) if uid else None
    if not u or not role_name: flash("Select a person and leadership role.","danger"); return redirect(url_for("leadership_view"))
    execute("UPDATE users SET leadership_role=?,leadership_level=?,department=? WHERE id=?",(role_name,level,dept,uid))
    execute("INSERT INTO leadership_assignments(user_id,leadership_role,level,department,appointed_by) VALUES(?,?,?,?,?)",(uid,role_name,level,dept,current_user()["id"]))
    notify_user(uid,"Leadership responsibility assigned",f"You have been assigned: {role_name}" ,url_for("leadership_view"))
    flash("Leadership structure updated.","success"); return redirect(url_for("leadership_view"))

@app.route("/assignments/create", methods=["POST"])
@login_required
@role_required("Teacher", "Staff")
def create_assignment():
    title=request.form.get("title", "").strip(); subject=request.form.get("subject", "").strip(); grade=request.form.get("grade", "").strip()
    if not title or not subject or not grade:
        flash("Title, subject and grade are required.", "danger"); return redirect(url_for("teacher_dashboard"))
    attachment_path=""
    file=request.files.get("attachment")
    allowed_types=request.form.get("allowed_types","pdf,doc,docx,xls,xlsx,ppt,pptx,csv,txt,png,jpg,jpeg,webp,zip").strip()
    max_submissions=max(1,min(5,request.form.get("max_submissions",2,type=int) or 2))
    allow_any=1 if request.form.get("allow_any_file") else 0
    if file and file.filename:
        ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
        allowed={x.strip().lower() for x in allowed_types.split(',') if x.strip()}
        if not allow_any and ext not in allowed:
            flash("This assignment's configured file types do not include that attachment.", "danger"); return redirect(url_for("teacher_dashboard"))
        filename=f"assignment-{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{ext}"
        file.save(UPLOAD_DIR/filename); attachment_path="uploads/"+filename
    assignment_id=execute("INSERT INTO assignments(title,subject,grade,description,deadline,attachment_path,posted_by,allowed_types,max_submissions,allow_any_file) VALUES(?,?,?,?,?,?,?,?,?,?)", (title,subject,grade,request.form.get("description",""),request.form.get("deadline",""),attachment_path,current_user()["id"],allowed_types,max_submissions,allow_any))
    student_users=q("SELECT u.id FROM users u JOIN students s ON s.id=u.student_id WHERE u.active=1 AND u.role='Student' AND lower(s.grade)=lower(?)",(grade,))
    notify_users([r["id"] for r in student_users],"New assignment",f"{title} — {subject}",url_for("student_dashboard"))
    audit(current_user()["id"], current_user()["full_name"], "Post Assignment", f"{title} posted to {grade}.")
    flash("Assignment posted to the selected class.", "success")
    return redirect(url_for("teacher_dashboard", grade=grade))


@app.route("/assignments/<int:assignment_id>/submit", methods=["POST"])
@login_required
@role_required("Student")
def submit_assignment(assignment_id):
    student=portal_student(request.form.get("student_id", type=int))
    assignment=q("SELECT * FROM assignments WHERE id=?", (assignment_id,), one=True)
    if not student or not assignment or str(assignment["grade"]).lower() != str(student["grade"]).lower(): abort(404)
    deadline=assignment["deadline"] or ""
    if deadline and datetime.now().strftime('%Y-%m-%dT%H:%M') > deadline[:16]:
        flash("The submission deadline has passed.","danger"); return redirect(url_for("student_dashboard", student_id=student["id"])+'#learning')
    attempts=q("SELECT COUNT(*) AS n FROM submissions WHERE assignment_id=? AND student_id=?",(assignment_id,student["id"]),one=True)["n"]
    limit=max(1,int(assignment["max_submissions"] or 2))
    if attempts >= limit:
        flash(f"This assignment allows {limit} submissions only.","danger"); return redirect(url_for("student_dashboard", student_id=student["id"])+'#learning')
    file=request.files.get('submission'); path=""
    if not file or not file.filename:
        flash("Choose a file before submitting.","danger"); return redirect(url_for("student_dashboard", student_id=student["id"])+'#learning')
    ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
    allowed={x.strip().lower() for x in (assignment["allowed_types"] or '').split(',') if x.strip()}
    if not assignment["allow_any_file"] and ext not in allowed:
        flash("That file type is not accepted for this assignment.","danger"); return redirect(url_for("student_dashboard", student_id=student["id"])+'#learning')
    filename=f"submission-{student['id']}-{assignment_id}-{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{ext}"
    file.save(UPLOAD_DIR/filename); path="uploads/"+filename
    execute("INSERT INTO submissions(assignment_id,student_id,attachment_path,note) VALUES(?,?,?,?)", (assignment_id,student["id"],path,request.form.get("note", "").strip()))
    flash(f"Submission {attempts+1} of {limit} recorded successfully.", "success")
    return redirect(url_for("student_dashboard", student_id=student["id"])+'#learning')


# -------------------------------------------------------------------
# Elections
# -------------------------------------------------------------------
def election_now_active(election):
    if not election or not election["active"] or not election["visible"]:
        return False
    now=datetime.now().strftime("%Y-%m-%dT%H:%M")
    return (not election["start_at"] or now >= election["start_at"]) and (not election["end_at"] or now <= election["end_at"])

@app.route("/elections/vote", methods=["POST"])
@login_required
@role_required("Student")
def cast_vote():
    if not school_settings()["elections_enabled"]:
        abort(404)
    election_id=request.form.get("election_id", type=int)
    candidate_id=request.form.get("candidate_id", type=int)
    election=q("SELECT * FROM elections WHERE id=?", (election_id,), one=True)
    candidate=q("SELECT * FROM election_candidates WHERE id=? AND election_id=? AND active=1", (candidate_id,election_id), one=True)
    if not election or not candidate or not election_now_active(election):
        flash("This election is not currently open for voting.", "warning")
        return redirect(url_for("student_dashboard")+"#elections")
    position=candidate["position"]
    already=q("SELECT id FROM election_votes WHERE election_id=? AND voter_user_id=? AND position=?",(election_id,current_user()["id"],position),one=True)
    if already:
        flash(f"You have already voted for {position} in this election.", "warning")
        return redirect(url_for("student_dashboard")+"#elections")
    execute("INSERT INTO election_votes(election_id,candidate_id,voter_user_id,position) VALUES(?,?,?,?)", (election_id,candidate_id,current_user()["id"],position))
    audit(current_user()["id"], current_user()["full_name"], "Cast Election Vote", f"Voted in election {election_id}.")
    flash("Your vote was recorded.", "success")
    return redirect(url_for("student_dashboard")+"#elections")

@app.route("/elections/create", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def create_election():
    title=request.form.get("title","").strip(); description=request.form.get("description","").strip(); start_at=request.form.get("start_at") or None; end_at=request.form.get("end_at") or None
    if not title:
        flash("Election title is required.", "danger"); return redirect(request.referrer or url_for("ict_dashboard"))
    eid=execute("INSERT INTO elections(title,description,start_at,end_at,visible,active,created_by) VALUES(?,?,?,?,0,0,?)", (title,description,start_at,end_at,current_user()["id"]))
    audit(current_user()["id"], current_user()["full_name"], "Create Election", title)
    flash("Election created. Add participants, then activate it when ready.", "success")
    return redirect((url_for("admin_dashboard") if current_user()["role"]=="Admin" else url_for("ict_dashboard"))+"#elections")

@app.route("/elections/<int:election_id>/toggle", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def toggle_election(election_id):
    election=q("SELECT * FROM elections WHERE id=?",(election_id,),one=True)
    if not election: abort(404)
    visible=1 if request.form.get("visible") else 0
    active=1 if request.form.get("active") else 0
    execute("UPDATE elections SET visible=?, active=? WHERE id=?",(visible,active,election_id))
    flash("Election visibility/status updated.","success")
    return redirect((url_for("admin_dashboard") if current_user()["role"]=="Admin" else url_for("ict_dashboard"))+"#elections")

@app.route("/elections/candidate/add", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def add_election_candidate():
    election_id=request.form.get("election_id",type=int); name=request.form.get("name","").strip(); position=request.form.get("position","").strip(); manifesto=request.form.get("manifesto","").strip()
    if not election_id or not name or not position:
        flash("Election, participant name and position are required.","danger"); return redirect(request.referrer or url_for("ict_dashboard"))
    file=request.files.get("image"); image_path=""
    if file and file.filename:
        fname=secure_filename(file.filename); dest=UPLOAD_DIR/"elections"; dest.mkdir(exist_ok=True); out=dest/f"{uuid.uuid4().hex}-{fname}"; file.save(out); image_path="uploads/elections/"+out.name
    execute("INSERT INTO election_candidates(election_id,name,position,manifesto,image_path) VALUES(?,?,?,?,?)",(election_id,name,position,manifesto,image_path))
    flash("Election participant added.","success")
    return redirect((url_for("admin_dashboard") if current_user()["role"]=="Admin" else url_for("ict_dashboard"))+"#elections")

@app.route("/elections/candidates/upload", methods=["POST"])
@login_required
@role_required("Admin","ICT")
def upload_election_candidates():
    election_id=request.form.get("election_id",type=int); file=request.files.get("participants")
    if not election_id or not file or not file.filename:
        flash("Select an election and CSV participant file.","danger"); return redirect(request.referrer or url_for("ict_dashboard"))
    try:
        text=file.read().decode("utf-8-sig")
        reader=csv.DictReader(io.StringIO(text)); count=0
        for row in reader:
            name=(row.get("name") or row.get("Name") or "").strip(); position=(row.get("position") or row.get("Position") or "").strip(); manifesto=(row.get("manifesto") or row.get("Manifesto") or "").strip()
            if not name or not position: continue
            execute("INSERT INTO election_candidates(election_id,name,position,manifesto) VALUES(?,?,?,?)",(election_id,name,position,manifesto)); count+=1
        flash(f"Imported {count} election participants.","success")
    except Exception as exc:
        flash(f"Could not read participant file: {exc}","danger")
    return redirect((url_for("admin_dashboard") if current_user()["role"]=="Admin" else url_for("ict_dashboard"))+"#elections")

# -------------------------------------------------------------------
# Library
# -------------------------------------------------------------------
@app.route("/library/add", methods=["POST"])
@login_required
@role_required("Admin","ICT","Librarian")
def add_library_item():
    title=request.form.get("title","").strip(); category=request.form.get("category","Book").strip() or "Book"
    author=request.form.get("author","").strip(); code=request.form.get("item_code","").strip(); location=request.form.get("location","").strip()
    resource_type=request.form.get("resource_type","Physical").strip(); qty=max(1,request.form.get("quantity",type=int) or 1)
    description=request.form.get("description","").strip(); external_url=request.form.get("external_url","").strip()
    youtube_url=request.form.get("youtube_url","").strip(); source_url=request.form.get("source_url","").strip(); source_name=request.form.get("source_name","").strip()
    class_level=request.form.get("class_level","").strip(); subject=request.form.get("subject","").strip()
    if not title: flash("Library title is required.","danger"); return redirect(request.referrer or url_for("librarian_dashboard"))
    file=request.files.get("resource"); file_path=""
    if file and file.filename:
        fname=secure_filename(file.filename); dest=UPLOAD_DIR/"library"; dest.mkdir(exist_ok=True); out=dest/f"{uuid.uuid4().hex}-{fname}"; file.save(out); file_path="uploads/library/"+out.name; resource_type="Digital"
    image=request.files.get("resource_image"); image_path=""
    if image and image.filename:
        ext=image.filename.rsplit('.',1)[-1].lower() if '.' in image.filename else ''
        if ext not in {"png","jpg","jpeg","webp"}: flash("Resource image must be PNG/JPG/JPEG/WEBP.","danger"); return redirect(request.referrer or url_for("librarian_dashboard"))
        dest=UPLOAD_DIR/"library"; dest.mkdir(exist_ok=True); out=dest/f"{uuid.uuid4().hex}-cover.{ext}"; image.save(out); image_path="uploads/library/"+out.name
    execute("""INSERT INTO library_items(title,category,author,item_code,quantity,available_quantity,location,resource_type,file_path,external_url,description,created_by,class_level,subject,image_path,youtube_url,source_url,source_name) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(title,category,author,code,qty,qty,location,resource_type,file_path,external_url,description,current_user()["id"],class_level,subject,image_path,youtube_url,source_url,source_name))
    audit(current_user()["id"],current_user()["full_name"],"Library Resource Added",f"Added {title} for class '{class_level or 'All'}' and subject '{subject or 'General'}'.")
    flash("Library resource added to the institutional catalogue.","success")
    return redirect((url_for("ict_dashboard") if current_user()["role"]=="ICT" else url_for("librarian_dashboard") if current_user()["role"]=="Librarian" else url_for("admin_dashboard"))+"#library")

@app.route("/library/<int:item_id>/remove", methods=["POST"])
@login_required
@role_required("Admin","ICT","Librarian")
def remove_library_item(item_id):
    execute("UPDATE library_items SET active=0 WHERE id=?",(item_id,)); flash("Library item removed from the catalogue.","success"); return redirect(request.referrer or url_for("librarian_dashboard"))

@app.route("/library/loan", methods=["POST"])
@login_required
@role_required("Admin","ICT","Librarian")
def issue_library_item():
    item_id=request.form.get("item_id",type=int); student_id=request.form.get("student_id",type=int); due=request.form.get("due_date") or None
    item=q("SELECT * FROM library_items WHERE id=? AND active=1",(item_id,),one=True); student=q("SELECT * FROM students WHERE id=? AND active=1",(student_id,),one=True)
    if not item or not student or item["available_quantity"]<1:
        flash("Item unavailable or student invalid.","danger"); return redirect(request.referrer or url_for("librarian_dashboard"))
    execute("INSERT INTO library_loans(item_id,student_id,issued_by,due_date) VALUES(?,?,?,?)",(item_id,student_id,current_user()["id"],due))
    execute("UPDATE library_items SET available_quantity=available_quantity-1 WHERE id=?",(item_id,)); flash("Library item issued.","success"); return redirect(request.referrer or url_for("librarian_dashboard"))

@app.route("/library/loan/<int:loan_id>/return", methods=["POST"])
@login_required
@role_required("Admin","ICT","Librarian")
def return_library_item(loan_id):
    loan=q("SELECT * FROM library_loans WHERE id=? AND status='Issued'",(loan_id,),one=True)
    if not loan: abort(404)
    execute("UPDATE library_loans SET returned_at=CURRENT_TIMESTAMP,status='Returned' WHERE id=?",(loan_id,)); execute("UPDATE library_items SET available_quantity=available_quantity+1 WHERE id=?",(loan["item_id"],)); flash("Library item returned.","success"); return redirect(request.referrer or url_for("librarian_dashboard"))

@app.route("/messages/send", methods=["POST"])
@login_required
def send_message():
    user=current_user(); recipient_role=request.form.get("recipient_role", "Teacher").strip()
    student_id=request.form.get("recipient_student_id", type=int)
    body=request.form.get("body", "").strip()
    allowed_pairs={("Teacher","Parent"),("Parent","Teacher")}
    if not body or (user["role"], recipient_role) not in allowed_pairs:
        flash("That messaging route is not available for this account.", "danger")
        return redirect(request.referrer or url_for("index"))
    student=q("SELECT id FROM students WHERE id=? AND active=1", (student_id,), one=True) if student_id else None
    if not student:
        flash("Select a valid pupil for this message.", "danger")
        return redirect(request.referrer or url_for("index"))
    if user["role"] == "Parent" and not can_access_student(student_id):
        abort(403)
    execute("INSERT INTO portal_messages(sender_role,sender_name,recipient_role,recipient_student_id,body) VALUES(?,?,?,?,?)", (user["role"], user["full_name"], recipient_role, student_id, body))
    flash("Message sent.", "success")
    return redirect(request.referrer or url_for("index"))


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    user = current_user()
    if request.method == "POST":
        full_name = request.form.get("full_name", "").strip()
        password = request.form.get("password", "")
        photo = request.files.get("profile_photo")
        if not full_name:
            flash("Full name is required.", "danger")
        else:
            if password:
                if len(password) < 4:
                    flash("Password must be at least 4 characters.", "danger")
                    return redirect(url_for("profile"))
                execute("UPDATE users SET full_name = ?, password_hash = ? WHERE id = ?", (full_name, generate_password_hash(password), user["id"]))
            else:
                execute("UPDATE users SET full_name = ? WHERE id = ?", (full_name, user["id"]))
            if photo and photo.filename:
                ext=photo.filename.rsplit('.',1)[-1].lower() if '.' in photo.filename else ''
                if ext not in {'png','jpg','jpeg','webp'}:
                    flash("Profile photo must be PNG, JPG, JPEG or WEBP.", "danger")
                    return redirect(url_for("profile"))
                folder=UPLOAD_DIR/"profile_photos"; folder.mkdir(exist_ok=True)
                out=folder/f"user-{user['id']}-{uuid.uuid4().hex[:12]}.{ext}"
                photo.save(out)
                execute("UPDATE users SET profile_photo=? WHERE id=?", ("uploads/profile_photos/"+out.name,user["id"]))
            session["user_id"] = user["id"]
            audit(user["id"], full_name, "Profile Update", "Profile details, password and/or profile photo updated.")
            flash("Profile updated successfully.", "success")
            return redirect(url_for("profile"))
    profile_user = q("SELECT id, full_name, username, role, created_at, profile_photo, student_id FROM users WHERE id = ?", (user["id"],), one=True)
    exam_card = q("SELECT id,qr_token,created_at FROM portal_documents WHERE document_type='Exam Card' AND student_id=? AND status='Valid' ORDER BY created_at DESC,id DESC LIMIT 1", (profile_user["student_id"],), one=True) if profile_user and profile_user["student_id"] else None
    return render_template("profile.html", profile_user=profile_user, workspace=workspace_for(profile_user["role"]), exam_card=exam_card)


@app.route("/settings", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def save_settings():
    school_name = request.form.get("school_name", "").strip() or "School"
    admission_prefix = request.form.get("admission_prefix", "").strip() or "ADM-"
    admission_suffix = request.form.get("admission_suffix", "").strip()
    student_name_prefix = request.form.get("student_name_prefix", "").strip()
    student_name_suffix = request.form.get("student_name_suffix", "").strip()
    currency_code = request.form.get("currency_code", "").strip() or "KES"
    try:
        school_fee = float(request.form.get("school_fee", "0") or 0)
    except ValueError:
        school_fee = 0.0
    student_email_domain = (request.form.get("student_email_domain") or "school.ac.ke").strip().lstrip("@").lower() or "school.ac.ke"
    public_location_title = (request.form.get("public_location_title") or "Visit us today at this location").strip()
    public_address = (request.form.get("public_address") or "").strip()
    public_location_notes = (request.form.get("public_location_notes") or "").strip()
    public_map_query = (request.form.get("public_map_query") or "").strip()
    execute(
        """
        UPDATE school_settings
        SET school_name = ?, admission_prefix = ?, admission_suffix = ?, student_name_prefix = ?, student_name_suffix = ?, currency_code = ?, school_fee = ?, student_email_domain = ?, public_location_title = ?, public_address = ?, public_location_notes = ?, public_map_query = ?
        WHERE id = 1
        """,
        (school_name, admission_prefix, admission_suffix, student_name_prefix, student_name_suffix, currency_code, school_fee, student_email_domain, public_location_title, public_address, public_location_notes, public_map_query),
    )
    audit(current_user()["id"], current_user()["full_name"], "Update Settings", f"School settings updated for {school_name}.")
    flash("School settings updated.", "success")
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/students/add", methods=["GET", "POST"])
@login_required
@role_required("Admin")
def add_student():
    """Single, administrator-only learner intake path.

    Only the identity fields are required. Everything else is optional enrichment;
    a missing optional field can never stop the learner from being created.
    """
    if request.method == "GET":
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    actor = current_user()
    admission_no = (request.form.get("admission_no") or "").strip()
    full_name = (request.form.get("full_name") or "").strip()
    grade = normalize_grade(request.form.get("grade") or "")
    age = (request.form.get("age") or "").strip()
    fields = {name: (request.form.get(name) or "").strip() for name in (
        "guardian_name","guardian_phone","guardian_email","alt_guardian_name","alt_guardian_phone","alt_guardian_email",
        "student_phone","student_email","medical_condition","allergies","special_info","notes",
        "address","date_of_birth","gender","id_reference","emergency_contact","blood_group","medical_notes")}
    transport_zone=(request.form.get('transport_zone') or '').strip()
    uses_bus=1 if request.form.get('uses_school_bus') in {'1','on','yes'} else 0
    meal_plan=(request.form.get('meal_plan') or 'None').strip() or 'None'
    fee_override=request.form.get('fee_override_total','').strip()
    fee_override_enabled=bool(fee_override)

    if not full_name or not grade:
        missing=[]
        if not full_name: missing.append("full name")
        if not grade: missing.append("grade / class")
        flash("Please add " + " and ".join(missing) + ". Everything else can be left blank.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    if fee_override:
        try:
            manual_fee=float(fee_override)
            if manual_fee < 0:
                raise ValueError
        except (TypeError, ValueError):
            flash("The exact fee total must be a valid amount of 0 or more. Example: 18500 or 18500.00. The student was not submitted.", "danger")
            return redirect(url_for("admin_dashboard") + "#admin-add-student")

    if uses_bus and not transport_zone:
        flash("School bus is set to Yes, but no transport zone was selected. Choose a transport zone or change School bus to No. The student was not submitted.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    try:
        meal_charge_value=float(request.form.get('meal_charge','0') or 0)
        if meal_charge_value < 0:
            raise ValueError
    except (TypeError, ValueError):
        flash("Meal charge must be a valid amount of 0 or more, or leave it at 0 when meals are not being charged. The student was not submitted.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    settings = school_settings()
    if not admission_no:
        admission_no = next_admission_no()
    if settings["student_name_prefix"]:
        full_name = f"{settings['student_name_prefix']} {full_name}".strip()
    if settings["student_name_suffix"]:
        full_name = f"{full_name} {settings['student_name_suffix']}".strip()

    if q("SELECT id FROM students WHERE admission_no=?", (admission_no,), one=True):
        flash("That admission number already exists. Open the existing learner instead of creating a duplicate.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    fee = float(settings["school_fee"] or 0)
    manual_fee=float(fee_override) if fee_override else None
    if manual_fee is not None:
        fee=manual_fee
    transport_charge=0.0
    if uses_bus and transport_zone:
        tr=q('SELECT amount FROM transport_rates WHERE zone_name=? AND active=1 LIMIT 1',(transport_zone,),one=True)
        transport_charge=float(tr['amount'] or 0) if tr else 0.0
    meal_charge=meal_charge_value if meal_plan!='None' else 0.0
    if fee_override_enabled:
        assessed=fee
    else:
        assessed=fee + transport_charge + meal_charge
    generated_email=student_email_for(admission_no, fields['student_email'])
    student_id = None
    warnings = []

    # The learner itself is authoritative. Build the INSERT from columns that are
    # actually present in the deployed database so an older schema cannot block intake.
    learner_values = {
        "admission_no": admission_no, "full_name": full_name, "grade": grade, "age": age,
        "guardian_name": fields["guardian_name"], "guardian_phone": fields["guardian_phone"],
        "guardian_email": fields["guardian_email"], "alt_guardian_name": fields["alt_guardian_name"],
        "alt_guardian_phone": fields["alt_guardian_phone"], "alt_guardian_email": fields["alt_guardian_email"],
        "student_phone": fields["student_phone"], "student_email": generated_email,
        "address": fields["address"], "date_of_birth": fields["date_of_birth"], "gender": fields["gender"],
        "id_reference": fields["id_reference"], "emergency_contact": fields["emergency_contact"],
        "blood_group": fields["blood_group"], "medical_notes": fields["medical_notes"],
        "medical_condition": fields["medical_condition"], "allergies": fields["allergies"],
        "special_info": fields["special_info"], "notes": fields["notes"],
        "payment_status": "Pending", "balance": assessed, "fee_assessed_total": assessed,
        "fee_override_enabled": 1 if fee_override_enabled else 0, "transport_zone": transport_zone,
        "uses_school_bus": uses_bus, "meal_plan": meal_plan, "transport_charge": transport_charge, "active": 1,
    }

    try:
        available_columns = table_columns(get_db(), "students")
        insert_columns = [c for c in learner_values if c in available_columns]
        if not {"admission_no", "full_name", "grade"}.issubset(insert_columns):
            raise sqlite3.OperationalError("The deployed students table is missing a required identity column. Please run the database migration.")
        placeholders = ",".join("?" for _ in insert_columns)
        student_id = execute(
            f"INSERT INTO students({','.join(insert_columns)}) VALUES({placeholders})",
            tuple(learner_values[c] for c in insert_columns),
        )
    except sqlite3.IntegrityError as exc:
        app.logger.exception("Student registration rejected at learner insert: %s", exc)
        message = str(exc).lower()
        if "unique" in message or "admission_no" in message:
            flash("That admission number already exists. Choose a different number, or leave Admission No. blank so the system can generate one.", "danger")
        else:
            flash("The learner could not be saved because the database rejected one of the submitted fields. Check the learner details and try again.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")
    except Exception as exc:
        app.logger.exception("Student registration failed at learner insert: %s", exc)
        flash("The learner could not be saved. Please check the full name, class and admission number, then submit again.", "danger")
        return redirect(url_for("admin_dashboard") + "#admin-add-student")

    # Everything below this point is optional enrichment. A failure here must not
    # undo the already-created learner.
    if not fee_override_enabled:
        try:
            if fee > 0:
                execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(student_id,None,fee,'Tuition / core school fee',actor['id']))
            if transport_charge > 0:
                execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(student_id,None,transport_charge,'School transport — '+transport_zone,actor['id']))
            if meal_charge > 0:
                execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(student_id,None,meal_charge,'Meal plan — '+meal_plan,actor['id']))
        except Exception as exc:
            warnings.append("fee breakdown could not be posted automatically")
            app.logger.exception("Student %s saved but fee charge setup failed: %s", student_id, exc)

    try:
        username = admission_no.lower()
        execute("INSERT OR IGNORE INTO users(full_name,username,password_hash,role,student_id,active,workspace_type,title,email) VALUES(?,?,?,?,?,1,?,?,?)",
                (full_name,username,generate_password_hash(admission_no),"Student",student_id,"Student","Student",generated_email))
    except Exception as exc:
        warnings.append("student portal account setup skipped")
        app.logger.warning("Student %s saved but portal account setup was skipped: %s", student_id, exc)

    try:
        placement = auto_place_new_student(student_id, grade, actor["id"])
    except Exception as exc:
        placement = {"class_teacher": None, "subjects": 0, "subject_teachers": 0}
        warnings.append("automatic class/subject allocation skipped")
        app.logger.warning("Student %s saved but automatic placement was skipped: %s", student_id, exc)

    try:
        audit(actor["id"], actor["full_name"], "Add Student",
              f"{full_name} ({admission_no}) created in school.db; grade={grade}; class_teacher={placement['class_teacher'] or 'pending'}; subjects={placement['subjects']}; subject_teachers={placement['subject_teachers']}.")
    except Exception as exc:
        app.logger.warning("Student %s saved but audit logging failed: %s", student_id, exc)

    if warnings:
        flash("Student added successfully. The learner record is saved; " + "; ".join(warnings) + ".", "warning")
    else:
        flash("Student added successfully. The learner record is ready.", "success")
    return redirect(url_for("admin_dashboard") + "#admin-add-student")

@app.route("/students/<int:student_id>/subjects", methods=["GET","POST"])
@login_required
def student_subjects_manage(student_id:int):
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True)
    if not student: abort(404)
    user=current_user(); role=user["role"]
    allowed_editor=role in {"Admin","ICT"} or is_reception_user(user)
    is_self=(role=="Student" and user["student_id"]==student_id)
    if not allowed_editor and not is_self: abort(403)
    subjects=q("SELECT * FROM subjects_catalog WHERE active=1 ORDER BY department,subject")
    departments=q("SELECT id,name,category FROM departments WHERE active=1 ORDER BY name")
    current_departments=q("SELECT department_id,status FROM student_departments WHERE student_id=? AND status!='Dropped'",(student_id,))
    current_department_ids={r['department_id'] for r in current_departments}
    current=q("SELECT ss.id,ss.subject_id,ss.status,sc.subject,sc.department FROM student_subjects ss JOIN subjects_catalog sc ON sc.id=ss.subject_id WHERE ss.student_id=? AND ss.status!='Dropped' ORDER BY sc.subject",(student_id,))
    current_ids={r['subject_id'] for r in current}
    if request.method=='POST':
        status='Approved' if allowed_editor else 'Pending'
        selected_departments=[]
        for value in request.form.getlist('department_ids'):
            try: selected_departments.append(int(value))
            except Exception: pass
        selected_departments=set(selected_departments)
        for did in selected_departments:
            if did in {x['id'] for x in departments}:
                execute("INSERT INTO student_departments(student_id,department_id,status,selected_by) VALUES(?,?,?,?) ON CONFLICT(student_id,department_id) DO UPDATE SET status=excluded.status,selected_by=excluded.selected_by,updated_at=CURRENT_TIMESTAMP",(student_id,did,status,user['id']))
        if selected_departments:
            placeholders=','.join('?'*len(selected_departments)); execute(f"UPDATE student_departments SET status='Dropped',updated_at=CURRENT_TIMESTAMP WHERE student_id=? AND department_id NOT IN ({placeholders})",(student_id,*selected_departments))
        else:
            execute("UPDATE student_departments SET status='Dropped',updated_at=CURRENT_TIMESTAMP WHERE student_id=?",(student_id,))
        selected=[]
        for value in request.form.getlist('subject_ids'):
            try: selected.append(int(value))
            except Exception: pass
        selected=set(selected)
        for sid in selected:
            if sid in {x['id'] for x in subjects}:
                execute("INSERT INTO student_subjects(student_id,subject_id,status,selected_by) VALUES(?,?,?,?) ON CONFLICT(student_id,subject_id) DO UPDATE SET status=excluded.status,selected_by=excluded.selected_by,updated_at=CURRENT_TIMESTAMP",(student_id,sid,status,user['id']))
        if selected:
            placeholders=','.join('?'*len(selected))
            params=(student_id,*selected)
            execute(f"UPDATE student_subjects SET status='Dropped',updated_at=CURRENT_TIMESTAMP WHERE student_id=? AND subject_id NOT IN ({placeholders})",params)
        else:
            execute("UPDATE student_subjects SET status='Dropped',updated_at=CURRENT_TIMESTAMP WHERE student_id=?",(student_id,))
        if is_self:
            flash("Subject selection submitted for approval.","success")
        else:
            flash("Student subject enrolment updated.","success")
        return redirect(url_for('student_subjects_manage',student_id=student_id))
    return render_template('student_subjects.html',settings=school_settings(),student=student,subjects=subjects,current=current,current_ids=current_ids,departments=departments,current_department_ids=current_department_ids,role=role,actor_name=user['full_name'],self_registration=is_self)

@app.route("/student/change-password", methods=["GET", "POST"])
@login_required
@role_required("Student")
def student_change_password():
    user=current_user()
    stored=q("SELECT password_hash,full_name,username,role FROM users WHERE id=?", (user["id"],), one=True)
    error=None
    if request.method=="POST":
        current=request.form.get("current_password", "")
        new=request.form.get("password", "")
        confirm=request.form.get("confirm_password", "")
        try: current_ok=check_password_hash(stored["password_hash"], current)
        except Exception: current_ok=False
        if not current_ok: error="Current password is incorrect."
        elif len(new)<8: error="Choose a password with at least 8 characters."
        elif new!=confirm: error="The two new passwords do not match."
        else:
            execute("UPDATE users SET password_hash=? WHERE id=?", (generate_password_hash(new), user["id"]))
            audit(user["id"], user["full_name"], "Student Password Changed", "Student changed their portal password.")
            flash("Password changed successfully. Your admission number is no longer your password.", "success")
            return redirect(url_for("student_dashboard"))
    return render_template("reset_password.html", settings=school_settings(), invalid=False, token="", user=stored, error=error, completed=False, student_self_change=True)

@app.route("/students/<int:student_id>")
@login_required
def student_profile(student_id:int):
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True)
    if not student: abort(404)
    user=current_user(); role=user["role"]
    if role in {"Student","Parent"} and not can_access_student(student_id): abort(403)
    if role not in {"Admin","ICT","Teacher","Finance","Librarian","Student","Parent"}: abort(403)
    guardians=q("""SELECT gu.id,gu.full_name,gu.username,gu.phone,gu.email,gl.relationship,gl.is_primary FROM guardian_links gl JOIN users gu ON gu.id=gl.guardian_user_id WHERE gl.student_id=? AND gl.active=1 ORDER BY gl.is_primary DESC,gu.full_name""",(student_id,))
    payments=q("SELECT p.*,u.full_name AS recorder FROM payments p LEFT JOIN users u ON u.id=p.recorded_by WHERE p.student_id=? ORDER BY p.created_at DESC,p.id DESC",(student_id,))
    records=q("SELECT r.*,u.full_name AS author,u.role AS author_role FROM student_records r JOIN users u ON u.id=r.author_user_id WHERE r.student_id=? AND (r.visible_to_parent=1 OR ? IN ('Admin','Teacher','ICT') OR r.author_user_id=?) ORDER BY r.created_at DESC,r.id DESC",(student_id,role,user["id"]))
    results=q("SELECT * FROM exam_results WHERE student_id=? ORDER BY term DESC,subject",(student_id,))
    student_departments=q("SELECT d.name,d.category,sd.status FROM student_departments sd JOIN departments d ON d.id=sd.department_id WHERE sd.student_id=? AND sd.status!='Dropped' ORDER BY d.name",(student_id,))
    subjects=q("SELECT sc.subject,sc.department,ss.status FROM student_subjects ss JOIN subjects_catalog sc ON sc.id=ss.subject_id WHERE ss.student_id=? AND ss.status!='Dropped' ORDER BY sc.department,sc.subject",(student_id,))
    class_attendance=q("SELECT ca.*,u.full_name AS teacher_name FROM class_attendance ca JOIN users u ON u.id=ca.teacher_user_id WHERE ca.student_id=? ORDER BY ca.attendance_date DESC,ca.id DESC LIMIT 120",(student_id,))
    awards=[r for r in records if r['category']=='Award']; discipline=[r for r in records if r['category']=='Indiscipline']
    return render_template('student_profile.html',student=student,guardians=guardians,payments=payments,records=records,results=results,class_attendance=class_attendance,awards=awards,discipline=discipline,settings=school_settings(),actor_name=user['full_name'],role=role,can_edit=role in {'Admin','ICT'},can_write_record=role in {'Admin','Teacher'},parent_visible_count=sum(1 for r in records if r['visible_to_parent']),subjects=subjects,student_departments=student_departments)

@app.route("/students/<int:student_id>/record", methods=["POST"])
@login_required
@role_required("Admin","Teacher")
def add_student_record(student_id:int):
    student=q("SELECT id,full_name FROM students WHERE id=?",(student_id,),one=True)
    if not student: abort(404)
    category=request.form.get('category','General').strip() or 'General'
    title=request.form.get('title','').strip(); content=request.form.get('content','').strip()
    if not title or not content:
        flash('A record title and details are required.','danger'); return redirect(url_for('student_profile',student_id=student_id))
    visible=1 if request.form.get('visible_to_parent')=='1' and current_user()['role'] in {'Admin','Teacher'} else 0
    execute("INSERT INTO student_records(student_id,author_user_id,category,title,content,visible_to_parent) VALUES(?,?,?,?,?,?)",(student_id,current_user()['id'],category,title,content,visible))
    audit(current_user()['id'],current_user()['full_name'],'Student Record',f"Added {category} record for {student['full_name']}; parent-visible={visible}.")
    flash('Student record added.','success'); return redirect(url_for('student_profile',student_id=student_id))

@app.route("/students/<int:student_id>/record/<int:record_id>/visibility", methods=["POST"])
@login_required
@role_required("Admin","Teacher")
def toggle_student_record_visibility(student_id:int,record_id:int):
    row=q("SELECT * FROM student_records WHERE id=? AND student_id=?",(record_id,student_id),one=True)
    if not row: abort(404)
    visible=1 if request.form.get('visible_to_parent')=='1' else 0
    execute("UPDATE student_records SET visible_to_parent=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",(visible,record_id))
    audit(current_user()['id'],current_user()['full_name'],'Student Record Visibility',f"Record #{record_id} visibility set to {visible}.")
    flash('Parent visibility updated.','success'); return redirect(url_for('student_profile',student_id=student_id))

@app.route("/students/search")
@login_required
@role_required("Admin","Teacher","ICT")
def student_search():
    term=request.args.get('q','').strip()
    rows=q("SELECT id,full_name,admission_no,grade,payment_status,balance,active FROM students WHERE full_name LIKE ? OR admission_no LIKE ? ORDER BY full_name LIMIT 100",(f'%{term}%',f'%{term}%')) if term else q("SELECT id,full_name,admission_no,grade,payment_status,balance,active FROM students WHERE active=1 ORDER BY full_name LIMIT 100")
    return render_template('student_search.html',rows=rows,term=term,settings=school_settings(),role=current_user()['role'],actor_name=current_user()['full_name'])

@app.route("/students/<int:student_id>/update", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def update_student(student_id: int):
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True)
    if not student: abort(404)
    def val(name, fallback=""):
        return request.form.get(name, fallback if fallback is not None else "").strip()
    full_name=val("full_name",student["full_name"]) or student["full_name"]
    admission_no=val("admission_no",student["admission_no"]) or student["admission_no"]
    grade=normalize_grade(val("grade",student["grade"]) or student["grade"])
    age=val("age",student["age"])
    fields={k:val(k,student[k] if k in student.keys() else "") for k in [
        "guardian_name","guardian_phone","guardian_email","alt_guardian_name","alt_guardian_phone","alt_guardian_email",
        "student_phone","student_email","address","date_of_birth","gender","id_reference","emergency_contact","blood_group",
        "medical_condition","medical_notes","allergies","special_info","notes"]}
    active=1 if request.form.get("active","1" if student["active"] else "0")=="1" else 0
    transport_zone=val("transport_zone",student["transport_zone"] if "transport_zone" in student.keys() else "")
    uses_bus=1 if request.form.get("uses_school_bus","1" if int(student["uses_school_bus"] or 0) else "0")=="1" else 0
    meal_plan=val("meal_plan",student["meal_plan"] if "meal_plan" in student.keys() else "None") or "None"
    try: meal_charge=max(0,float(request.form.get("meal_charge",0) or 0))
    except ValueError: meal_charge=float(student["meal_charge"] or 0) if "meal_charge" in student.keys() else 0.0
    try: fee_total=max(0,float(request.form.get("fee_assessed_total",student["fee_assessed_total"] or 0) or 0))
    except ValueError: fee_total=float(student["fee_assessed_total"] or 0)
    override=1 if request.form.get("fee_override_enabled","1" if int(student["fee_override_enabled"] or 0) else "0")=="1" else 0
    tr=q("SELECT amount FROM transport_rates WHERE zone_name=? AND active=1 LIMIT 1",(transport_zone,),one=True) if uses_bus and transport_zone else None
    transport_charge=float(tr["amount"] or 0) if tr else 0.0
    old_admission=student["admission_no"] or ""
    if not fields["student_email"]:
        fields["student_email"]=student_email_for(admission_no)
    try:
        execute("""UPDATE students SET admission_no=?,full_name=?,grade=?,age=?,grade_category=?,guardian_name=?,guardian_phone=?,guardian_email=?,alt_guardian_name=?,alt_guardian_phone=?,alt_guardian_email=?,student_phone=?,student_email=?,address=?,date_of_birth=?,gender=?,id_reference=?,emergency_contact=?,blood_group=?,medical_condition=?,medical_notes=?,allergies=?,special_info=?,notes=?,active=?,fee_assessed_total=?,fee_override_enabled=?,transport_zone=?,uses_school_bus=?,meal_plan=?,transport_charge=?,updated_at=CURRENT_TIMESTAMP WHERE id=?""",
            (admission_no,full_name,grade,age,grade,fields["guardian_name"],fields["guardian_phone"],fields["guardian_email"],fields["alt_guardian_name"],fields["alt_guardian_phone"],fields["alt_guardian_email"],fields["student_phone"],fields["student_email"],fields["address"],fields["date_of_birth"],fields["gender"],fields["id_reference"],fields["emergency_contact"],fields["blood_group"],fields["medical_condition"],fields["medical_notes"],fields["allergies"],fields["special_info"],fields["notes"],active,fee_total,override,transport_zone,uses_bus,meal_plan,transport_charge,student_id))
        if not override:
            # Replace the current auto-generated transport/meal charges, while leaving manually-added fee items intact.
            execute("DELETE FROM fee_charges WHERE student_id=? AND (description LIKE 'School transport — %' OR description LIKE 'Meal plan — %')",(student_id,))
            if transport_charge>0: execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(student_id,None,transport_charge,'School transport — '+transport_zone,current_user()["id"]))
            if meal_plan!="None" and meal_charge>0: execute("INSERT INTO fee_charges(student_id,fee_structure_id,amount,description,created_by) VALUES(?,?,?,?,?)",(student_id,None,meal_charge,'Meal plan — '+meal_plan,current_user()["id"]))
        else:
            # In override mode the supplied total is authoritative; no synthetic charges are needed.
            execute("UPDATE students SET fee_assessed_total=? WHERE id=?",(fee_total,student_id))
        recalculate_student_balance(student_id)
        linked=q("SELECT id,password_hash FROM users WHERE student_id=? AND role='Student' AND active=1 LIMIT 1",(student_id,),one=True)
        if linked:
            try: bootstrap=bool(old_admission) and check_password_hash(linked["password_hash"],old_admission)
            except Exception: bootstrap=False
            if bootstrap: execute("UPDATE users SET username=?,full_name=?,email=?,password_hash=? WHERE id=?",(admission_no,full_name,fields["student_email"],generate_password_hash(admission_no),linked["id"]))
            else: execute("UPDATE users SET full_name=?,email=? WHERE id=?",(full_name,fields["student_email"],linked["id"]))
        else:
            execute("INSERT OR IGNORE INTO users(full_name,username,password_hash,role,student_id,active,workspace_type,title,email) VALUES(?,?,?,?,?,1,?,?,?)",(full_name,admission_no,generate_password_hash(admission_no),"Student",student_id,"Student","Student",fields["student_email"]))
        audit(current_user()["id"],current_user()["full_name"],"Update Student",f"Updated {full_name} ({admission_no}); balance recalculated from charges/payments.")
        flash("Learner updated. Fees, payments and transport have been reconciled automatically.","success")
    except Exception as exc:
        app.logger.exception("Student update failed: %s",exc); flash("Learner could not be updated safely.","danger")
    return redirect(url_for("student_profile",student_id=student_id)+"#edit")

@app.route("/students/<int:student_id>/delete", methods=["POST"])
@login_required
@role_required("Admin")
def delete_student(student_id: int):
    student = q("SELECT * FROM students WHERE id = ?", (student_id,), one=True)
    if not student:
        abort(404)
    execute("DELETE FROM students WHERE id = ?", (student_id,))
    audit(current_user()["id"], current_user()["full_name"], "Delete Student", f"Student {student['admission_no']} deleted.")
    flash("Pupil deleted.", "success")
    return redirect(request.referrer or url_for("admin_dashboard"))


def notify_teachers_of_payment(student, amount, balance):
    teacher_ids={int(r["id"]) for r in q("SELECT DISTINCT teacher_user_id AS id FROM student_teacher_assignments WHERE student_id=? AND active=1",(student["id"],)) if r["id"]}
    teacher_ids|={int(r["id"]) for r in q("SELECT DISTINCT teacher_user_id AS id FROM class_teacher_assignments WHERE class_name=?",(student["grade"],)) if r["id"]}
    teacher_ids|={int(r["id"]) for r in q("SELECT DISTINCT teacher_user_id AS id FROM teacher_assignments WHERE class_name=? AND active=1",(student["grade"],)) if r["id"]}
    if not teacher_ids: return
    assessed=float(student["fee_assessed_total"] or 0)
    status="Fully paid" if balance<=0 else ("Partially paid" if assessed>0 and balance<assessed else "Unpaid")
    notify_users(sorted(teacher_ids),"Fee payment received",f"{student['full_name']} ({student['admission_no']}) paid KES {float(amount):,.2f}. Current status: {status}; balance KES {max(balance,0):,.2f}.",url_for("student_profile",student_id=student["id"]))

@app.route("/payments/add", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def add_payment():
    admission_no = request.form.get("admission_no", "").strip()
    amount = request.form.get("amount", "").strip()
    method = request.form.get("method", "").strip()
    reference_no = request.form.get("reference_no", "").strip()
    student = q("SELECT * FROM students WHERE admission_no = ?", (admission_no,), one=True)
    if not student:
        flash("Student admission number not found.", "danger")
        return redirect(request.referrer or url_for("dashboard"))
    try:
        amount_f = float(amount)
        if amount_f <= 0:
            raise ValueError
    except ValueError:
        flash("Enter a valid payment amount.", "danger")
        return redirect(request.referrer or url_for("dashboard"))
    if method not in {"Cash", "M-Pesa", "Bank", "Cheque"}:
        flash("Select a valid payment method.", "danger")
        return redirect(request.referrer or url_for("dashboard"))
    execute(
        "INSERT INTO payments(student_id, amount, method, reference_no, recorded_by, status) VALUES (?, ?, ?, ?, ?, 'Posted')",
        (student["id"], amount_f, method, reference_no, current_user()["id"]),
    )
    recalculate_student_balance(student["id"])
    new_balance=float(q("SELECT balance FROM students WHERE id=?",(student['id'],),one=True)['balance'] or 0)
    notify_teachers_of_payment(q("SELECT * FROM students WHERE id=?",(student['id'],),one=True), amount_f, new_balance)
    audit(current_user()["id"], current_user()["full_name"], "Record Payment", f"{amount_f:.2f} recorded for {student['admission_no']} using {method}.")
    flash("Payment recorded.", "success")
    return redirect(request.referrer or url_for("dashboard", student_id=student["id"]))


@app.route("/payments/<int:payment_id>/reverse", methods=["POST"])
@login_required
@role_required("Admin", "Finance")
def reverse_payment(payment_id: int):
    payment = q("SELECT * FROM payments WHERE id = ?", (payment_id,), one=True)
    if not payment:
        abort(404)
    if payment["status"] == "Reversed":
        flash("Payment is already reversed.", "warning")
        return redirect(request.referrer or url_for("admin_dashboard"))
    execute("UPDATE payments SET status = 'Reversed', reversed_at = CURRENT_TIMESTAMP WHERE id = ?", (payment_id,))
    student = q("SELECT * FROM students WHERE id = ?", (payment["student_id"],), one=True)
    if student:
        recalculate_student_balance(student["id"])
        updated_balance=float(q("SELECT balance FROM students WHERE id=?",(student['id'],),one=True)['balance'] or 0)
    audit(current_user()["id"], current_user()["full_name"], "Reverse Payment", f"Payment #{payment_id} reversed.")
    flash("Payment reversed.", "success")
    return redirect(request.referrer or url_for("admin_dashboard"))

@app.route("/finance/payment", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def finance_record_payment():
    student=q("SELECT * FROM students WHERE id=?",(request.form.get("student_id",type=int),),one=True)
    if not student: abort(404)
    try: amount=float(request.form.get("amount","0") or 0)
    except ValueError: amount=0
    if amount<=0: flash("Enter a valid payment amount.","danger"); return redirect(url_for("finance_dashboard"))
    method=request.form.get("method","M-Pesa").strip(); reference=request.form.get("reference_no","").strip(); receipt_path=""
    file=request.files.get("receipt")
    if file and file.filename:
        ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
        if ext not in {"pdf","png","jpg","jpeg","webp"}: flash("Receipt must be PDF or image.","danger"); return redirect(url_for("finance_dashboard"))
        receipt_dir=UPLOAD_DIR/"receipts"; receipt_dir.mkdir(exist_ok=True)
        name=f"receipt-{student['id']}-{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{ext}"; file.save(receipt_dir/name); receipt_path="uploads/receipts/"+name
    payment_id=execute("INSERT INTO payments(student_id,amount,method,reference_no,recorded_by,status,receipt_path) VALUES(?,?,?,?,?,'Posted',?)",(student['id'],amount,method,reference,current_user()['id'],receipt_path))
    recalculate_student_balance(student['id']); new_balance=float(q("SELECT balance FROM students WHERE id=?",(student['id'],),one=True)['balance'] or 0)
    notify_teachers_of_payment(q("SELECT * FROM students WHERE id=?",(student['id'],),one=True), amount, new_balance)
    audit(current_user()['id'],current_user()['full_name'],'Finance Payment',f"Payment #{payment_id}: {amount:.2f} for {student['admission_no']}; balance {new_balance:.2f}.")
    flash("Payment posted and the student/parent balance has been updated.","success"); return redirect(url_for("finance_dashboard"))

@app.route("/finance/submit-day", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def finance_submit_day():
    closing_date=request.form.get("closing_date", datetime.utcnow().strftime("%Y-%m-%d")).strip()
    notes=request.form.get("notes","").strip()
    execute("INSERT INTO finance_closings(closing_date,submitted_by,notes,status) VALUES(?,?,?,'Submitted')",(closing_date,current_user()["id"],notes))
    audit(current_user()["id"],current_user()["full_name"],"Finance Day Submitted",f"Finance submitted the day {closing_date} for Administrator review.")
    flash("Finance records submitted to Administrator for review.","success"); return redirect(url_for("finance_dashboard"))

@app.route("/finance/results/<int:batch_id>/approve", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def finance_approve_results(batch_id):
    batch=q("SELECT * FROM exam_batches WHERE id=?",(batch_id,),one=True)
    if not batch: abort(404)
    settings=school_settings(); limit=float(settings['result_download_balance_limit'] or 0)
    balances=q("SELECT s.full_name,s.balance FROM students s JOIN exam_results r ON r.student_id=s.id WHERE r.batch_id=? GROUP BY s.id",(batch_id,))
    over=[r for r in balances if float(r['balance'] or 0)>limit]; override=request.form.get('override')=='1' and current_user()['role']=='Admin'
    if over and not override:
        flash(f"Approval blocked: {len(over)} student(s) exceed the result-release balance limit of {settings['currency_code']} {limit:,.0f}.","danger"); return redirect(url_for('finance_dashboard'))
    note=request.form.get('finance_note','').strip() or ("Approved within balance policy." if not override else "Approved by Admin override.")
    execute("UPDATE exam_batches SET finance_status='Approved',finance_note=?,approved_by=?,approved_at=CURRENT_TIMESTAMP,updated_at=CURRENT_TIMESTAMP WHERE id=?",(note,current_user()['id'],batch_id)); audit(current_user()['id'],current_user()['full_name'],'Approve Results',f"Exam batch #{batch_id} approved. {note}"); flash("Results approved for release.","success"); return redirect(url_for('finance_dashboard'))

@app.route("/finance/results/<int:batch_id>/reject", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def finance_reject_results(batch_id):
    batch=q("SELECT id FROM exam_batches WHERE id=?",(batch_id,),one=True)
    if not batch: abort(404)
    note=request.form.get('finance_note','').strip() or "Finance did not approve this result batch."
    execute("UPDATE exam_batches SET finance_status='Rejected',finance_note=?,approved_by=?,approved_at=CURRENT_TIMESTAMP,updated_at=CURRENT_TIMESTAMP WHERE id=?",(note,current_user()['id'],batch_id)); audit(current_user()['id'],current_user()['full_name'],'Reject Results',f"Exam batch #{batch_id} rejected. {note}"); flash("Results marked not approved for release.","warning"); return redirect(url_for('finance_dashboard'))

@app.route("/finance/policy", methods=["POST"])
@login_required
@role_required("Finance", "Admin")
def finance_policy():
    try: limit=max(0,float(request.form.get('result_download_balance_limit','0') or 0))
    except ValueError: limit=500
    execute("UPDATE school_settings SET result_download_balance_limit=? WHERE id=1",(limit,)); audit(current_user()['id'],current_user()['full_name'],'Finance Policy',f"Result download balance limit set to {limit:.2f}."); flash("Result download threshold updated.","success"); return redirect(url_for('finance_dashboard'))

@app.route("/report-card/<int:student_id>/<int:batch_id>")
@login_required
def report_card(student_id:int,batch_id:int):
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True); batch=q("SELECT * FROM exam_batches WHERE id=?",(batch_id,),one=True)
    if not student or not batch: abort(404)
    role=current_user()['role']
    if role in {'Student','Parent'} and not can_access_student(student_id): abort(403)
    if role=='Parent' and not parent_portal_enabled(): abort(403)
    limit=float(school_settings()['result_download_balance_limit'] or 0)
    eligible=batch['finance_status']=='Approved' and float(student['balance'] or 0)<=limit
    if role in {'Student','Parent'} and not eligible:
        flash('This online report card is awaiting Finance clearance.','warning'); return redirect(request.referrer or url_for('dashboard'))
    results=q("SELECT subject,term,mark,max_mark FROM exam_results WHERE student_id=? AND batch_id=? ORDER BY subject",(student_id,batch_id))
    return render_template('report_card.html',student=student,batch=batch,results=results,eligible=eligible,settings=school_settings())

@app.route("/results/<int:student_id>/<int:batch_id>/download")
@login_required
def download_result(student_id,batch_id):
    student=q("SELECT * FROM students WHERE id=?",(student_id,),one=True); batch=q("SELECT * FROM exam_batches WHERE id=?",(batch_id,),one=True)
    if not student or not batch: abort(404)
    if current_user()['role'] not in {'Admin','Finance','Student','Parent'}: abort(403)
    if current_user()['role'] in {'Student','Parent'} and not can_access_student(student_id): abort(403)
    if batch['finance_status']!='Approved' or float(student['balance'] or 0)>float(school_settings()['result_download_balance_limit'] or 0):
        flash("This result is not available for download under the current finance release policy.","warning"); return redirect(request.referrer or url_for('student_dashboard',student_id=student_id))
    results=q("SELECT subject,mark,max_mark FROM exam_results WHERE student_id=? AND batch_id=? ORDER BY subject",(student_id,batch_id))
    doc=q("SELECT * FROM portal_documents WHERE document_type='Result' AND student_id=? AND batch_id=? AND status='Valid' ORDER BY id DESC LIMIT 1",(student_id,batch_id),one=True)
    if not doc:
        token=make_qr_token('Result',student_id); execute("INSERT INTO portal_documents(document_type,student_id,batch_id,qr_token,issued_by) VALUES(?,?,?,?,?)",('Result',student_id,batch_id,token,current_user()['id'])); doc=q("SELECT * FROM portal_documents WHERE qr_token=?",(token,),one=True)
    pdf=generate_result_pdf(student,batch,results,doc,create_qr_file(doc['qr_token']))
    return send_file(pdf,as_attachment=True,download_name=f"Result-{student['admission_no']}-{batch['term'].replace(' ','_')}.pdf",mimetype='application/pdf')

@app.route("/exam-cards/create", methods=["POST"])
@login_required
@role_required("Admin", "Finance")
def create_exam_cards():
    student=q("SELECT * FROM students WHERE id=?",(request.form.get('student_id',type=int),),one=True)
    if not student: abort(404)
    batch_name=request.form.get('batch_name','Current Examination').strip() or 'Current Examination'; uploaded=request.files.get('exam_card'); token=make_qr_token('Exam Card',student['id']); file_path=''
    if uploaded and uploaded.filename:
        ext=uploaded.filename.rsplit('.',1)[-1].lower() if '.' in uploaded.filename else ''
        if ext not in {'pdf','png','jpg','jpeg','webp'}: flash("Exam card must be PDF or image.","danger"); return redirect(url_for('finance_dashboard'))
        name=f"exam-card-{student['id']}-{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{ext}"; uploaded.save(UPLOAD_DIR/name); raw_path=UPLOAD_DIR/name
        file_path=embed_qr_in_document(raw_path,create_qr_file(token),token) or ('uploads/'+name)
    doc_id=execute("INSERT INTO portal_documents(document_type,student_id,qr_token,file_path,issued_by) VALUES(?,?,?,?,?)",('Exam Card',student['id'],token,file_path,current_user()['id']))
    audit(current_user()['id'],current_user()['full_name'],'Exam Card',f"Exam card prepared for {student['admission_no']}.")
    if uploaded and file_path: flash("Exam card uploaded and registered. The verification QR is associated with this official document record.","success"); return redirect(url_for('finance_dashboard'))
    flash("Exam card generated and registered.","success"); return redirect(url_for('download_exam_card',document_id=doc_id,batch_name=batch_name))

@app.route("/exam-cards/<int:document_id>/download")
@login_required
def download_exam_card(document_id):
    doc=q("SELECT d.*,s.full_name,s.admission_no,s.grade,s.balance,u.profile_photo FROM portal_documents d JOIN students s ON s.id=d.student_id LEFT JOIN users u ON u.student_id=s.id AND u.role='Student' AND u.active=1 WHERE d.id=? AND d.document_type='Exam Card'",(document_id,),one=True)
    if not doc: abort(404)
    if current_user()["role"] in {"Student", "Parent"} and not can_access_student(doc["student_id"]): abort(403)
    if doc['file_path']: return send_file(BASE_DIR/doc['file_path'],as_attachment=True,download_name=Path(doc['file_path']).name)
    pdf=generate_exam_card_pdf(doc,request.args.get('batch_name','Current Examination'),doc,create_qr_file(doc['qr_token']))
    return send_file(pdf,as_attachment=True,download_name=f"Exam-Card-{doc['admission_no']}.pdf",mimetype='application/pdf')

@app.route("/verify/<token>")
def verify_document(token):
    doc=q("SELECT d.*,s.full_name,s.admission_no,s.grade FROM portal_documents d JOIN students s ON s.id=d.student_id WHERE d.qr_token=?",(token,),one=True)
    if not doc: return render_template('error.html',message='Verification token not found.'),404
    status='AUTHENTIC' if doc['status']=='Valid' else 'REVOKED'
    return render_template('error.html',message=f"Document {status}: {doc['document_type']} — {doc['full_name']} ({doc['admission_no']}). Token: {doc['qr_token']}")

@app.route("/exams/submit", methods=["POST"])
@login_required
@role_required("Teacher")
def submit_exams():
    grade = request.form.get("grade", "").strip()
    term = request.form.get("term", "").strip() or "Term 1"
    payload_raw = request.form.get("exam_payload", "").strip()
    if not grade:
        flash("Choose a grade first.", "danger")
        return redirect(request.referrer or url_for("dashboard"))
    if not payload_raw:
        flash("Add at least one subject and a few marks before submitting.", "danger")
        return redirect(request.referrer or url_for("dashboard", exam_grade=grade))
    try:
        payload = json.loads(payload_raw)
    except json.JSONDecodeError:
        flash("Could not read the marks submission.", "danger")
        return redirect(request.referrer or url_for("dashboard", exam_grade=grade))

    subjects = [str(item).strip() for item in payload.get("subjects", []) if str(item).strip()]
    rows = payload.get("rows", [])
    if not subjects or not rows:
        flash("Add at least one subject and one pupil mark.", "danger")
        return redirect(request.referrer or url_for("dashboard", exam_grade=grade))

    valid_students = {row["id"] for row in q("SELECT id FROM students WHERE grade = ?", (grade,))}
    batch_id = execute(
        "INSERT INTO exam_batches(grade, term, submitted_by, status) VALUES (?, ?, ?, 'Submitted')",
        (grade, term, current_user()["id"]),
    )

    saved = 0
    for row in rows:
        try:
            student_id = int(row.get("student_id"))
        except (TypeError, ValueError):
            continue
        if student_id not in valid_students:
            continue
        student = q("SELECT id, full_name, admission_no, grade FROM students WHERE id = ?", (student_id,), one=True)
        if not student:
            continue
        marks = row.get("marks", {}) or {}
        for subject in subjects:
            raw_mark = marks.get(subject, "")
            if raw_mark in (None, ""):
                continue
            try:
                mark = float(raw_mark)
            except (TypeError, ValueError):
                continue
            mark = max(0.0, min(100.0, mark))
            execute(
                """
                INSERT INTO exam_results(
                    batch_id, grade, term, subject, student_id, student_name, admission_no, mark, max_mark, submitted_by
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 100, ?)
                """,
                (batch_id, grade, term, subject, student_id, student["full_name"], student["admission_no"], mark, current_user()["id"]),
            )
            saved += 1

    audit(current_user()["id"], current_user()["full_name"], "Submit Exam Results", f"{grade} {term} submitted with {saved} mark entries.")
    flash(f"Exam marks submitted for {grade} {term}.", "success")
    return redirect(url_for("dashboard", exam_grade=grade, exam_term=term))


@app.route("/admin/institution-profile", methods=["POST"])
@login_required
@role_required("Admin")
def admin_institution_profile():
    institution_type=request.form.get("institution_type","Secondary School")
    if institution_type not in INSTITUTION_TYPES: institution_type="Secondary School"
    learner_label=request.form.get("learner_label","Student").strip() or "Student"
    staff_label=request.form.get("staff_label","Teacher").strip() or "Teacher"
    academic_period_label=request.form.get("academic_period_label","Term").strip() or "Term"
    class_label=request.form.get("class_label","Class").strip() or "Class"
    department_label=request.form.get("department_label","Department").strip() or "Department"
    parent_enabled = 0 if institution_type in {"TVET", "College", "University"} else (1 if request.form.get('parent_portal_enabled') == '1' else 0)
    type_orders={
        "Kindergarten":"Home,Students,Assignments,Messages,Library,Institution",
        "Primary School":"Home,Students,Assignments,Results,Messages,Library,Institution",
        "High School":"Home,Students,Classes,Assignments,Results,Messages,Library,Institution",
        "Secondary School":"Home,Students,Classes,Assignments,Results,Messages,Library,Institution",
        "TVET":"Home,Students,Departments,Courses,Assessments,Results,Finance,Library,Messages,Institution",
        "College":"Home,Students,Departments,Courses,Assessments,Results,Finance,Library,Messages,Institution",
        "University":"Home,Students,Faculties,Departments,Courses,Assessments,Results,Finance,Library,Messages,Institution",
        "Mixed Institution":"Home,Students,Classes,Departments,Assignments,Results,Finance,Library,Messages,Institution",
    }
    order=type_orders.get(institution_type,type_orders["High School"])
    default_staff="Teacher / Lecturer" if institution_type in {"TVET","College","University"} else "Teacher"
    default_period="Semester" if institution_type in {"TVET","College","University"} else "Term"
    default_class="Cohort" if institution_type in {"TVET","College","University"} else "Class"
    execute("""UPDATE school_settings SET institution_type=?, learner_label=?, staff_label=?, academic_period_label=?, class_label=?, department_label=?, parent_portal_enabled=?, menu_order=? WHERE id=1""",(institution_type,learner_label,staff_label or default_staff,academic_period_label or default_period,class_label or default_class,department_label,parent_enabled,order))
    if institution_type in {"TVET","College","University"}:
        for name in DEFAULT_DEPARTMENTS:
            execute("INSERT OR IGNORE INTO departments(name,category) VALUES(?, 'Academic')",(name,))
    audit(current_user()["id"],current_user()["full_name"],"Institution Profile","Updated institution type and terminology.")
    flash("Institution profile and terminology saved.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/admin/public-settings", methods=["POST"])
@login_required
@role_required("Admin")
def admin_public_settings():
    keys=["institution_history","institution_performance","institution_religion","institution_affiliations","institution_help","institution_contact","institution_owners","developer_name","developer_about","company_name","company_about"]
    values={k:request.form.get(k,"").strip() for k in keys}
    selected=[k for k in ["institution","history","achievements","owners","developer","company"] if request.form.get(f"show_{k}")]
    execute("""UPDATE school_settings SET institution_history=?, institution_performance=?, institution_religion=?, institution_affiliations=?, institution_help=?, institution_contact=?, institution_owners=?, developer_name=?, developer_about=?, company_name=?, company_about=?, prelogin_sections=?, landing_hero_title=?, landing_hero_text=?, landing_cta_primary=?, landing_cta_secondary=?, landing_announcement=?, landing_contact=?, landing_show_dates=?, landing_show_gallery=?, landing_show_roles=? WHERE id=1""", (values["institution_history"],values["institution_performance"],values["institution_religion"],values["institution_affiliations"],values["institution_help"],values["institution_contact"],values["institution_owners"],values["developer_name"],values["developer_about"],values["company_name"],values["company_about"],",".join(selected), request.form.get("landing_hero_title","").strip()[:240], request.form.get("landing_hero_text","").strip()[:2000], request.form.get("landing_cta_primary","").strip()[:80] or "Sign in to your workspace", request.form.get("landing_cta_secondary","").strip()[:80] or "View school information", request.form.get("landing_announcement","").strip()[:500], request.form.get("landing_contact","").strip()[:500], 1 if request.form.get("landing_show_dates") else 0, 1 if request.form.get("landing_show_gallery") else 0, 1 if request.form.get("landing_show_roles") else 0))
    audit(current_user()["id"], current_user()["full_name"], "Public Information", "Pre-login information sections updated.")
    flash("Public information settings saved.", "success")
    return redirect(url_for("admin_dashboard"))

@app.route("/admin/login-access", methods=["POST"])
@login_required
@role_required("Admin")
def admin_login_access():
    actor = current_user()
    desired = request.form.get("desired_state") == "1"
    confirmed = request.form.get("confirm_login_access") == "yes"
    password = request.form.get("admin_password", "")
    if not confirmed:
        flash("Please confirm the login-access change.", "warning")
        return redirect(url_for("admin_dashboard"))
    full = q("SELECT password_hash FROM users WHERE id=? AND role='Admin' AND active=1", (actor["id"],), one=True)
    if not full or not check_password_hash(full["password_hash"], password):
        flash("Administrator password is incorrect. Login-access setting was not changed.", "danger")
        return redirect(url_for("admin_dashboard"))
    execute("UPDATE school_settings SET auth_required=? WHERE id=1", (1 if desired else 0,))
    state = "enabled" if desired else "disabled"
    audit(actor["id"], actor["full_name"], "Login Access", f"Administrator {state} the login page for non-Administrator users.")
    flash(f"Login page is now {state}.", "success")
    return redirect(url_for("admin_dashboard"))

@app.route("/admin/parent/add", methods=["POST"])
@login_required
@role_required("Admin")
def add_parent_account():
    parent_name=request.form.get("parent_full_name", "").strip()
    parent_email=request.form.get("parent_email", "").strip()
    parent_phone=request.form.get("parent_phone", "").strip()
    password=request.form.get("parent_password", "")
    relationship=request.form.get("relationship", "Parent/Guardian").strip() or "Parent/Guardian"
    enable_portal=bool(request.form.get("enable_parent_portal"))
    student_ids=[]
    for raw in request.form.getlist("student_ids"):
        try: student_ids.append(int(raw))
        except Exception: pass
    student_ids=list(dict.fromkeys(student_ids))
    students=q(f"SELECT id,full_name,admission_no,grade FROM students WHERE active=1 AND id IN ({','.join('?'*len(student_ids))}) ORDER BY grade,full_name", tuple(student_ids)) if student_ids else []
    if not parent_name or not students:
        flash("Enter the parent/guardian name and select at least one existing student.", "danger")
        return redirect(url_for("admin_dashboard")+"#admin-add-student")
    if enable_portal and len(password) < 4:
        flash("A parent portal password must be at least 4 characters, or leave portal access disabled.", "danger")
        return redirect(url_for("admin_dashboard")+"#admin-add-student")
    try:
        if enable_portal:
            internal_username=f"parent-{uuid.uuid4().hex[:12]}"
            uid=execute("""INSERT INTO users(full_name,username,password_hash,role,student_id,active,phone,email,title,workspace_type)
                           VALUES(?,?,?,?,?,1,?,?,?,?)""",
                        (parent_name,internal_username,generate_password_hash(password),"Parent",students[0]["id"],parent_phone,parent_email,"Parent / Guardian","Parent"))
            for st in students:
                execute("INSERT OR IGNORE INTO guardian_links(guardian_user_id,student_id,relationship,is_primary) VALUES(?,?,?,?)",(uid,st["id"],relationship,1 if st["id"]==students[0]["id"] else 0))
                execute("UPDATE students SET guardian_name=COALESCE(NULLIF(guardian_name,''),?), guardian_phone=COALESCE(NULLIF(guardian_phone,''),?), guardian_email=COALESCE(NULLIF(guardian_email,''),?) WHERE id=?",(parent_name,parent_phone,parent_email,st["id"]))
            audit(current_user()["id"],current_user()["full_name"],"Add Parent",f"Parent {parent_name} linked to {len(students)} learner(s). Portal enabled.")
            flash(f"Parent account created and linked to {len(students)} student(s). Parent login uses the registered full name and password.","success")
        else:
            # Keep contact details on the selected learners without exposing a portal account.
            for st in students:
                execute("UPDATE students SET guardian_name=?,guardian_phone=?,guardian_email=? WHERE id=?",(parent_name,parent_phone,parent_email,st["id"]))
            audit(current_user()["id"],current_user()["full_name"],"Add Parent Contact",f"Parent contact {parent_name} linked to {len(students)} learner(s); portal disabled.")
            flash(f"Parent contact saved for {len(students)} student(s). Portal access was left disabled.","success")
    except sqlite3.IntegrityError as exc:
        flash(f"Parent could not be saved: {exc}","danger")
    return redirect(request.referrer or (url_for("admin_dashboard")+"#admin-add-student"))

@app.route("/users/add", methods=["GET", "POST"])
@login_required
def add_user():
    actor=current_user()
    if actor["role"] not in {"Admin","ICT"}: abort(403)
    # GET is an entry point retained for older Staff directory links.
    # Keep the actual form in the role dashboard so permissions and fields
    # stay centralized, but never expose a dead /users/add URL.
    if request.method == "GET":
        return redirect(url_for("all_employees", add=1))
    full_name=request.form.get("full_name", "").strip()
    username=request.form.get("username", "").strip().lower()
    password=request.form.get("password", "")
    role=request.form.get("role", "Teacher")
    if role in {"Student", "Parent", "System"}:
        flash("Student and parent records are not staff accounts. Use the administrator-only student intake.", "warning")
        return redirect(request.referrer or url_for("admin_dashboard"))
    title=request.form.get("title", "").strip()
    leadership_role=request.form.get("leadership_role", "").strip()
    department=request.form.get("department", "").strip()
    workspace_type=request.form.get("workspace_type", "Teaching").strip() or "Teaching"
    if workspace_type not in {"Teaching","Driver","Reception","Guard","Cook","Other Staff"}: workspace_type="Teaching"
    student_id=request.form.get("student_id") or None
    allowed=set(ALL_PORTAL_ROLES) - {SYSTEM_ROLE, "Student", "Parent"}
    # ICT is a technical operator, not a privilege escalator.
    if role not in allowed or (actor["role"]=="ICT" and role in {"Admin","ICT"}):
        flash("This account type cannot be created by your role.", "danger")
        return redirect(request.referrer or url_for("admin_dashboard"))
    if not full_name or (role != "Student" and (not username or len(password)<4)):
        flash("Name is required. Staff/portal accounts need a username and password of at least 4 characters.", "danger")
        return redirect(request.referrer or url_for("admin_dashboard"))
    student_id=None
    if q("SELECT id FROM users WHERE lower(username)=? LIMIT 1", (username,), one=True):
        flash("Username already exists. Choose a different username.", "danger")
        return redirect(request.referrer or url_for("admin_dashboard"))
    try:
        school_unit=(request.form.get("school_unit","").strip() or school_settings()["school_name"])
        school_location=request.form.get("school_location","").strip()
        position_code=staff_code_for(role, workspace_type) if role not in {"Student","Parent","System"} else ""
        reception_enabled=1 if workspace_type==RECEPTION_WORKSPACE else 0
        uid=execute("""INSERT INTO users(full_name, username, password_hash, role, student_id, active, title, department, phone, email, date_of_birth, gender, id_reference, address, emergency_contact, blood_group, medical_notes, accountability_notes, workspace_type, school_unit, school_location, reception_enabled, position_code, staff_code, qr_access_token, qr_login_enabled)
                      VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, lower(hex(randomblob(16))), 0)""",
                   (full_name,username,generate_password_hash(password),role,student_id,title,department,request.form.get("phone","").strip(),request.form.get("email","").strip(),request.form.get("date_of_birth","").strip(),request.form.get("gender","").strip(),request.form.get("id_reference","").strip(),request.form.get("address","").strip(),request.form.get("emergency_contact","").strip(),request.form.get("blood_group","").strip(),request.form.get("medical_notes","").strip(),request.form.get("accountability_notes","").strip(),workspace_type,school_unit,school_location,reception_enabled,position_code,position_code))
        if leadership_role and role not in {"Student","Parent","System"}:
            execute("UPDATE users SET leadership_role=?,leadership_level=? WHERE id=?",(leadership_role,1 if leadership_role in {"Dean","Deputy","Deputy Principal","HOD","Head of Department"} else 0,uid))
        audit(actor["id"],actor["full_name"],"Add User",f"{full_name} ({username}) added as {role}; title={title or '—'}; department={department or '—'}.")
    except sqlite3.IntegrityError:
        flash("Username already exists or the supplied learner link is invalid.", "danger")
        return redirect(request.referrer or url_for("admin_dashboard"))
    flash("Account created successfully.", "success")
    return redirect(request.referrer or url_for("admin_dashboard"))

@app.route("/admin/access/<int:user_id>")
@login_required
@role_required("Admin")
def admin_access_user(user_id: int):
    """Open any active institutional account directly from the Admin console.

    The selected user's password is neither requested nor changed. Access is
    represented by a short-lived, revocable portal-context token and the
    original Admin session is preserved so the administrator can keep the
    command centre open separately.
    """
    user = q("SELECT id, full_name, username, role, active FROM users WHERE id=? AND role!='System'", (user_id,), one=True)
    if not user:
        abort(404)
    if not user["active"]:
        flash("That account is archived. Restore it before using direct access.", "warning")
        return redirect(url_for("admin_dashboard"))
    actor = current_user()
    session["admin_impersonation"] = True
    context = _portal_context_for(user_id)
    audit(actor["id"], actor["full_name"], "Admin Direct Access", f"Opened {user['full_name']} ({user['username']}) as {user['role']} without using the account password.")
    target = specialized_dashboard_for(user)
    return redirect(target + "?" + urllib.parse.urlencode({"portal_context": context}))


@app.route("/users/<int:user_id>/edit", methods=["GET","POST"])
@login_required
def edit_user(user_id:int):
    actor=current_user()
    if actor["role"] not in {"Admin","ICT"}: abort(403)
    user=q("SELECT * FROM users WHERE id=?",(user_id,),one=True)
    if not user or user["role"] in {SYSTEM_ROLE, "Student"}: abort(404)
    if actor["role"]=="ICT" and user["role"] in {"Admin","ICT"}: abort(403)
    if request.method=="POST":
        role=request.form.get("role",user["role"])
        if user["id"] == actor["id"] and role != actor["role"]:
            flash("The currently signed-in Administrator account cannot be changed into another role. Create or edit another account instead.", "warning")
            return redirect(url_for("edit_user", user_id=user_id))
        workspace_type=request.form.get("workspace_type", user["workspace_type"] if "workspace_type" in user.keys() else "Teaching").strip() or "Teaching"
        if workspace_type not in {"Teaching","Driver","Reception","Guard","Cook","Other Staff"}: workspace_type="Teaching"
        if actor["role"]=="ICT" and role in {"Admin","ICT"}: abort(403)
        if role not in set(ALL_PORTAL_ROLES)-{SYSTEM_ROLE, "Student"}: abort(400)
        student_id=request.form.get("student_id") or None
        if role != "Parent":
            student_id=None
        elif student_id:
            try: student_id=int(student_id)
            except ValueError: student_id=None
        else:
            student_id=None
        new_username = request.form.get("username","").strip().lower()
        conflict = q("SELECT id FROM users WHERE lower(username)=? AND id!=? LIMIT 1", (new_username, user_id), one=True)
        if conflict:
            flash("Username already exists. Choose a different username.", "danger")
            return redirect(url_for("edit_user", user_id=user_id))
        new_title=request.form.get("title","").strip(); leadership_role=request.form.get("leadership_role","").strip(); leadership_level=1 if leadership_role in {"Dean","Deputy","Deputy Principal","HOD","Head of Department"} else 0; school_unit=request.form.get("school_unit","").strip() or school_settings()["school_name"]; school_location=request.form.get("school_location","").strip(); reception_enabled=1 if workspace_type==RECEPTION_WORKSPACE else 0
        existing_code=user["position_code"] or user["staff_code"] or (staff_code_for(role,workspace_type) if role not in {"Student","Parent","System"} else "")
        execute("""UPDATE users SET full_name=?, username=?, role=?, student_id=?, title=?, department=?, phone=?, email=?, date_of_birth=?, gender=?, id_reference=?, address=?, emergency_contact=?, blood_group=?, medical_notes=?, accountability_notes=?, workspace_type=?, school_unit=?, school_location=?, leadership_role=?, leadership_level=?, reception_enabled=?, position_code=?, staff_code=? WHERE id=?""",
               (request.form.get("full_name","").strip(),new_username,role,student_id,new_title,request.form.get("department","").strip(),request.form.get("phone","").strip(),request.form.get("email","").strip(),request.form.get("date_of_birth","").strip(),request.form.get("gender","").strip(),request.form.get("id_reference","").strip(),request.form.get("address","").strip(),request.form.get("emergency_contact","").strip(),request.form.get("blood_group","").strip(),request.form.get("medical_notes","").strip(),request.form.get("accountability_notes","").strip(),workspace_type,school_unit,school_location,leadership_role,leadership_level,reception_enabled,existing_code,existing_code,user_id))
        if role=="Parent" and student_id:
            execute("INSERT OR IGNORE INTO guardian_links(guardian_user_id,student_id,relationship,is_primary) VALUES(?,?,?,?)",(user_id,student_id,request.form.get("relationship","Guardian").strip() or "Guardian",1))
        audit(actor["id"],actor["full_name"],"Edit User",f"Updated {user['username']} ({user['role']}) -> {request.form.get('username','').strip()} ({role}).")
        flash("Person profile updated.","success")
        return redirect(url_for("admin_dashboard"))
    students=q("SELECT id, full_name, admission_no FROM students WHERE active=1 ORDER BY full_name")
    depts=q("SELECT name FROM departments WHERE active=1 ORDER BY name")
    return render_template("user_edit.html", user=user, students=students, departments=depts, role_options=tuple(r for r in ALL_PORTAL_ROLES if r != "Student"), guardian_links=q("SELECT * FROM guardian_links WHERE guardian_user_id=? AND active=1",(user_id,)))

@app.route("/users/<int:user_id>/reset-password", methods=["POST"])
@login_required
def reset_user_password(user_id:int):
    actor=current_user()
    if actor["role"]!="Admin": abort(403)
    user=q("SELECT * FROM users WHERE id=?",(user_id,),one=True)
    if not user or user["role"]==SYSTEM_ROLE: abort(404)
    password=request.form.get("password","")
    if len(password)<4:
        flash("Temporary password must be at least 4 characters.","danger")
        return redirect(url_for("admin_dashboard"))
    execute("UPDATE users SET password_hash=? WHERE id=?",(generate_password_hash(password),user_id))
    audit(actor["id"],actor["full_name"],"Password Reset",f"Administrator reset the password for {user['username']}.")
    flash("Password reset successfully. Give the temporary password directly to the account holder.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/users/<int:user_id>/archive", methods=["POST"])
@login_required
def archive_user(user_id:int):
    actor=current_user()
    if actor["role"]!="Admin": abort(403)
    user=q("SELECT * FROM users WHERE id=?",(user_id,),one=True)
    if not user or user["role"]==SYSTEM_ROLE: abort(404)
    if user["id"]==actor["id"]:
        flash("You cannot archive your own Administrator account.","warning")
        return redirect(url_for("admin_dashboard"))
    if user["role"]=="Admin" and q("SELECT COUNT(*) AS c FROM users WHERE role='Admin' AND active=1",one=True)["c"]<=1:
        flash("The last active Administrator cannot be archived.","warning")
        return redirect(url_for("admin_dashboard"))
    execute("UPDATE users SET active=0, archived_at=CURRENT_TIMESTAMP WHERE id=?",(user_id,))
    audit(actor["id"],actor["full_name"],"Archive User",f"{user['username']} ({user['role']}) archived; historical records retained.")
    flash("Account archived. Historical records remain intact.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/users/<int:user_id>/restore", methods=["POST"])
@login_required
def restore_user(user_id:int):
    actor=current_user()
    if actor["role"]!="Admin": abort(403)
    user=q("SELECT * FROM users WHERE id=?",(user_id,),one=True)
    if not user or user["role"]==SYSTEM_ROLE: abort(404)
    execute("UPDATE users SET active=1, archived_at=NULL WHERE id=?",(user_id,))
    audit(actor["id"],actor["full_name"],"Restore User",f"{user['username']} ({user['role']}) restored.")
    flash("Account restored.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/users/<int:user_id>/guardian-link", methods=["POST"])
@login_required
def guardian_link(user_id:int):
    actor=current_user()
    if actor["role"] not in {"Admin","ICT"}: abort(403)
    user=q("SELECT * FROM users WHERE id=?",(user_id,),one=True)
    if not user or user["role"]!="Parent": abort(404)
    student_id=request.form.get("student_id")
    try: student_id=int(student_id)
    except (TypeError,ValueError): abort(400)
    execute("INSERT OR REPLACE INTO guardian_links(guardian_user_id,student_id,relationship,is_primary,notes,active) VALUES(?,?,?,?,?,1)",(user_id,student_id,request.form.get("relationship","Guardian").strip() or "Guardian",1,request.form.get("notes","").strip()))
    execute("UPDATE users SET student_id=? WHERE id=? AND student_id IS NULL",(student_id,user_id))
    audit(actor["id"],actor["full_name"],"Guardian Link",f"Linked {user['username']} to student {student_id}.")
    flash("Guardian link saved.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/users/<int:user_id>/delete", methods=["POST"])
@login_required
def delete_user(user_id: int):
    actor=current_user()
    if actor["role"] != "Admin": abort(403)
    user=q("SELECT * FROM users WHERE id=? AND role!=?",(user_id,SYSTEM_ROLE),one=True)
    if not user: abort(404)
    if user["id"] == actor["id"]:
        flash("You cannot delete your own Administrator account.","warning")
        return redirect(url_for("admin_dashboard"))
    if user["role"] == "Admin" and q("SELECT COUNT(*) AS c FROM users WHERE role='Admin' AND active=1",one=True)["c"] <= 1:
        flash("The last active Administrator cannot be deleted.","warning")
        return redirect(url_for("admin_dashboard"))
    try:
        conn=get_db()
        # Preserve nullable historical references; remove dependent records that
        # cannot survive an account deletion. SQLite then performs normal CASCADEs.
        refs=[]
        tables=conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchall()
        for row in tables:
            table=row[0]
            for fk in conn.execute(f"PRAGMA foreign_key_list([{table}])").fetchall():
                if fk[2] != "users": continue
                refs.append((table,fk[3],fk[6]))
        for table,col,on_delete in refs:
            if table in {"users"}: continue
            try:
                notnull=next((r[3] for r in conn.execute(f"PRAGMA table_info([{table}])").fetchall() if r[1]==col),0)
                if str(on_delete).upper() in {"SET NULL","CASCADE"}:
                    continue
                # RESTRICT/NO ACTION on a non-nullable child must be removed.
                conn.execute(f"DELETE FROM [{table}] WHERE [{col}]=?",(user_id,))
            except Exception:
                pass
        conn.execute("DELETE FROM users WHERE id=?",(user_id,))
        conn.commit()
    except Exception as exc:
        try: conn.rollback()
        except Exception: pass
        flash(f"The account could not be deleted safely: {exc}","danger")
        return redirect(url_for("admin_dashboard"))
    audit(actor["id"],actor["full_name"],"Delete User",f"Permanently deleted {user['username']} ({user['role']}).")
    flash("Account permanently deleted.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/export/<kind>/<fmt>")
@login_required
@role_required("Admin","ICT")
def export_formatted(kind,fmt):
    if kind not in {"students","employees","attendance"} or fmt.lower() not in {"csv","pdf","docx"}: abort(404)
    fmt=fmt.lower()
    if kind=="students":
        headers=["Name","Admission No.","Class","Status"]; rows=[[r["full_name"],r["admission_no"],r["grade"],"Active" if r["active"] else "Archived"] for r in q("SELECT full_name,admission_no,grade,active FROM students ORDER BY grade,full_name")]
    elif kind=="employees":
        headers=["Name","Role","Department","Status"]; rows=[[r["full_name"],r["role"],r["department"],"Active" if r["active"] else "Archived"] for r in q("SELECT full_name,role,department,active FROM users WHERE role!='System' ORDER BY role,full_name")]
    else:
        headers=["Name","Role","Date","Action","Time","Location"]; rows=[]
        for r in q("SELECT u.full_name,u.role,a.action,a.event_at,a.latitude,a.longitude FROM attendance_events a JOIN users u ON u.id=a.user_id ORDER BY a.event_at DESC LIMIT 5000"):
            dt=str(r["event_at"] or ""); loc=f"{float(r['latitude']):.6f}, {float(r['longitude']):.6f}" if r["latitude"] is not None and r["longitude"] is not None else ""; rows.append([r["full_name"],r["role"],dt[:10],r["action"],dt[11:19],loc])
    filename=f"prime-{kind}-export-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}"
    if fmt=="csv":
        out=io.StringIO(); csv.writer(out).writerows([headers,*rows]); return send_file(io.BytesIO(out.getvalue().encode('utf-8-sig')),mimetype='text/csv',as_attachment=True,download_name=filename+'.csv')
    if fmt=="docx":
        from docx import Document
        doc=Document(); doc.add_heading(f"{school_settings()['school_name']} — {kind.title()}",0); table=doc.add_table(rows=1,cols=len(headers))
        for i,h in enumerate(headers): table.rows[0].cells[i].text=str(h)
        for row in rows:
            cells=table.add_row().cells
            for i,val in enumerate(row): cells[i].text=str(val)
        bio=io.BytesIO(); doc.save(bio); bio.seek(0); return send_file(bio,mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',as_attachment=True,download_name=filename+'.docx')
    from reportlab.lib.pagesizes import A4,landscape
    from reportlab.platypus import SimpleDocTemplate,Table,TableStyle,Paragraph
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    bio=io.BytesIO(); doc=SimpleDocTemplate(bio,pagesize=landscape(A4),leftMargin=18,rightMargin=18,topMargin=22,bottomMargin=22); styles=getSampleStyleSheet(); t=Table([headers]+[[str(v) for v in row] for row in rows],repeatRows=1); t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#1f2937')),('TEXTCOLOR',(0,0),(-1,0),colors.white),('GRID',(0,0),(-1,-1),.25,colors.HexColor('#cbd5e1')),('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),7)])); doc.build([Paragraph(f"{school_settings()['school_name']} — {kind.title()}",styles['Title']),t]); bio.seek(0); return send_file(bio,mimetype='application/pdf',as_attachment=True,download_name=filename+'.pdf')

@app.route("/export/<kind>")
@login_required
@role_required("Admin")
def export_data(kind: str):
    mapping = {
        "students": (
            "Student Export",
            [
                "admission_no", "full_name", "grade", "guardian_name", "guardian_phone", "guardian_email",
                "alt_guardian_name", "alt_guardian_phone", "alt_guardian_email", "student_phone", "student_email",
                "medical_condition", "allergies", "special_info", "notes", "payment_status", "balance", "active",
            ],
            "SELECT admission_no, full_name, grade, guardian_name, guardian_phone, guardian_email, alt_guardian_name, alt_guardian_phone, alt_guardian_email, student_phone, student_email, medical_condition, allergies, special_info, notes, payment_status, balance, active FROM students ORDER BY id",
        ),
        "users": (
            "Employee Export",
            ["full_name", "username", "role", "created_at"],
            "SELECT full_name, username, role, created_at FROM users ORDER BY id",
        ),
        "payments": (
            "Payments Export",
            ["student_id", "amount", "method", "reference_no", "recorded_by", "status", "created_at"],
            "SELECT student_id, amount, method, reference_no, recorded_by, status, created_at FROM payments ORDER BY id",
        ),
        "audit": (
            "Audit Export",
            ["actor_name", "action", "details", "created_at"],
            "SELECT actor_name, action, details, created_at FROM audit_log ORDER BY id",
        ),
    }
    if kind not in mapping:
        abort(404)
    title, headers, sql = mapping[kind]
    rows = q(sql)
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([row[h] for h in headers])
    data = io.BytesIO(buffer.getvalue().encode("utf-8-sig"))
    data.seek(0)
    filename = f"{kind}_export.csv"
    audit(current_user()["id"], current_user()["full_name"], "Export Data", f"{title} downloaded.")
    return send_file(data, as_attachment=True, download_name=filename, mimetype="text/csv")



@app.route("/system-help")
@login_required
def system_help():
    user=current_user(); settings=school_settings(); rows=q("SELECT * FROM system_help WHERE active=1 ORDER BY category,sort_order,title")
    visible=[]
    for row in rows:
        scope=(row["role_scope"] or "All").split(",")
        if "All" in scope or user["role"] in [x.strip() for x in scope]: visible.append(row)
    role_guides={
        "Admin": settings["institution_admin_guide"],
        "ICT": settings["institution_ict_guide"],
        "Finance": settings["institution_finance_guide"],
        "Teacher": settings["institution_portal_guide"],
        "Student": settings["institution_portal_guide"],
        "Parent": settings["institution_portal_guide"],
    }
    return render_template("system_help.html", role=user["role"], workspace=workspace_for(user["role"]), help_rows=visible, settings=settings, actor_name=user["full_name"], role_guide=role_guides.get(user["role"], settings["institution_portal_guide"]))

@app.route("/ai-assistant")
@login_required
def ai_assistant():
    return render_template("ai_assistant.html", role=current_user()["role"], workspace=workspace_for(current_user()["role"]), settings=school_settings(), actor_name=current_user()["full_name"])

@app.route("/ai/ask", methods=["POST"])
@login_required
def ai_ask():
    settings=school_settings(); user=current_user()
    if not int(settings["ai_enabled"] or 0): return jsonify({"error":"AI assistance is disabled by the administrator."}),403
    prompt=(request.form.get("prompt") or (request.get_json(silent=True) or {}).get("prompt") or "").strip()
    if not prompt: return jsonify({"error":"Enter a question."}),400
    if len(prompt)>8000: return jsonify({"error":"Question is too long."}),400
    provider=settings["ai_provider"] or "openai_responses"; model=settings["ai_model"] or "gpt-5.6"
    api_key=os.environ.get("OPENAI_API_KEY","").strip()
    if not api_key: return jsonify({"error":"OpenAI API is not configured on this server yet. Set OPENAI_API_KEY in the deployment environment."}),503
    system_prompt=f"You are the institutional AI assistant for {settings['school_name']}. User role: {user['role']}. Be practical, concise, safe, and never invent private institutional data. If data is not supplied in the conversation, say so."
    try:
        if provider=="openai_chat_completions":
            payload=json.dumps({"model":model,"messages":[{"role":"system","content":system_prompt},{"role":"user","content":prompt}],"temperature":0.2}).encode()
            req=urllib.request.Request("https://api.openai.com/v1/chat/completions",data=payload,headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"},method="POST")
            with urllib.request.urlopen(req,timeout=45) as resp: data=json.loads(resp.read().decode())
            answer=data.get("choices",[{}])[0].get("message",{}).get("content","")
        else:
            payload=json.dumps({"model":model,"input":[{"role":"system","content":system_prompt},{"role":"user","content":prompt}]}).encode()
            req=urllib.request.Request("https://api.openai.com/v1/responses",data=payload,headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"},method="POST")
            with urllib.request.urlopen(req,timeout=45) as resp: data=json.loads(resp.read().decode())
            answer=(data.get("output_text") or "").strip()
            if not answer:
                chunks=[]
                for item in data.get("output",[]) or []:
                    for content in item.get("content",[]) or []:
                        if content.get("type") in {"output_text","text"}: chunks.append(content.get("text","") or content.get("output_text",""))
                answer="\n".join(chunks).strip()
        if not answer: raise RuntimeError("AI returned an empty response.")
        execute("INSERT INTO ai_usage_log(user_id,role,provider,model,prompt_preview,response_preview,status) VALUES(?,?,?,?,?,?,?)",(user["id"],user["role"],provider,model,prompt[:300],answer[:500],"Success"))
        return jsonify({"answer":answer,"provider":provider,"model":model})
    except Exception as exc:
        execute("INSERT INTO ai_usage_log(user_id,role,provider,model,prompt_preview,response_preview,status) VALUES(?,?,?,?,?,?,?)",(user["id"],user["role"],provider,model,prompt[:300],str(exc)[:500],"Failed"))
        return jsonify({"error":f"AI request failed: {exc}"}),502

@app.route("/admin/ai-settings", methods=["POST"])
@login_required
@role_required("Admin")
def admin_ai_settings():
    enabled=1 if request.form.get("ai_enabled") else 0
    provider=request.form.get("ai_provider","openai_responses")
    if provider not in {"openai_responses","openai_chat_completions"}: provider="openai_responses"
    model=request.form.get("ai_model","gpt-5.6").strip() or "gpt-5.6"
    execute("UPDATE school_settings SET ai_enabled=?, ai_provider=?, ai_model=? WHERE id=1",(enabled,provider,model))
    audit(current_user()["id"],current_user()["full_name"],"AI Settings",f"AI {'enabled' if enabled else 'disabled'} using {provider} / {model}.")
    flash("AI configuration saved. API keys remain server-side only.","success")
    return redirect(url_for("admin_dashboard"))

@app.route("/admin/help/save", methods=["POST"])
@login_required
@role_required("Admin")
def admin_help_save():
    title=request.form.get("title","").strip(); category=request.form.get("category","Getting Started").strip(); content=request.form.get("content","").strip(); scope=request.form.get("role_scope","All").strip() or "All"
    if not title or not content: flash("Help title and content are required.","danger"); return redirect(url_for("admin_dashboard"))
    execute("INSERT INTO system_help(title,category,content,role_scope,sort_order) VALUES(?,?,?,?,?)",(title,category,content,scope,100))
    audit(current_user()["id"],current_user()["full_name"],"System Help",f"Added help article: {title}.")
    flash("System help article published.","success"); return redirect(url_for("admin_dashboard"))

@app.route("/admin/learners")
@login_required
@role_required("Admin","ICT")
def all_learners():
    rows=q("SELECT * FROM students ORDER BY active DESC,grade,full_name")
    guardians=q("""SELECT gl.student_id, GROUP_CONCAT(gu.full_name || ' (' || gl.relationship || ')', ', ') AS guardians FROM guardian_links gl JOIN users gu ON gu.id=gl.guardian_user_id WHERE gl.active=1 GROUP BY gl.student_id""")
    guardian_map={r["student_id"]:r["guardians"] for r in guardians}
    return render_template("directory.html", directory_type="Learners", rows=rows, settings=school_settings(), role=current_user()["role"], actor_name=current_user()["full_name"], guardian_map=guardian_map)

@app.route("/admin/employees")
@login_required
@role_required("Admin","ICT")
def all_employees():
    today=(datetime.utcnow()+KENYA_TZ_OFFSET).date().isoformat()
    start_utc,end_utc=attendance_day_bounds_utc(today)
    rows=q("""
      SELECT u.*,
        (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS check_in_at,
        (SELECT a.event_at FROM attendance_events a WHERE a.user_id=u.id AND a.action='OUT' AND a.event_at>=? AND a.event_at<? ORDER BY a.event_at DESC,a.id DESC LIMIT 1) AS check_out_at,
        (SELECT a.location_label FROM attendance_events a WHERE a.user_id=u.id AND a.action='IN' AND a.event_at>=? AND a.event_at<? AND a.location_label!='' ORDER BY a.event_at ASC,a.id ASC LIMIT 1) AS check_in_location
      FROM users u WHERE u.role NOT IN ('Student','Parent','System') ORDER BY u.active DESC, u.full_name
    """,(start_utc,end_utc,start_utc,end_utc,start_utc,end_utc))
    rows=[dict(r) for r in rows]
    for r in rows:
        r['check_in_local']=_local_iso(_parse_stored_event(r.get('check_in_at'))) if r.get('check_in_at') else None
        r['check_out_local']=_local_iso(_parse_stored_event(r.get('check_out_at'))) if r.get('check_out_at') else None
    return render_template("directory.html", directory_type="Employees", rows=rows, settings=school_settings(), role=current_user()["role"], actor_name=current_user()["full_name"], guardian_map={}, today=today, add_mode=bool(request.args.get("add")), all_roles=ALL_PORTAL_ROLES, departments=q("SELECT id,name,category FROM departments WHERE active=1 ORDER BY name"))

@app.route("/users/<int:user_id>/qr")
@login_required
def user_qr(user_id:int):
    user=q("SELECT * FROM users WHERE id=? AND role!='System'",(user_id,),one=True)
    if not user: abort(404)
    if user["role"] in {"Student","Parent","System"}: abort(404)
    if current_user()["role"] not in {"Admin","ICT"} and current_user()["id"]!=user_id: abort(403)
    token=user["qr_access_token"]
    if not token:
        token=uuid.uuid4().hex; execute("UPDATE users SET qr_access_token=? WHERE id=?",(token,user_id))
    payload=url_for("qr_landing", token=token, _external=True)
    qr=qrcode.QRCode(version=1,box_size=8,border=3); qr.add_data(payload); qr.make(fit=True)
    buf=io.BytesIO(); qr.make_image().save(buf,format="PNG"); buf.seek(0)
    return send_file(buf,mimetype="image/png",download_name=f"{secure_filename(user['username'])}-portal-qr.png")

@app.route("/users/<int:user_id>/attendance-qr")
@login_required
def staff_attendance_qr(user_id:int):
    actor=current_user()
    if actor['role'] not in {'Admin','ICT'} and not is_reception_user(actor): abort(403)
    user=q("SELECT * FROM users WHERE id=? AND active=1 AND role NOT IN ('System','Admin')",(user_id,),one=True)
    if not user: abort(404)
    token=user['qr_access_token'] or uuid.uuid4().hex
    execute("UPDATE users SET qr_access_token=? WHERE id=?",(token,user_id))
    payload='STAFF|'+token+'|'+(user['full_name'] or '')+'|'+(user['title'] or user['role'])+'|'+(user['position_code'] or user['staff_code'] or '')
    code=qrcode.QRCode(version=3,box_size=9,border=3); code.add_data(payload); code.make(fit=True); buf=io.BytesIO(); code.make_image().save(buf,format='PNG'); buf.seek(0)
    return send_file(buf,mimetype='image/png',download_name=f"{secure_filename(user['full_name'])}-attendance-qr.png",as_attachment=False)

@app.route("/qr/<token>")
def qr_landing(token:str):
    user=q("SELECT * FROM users WHERE qr_access_token=? AND active=1 AND role NOT IN ('Student','Parent','System')",(token,),one=True)
    if not user: return render_template('error.html',message='This QR code is not linked to an active staff account.'),404
    if not qr_login_allowed(user): return render_template('error.html',message='QR sign-in is not enabled for this staff account yet. Use your normal username and password first.'),403
    return render_template('qr_staff_landing.html',user_name=user['full_name'],token=token,dashboard_url=specialized_dashboard_for(user),next_action=next_attendance_action(user['id']))

@app.route("/qr/<token>/consume",methods=['POST'])
def qr_consume(token:str):
    user=q("SELECT * FROM users WHERE qr_access_token=? AND active=1 AND role NOT IN ('Student','Parent','System')",(token,),one=True)
    if not user or not qr_login_allowed(user): return jsonify({'ok':False,'message':'This staff QR is no longer valid.'}),403
    payload=request.get_json(silent=True) or request.form
    action=str(payload.get('action') or next_attendance_action(user['id'])).upper()
    result=record_account_attendance(user,action,payload.get('event_at') or None,payload.get('source','online'),'QR',_payload_float(payload,'latitude'),_payload_float(payload,'longitude'),_payload_float(payload,'accuracy'),payload.get('device_note',''),payload.get('location_label',''))
    if not result['ok']: return jsonify(result),409
    login_id=record_login_event(user,'QR',_payload_float(payload,'latitude'),_payload_float(payload,'longitude'),_payload_float(payload,'accuracy'))
    session.clear(); session.permanent=True
    session['user_id']=user['id']; session['active_portal_role']=user['role']; session['login_event_id']=login_id; session['login_location_pending']=1
    audit(user['id'],user['full_name'],'QR Login + Attendance',f'{user["full_name"]} signed in with personal staff QR and was marked {action}.')
    return jsonify({'ok':True,'message':result['message'],'dashboard':result['dashboard'],'action':action})

@app.route('/qr/<token>/offline-dashboard')
def qr_offline_dashboard(token:str):
    user=q("SELECT full_name,role,title,department,workspace_type FROM users WHERE qr_access_token=? AND active=1 AND role NOT IN ('Student','Parent','System')",(token,),one=True)
    if not user: abort(404)
    return render_template('qr_offline_dashboard.html',settings=school_settings(),user=user,dashboard_url=specialized_dashboard_for(user))

@app.route("/qr/offline-sync",methods=['POST'])
def qr_offline_sync():
    payload=request.get_json(silent=True) or {}; token=str(payload.get('token') or '').strip()
    user=q("SELECT * FROM users WHERE qr_access_token=? AND active=1 AND role NOT IN ('Student','Parent','System')",(token,),one=True)
    if not user or not qr_login_allowed(user): return jsonify({'ok':False,'message':'Invalid staff QR.'}),403
    saved=0
    for item in (payload.get('events') or [])[:20]:
        result=record_account_attendance(user,str(item.get('action') or next_attendance_action(user['id'])).upper(),item.get('event_at'),'offline-sync','QR',item.get('latitude'),item.get('longitude'),item.get('accuracy'),item.get('device_note','offline qr'),item.get('location_label',''))
        if result.get('ok'): saved+=1
    return jsonify({'ok':True,'saved':saved,'dashboard':specialized_dashboard_for(user)})

@app.route("/finance/ledger", methods=["POST"])
@login_required
@role_required("Finance","Admin")
def finance_post_ledger():
    entry_type=request.form.get("entry_type","Expense"); category=request.form.get("category","General").strip() or "General"; amount=float(request.form.get("amount",0) or 0)
    description=request.form.get("description","").strip(); reference=request.form.get("reference_no","").strip(); payee=request.form.get("payee_user_id",type=int) or None
    if entry_type not in {"Income","Expense","Payroll","Adjustment"} or amount<=0 or not description:
        flash("Enter a valid ledger transaction.","danger"); return redirect(url_for("finance_dashboard"))
    if entry_type=="Payroll" and (not payee or not q("SELECT id FROM users WHERE id=? AND active=1 AND role NOT IN ('Student','Parent','System')",(payee,),one=True)):
        flash("Select a valid active employee for payroll.","danger"); return redirect(url_for("finance_dashboard"))
    receipt_path=""
    file=request.files.get("receipt")
    if file and file.filename:
        ext=file.filename.rsplit('.',1)[-1].lower() if '.' in file.filename else ''
        if ext in {"pdf","png","jpg","jpeg","webp"}:
            rd=UPLOAD_DIR/"finance_receipts"; rd.mkdir(exist_ok=True)
            name=f"ledger-{datetime.now().strftime('%Y%m%d%H%M%S%f')}.{ext}"; file.save(rd/name); receipt_path="uploads/finance_receipts/"+name
    execute("INSERT INTO finance_ledger(entry_type,category,payee_user_id,amount,description,reference_no,posted_by,receipt_path) VALUES(?,?,?,?,?,?,?,?)",(entry_type,category,payee,amount,description,reference,current_user()["id"],receipt_path))
    audit(current_user()["id"],current_user()["full_name"],"Finance Transaction",f"Posted {entry_type} of {amount:.2f} ({reference or 'no reference'}).")
    flash("Transaction posted and locked. Only an Administrator can reverse it.","success"); return redirect(url_for("finance_dashboard"))

@app.route("/finance/ledger/<int:entry_id>/reverse", methods=["POST"])
@login_required
@role_required("Admin")
def finance_reverse_ledger(entry_id:int):
    row=q("SELECT * FROM finance_ledger WHERE id=?",(entry_id,),one=True)
    if not row: abort(404)
    if row["status"]=="Reversed": flash("Transaction is already reversed.","warning"); return redirect(url_for("finance_dashboard"))
    execute("UPDATE finance_ledger SET status='Reversed',reversed_by=?,reversed_at=CURRENT_TIMESTAMP WHERE id=? AND status='Posted'",(current_user()["id"],entry_id))
    audit(current_user()["id"],current_user()["full_name"],"Reverse Finance Transaction",f"Reversed ledger entry #{entry_id}.")
    flash("Transaction reversed by Administrator. Original record remains visible for audit.","success"); return redirect(url_for("finance_dashboard"))

@app.route("/admin/theme/restore/<snapshot_type>", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def restore_previous_theme(snapshot_type):
    if snapshot_type not in {"workspace","landing"}: abort(400)
    snap=q("SELECT * FROM theme_snapshots WHERE snapshot_type=? ORDER BY id DESC LIMIT 1",(snapshot_type,),one=True)
    if not snap:
        flash("There is no previous saved setting to restore.","warning")
        return redirect(request.referrer or url_for("ict_dashboard"))
    try: data=json.loads(snap["settings_json"])
    except Exception:
        flash("The previous setting could not be read.","danger")
        return redirect(request.referrer or url_for("ict_dashboard"))
    cols=set(table_columns(get_db(),"school_settings"))
    allowed={k:v for k,v in data.items() if k in cols}
    if not allowed:
        flash("No compatible setting was found in the restore point.","danger")
        return redirect(request.referrer or url_for("ict_dashboard"))
    sets=", ".join(f"{k}=?" for k in allowed)
    execute(f"UPDATE school_settings SET {sets} WHERE id=1",tuple(allowed.values()))
    execute("DELETE FROM theme_snapshots WHERE id=?",(snap["id"],))
    audit(current_user()["id"],current_user()["full_name"],"Theme Restore",f"Restored previous {snapshot_type} presentation settings.")
    flash(f"Previous {snapshot_type} presentation restored.","success")
    return redirect(request.referrer or url_for("ict_dashboard"))

@app.route("/admin/theme", methods=["POST"])
@login_required
@role_required("Admin", "ICT")
def admin_theme():
    return ict_settings()


def _backup_tables():
    names=[r[0] for r in get_db().execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name").fetchall()]
    return names

def _json_backup_payload(include_assets=True):
    conn=get_db(); tables={}
    for name in _backup_tables():
        rows=conn.execute(f"SELECT * FROM [{name}]").fetchall()
        tables[name]={"columns":list(rows[0].keys()) if rows else list(table_columns(conn,name)),"rows":[dict(r) for r in rows]}
    assets=[]
    if include_assets and UPLOAD_DIR.exists():
        for path in UPLOAD_DIR.rglob('*'):
            if not path.is_file() or path.name.endswith('.bak'): continue
            try:
                data=path.read_bytes()
                if len(data)<=30*1024*1024:
                    assets.append({"path":str(path.relative_to(DATA_DIR)).replace('\\','/'),"mime":mimetypes.guess_type(str(path))[0] or 'application/octet-stream',"data_base64":base64.b64encode(data).decode('ascii')})
            except OSError:
                continue
    settings=dict(q("SELECT * FROM school_settings WHERE id=1",one=True) or {})
    return {"format":"Prime Institution OS Full System JSON","version":2,"created_at":datetime.utcnow().isoformat(timespec='seconds')+'Z',"settings":settings,"tables":tables,"assets":assets,"notes":"Restore through Administration > Backup & Recovery. Password hashes are preserved; plaintext passwords and API keys are never exported."}

def admin_root_user():
    uid=flask_session.get("user_id")
    if not uid:
        return None
    return q("SELECT id,full_name,role,active FROM users WHERE id=? AND role='Admin' AND active=1 LIMIT 1", (uid,), one=True)

def admin_root_required(view):
    @wraps(view)
    def wrapper(*args, **kwargs):
        actor=admin_root_user()
        if not actor:
            abort(403)
        g.admin_actor=actor
        return view(*args, **kwargs)
    return wrapper

@app.route("/backup/json")
@login_required
@admin_root_required
def backup_json_download():
    payload=_json_backup_payload(include_assets=True)
    raw=json.dumps(payload,ensure_ascii=False,indent=2).encode('utf-8')
    name=f"prime-institution-full-backup-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}.json"
    execute("INSERT INTO backup_registry(backup_type,file_name,row_count,created_by) VALUES(?,?,?,?)",("JSON",name,sum(len(v.get('rows') or []) for v in payload.get('tables',{}).values() if isinstance(v,dict)),current_user()["id"]))
    return send_file(io.BytesIO(raw),mimetype='application/json',as_attachment=True,download_name=name)

@app.route("/backup/json/restore", methods=["POST"])
@login_required
@admin_root_required
def backup_json_restore():
    file=request.files.get('json_backup')
    if not file or not file.filename or not file.filename.lower().endswith('.json'):
        flash('Choose a .json full system backup first.','danger'); return redirect(request.referrer or url_for('admin_dashboard'))
    try: payload=json.load(file.stream)
    except Exception as exc:
        flash(f'Backup JSON could not be read: {exc}','danger'); return redirect(request.referrer or url_for('admin_dashboard'))
    if payload.get('format')!='Prime Institution OS Full System JSON' or not isinstance(payload.get('tables'),dict):
        flash('This file is not a valid Prime Institution OS backup.','danger'); return redirect(request.referrer or url_for('admin_dashboard'))
    conn=get_db(); old_autocommit=conn.isolation_level
    try:
        session['_restore_actor_username']=current_user()['username']
        conn.execute('PRAGMA foreign_keys=OFF')
        conn.execute('BEGIN')
        existing=_backup_tables()
        for table in existing:
            conn.execute(f'DELETE FROM [{table}]')
        for table,data in payload['tables'].items():
            if table.startswith('sqlite_') or table not in existing: continue
            columns=data.get('columns') or []
            for row in data.get('rows') or []:
                cols=[c for c in columns if c in table_columns(conn,table)]
                if not cols: continue
                vals=[row.get(c) for c in cols]
                placeholders=','.join('?' for _ in cols)
                conn.execute(f"INSERT INTO [{table}] ({','.join('['+c+']' for c in cols)}) VALUES ({placeholders})",vals)
        for asset in payload.get('assets') or []:
            rel=str(asset.get('path',''))
            if not rel.startswith('uploads/'): continue
            target=(DATA_DIR/rel).resolve(); root=UPLOAD_DIR.resolve()
            if root not in target.parents: continue
            target.parent.mkdir(parents=True,exist_ok=True)
            target.write_bytes(base64.b64decode(asset.get('data_base64','')))
        conn.commit()
        conn.execute('PRAGMA foreign_keys=ON')
        init_db()
        restored_actor=q('SELECT id,full_name FROM users WHERE username=? AND active=1 LIMIT 1',(session.get('_restore_actor_username',''),),one=True)
        if not restored_actor:
            restored_actor=q("SELECT id,full_name FROM users WHERE role='Admin' AND active=1 ORDER BY id LIMIT 1",one=True)
        if restored_actor:
            audit(restored_actor['id'],restored_actor['full_name'],'Restore JSON Backup',f"Full system JSON backup restored: {file.filename}.")
        flash('Full system JSON restored successfully. Database records, settings and included uploaded assets are back in place.','success')
    except Exception as exc:
        conn.rollback(); conn.execute('PRAGMA foreign_keys=ON'); flash(f'JSON restore failed safely: {exc}','danger')
    finally:
        conn.isolation_level=old_autocommit
    return redirect(request.referrer or url_for('admin_dashboard'))

@app.route("/backup/download")
@login_required
@admin_root_required
def backup_download():
    if not DB_PATH.exists():
        abort(404)
    return send_file(DB_PATH, as_attachment=True, download_name="school_backup.sqlite3")


@app.route("/backup/restore", methods=["POST"])
@login_required
@admin_root_required
def backup_restore():
    file=request.files.get('backup_file')
    if not file or not file.filename:
        flash('Choose a backup file first.','danger'); return redirect(request.referrer or url_for('admin_dashboard'))
    if file.filename.lower().endswith('.json'):
        return backup_json_restore()
    if not allowed_filename(file.filename):
        flash('Only .db, .sqlite or .sqlite3 database backups are allowed.','danger'); return redirect(request.referrer or url_for('admin_dashboard'))
    safe_name=secure_filename(file.filename)
    temp=UPLOAD_DIR/(safe_name+'.incoming')
    file.save(temp)
    old_db=DB_PATH.with_suffix('.pre-restore.bak')
    try:
        with sqlite3.connect(temp,timeout=30) as test:
            if (test.execute('PRAGMA integrity_check').fetchone() or [''])[0].lower()!='ok': raise ValueError('SQLite integrity check failed.')
            present={r[0] for r in test.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()}
            required={'users','students','school_settings','payments','exam_batches','exam_results'}
            if not required.issubset(present): raise ValueError('The selected database is missing required Prime tables.')
        # Never replace an SQLite file while a Flask request-scoped connection still owns it.
        db=g.pop('db',None)
        if db is not None: db.close()
        for suffix in ('-wal','-shm'):
            DB_PATH.with_name(DB_PATH.name+suffix).unlink(missing_ok=True)
        if old_db.exists(): old_db.unlink()
        if DB_PATH.exists(): DB_PATH.replace(old_db)
        temp.replace(DB_PATH)
        try:
            init_db()
            audit(current_user()['id'],current_user()['full_name'],'Restore Backup',f'Backup restored from {safe_name}.')
        except Exception:
            if DB_PATH.exists(): DB_PATH.unlink()
            if old_db.exists(): old_db.replace(DB_PATH)
            init_db()
            raise
        flash('Backup restored safely. The previous database was retained as a rollback copy.','success')
    except Exception as exc:
        temp.unlink(missing_ok=True)
        flash(f'Restore failed safely: {exc}','danger')
    return redirect(request.referrer or url_for('admin_dashboard'))


@app.route("/uploads/<path:filename>")
def uploaded_file(filename: str):
    # Public assets (logo, library, election and institution media) can be served.
    # Private finance/document artifacts must never be anonymously retrievable.
    rel_parts = [secure_filename(part) for part in Path(filename).parts if part not in {"", ".", ".."}]
    relative_name = "/".join(rel_parts)
    first_part = rel_parts[0] if rel_parts else ""
    private = first_part in {"documents", "qr", "receipts"} or (len(rel_parts) == 1 and rel_parts[0].startswith("receipt-"))
    if private and not current_user():
        abort(403)
    if private and current_user()["role"] in {"Student", "Parent"}:
        abort(403)
    if not rel_parts:
        abort(404)
    target = (UPLOAD_DIR / Path(*rel_parts)).resolve()
    upload_root = UPLOAD_DIR.resolve()
    if upload_root not in target.parents and target != upload_root:
        abort(403)
    if not target.exists() or not target.is_file():
        abort(404)
    return send_file(target)


@app.route("/api/student/<int:student_id>")
@login_required
def api_student(student_id: int):
    student = q("SELECT * FROM students WHERE id = ?", (student_id,), one=True)
    if not student:
        return jsonify({"error": "Not found"}), 404
    if not can_access_student(student_id):
        return jsonify({"error": "Forbidden"}), 403
    payments = q(
        """
        SELECT p.*, u.full_name AS recorded_by_name
        FROM payments p
        JOIN users u ON u.id = p.recorded_by
        WHERE p.student_id = ?
        ORDER BY p.created_at DESC, p.id DESC
        """,
        (student_id,),
    )
    return jsonify({"student": dict(student), "payments": [dict(p) for p in payments]})


@app.errorhandler(403)
def forbidden(_):
    return ("<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Access denied</title></head><body><h1>Access denied</h1><p>You do not have permission to access this area.</p><p><a href='/'>Return to portal</a></p></body></html>"), 403


@app.errorhandler(500)
def internal_error(error):
    app.logger.exception("Unhandled Prime application error: %s", error)
    # Keep the last-resort 500 response dependency-free. It must never render a
    # template that itself depends on static routes or another failing service.
    return ("<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Prime Portal Error</title>"
            "<style>body{font-family:system-ui,sans-serif;margin:0;padding:40px;background:#f4f7fb;color:#152033}main{max-width:720px;margin:auto;background:white;border-radius:18px;padding:28px;box-shadow:0 10px 30px rgba(0,0,0,.08)}"
            "a{display:inline-block;margin-top:16px;padding:11px 15px;border-radius:10px;background:#2457d6;color:#fff;text-decoration:none}</style></head>"
            "<body><main><h1>Something went wrong</h1><p>The portal could not complete that request. Your learner data is protected from partial writes.</p><a href='/'>Return to portal</a></main></body></html>"), 500


@app.errorhandler(404)
def not_found(_):
    return ("<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>Not found</title></head><body><h1>Page not found</h1><p>The page you requested could not be found.</p><p><a href='/'>Return to portal</a></p></body></html>"), 404


@app.errorhandler(413)
def too_large(_):
    return ("<!doctype html><html><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>File too large</title></head><body><h1>File too large</h1><p>The uploaded file is too large.</p><p><a href='/'>Return to portal</a></p></body></html>"), 413


with app.app_context():
    init_db()
    migrate_legacy_student_store()
    backfill_student_allocations()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
